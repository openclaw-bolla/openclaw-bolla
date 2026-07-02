#!/usr/bin/env python3
"""
AURORA Satz-Kur — Review-Dokument (Word, Aptos).

Baut aus allen data/satzkur/batch_*.json EIN übersichtliches .docx für Chris:
  Teil A: KONSISTENZ-FUNDE (brauchen Chris' Entscheidung) — pro Kapitel:
          Typ · Fund/Zitat · Problem · Vorschlag.
  Teil B: SATZ-KUR (Lesbarkeit) — pro Kapitel: Vorher → Nachher (+ kurzer Grund).

Format nach Chris-Präferenz: Aptos, Absatzabstand 0pt, Zeilenabstand 1,15, kompakt.
Ausgabe auf den OneDrive-Desktop.
"""
import json, os, glob, re, datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_LINE_SPACING

OUTDIR = "/home/bolla/workspace/data/satzkur"
DESKTOP = "/mnt/d/OneDrive/Desktop"
FONT = "Aptos"


def chapnum_key(nr):
    return (nr is None, nr if nr is not None else 999)


def style_para(p):
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.15


def run(p, text, bold=False, italic=False, size=11, color=None):
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = RGBColor(*color)
    return r


def load_all():
    rewrites = {}   # nr -> [ {original,revised,grund} ]
    konsist = {}    # nr -> [ {typ,fund,problem,vorschlag} ]
    nrw = nks = 0
    for path in sorted(glob.glob(f"{OUTDIR}/batch_*.json")):
        d = json.load(open(path, encoding="utf-8"))
        for kap in d.get("kapitel", []):
            nr = kap.get("nr")
            for r in kap.get("rewrites", []):
                rewrites.setdefault(nr, []).append(r); nrw += 1
        for ks in d.get("konsistenz", []):
            konsist.setdefault(ks.get("nr"), []).append(ks); nks += 1
    return rewrites, konsist, nrw, nks


def main():
    rewrites, konsist, nrw, nks = load_all()
    doc = Document()
    # Standardschrift
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(11)

    h = doc.add_paragraph(); style_para(h)
    run(h, "AURORA — Fable Satz-Kur & Konsistenz: Review", bold=True, size=16)
    d = doc.add_paragraph(); style_para(d)
    run(d, datetime.datetime.now().strftime("Stand %d.%m.%Y %H:%M") +
        f"  ·  {nks} Konsistenz-Funde  ·  {nrw} Satz-Umbauten", italic=True, size=10,
        color=(0x66, 0x66, 0x66))
    intro = doc.add_paragraph(); style_para(intro)
    run(intro, "Nichts ist bisher am Buch geändert. Sag mir, was raus soll — den Rest spiele ich "
               "mit Backup + Verifikation ein.", size=10, italic=True)

    # ---- TEIL A: KONSISTENZ ----
    a = doc.add_paragraph(); style_para(a)
    run(a, "Teil A — Konsistenz-Funde (deine Entscheidung)", bold=True, size=14,
        color=(0xB0, 0x30, 0x30))
    for nr in sorted(konsist, key=chapnum_key):
        kh = doc.add_paragraph(); style_para(kh)
        run(kh, f"Kapitel {nr}", bold=True, size=12)
        for ks in konsist[nr]:
            p = doc.add_paragraph(); style_para(p)
            p.paragraph_format.left_indent = Pt(12)
            run(p, f"[{ks.get('typ','?').upper()}] ", bold=True, size=10, color=(0xB0, 0x30, 0x30))
            run(p, "Fund: ", bold=True, size=10)
            run(p, ks.get("fund", "")[:400], size=10)
            p2 = doc.add_paragraph(); style_para(p2); p2.paragraph_format.left_indent = Pt(12)
            run(p2, "Problem: ", bold=True, size=10)
            run(p2, ks.get("problem", ""), size=10)
            p3 = doc.add_paragraph(); style_para(p3); p3.paragraph_format.left_indent = Pt(12)
            run(p3, "Vorschlag: ", bold=True, size=10, color=(0x1E, 0x7A, 0x1E))
            run(p3, ks.get("vorschlag", ""), size=10, color=(0x1E, 0x7A, 0x1E))
            sp = doc.add_paragraph(); style_para(sp)  # kleine Lücke

    # ---- TEIL B: SATZ-KUR ----
    doc.add_page_break()
    b = doc.add_paragraph(); style_para(b)
    run(b, "Teil B — Satz-Kur (Lesbarkeit: Vorher → Nachher)", bold=True, size=14,
        color=(0x30, 0x50, 0xB0))
    for nr in sorted(rewrites, key=chapnum_key):
        kh = doc.add_paragraph(); style_para(kh)
        run(kh, f"Kapitel {nr}  ({len(rewrites[nr])} Umbauten)", bold=True, size=12)
        for r in rewrites[nr]:
            pv = doc.add_paragraph(); style_para(pv); pv.paragraph_format.left_indent = Pt(12)
            run(pv, "Vorher: ", bold=True, size=10, color=(0x99, 0x55, 0x00))
            run(pv, r.get("original", ""), size=10)
            pn = doc.add_paragraph(); style_para(pn); pn.paragraph_format.left_indent = Pt(12)
            run(pn, "Nachher: ", bold=True, size=10, color=(0x1E, 0x7A, 0x1E))
            run(pn, r.get("revised", ""), size=10)
            if r.get("grund"):
                pg = doc.add_paragraph(); style_para(pg); pg.paragraph_format.left_indent = Pt(12)
                run(pg, "→ " + r["grund"], italic=True, size=9, color=(0x66, 0x66, 0x66))
            sp = doc.add_paragraph(); style_para(sp)

    os.makedirs(DESKTOP, exist_ok=True)
    out = f"{DESKTOP}/AURORA_Satzkur_Review.docx"
    doc.save(out)
    print(f"✅ Review-Dokument: {out}")
    print(f"   {nks} Konsistenz-Funde, {nrw} Satz-Umbauten")


if __name__ == "__main__":
    main()
