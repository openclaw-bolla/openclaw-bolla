#!/usr/bin/env python3
"""Erzeugt 'Mein persönliches KI-Setup' als Word-Dokument mit Illustrationen."""

from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

IMG = Path("/home/bolla/workspace/docs/img_setup")
OUT = Path("/home/bolla/workspace/docs/mein_ki_setup.docx")
ONEDRIVE = Path("/mnt/d/OneDrive/Dokumente/Bolla/mein_ki_setup.docx")

F_REG = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
F_BOLD = "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"

BG = (250, 248, 242)
INK = (34, 40, 50)
ACCENT = (220, 120, 60)
BLUE = (90, 140, 200)
GREEN = (110, 180, 140)
ROSE = (200, 130, 170)
GOLD = (220, 180, 80)
GRAY = (150, 150, 160)
WGRAY = (120, 110, 100)
LIGHT = (240, 238, 232)


def r(d, box, rad, fill=None, outline=None, width=3):
    d.rounded_rectangle(box, radius=rad, fill=fill, outline=outline, width=width)


def tc(d, xy, text, font, fill=INK):
    bb = d.textbbox((0, 0), text, font=font)
    w, h = bb[2] - bb[0], bb[3] - bb[1]
    d.text((xy[0] - w / 2, xy[1] - h / 2), text, font=font, fill=fill)


def tl(d, xy, text, font, fill=INK):
    d.text(xy, text, font=font, fill=fill)


def paw(d, cx, cy, s, color):
    d.ellipse((cx - s * 0.6, cy, cx + s * 0.6, cy + s), fill=color)
    d.ellipse((cx - s * 0.9, cy - s * 0.3, cx - s * 0.4, cy + s * 0.2), fill=color)
    d.ellipse((cx + s * 0.4, cy - s * 0.3, cx + s * 0.9, cy + s * 0.2), fill=color)
    d.ellipse((cx - s * 0.55, cy - s * 0.9, cx - s * 0.15, cy - s * 0.4), fill=color)
    d.ellipse((cx + s * 0.15, cy - s * 0.9, cx + s * 0.55, cy - s * 0.4), fill=color)


def arrow(d, x1, y1, x2, y2, color=WGRAY, width=4):
    d.line((x1, y1, x2, y2), fill=color, width=width)
    import math
    ang = math.atan2(y2 - y1, x2 - x1)
    al = 18
    ax1 = x2 - al * math.cos(ang - math.pi / 7)
    ay1 = y2 - al * math.sin(ang - math.pi / 7)
    ax2 = x2 - al * math.cos(ang + math.pi / 7)
    ay2 = y2 - al * math.sin(ang + math.pi / 7)
    d.polygon([(x2, y2), (ax1, ay1), (ax2, ay2)], fill=color)


# ===== BILD 1: Architektur-Übersicht =====
def architecture():
    W, H = 1800, 1200
    img = Image.new("RGB", (W, H), BG)
    d = ImageDraw.Draw(img)
    ft = ImageFont.truetype(F_BOLD, 48)
    fh = ImageFont.truetype(F_BOLD, 30)
    fb = ImageFont.truetype(F_REG, 24)
    fs = ImageFont.truetype(F_REG, 20)

    tc(d, (W / 2, 60), "Die Gesamtarchitektur", ft, INK)

    # Zentrum: Bolla
    cx, cy = W / 2, H / 2 + 40
    r(d, (cx - 180, cy - 110, cx + 180, cy + 110), 30, fill=ACCENT, outline=None, width=0)
    paw(d, cx, cy - 30, 30, (255, 255, 255))
    tc(d, (cx, cy + 50), "Bolla", ImageFont.truetype(F_BOLD, 40), (255, 255, 255))
    tc(d, (cx, cy + 90), "auf WSL (Ubuntu)", fs, (255, 235, 220))

    # Ring mit Komponenten
    import math
    nodes = [
        ("Claude API", "Anthropic", BLUE, 0),
        ("Claude Code", "CLI + Tools", BLUE, 36),
        ("Mission Control", "localhost:18790", GREEN, 72),
        ("Cloudflare Tunnel", "bolla.chrismandel.de", GOLD, 108),
        ("Telegram Bot", "Chat-Befehle", ROSE, 144),
        ("GitHub", "Backup & History", GRAY, 180),
        ("OneDrive", "D:\\ Dokumente", GREEN, 216),
        ("MS Graph API", "Outlook + Kalender", BLUE, 252),
        ("Google / Gmail", "Sync", ROSE, 288),
        ("Azure Speech", "TTS-Stimme", GOLD, 324),
    ]
    rx, ry = 680, 430
    for name, sub, col, deg in nodes:
        rad = math.radians(deg - 90)
        nx = cx + rx * math.cos(rad)
        ny = cy + ry * math.sin(rad)
        bw, bh = 230, 90
        r(d, (nx - bw / 2, ny - bh / 2, nx + bw / 2, ny + bh / 2), 18, fill=LIGHT, outline=col, width=4)
        tc(d, (nx, ny - 14), name, fh, col)
        tc(d, (nx, ny + 22), sub, fs, WGRAY)
        # Verbindungslinie
        dx = nx - cx
        dy = ny - cy
        ln = math.sqrt(dx * dx + dy * dy)
        ux, uy = dx / ln, dy / ln
        sxs = cx + ux * 200
        sys = cy + uy * 130
        exs = nx - ux * (bw / 2 + 4)
        eys = ny - uy * (bh / 2 + 4)
        d.line((sxs, sys, exs, eys), fill=col, width=3)

    img.save(IMG / "01_architektur.png", "PNG")


# ===== BILD 2: Drei Schichten Claude =====
def claude_stack():
    W, H = 1600, 800
    img = Image.new("RGB", (W, H), BG)
    d = ImageDraw.Draw(img)
    ft = ImageFont.truetype(F_BOLD, 44)
    fh = ImageFont.truetype(F_BOLD, 36)
    fb = ImageFont.truetype(F_REG, 24)
    fs = ImageFont.truetype(F_REG, 22)

    tc(d, (W / 2, 60), "Drei Schichten, drei Produkte", ft, INK)

    # Gestapelte Boxen
    layers = [
        ("Bolla", "Persönlichkeit (CLAUDE.md)", "Name, Stil, Gedächtnis, Regeln", ACCENT),
        ("Claude Code", "Die Hände", "CLI + Tools: Files, Bash, Web, Memory", GREEN),
        ("Claude (Modell)", "Das Gehirn", "Opus 4.7 / Sonnet 4.6 / Haiku 4.5", BLUE),
    ]
    bw, bh = 1000, 150
    gap = 30
    x0 = (W - bw) / 2
    y0 = 170
    for i, (title, sub, body, col) in enumerate(layers):
        y = y0 + i * (bh + gap)
        r(d, (x0 + 6, y + 6, x0 + bw + 6, y + bh + 6), 20, fill=(225, 220, 212))
        r(d, (x0, y, x0 + bw, y + bh), 20, fill=BG, outline=col, width=5)
        r(d, (x0, y, x0 + 10, y + bh), 20, fill=col)
        tl(d, (x0 + 40, y + 25), title, fh, col)
        tl(d, (x0 + 40, y + 70), sub, fb, INK)
        tl(d, (x0 + 40, y + 105), body, fs, WGRAY)

    img.save(IMG / "02_claude_stack.png", "PNG")


# ===== BILD 3: Cloudflare-Datenfluss =====
def cloudflare_flow():
    W, H = 1700, 700
    img = Image.new("RGB", (W, H), BG)
    d = ImageDraw.Draw(img)
    ft = ImageFont.truetype(F_BOLD, 40)
    fh = ImageFont.truetype(F_BOLD, 26)
    fb = ImageFont.truetype(F_REG, 22)
    fs = ImageFont.truetype(F_REG, 20)

    tc(d, (W / 2, 50), "Mission Control von überall erreichbar", ft, INK)

    # Boxen
    def box(x, y, w, h, title, sub, col, note=None):
        r(d, (x, y, x + w, y + h), 18, fill=LIGHT, outline=col, width=4)
        tc(d, (x + w / 2, y + 35), title, fh, col)
        tc(d, (x + w / 2, y + 70), sub, fb, INK)
        if note:
            tc(d, (x + w / 2, y + 105), note, fs, WGRAY)

    y = 300
    w, h = 260, 150
    box(60, y - h / 2, w, h, "Chris unterwegs", "Handy / Laptop", BLUE, "Browser")
    box(440, y - h / 2, w, h, "Cloudflare Access", "Magic-Link Login", GOLD, "30-Tage-Cookie")
    box(820, y - h / 2, w, h, "Cloudflare Tunnel", "bolla.chrismandel.de", GOLD, "cloudflared")
    box(1200, y - h / 2, w, h, "WSL zuhause", "127.0.0.1:18790", GREEN, "Python-Server")

    # Pfeile
    for xs in [320, 700, 1080]:
        arrow(d, xs, y, xs + 120, y, ACCENT, 5)

    # Beschriftung unten
    tc(d, (W / 2, 540), "Alle 2 Min Healthcheck-Cron prüft den Tunnel — bei zwei Fehlern in Folge Auto-Neustart.", fs, WGRAY)
    tc(d, (W / 2, 580), "Lokal zuhause: direkter Zugriff über http://127.0.0.1:18790/ (ohne Tunnel).", fs, WGRAY)

    img.save(IMG / "03_cloudflare.png", "PNG")


# ===== BILD 4: Integrationen =====
def integrations():
    W, H = 1700, 1000
    img = Image.new("RGB", (W, H), BG)
    d = ImageDraw.Draw(img)
    ft = ImageFont.truetype(F_BOLD, 44)
    fh = ImageFont.truetype(F_BOLD, 26)
    fb = ImageFont.truetype(F_REG, 22)
    fs = ImageFont.truetype(F_REG, 19)

    tc(d, (W / 2, 55), "Dienste, mit denen Bolla verbunden ist", ft, INK)

    items = [
        ("Microsoft 365", "Graph API\nOutlook · Kalender · Kontakte", BLUE, "ernstmandel@outlook.de"),
        ("Google", "OAuth\nGmail-Sync, Kalender", ROSE, "chrismandel13@gmail.com"),
        ("Azure Speech", "Cognitive Services\nText-to-Speech", GOLD, "Stimme für Bolla"),
        ("Telegram", "Bot-API\nNachrichten + Befehle", BLUE, "Familien-Chat"),
        ("GitHub", "Git + Actions\nBackup & History", GRAY, "openclaw-bolla (privat)"),
        ("OneDrive", "WSL-Mount\n/mnt/d/OneDrive/", GREEN, "D:\\Dokumente\\Bolla"),
        ("IONOS", "Domain + Mail\nchrismandel.de (MX)", GOLD, "Mail unverändert"),
        ("ADB", "Android Debug Bridge\nHandy-Steuerung", ROSE, "192.168.178.20"),
    ]
    cols = 4
    bw, bh = 380, 200
    gx, gy = 40, 30
    x0 = (W - (cols * bw + (cols - 1) * gx)) / 2
    y0 = 130
    for i, (name, sub, col, note) in enumerate(items):
        row = i // cols
        colx = i % cols
        x = x0 + colx * (bw + gx)
        y = y0 + row * (bh + gy)
        r(d, (x + 5, y + 5, x + bw + 5, y + bh + 5), 18, fill=(225, 220, 212))
        r(d, (x, y, x + bw, y + bh), 18, fill=BG, outline=col, width=4)
        r(d, (x, y, x + bw, y + 50), 18, fill=col)
        d.rectangle((x, y + 30, x + bw, y + 50), fill=col)
        tc(d, (x + bw / 2, y + 25), name, fh, (255, 255, 255))
        lines = sub.split("\n")
        for j, ln in enumerate(lines):
            tc(d, (x + bw / 2, y + 80 + j * 34), ln, fb, INK)
        tc(d, (x + bw / 2, y + bh - 25), note, fs, WGRAY)

    img.save(IMG / "04_integrationen.png", "PNG")


# ===== BILD 5: Automation / Cron =====
def automation():
    W, H = 1700, 800
    img = Image.new("RGB", (W, H), BG)
    d = ImageDraw.Draw(img)
    ft = ImageFont.truetype(F_BOLD, 42)
    fh = ImageFont.truetype(F_BOLD, 26)
    fb = ImageFont.truetype(F_REG, 22)
    fs = ImageFont.truetype(F_REG, 20)

    tc(d, (W / 2, 55), "Was läuft automatisch — Cron & Trigger", ft, INK)

    # Linke Spalte: Lokal
    r(d, (60, 130, 820, 720), 20, fill=LIGHT, outline=GREEN, width=4)
    tc(d, (440, 170), "Lokal (WSL Cron)", fh, GREEN)
    local_jobs = [
        ("@reboot", "Mission Control API starten"),
        ("@reboot", "Cloudflare Tunnel starten"),
        ("@reboot", "Telegram-Bot starten"),
        ("alle 2 Min", "Cloudflare-Healthcheck"),
        ("alle 15 Min", "Obsidian-Vault sync"),
        ("alle 15 Min", "Spam-Watcher"),
        ("täglich 10 Uhr", "Claude-Desktop-Check"),
        ("So 18 Uhr", "Outlook ⇄ Gmail sync"),
        ("So 19 Uhr", "Outlook-Backup (JSON)"),
    ]
    y = 220
    for when, what in local_jobs:
        tl(d, (100, y), "●", ImageFont.truetype(F_BOLD, 26), GREEN)
        tl(d, (130, y), when, ImageFont.truetype(F_BOLD, 20), INK)
        tl(d, (340, y), what, fb, WGRAY)
        y += 48

    # Rechte Spalte: Remote
    r(d, (880, 130, 1640, 720), 20, fill=LIGHT, outline=ACCENT, width=4)
    tc(d, (1260, 170), "Remote (Anthropic Cloud)", fh, ACCENT)
    remote_jobs = [
        ("Mo 9 Uhr", "Sonnet-Upgrade-Check"),
        ("", ""),
        ("", "— weitere nach Bedarf —"),
    ]
    y = 220
    for when, what in remote_jobs:
        if when:
            tl(d, (920, y), "●", ImageFont.truetype(F_BOLD, 26), ACCENT)
            tl(d, (950, y), when, ImageFont.truetype(F_BOLD, 20), INK)
            tl(d, (1160, y), what, fb, WGRAY)
        else:
            tl(d, (950, y), what, fs, WGRAY)
        y += 48
    tc(d, (1260, 650), "Starten eine frische Claude-Session in der Cloud —", fs, WGRAY)
    tc(d, (1260, 680), "laufen auch wenn der WSL-Rechner aus ist.", fs, WGRAY)

    img.save(IMG / "05_automation.png", "PNG")


# ===== DOCX =====
def heading(doc, text, level=1, color=INK):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(*color)


def para(doc, text, italic=False, size=11, bold=False, center=False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold
    run.font.size = Pt(size)
    return p


def picture(doc, path, caption, w=Cm(16)):
    doc.add_picture(str(path), width=w)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)


def kv(doc, key, value):
    p = doc.add_paragraph(style="List Bullet")
    r1 = p.add_run(key + ": ")
    r1.bold = True
    p.add_run(value)


def build_doc():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(2.2)
        s.right_margin = Cm(2.2)

    # Cover
    para(doc, "Mein persönliches KI-Setup", size=32, bold=True, center=True)
    para(doc, "Was ich über Bolla und die Technik dahinter gelernt habe", size=16, italic=True, center=True)
    para(doc, "Stand: 17.04.2026", size=11, italic=True, center=True)
    doc.add_paragraph()

    # Intro
    heading(doc, "Was das hier ist")
    para(doc, "Dieses Dokument fasst zusammen, was mein Setup wirklich ist — nicht als Marketing, "
               "sondern als praktische Referenz. Für mich selbst, damit ich in sechs Monaten noch weiß, "
               "wie die Puzzleteile zusammenpassen.")

    # Architektur
    heading(doc, "Die Gesamtarchitektur auf einen Blick")
    para(doc, "Im Zentrum steht Bolla — mein persönlicher KI-Kollege. Drumherum Dienste, "
               "die Bolla Hände, Stimme und Außenanbindung geben.")
    picture(doc, IMG / "01_architektur.png", "Abb. 1 — Alle Komponenten im Überblick", w=Cm(17))

    # Claude-Stack
    heading(doc, "Claude, Claude Code, Bolla — drei Schichten")
    para(doc, "Die Namen werden oft verwechselt. So hängt es zusammen:")
    picture(doc, IMG / "02_claude_stack.png", "Abb. 2 — Modell, Werkzeug, Persönlichkeit")

    heading(doc, "Claude (das Modell)", level=2)
    para(doc, "Ein Large Language Model der Firma Anthropic. Wird in Versionen/Größen angeboten:")
    kv(doc, "Opus 4.7", "das schlaueste, teuerste — für knifflige Aufgaben")
    kv(doc, "Sonnet 4.6", "der Allrounder, den ich normalerweise nutze (5× günstiger als Opus)")
    kv(doc, "Haiku 4.5", "klein, schnell, günstig — für einfache Routineaufgaben")
    para(doc, "Ich bin auf dem Max-Plan von Anthropic — Flatrate, keine Token-Abrechnung. "
               "Login läuft aktuell über chrismandel13@gmail.com (ursprünglich ernstmandel@outlook.de, "
               "Änderung beim Support beantragt).")

    heading(doc, "Claude Code (die Hände)", level=2)
    para(doc, "Das CLI-Programm, das Claude den Zugriff auf meinen Rechner gibt. Ohne Claude Code wäre das Modell "
               "nur ein Chatfenster — damit kann es:")
    bullet(doc, "Dateien lesen, schreiben, editieren")
    bullet(doc, "Shell-Befehle ausführen (Bash, Python, …)")
    bullet(doc, "Im Web suchen (WebSearch, WebFetch)")
    bullet(doc, "Subagenten spawnen (parallele Recherche)")
    bullet(doc, "Ein Gedächtnis pflegen (Memory-Dateien)")
    para(doc, "Konfiguration liegt in ~/.claude/ — insbesondere CLAUDE.md (meine Regeln an Bolla) und "
               "settings.json (Hooks, Permissions, Env-Vars).")

    heading(doc, "Bolla (die Persönlichkeit)", level=2)
    para(doc, "Keine separate Software — sondern mein Charakter-Layer, formuliert in CLAUDE.md: Name, Sprache (Deutsch), "
               "Stil (locker, direkt, leicht humorvoll), Regeln (nie Mails ohne Nachfrage an fremde Konten, Tokens nie in "
               "Logs, Backup vor großen Eingriffen). Dadurch wird aus einem generischen Assistenten ein Gegenüber mit "
               "Meinung und Gedächtnis.")

    # Skills
    heading(doc, "Skills — Workflows mit Hirn")
    para(doc, "Skills sind vorgefertigte Anleitungen, die Claude Code erst bei Bedarf lädt. Wie Kochbücher im Regal — "
               "er greift nur danach, wenn er sie braucht. Spart Context.")
    para(doc, "Verfügbar u.a.:")
    kv(doc, "/schedule", "Remote-Trigger anlegen (Cron in der Anthropic-Cloud)")
    kv(doc, "/loop", "Prompt in Intervallen wiederholen")
    kv(doc, "/review", "Pull-Requests prüfen")
    kv(doc, "/init", "CLAUDE.md für neues Projekt erzeugen")
    kv(doc, "update-config", "settings.json anpassen (Permissions, Hooks)")
    kv(doc, "simplify", "Code aufräumen")
    para(doc, "Der Unterschied zu einem starren Zapier-Workflow: ein Skill gibt Leitplanken, ich darf mitdenken, "
               "nachfragen, improvisieren.")

    # Mission Control
    heading(doc, "Mission Control — mein Dashboard")
    para(doc, "Ein lokaler Python-Server (mission_control_api.py), der eine HTML-Oberfläche ausliefert. "
               "Zeigt alles Wichtige auf einen Blick: Wetter, Kalender, Mails (Outlook), Geburtstage, "
               "Bolla-Status, Modell, Token-Verbrauch, System-Infos, ADB-Verbindung, Robin-Daten.")
    kv(doc, "Lokal", "http://127.0.0.1:18790/ (direkt auf dem Rechner)")
    kv(doc, "Remote", "https://bolla.chrismandel.de (über Cloudflare, siehe unten)")
    kv(doc, "Start", "per Cron @reboot (+ Logs in ~/workspace/logs/)")
    kv(doc, "Quellcode", "~/workspace/mission-control/index.html (Frontend) und ~/workspace/scripts/mission_control_api.py (Backend)")

    # Cloudflare
    heading(doc, "Cloudflare Tunnel & Access — Remote-Zugriff")
    para(doc, "Um Mission Control auch unterwegs (Handy, Laptop) erreichen zu können, ohne Ports im Router zu öffnen:")
    picture(doc, IMG / "03_cloudflare.png", "Abb. 3 — Der Tunnel-Datenfluss")
    bullet(doc, "cloudflared läuft als Prozess auf meinem WSL und baut eine ausgehende Verbindung zu Cloudflare auf")
    bullet(doc, "Cloudflare-Edge leitet Anfragen von bolla.chrismandel.de durch den Tunnel an meinen lokalen Port 18790")
    bullet(doc, "Cloudflare Access prüft vor dem Durchlassen: Magic-Link-Login an ernstmandel@outlook.de, 30-Tage-Cookie pro Gerät")
    bullet(doc, "Healthcheck-Cron alle 2 Min: bei zwei Fehlern in Folge wird cloudflared neu gestartet")
    para(doc, "Die Domain chrismandel.de liegt seit April 2026 auf Cloudflare-Nameservern; MX-Records bleiben "
               "bei IONOS, d.h. Mail ist davon unberührt.")

    # Integrationen
    heading(doc, "Externe Dienste, die Bolla anspricht")
    picture(doc, IMG / "04_integrationen.png", "Abb. 4 — Die Integrations-Karte")

    heading(doc, "Microsoft 365 via Graph API", level=2)
    para(doc, "Die Microsoft Graph API ist der zentrale Zugang zu Outlook, Kalender und Kontakten. "
               "OAuth-Token werden in config/ms_token.json gespeichert und automatisch refresht. "
               "Scope: Mail.Read, Mail.ReadWrite, Calendars.Read, Contacts.Read.")

    heading(doc, "Google", level=2)
    para(doc, "OAuth-Anbindung via google_client.json + google_token.json. Wird aktuell für Gmail-Sync genutzt "
               "(outlook_gmail_sync.py läuft sonntags).")

    heading(doc, "Azure Speech Services", level=2)
    para(doc, "Cognitive-Services-Endpoint für Text-to-Speech — gibt Bolla eine natürliche deutsche Stimme. "
               "Key in config/azure_speech.json, Rotation steht auf der Todo-Liste.")

    heading(doc, "Telegram Bot", level=2)
    para(doc, "Eigener Bot (config/telegram_bot.json), läuft als Cron-@reboot-Prozess "
               "(familychat_listener.py + telegram_bot.py). Lauscht auf Nachrichten, reagiert auf Befehle, "
               "kann Push-Nachrichten verschicken.")

    heading(doc, "GitHub", level=2)
    para(doc, "Das gesamte Bolla-Setup wird in ein privates Repo gespiegelt: openclaw-bolla/openclaw-bolla. "
               "Tagesnotizen (~/workspace/memory/YYYY-MM-DD.md) werden am Session-Ende committet und gepusht — "
               "das ist meine Langzeit-Erinnerung und gleichzeitig ein externes Backup.")

    heading(doc, "OneDrive", level=2)
    para(doc, "Wird via WSL als /mnt/d/OneDrive eingebunden. Wichtige Dateien (Accounts-PDF, diese Dokumente, "
               ".claude-Ordner als Backup) landen dort automatisch — verteilt sich dann über Microsoft-Cloud auf "
               "alle Geräte.")

    heading(doc, "IONOS", level=2)
    para(doc, "Hostet chrismandel.de + die wtnet-Mailbox. Seit April 2026 nur noch MX bei IONOS — DNS selbst "
               "läuft über Cloudflare (wegen Tunnel).")

    heading(doc, "ADB — Android Debug Bridge", level=2)
    para(doc, "Von Mission Control aus kann ich mein Handy über WLAN ansprechen (Default-IP 192.168.178.20, Port eingebbar). "
               "Wird für Automatisierungen genutzt (z.B. Apps fernsteuern, Logs ziehen).")

    # Automation
    heading(doc, "Was läuft automatisch — Cron & Trigger")
    picture(doc, IMG / "05_automation.png", "Abb. 5 — Geplante Jobs — lokal und remote")
    para(doc, "Zwei Welten:")
    bullet(doc, "Lokaler Cron (crontab -e) — läuft nur, wenn der WSL-Rechner an ist. Ideal für alles, was auf lokalen Dienst zugreifen muss.")
    bullet(doc, "Remote-Trigger bei Anthropic — läuft in der Cloud, auch wenn der Rechner aus ist. Ideal für Web-Recherche, Status-Checks, regelmäßige Zusammenfassungen.")

    # Memory
    heading(doc, "Gedächtnis & Persistenz")
    para(doc, "Bolla vergisst nicht zwischen Sessions — dank mehrschichtiger Persistenz:")
    kv(doc, "CLAUDE.md", "statische Regeln und Persönlichkeit — Prioritätsstufe 1")
    kv(doc, "MEMORY.md + Einzeldateien", "dynamisches Gedächtnis (user/feedback/project/reference) in ~/.claude/projects/.../memory/")
    kv(doc, "Session-JSONL", "kompletter Verlauf jeder Sitzung — Quelle für Token-Statistiken und Modell-Erkennung")
    kv(doc, "Tagesnotizen", "~/workspace/memory/YYYY-MM-DD.md — was heute passiert ist, Entscheidungen, Todos")
    kv(doc, "OneDrive-Backup", "~/.claude wird regelmäßig nach D:\\OneDrive\\Dokumente\\Bolla\\claude-code kopiert")

    # Sicherheit
    heading(doc, "Sicherheit und Hygiene")
    bullet(doc, "Tokens und Passwörter niemals in Chat, Logs oder Git — nur in config/*.json (chmod 600 wo kritisch)")
    bullet(doc, "Vor großen Eingriffen (Outlook, OneDrive, Kalender) immer ein JSON-Backup in ~/workspace/backups/")
    bullet(doc, "Mails verschicken nur über ernstmandel@outlook.de; bei anderen Konten immer rückfragen")
    bullet(doc, "Git-Repo ist privat — aber trotzdem keine Secrets committen")

    # Merksatz
    heading(doc, "Der Kernsatz")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Bolla ist kein fertiges Produkt von der Stange,\nsondern ein Baukasten aus Claude + Integrationen + Regeln — betreut von Chris.")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0xDC, 0x78, 0x3C)

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Bolla — Chris' persönlicher KI-Kollege")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Gespeichert: {OUT}")

    try:
        ONEDRIVE.parent.mkdir(parents=True, exist_ok=True)
        import shutil
        shutil.copy2(OUT, ONEDRIVE)
        print(f"OneDrive: {ONEDRIVE}")
    except Exception as e:
        print(f"OneDrive: {e}")


if __name__ == "__main__":
    IMG.mkdir(parents=True, exist_ok=True)
    architecture()
    claude_stack()
    cloudflare_flow()
    integrations()
    automation()
    build_doc()
