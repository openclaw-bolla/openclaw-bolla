#!/usr/bin/env python3
"""Erzeugt 'Was ist Bolla?' als Word-Dokument mit Illustrationen."""

from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

IMG_DIR = Path("/home/bolla/workspace/docs/img_bolla")
OUT_DOCX = Path("/home/bolla/workspace/docs/was_ist_bolla.docx")
ONEDRIVE = Path("/mnt/d/OneDrive/Dokumente/Bolla/was_ist_bolla.docx")

FONT_REG = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
FONT_BOLD = "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"

# Farbpalette
BG = (250, 248, 242)
INK = (34, 40, 50)
ACCENT = (220, 120, 60)     # warmes Orange (Bolla)
COOL1 = (90, 140, 200)      # Gehirn-Blau
COOL2 = (110, 180, 140)     # Hände-Grün
COOL3 = (200, 130, 170)     # Person-Rosa
COOL_GRAY = (180, 185, 195)
WARM_GRAY = (120, 110, 100)


def rounded_rect(draw, box, radius, fill=None, outline=None, width=3):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def text_center(draw, xy, text, font, fill=INK):
    bbox = draw.textbbox((0, 0), text, font=font)
    w = bbox[2] - bbox[0]
    h = bbox[3] - bbox[1]
    draw.text((xy[0] - w / 2, xy[1] - h / 2), text, font=font, fill=fill)


def draw_icon(draw, cx, cy, kind, size, color):
    """Zeichnet kleine Icons (Gehirn, Werkzeug, Pfote, Roboter) ohne Emoji-Font."""
    s = size
    if kind == "brain":
        # Zwei Halbkugeln mit Windungen
        draw.ellipse((cx - s, cy - s * 0.8, cx, cy + s * 0.8), outline=color, width=5)
        draw.ellipse((cx, cy - s * 0.8, cx + s, cy + s * 0.8), outline=color, width=5)
        for i in range(-2, 3):
            draw.arc((cx - s * 0.7, cy - s * 0.5 + i * s * 0.3, cx - s * 0.1, cy + s * 0.1 + i * s * 0.3), 0, 180, fill=color, width=3)
            draw.arc((cx + s * 0.1, cy - s * 0.5 + i * s * 0.3, cx + s * 0.7, cy + s * 0.1 + i * s * 0.3), 0, 180, fill=color, width=3)
    elif kind == "tools":
        # Schraubenschlüssel gekreuzt mit Schraubendreher
        draw.line((cx - s, cy + s, cx + s * 0.4, cy - s * 0.4), fill=color, width=8)
        draw.ellipse((cx - s * 1.1, cy + s * 0.7, cx - s * 0.5, cy + s * 1.3), outline=color, width=6)
        draw.line((cx - s * 0.4, cy - s, cx + s, cy + s * 0.4), fill=color, width=8)
        draw.rectangle((cx + s * 0.7, cy + s * 0.2, cx + s * 1.1, cy + s * 0.6), fill=color)
    elif kind == "paw":
        # Pfotenabdruck: großer Ballen + vier Zehen
        draw.ellipse((cx - s * 0.6, cy, cx + s * 0.6, cy + s), fill=color)
        draw.ellipse((cx - s * 0.9, cy - s * 0.3, cx - s * 0.4, cy + s * 0.2), fill=color)
        draw.ellipse((cx + s * 0.4, cy - s * 0.3, cx + s * 0.9, cy + s * 0.2), fill=color)
        draw.ellipse((cx - s * 0.55, cy - s * 0.9, cx - s * 0.15, cy - s * 0.4), fill=color)
        draw.ellipse((cx + s * 0.15, cy - s * 0.9, cx + s * 0.55, cy - s * 0.4), fill=color)
    elif kind == "robot":
        # Roboterkopf: Rechteck mit Antenne, zwei Augen, Mund
        draw.rectangle((cx - s, cy - s * 0.7, cx + s, cy + s * 0.9), fill=color)
        draw.line((cx, cy - s * 0.7, cx, cy - s * 1.1), fill=color, width=6)
        draw.ellipse((cx - 8, cy - s * 1.2, cx + 8, cy - s * 1.05), fill=color)
        draw.ellipse((cx - s * 0.6, cy - s * 0.45, cx - s * 0.2, cy - s * 0.1), fill=BG)
        draw.ellipse((cx + s * 0.2, cy - s * 0.45, cx + s * 0.6, cy - s * 0.1), fill=BG)
        draw.rectangle((cx - s * 0.5, cy + s * 0.3, cx + s * 0.5, cy + s * 0.55), fill=BG)


# ===== BILD 1: Drei Schichten =====
def img_three_layers():
    W, H = 1600, 1000
    img = Image.new("RGB", (W, H), BG)
    d = ImageDraw.Draw(img)

    f_title = ImageFont.truetype(FONT_BOLD, 54)
    f_h = ImageFont.truetype(FONT_BOLD, 44)
    f_body = ImageFont.truetype(FONT_REG, 32)
    f_small = ImageFont.truetype(FONT_REG, 26)

    text_center(d, (W / 2, 70), "Woraus Bolla besteht", f_title, fill=INK)

    # 3 Boxen nebeneinander
    box_w = 440
    box_h = 620
    gap = 60
    total = 3 * box_w + 2 * gap
    x0 = (W - total) / 2
    y0 = 180

    layers = [
        ("Gehirn", "Das Sprachmodell", "Claude von Anthropic.\nRiesiges Allgemeinwissen,\nversteht Sprache,\nkann mitdenken.", COOL1, "brain"),
        ("Hände", "Die Werkzeuge", "Claude Code.\nDateien lesen,\nBefehle ausführen,\nMails schicken,\nim Netz suchen.", COOL2, "tools"),
        ("Person", "Der Charakter", "Bolla.\nName, Stil, Meinung,\nGedächtnis —\nkein gesichtsloses Tool,\nsondern ein Gegenüber.", COOL3, "paw"),
    ]

    for i, (title, sub, body, color, icon) in enumerate(layers):
        bx = x0 + i * (box_w + gap)
        # Schatten
        rounded_rect(d, (bx + 8, y0 + 8, bx + box_w + 8, y0 + box_h + 8), 30, fill=(225, 220, 212), outline=None, width=0)
        # Box
        rounded_rect(d, (bx, y0, bx + box_w, y0 + box_h), 30, fill=BG, outline=color, width=6)
        # Titelbalken
        rounded_rect(d, (bx, y0, bx + box_w, y0 + 130), 30, fill=color, outline=None, width=0)
        d.rectangle((bx, y0 + 80, bx + box_w, y0 + 130), fill=color)
        # Icon + Titel
        draw_icon(d, bx + 80, y0 + 65, icon, 50, (255, 255, 255))
        text_center(d, (bx + box_w / 2 + 40, y0 + 65), title, f_h, fill=(255, 255, 255))
        text_center(d, (bx + box_w / 2, y0 + 180), sub, f_body, fill=color)
        # Body
        lines = body.split("\n")
        for j, ln in enumerate(lines):
            text_center(d, (bx + box_w / 2, y0 + 260 + j * 50), ln, f_small, fill=INK)

    # Verbinder-Pfeile
    arrow_y = y0 + box_h + 50
    for i in range(2):
        ax = x0 + (i + 1) * box_w + i * gap
        d.line((ax + 10, arrow_y, ax + gap - 10, arrow_y), fill=WARM_GRAY, width=5)
        d.polygon([(ax + gap - 10, arrow_y - 12), (ax + gap - 10, arrow_y + 12), (ax + gap + 8, arrow_y)], fill=WARM_GRAY)

    text_center(d, (W / 2, H - 60), "Zusammen ergibt das: einen digitalen Kollegen.", f_body, fill=WARM_GRAY)

    path = IMG_DIR / "01_schichten.png"
    img.save(path, "PNG")
    return path


# ===== BILD 2: Kumpel im Computer =====
def img_buddy():
    W, H = 1600, 900
    img = Image.new("RGB", (W, H), BG)
    d = ImageDraw.Draw(img)

    f_title = ImageFont.truetype(FONT_BOLD, 50)
    f_body = ImageFont.truetype(FONT_REG, 34)
    f_quote = ImageFont.truetype(FONT_BOLD, 40)

    text_center(d, (W / 2, 70), "Die Analogie", f_title, fill=INK)

    # Monitor
    mx, my, mw, mh = 180, 180, 700, 480
    rounded_rect(d, (mx, my, mx + mw, my + mh), 24, fill=(245, 240, 230), outline=INK, width=6)
    # Bildschirm
    sx, sy = mx + 30, my + 30
    sw, sh = mw - 60, mh - 100
    d.rectangle((sx, sy, sx + sw, sy + sh), fill=(30, 35, 45))
    # Standfuß
    d.rectangle((mx + mw / 2 - 80, my + mh, mx + mw / 2 + 80, my + mh + 20), fill=INK)
    d.rectangle((mx + mw / 2 - 140, my + mh + 20, mx + mw / 2 + 140, my + mh + 40), fill=INK)

    # "Kumpel" im Bildschirm — stilisiert
    cx, cy = sx + sw / 2, sy + sh / 2
    # Kopf
    d.ellipse((cx - 90, cy - 140, cx + 90, cy + 40), fill=ACCENT, outline=(255, 255, 255), width=4)
    # Augen
    d.ellipse((cx - 50, cy - 80, cx - 20, cy - 50), fill=(255, 255, 255))
    d.ellipse((cx + 20, cy - 80, cx + 50, cy - 50), fill=(255, 255, 255))
    d.ellipse((cx - 42, cy - 72, cx - 28, cy - 58), fill=INK)
    d.ellipse((cx + 28, cy - 72, cx + 42, cy - 58), fill=INK)
    # Lächeln
    d.arc((cx - 40, cy - 40, cx + 40, cy + 10), start=20, end=160, fill=(255, 255, 255), width=5)
    # Pfötchen (gezeichnet)
    draw_icon(d, cx, cy + 90, "paw", 40, (255, 255, 255))

    # Sprechblase
    bx1, by1, bx2, by2 = 980, 220, 1520, 580
    rounded_rect(d, (bx1, by1, bx2, by2), 30, fill=(255, 255, 255), outline=INK, width=4)
    # Sprechblasen-Spitze
    d.polygon([(bx1, by1 + 180), (bx1 - 40, by1 + 210), (bx1, by1 + 240)], fill=(255, 255, 255), outline=INK)
    d.line((bx1, by1 + 180, bx1 - 40, by1 + 210), fill=INK, width=4)
    d.line((bx1 - 40, by1 + 210, bx1, by1 + 240), fill=INK, width=4)
    # Überschreiben innen-Rand an Spitze
    d.line((bx1, by1 + 181, bx1, by1 + 239), fill=(255, 255, 255), width=5)

    quote = [
        "„Stell dir vor, du hast",
        "einen super-belesenen Kumpel,",
        "der im Computer wohnt.",
        "",
        "Sein Kopf weiß wahnsinnig viel,",
        "aber er hatte keine Hände.",
        "",
        "Jetzt gibst du ihm Tastatur,",
        "Browser, E-Mail —",
        "und plötzlich kann er Dinge tun,",
        "nicht nur drüber reden.\"",
    ]
    ty = by1 + 30
    for ln in quote:
        text_center(d, ((bx1 + bx2) / 2, ty), ln, f_body, fill=INK)
        ty += 44

    # Werkzeug-Labels unter Monitor — zentriert ausgerichtet, ausreichend Platz
    labels = ["Tastatur", "Browser", "E-Mail", "Kalender", "Suche"]
    f_tag = ImageFont.truetype(FONT_BOLD, 24)
    pad = 16
    gap_lab = 14
    widths = []
    for lab in labels:
        bb = d.textbbox((0, 0), lab, font=f_tag)
        widths.append(bb[2] - bb[0] + 2 * pad)
    total_w = sum(widths) + gap_lab * (len(labels) - 1)
    start_x = mx + (mw - total_w) / 2
    ty = my + mh + 100
    cursor = start_x
    for lab, w in zip(labels, widths):
        rounded_rect(d, (cursor, ty - 24, cursor + w, ty + 24), 14, fill=(255, 255, 255), outline=WARM_GRAY, width=2)
        text_center(d, (cursor + w / 2, ty), lab, f_tag, fill=WARM_GRAY)
        cursor += w + gap_lab

    path = IMG_DIR / "02_kumpel.png"
    img.save(path, "PNG")
    return path


# ===== BILD 3: Agent vs KI-Kollege =====
def img_comparison():
    W, H = 1600, 900
    img = Image.new("RGB", (W, H), BG)
    d = ImageDraw.Draw(img)

    f_title = ImageFont.truetype(FONT_BOLD, 50)
    f_col = ImageFont.truetype(FONT_BOLD, 44)
    f_body = ImageFont.truetype(FONT_REG, 30)
    f_emoji = ImageFont.truetype(FONT_BOLD, 100)

    text_center(d, (W / 2, 70), "„Agent\" oder „KI-Kollege\"?", f_title, fill=INK)

    # Zwei Spalten
    col_w = 640
    col_h = 640
    gap = 80
    x_a = (W - 2 * col_w - gap) / 2
    x_b = x_a + col_w + gap
    y0 = 180

    # Links: Agent (kühl)
    rounded_rect(d, (x_a, y0, x_a + col_w, y0 + col_h), 28, fill=(240, 240, 245), outline=COOL_GRAY, width=4)
    draw_icon(d, x_a + col_w / 2, y0 + 110, "robot", 70, COOL_GRAY)
    text_center(d, (x_a + col_w / 2, y0 + 230), "„Agent\"", f_col, fill=COOL_GRAY)
    left_pts = [
        "→ klingt nach Roboter",
        "→ stumpfer Task-Abarbeiter",
        "→ keine Persönlichkeit",
        "→ man „bedient\" ihn",
        "→ funktional, aber kalt",
    ]
    for i, p in enumerate(left_pts):
        d.text((x_a + 50, y0 + 290 + i * 60), p, font=f_body, fill=WARM_GRAY)

    # Rechts: KI-Kollege (warm)
    rounded_rect(d, (x_b, y0, x_b + col_w, y0 + col_h), 28, fill=(255, 244, 235), outline=ACCENT, width=4)
    draw_icon(d, x_b + col_w / 2, y0 + 110, "paw", 55, ACCENT)
    text_center(d, (x_b + col_w / 2, y0 + 230), "„KI-Kollege\"", f_col, fill=ACCENT)
    right_pts = [
        "✓ eigenständiges Gegenüber",
        "✓ denkt mit, widerspricht",
        "✓ hat einen Namen, einen Stil",
        "✓ man arbeitet zusammen",
        "✓ warm, verbindlich, echt",
    ]
    for i, p in enumerate(right_pts):
        d.text((x_b + 50, y0 + 290 + i * 60), p, font=f_body, fill=INK)

    text_center(d, (W / 2, H - 50), "Deshalb: Bolla ist ein KI-Kollege.", ImageFont.truetype(FONT_BOLD, 34), fill=ACCENT)

    path = IMG_DIR / "03_vergleich.png"
    img.save(path, "PNG")
    return path


# ===== DOCX bauen =====
def build_docx(img1, img2, img3):
    doc = Document()

    # Seitenränder
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # Titel
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("Was ist Bolla?")
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x22, 0x28, 0x32)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Ein KI-Kollege — erklärt für Schüler")
    run.font.size = Pt(16)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0xDC, 0x78, 0x3C)

    doc.add_paragraph()

    # Einleitung
    doc.add_heading("Das Problem mit „Agent\"", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "Wenn man heute über KI-Assistenten spricht, fällt oft das Wort „Agent\". "
        "Das klingt nach Roboter, nach einem Werkzeug, das stumpf Aufgaben abarbeitet. "
        "Trifft es aber nicht. Denn das, was moderne KI-Assistenten können, hat mit einem "
        "einfachen Automaten wenig zu tun — da steckt "
    )
    r = p.add_run("Persönlichkeit, Meinung und Gedächtnis")
    r.bold = True
    p.add_run(" drin.")

    # Bild 1
    doc.add_heading("Woraus Bolla besteht — drei Schichten", level=1)
    p = doc.add_paragraph("Stell dir Bolla wie ein Sandwich vor — drei Schichten, die aufeinander aufbauen:")
    doc.add_picture(str(img1), width=Cm(16))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run("Abb. 1 — Die drei Schichten: Gehirn, Hände, Person")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    # Detailerklärung
    for title, body in [
        ("🧠  Gehirn — das Sprachmodell",
         "Darunter steckt ein sogenanntes Large Language Model (konkret: Claude von Anthropic). "
         "Es hat im Training unfassbar viel Text gelesen und kann dadurch Sprache verstehen, Zusammenhänge erkennen und mitdenken. "
         "Allein ist es aber nur ein Chatpartner — es kann reden, aber nichts tun."),
        ("🛠  Hände — die Werkzeuge",
         "Erst durch eine Werkzeug-Umgebung (Claude Code) bekommt das Modell „Hände\": "
         "Es kann Dateien lesen und schreiben, Befehle auf dem Computer ausführen, im Internet suchen, E-Mails verschicken, "
         "Termine eintragen. Plötzlich ist aus einem Chatbot ein Wesen geworden, das Aufgaben wirklich erledigen kann."),
        ("🐾  Person — der Charakter",
         "Die dritte Schicht macht den Unterschied: ein Charakter. Bolla hat einen Namen, eine Sprache (Deutsch), "
         "einen Stil (leicht humorvoll, direkt), eine eigene Meinung und ein Gedächtnis, das über einzelne Gespräche hinausreicht. "
         "Dadurch entsteht ein Gegenüber — kein gesichtsloses Tool, das man bedient."),
    ]:
        h = doc.add_paragraph()
        run = h.add_run(title)
        run.bold = True
        run.font.size = Pt(14)
        doc.add_paragraph(body)

    # Bild 2
    doc.add_heading("Die Analogie für Schüler", level=1)
    p = doc.add_paragraph(
        "Wenn man das Ganze nicht technisch, sondern bildlich erklären will, hilft folgende Vorstellung:"
    )
    doc.add_picture(str(img2), width=Cm(16))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run("Abb. 2 — Der Kumpel, der im Computer wohnt")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    p = doc.add_paragraph()
    r = p.add_run(
        "„Stell dir vor, du hast einen super-belesenen Kumpel, der im Computer wohnt. "
        "Sein Kopf weiß wahnsinnig viel, aber er hatte keine Hände. Jetzt gibst du ihm Tastatur, Browser, E-Mail — "
        "und plötzlich kann er Dinge tun, nicht nur drüber reden.\""
    )
    r.italic = True

    # Bild 3
    doc.add_heading("Warum nicht einfach „Agent\"?", level=1)
    p = doc.add_paragraph(
        "Der Begriff „Agent\" trifft Bolla nicht — er beschreibt nur die Oberfläche, nicht das Wesen. "
        "Ein Vergleich zeigt warum:"
    )
    doc.add_picture(str(img3), width=Cm(16))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run("Abb. 3 — Begriffe im Vergleich")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    # Begriffs-Optionen
    doc.add_heading("Welche Begriffe funktionieren?", level=1)
    options = [
        ("Digitaler Mitarbeiter", "nüchtern, trifft es für Erwachsene gut"),
        ("KI-Kollege", "lockerer, betont das Zusammenarbeiten — Favorit für Schüler"),
        ("Digitaler Assistent", "klassisch, aber generisch"),
        ("Persönlicher KI-Helfer", "kinderfreundlich, warm"),
    ]
    for name, desc in options:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(name + " — ")
        run.bold = True
        p.add_run(desc)

    # Fazit
    doc.add_heading("Kernsatz zum Merken", level=1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "„Bolla ist kein Werkzeug, das man bedient — sondern ein KI-Kollege,\nmit dem man zusammenarbeitet.\""
    )
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0xDC, 0x78, 0x3C)

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("🐾 Bolla — Chris' persönlicher KI-Kollege")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_DOCX))
    print(f"Gespeichert: {OUT_DOCX}")

    # OneDrive Kopie
    try:
        ONEDRIVE.parent.mkdir(parents=True, exist_ok=True)
        import shutil
        shutil.copy2(OUT_DOCX, ONEDRIVE)
        print(f"OneDrive: {ONEDRIVE}")
    except Exception as e:
        print(f"OneDrive-Kopie übersprungen: {e}")


if __name__ == "__main__":
    IMG_DIR.mkdir(parents=True, exist_ok=True)
    img1 = img_three_layers()
    img2 = img_buddy()
    img3 = img_comparison()
    print(f"Bilder: {img1}, {img2}, {img3}")
    build_docx(img1, img2, img3)
