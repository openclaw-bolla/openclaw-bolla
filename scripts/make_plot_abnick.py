#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""AURORA Plot-Lücken + Brandt — Abnick-Doc aus plot_vorschlaege.json."""
import json, os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

V = json.load(open("/home/bolla/workspace/data/satzkur/plot_vorschlaege.json", encoding="utf-8"))["vorschlaege"]

# Punkt-Titel je id
TITEL = {
    "43": "① Sicherungstausch / Notstrom (Kap 27)",
    "44": "② Herkunft der Versiegelung (Kap 27)",
    "47": "③ Verschwundener Notartermin (Kap 29)",
    "77": "④ Wagen-Doppelung (Kap 45)",
    "78": "⑤ Koffer-Übergabe (Kap 45)",
    "83": "⑥ Abriss-Enthüllung (Kap 46)",
    "10": "⑦ Doppelt erzählter Anruf — Entdopplung (Kap 3)",
    "brandt": "⑧ Günter Brandt → Wachmann (Kap 27/28/29)",
}
ORDER = ["43", "44", "47", "77", "78", "83", "10", "brandt"]

doc = Document()
st = doc.styles["Normal"]; st.font.name = "Aptos"; st.font.size = Pt(10.5)
st.element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
st.paragraph_format.space_after = Pt(0); st.paragraph_format.line_spacing = 1.13

def para(text="", size=10.5, bold=False, color=None, after=0, before=0, italic=False):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.line_spacing = 1.13
    r = p.add_run(text); r.font.name = "Aptos"; r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = RGBColor(*color)
    return p

para("AURORA — Plot-Lücken & Brandt zum Abnicken", 19, True, (0x1a,0x1a,0x3a), after=2)
para("Fable-Vorschläge · nichts davon ist im Buch, bis du es freigibst · 3. Juli 2026", 11, False, (0x66,0x66,0x66), after=6, italic=True)
para("Pro Punkt: sag „ja“, „nein/offen lassen“ oder „so aber ändere …“. Der TEXT in Grün ist der konkrete Vorschlag.", 10.5, bold=True, after=8)

ACT = {"einfuegen_nach": "EINFÜGEN nach dieser Stelle", "ersetzen": "ERSETZEN", "streichen": "STREICHEN"}

for key in ORDER:
    items = [x for x in V if str(x["id"]) == key]
    if not items: continue
    para(TITEL[key], 14, True, (0x1a,0x1a,0x3a), after=2, before=7)
    multi = len(items) > 1
    for j, x in enumerate(items, 1):
        if multi:
            para(f"   Teil {j}/{len(items)} — {ACT.get(x['aktion'], x['aktion'])}", 10, bold=True, color=(0x2a,0x4a,0x8a), after=1, before=2)
        else:
            para(ACT.get(x["aktion"], x["aktion"]), 10, bold=True, color=(0x2a,0x4a,0x8a), after=1)
        para("Warum: " + x.get("begruendung", ""), 9.5, italic=True, color=(0x55,0x55,0x55), after=1)
        anker = (x.get("anker") or "").replace("\n", " ").strip()
        if len(anker) > 150: anker = anker[:150] + "…"
        label = "Stelle" if x["aktion"] == "einfuegen_nach" else ("Alt" if x["aktion"] == "ersetzen" else "Streichen")
        para(f"{label}: „{anker}“", 9.5, color=(0x70,0x40,0x40), after=1)
        if x.get("vorschlag"):
            lbl = "► Neuer Text: " if x["aktion"] == "einfuegen_nach" else "► Neu: "
            p = para("", after=3)
            r1 = p.add_run(lbl); r1.bold = True; r1.font.name = "Aptos"; r1.font.size = Pt(10)
            r2 = p.add_run(x["vorschlag"]); r2.font.name = "Aptos"; r2.font.size = Pt(10.5); r2.font.color.rgb = RGBColor(0x12,0x5a,0x28)

out = "/mnt/d/OneDrive/Desktop/AURORA_PlotLuecken_Abnicken.docx"
doc.save(out)
print("Gespeichert:", out, "|", os.path.getsize(out), "bytes |", len(V), "Einträge")
