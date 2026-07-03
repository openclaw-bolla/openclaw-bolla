#!/usr/bin/env python3
"""Baut das AURORA Konsistenz-Review .docx aus review_B.json (Aptos, kompakt)."""
import json
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

FUNDE = json.load(open("/home/bolla/workspace/data/satzkur/review_B.json", encoding="utf-8"))

doc = Document()
# Basisschrift Aptos
style = doc.styles["Normal"]
style.font.name = "Aptos"
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
pf = style.paragraph_format
pf.space_after = Pt(0); pf.space_before = Pt(0); pf.line_spacing = 1.12

def kap_sort(f):
    n = f.get("nr")
    return (999 if n is None else n)

def p(text="", size=10.5, bold=False, color=None, after=0, before=0, italic=False, align=None):
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(after)
    par.paragraph_format.space_before = Pt(before)
    par.paragraph_format.line_spacing = 1.12
    if align: par.alignment = align
    r = par.add_run(text)
    r.font.name = "Aptos"; r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = RGBColor(*color)
    return par, r

# --- Titel ---
p("AURORA — Konsistenz-Review", 20, True, (0x1a,0x1a,0x3a), after=2)
p("86 Funde zum Durchwinken · Stand 3. Juli 2026", 11, False, (0x66,0x66,0x66), after=8, italic=True)

# --- Bereits erledigt (grün) ---
p("✅ Schon erledigt (autonom, mit Backup):", 12, True, (0x15,0x80,0x3d), after=2)
for t in [
    "83 von 94 Fable-Satz-Umbauten angewendet (11 nach Fable-Kurationspass verworfen).",
    "2 glasklare Fixes: „Frau Hoffmann“→„Frau Yilmaz“ (Kap 29), „Siebenjährigen“→„Siebenjährige“ (Mia, Kap 46).",
    "Auftakt-Datum kalendarisch korrigiert: Prolog jetzt Donnerstag, 1. März 2035 (+ Folgetag 2. März).",
]:
    p("   •  " + t, 10.5, after=1)
p("", 4)

# --- Timeline-Frage prominent ---
p("⚠️ EINE Entscheidung vorab (Timeline):", 12, True, (0xb4,0x53,0x09), after=2)
p("Der Kapitel-Header „Freitag, 2. März 2035, 06:12 Uhr“ (Marlies Heimweg) gehört laut Konsistenz-Prüfung "
  "eigentlich in DIESELBE Prolognacht — also müsste er „Donnerstag, 1. März“ heißen, nicht Freitag. "
  "Kalendarisch stimmt Freitag jetzt zwar, erzählerisch ist es aber dieselbe Nacht.", 10.5, after=2)
p("   → Deine Wahl: (a) so lassen (Freitag 2. März) oder (b) auf „Donnerstag, 1. März 06:12“ ziehen.",
  10.5, True, after=8)

# --- Anleitung ---
p("So nutzt du die Liste:", 11, True, after=2)
p("Jeder Fund hat eine Nummer. Sag mir einfach welche RAUS sollen (z. B. „5, 12, 30 weglassen“) — "
  "oder „alle übernehmen“. Der ✏️-Vorschlag ist jeweils die geplante Korrektur.", 10.5, after=8)

# --- Gruppen ---
GROUPS = [("datum","📅 Datum & Uhrzeit"), ("zahl","🔢 Zahlen & Zeitspannen"),
          ("ort","📍 Orte"), ("name","👤 Namen & Zuordnung")]
counter = 0
for typ, label in GROUPS:
    items = sorted([f for f in FUNDE if f["typ"] == typ], key=kap_sort)
    if not items: continue
    p("", 6)
    p(f"{label}  ({len(items)})", 14, True, (0x1a,0x1a,0x3a), after=3, before=4)
    for f in items:
        counter += 1
        kap = f.get("nr")
        kaplabel = "Prolog" if kap == 0 else (f"Kap {kap}" if kap is not None else "—")
        # Kopfzeile: Nummer + Kapitel
        par = doc.add_paragraph(); par.paragraph_format.space_after = Pt(0); par.paragraph_format.space_before = Pt(3)
        r1 = par.add_run(f"{counter}.  "); r1.bold = True; r1.font.name="Aptos"; r1.font.size=Pt(10.5)
        r2 = par.add_run(f"[{kaplabel}]  "); r2.bold = True; r2.font.name="Aptos"; r2.font.size=Pt(10.5); r2.font.color.rgb=RGBColor(0x2a,0x4a,0x8a)
        fund = (f.get("fund") or "").strip().replace("\n"," ")
        if len(fund) > 160: fund = fund[:160] + "…"
        r3 = par.add_run(fund); r3.italic = True; r3.font.name="Aptos"; r3.font.size=Pt(10)
        # Problem
        prob = (f.get("problem") or "").strip().replace("\n"," ")
        p("      Problem: " + prob, 10, after=0)
        # Vorschlag
        vor = (f.get("vorschlag") or "").strip().replace("\n"," ")
        pv, rv = p("      ✏️ " + vor, 10, after=0)
        rv.font.color.rgb = RGBColor(0x15,0x60,0x2d)

out = "/mnt/d/OneDrive/Desktop/AURORA_Konsistenz_Review.docx"
import os
try:
    doc.save(out)
    print("Gespeichert:", out, "|", os.path.getsize(out), "bytes |", counter, "Funde")
except Exception as e:
    alt = "/home/bolla/workspace/AURORA_Konsistenz_Review.docx"
    doc.save(alt); print("OneDrive-Desktop nicht erreichbar, gespeichert:", alt, "|", e)
