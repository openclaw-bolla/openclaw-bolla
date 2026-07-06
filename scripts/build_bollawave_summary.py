#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Baut die bollawave-Zusammenfassung als kompakte Word (Aptos, Ordner-Links)."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = "/mnt/d/OneDrive/Dokumente/Bolla/Suno_DistroKid/bollawave - Ueberblick.docx"

TEAL = RGBColor(0x0B, 0x76, 0x6E)
DARK = RGBColor(0x22, 0x22, 0x22)
GREY = RGBColor(0x55, 0x55, 0x55)
LINK = RGBColor(0x1A, 0x5F, 0xB4)

doc = Document()

# Standard-Font Aptos, kompakt
style = doc.styles["Normal"]
style.font.name = "Aptos"
style.font.size = Pt(10.5)
style.font.color.rgb = DARK
pf = style.paragraph_format
pf.space_before = Pt(0)
pf.space_after = Pt(0)
pf.line_spacing = 1.15

def set_aptos(run):
    run.font.name = "Aptos"
    r = run._element
    rpr = r.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(a), "Aptos")

def para(space_after=2, space_before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.15
    if align is not None:
        p.alignment = align
    return p

def run(p, text, size=10.5, bold=False, color=DARK, italic=False):
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = color
    set_aptos(r)
    return r

def add_hyperlink(p, url, text, color=LINK):
    part = p.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyper = OxmlElement("w:hyperlink")
    hyper.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(a), "Aptos")
    rpr.append(rfonts)
    col = OxmlElement("w:color"); col.set(qn("w:val"), "1A5FB4"); rpr.append(col)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rpr.append(u)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "21"); rpr.append(sz)
    new_run.append(rpr)
    t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text
    new_run.append(t)
    hyper.append(new_run)
    p._p.append(hyper)
    return hyper

def heading(text, emoji=""):
    p = para(space_after=3, space_before=8)
    run(p, f"{emoji} {text}".strip(), size=13, bold=True, color=TEAL)
    return p

def rule():
    p = para(space_after=4, space_before=2)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), "C9C9C9")
    pbdr.append(bottom); pPr.append(pbdr)

def bullet(text_parts, space_after=1):
    """text_parts: Liste von (text, bold, color) oder String."""
    p = para(space_after=space_after)
    p.paragraph_format.left_indent = Pt(12)
    run(p, "•  ", bold=True, color=TEAL)
    if isinstance(text_parts, str):
        text_parts = [(text_parts, False, DARK)]
    for tp in text_parts:
        if isinstance(tp, str):
            run(p, tp)
        else:
            txt, bold, color = tp
            run(p, txt, bold=bold, color=color)
    return p

# ---------- Titel ----------
p = para(space_after=1)
run(p, "🌊 bollawave", size=24, bold=True, color=TEAL)
p = para(space_after=8)
run(p, "Chris' Musik-Projekt mit Bolla — Überblick & Veröffentlichungs-Weg", size=11, color=GREY, italic=True)

# ---------- Was ist bollawave ----------
heading("Was ist bollawave?", "🎵")
p = para(space_after=3)
run(p, "bollawave", bold=True)
run(p, " ist Chris' Künstlername für die KI-Musik bei DistroKid / Spotify / Apple Music — "
        "das musikalische Pendant zum Buchprojekt AURORA (beides zusammen mit Bolla). "
        "Marken-Stimme: ")
run(p, "„Feel-Good mit Augenzwinkern.\"", bold=True, color=TEAL)
p = para(space_after=2)
run(p, "Vertrieb komplett über DistroKid. KI-Einsatz wird immer transparent deklariert.", color=GREY, size=9.5, italic=True)

# ---------- Wo liegt was ----------
heading("Wo liegt was?", "📁")
p = para(space_after=2)
run(p, "Master-Archiv (alle finalen Songs + Cover): ", bold=True)
p = para(space_after=3); p.paragraph_format.left_indent = Pt(12)
add_hyperlink(p, "file:///D:/OneDrive/Dokumente/Bolla/Suno_DistroKid/", "📂 Bolla\\Suno_DistroKid\\")

p = para(space_after=2)
run(p, "Aussortiert (off-brand / verworfen, soft-delete): ", bold=True)
p = para(space_after=3); p.paragraph_format.left_indent = Pt(12)
add_hyperlink(p, "file:///D:/OneDrive/Dokumente/Bolla/Suno_DistroKid/_aussortiert/", "📂 Suno_DistroKid\\_aussortiert\\")

p = para(space_after=2)
run(p, "Staging (pro Release ein Ordner auf dem Desktop — hochladen & danach löschen): ", bold=True)
p = para(space_after=3); p.paragraph_format.left_indent = Pt(12)
add_hyperlink(p, "file:///D:/OneDrive/Desktop/DistroKid/", "📂 Desktop\\DistroKid\\")

p = para(space_after=2)
run(p, "Upload-Checkliste (Feld-für-Feld für DistroKid): ", bold=True)
p = para(space_after=3); p.paragraph_format.left_indent = Pt(12)
add_hyperlink(p, "file:///D:/OneDrive/Dokumente/Bolla/Suno_DistroKid/DistroKid_Upload_Checkliste.docx", "📄 DistroKid_Upload_Checkliste.docx")

# ---------- Der Veröffentlichen-Knopf ----------
heading("Der „Veröffentlichen\"-Knopf macht jetzt alles", "🚀")
p = para(space_after=3)
run(p, "In ")
run(p, "Bolla Songs → „Song herunterladen + Cover erstellen\"", bold=True)
run(p, " sitzt die grüne Karte ")
run(p, "„🚀 Für DistroKid fertig machen\"", bold=True, color=TEAL)
run(p, ". Titel eintippen (wie in Suno), Sprache + Genre wählen, Liedtext einfügen, auf den Knopf — und der Server erledigt in einem Rutsch:")
bullet([("Aufräumen: ", True, DARK), ("bereits veröffentlichte Songs wandern vom Desktop ins Backup-Archiv.", False, DARK)])
bullet([("MP3 + Cover: ", True, DARK), ("holt den Song aus Suno, macht 320 kbps + textfreies 3000²-Cover (abschaltbar, falls schon vorhanden).", False, DARK)])
bullet([("Staging-Ordner: ", True, DARK), ("MP3, Cover, Lyrics.txt und eine DistroKid-Upload-Checkliste — alles im richtigen Format.", False, DARK)])
bullet([("Social-Paket: ", True, DARK), ("TikTok/Insta-Video (Hochkant, mit Ton), Captions und Poster-Anleitung.", False, DARK)])
bullet([("Lyrics-Scan: ", True, DARK), ("warnt automatisch, falls Namen/Schulbezug/Künstlername im Text stecken.", False, DARK)])
p = para(space_after=2, space_before=3)
run(p, "Du machst danach nur noch den DistroKid-Upload selbst (die Checkliste liegt im Ordner). Videos/Captions erst NACH dem Release posten.", italic=True, color=GREY, size=9.5)

# ---------- Katalog ----------
heading("Aktueller Katalog & Kalender", "📅")
bullet([("LIVE: ", True, TEAL), ("Pausenklingel-Magie (DE Rap) · Summertime's Callin' (EN Pop) · Läuft, Sommer! (DE) · Sun, Sand & Baluhai! (EN)", False, DARK)])
bullet([("Als Nächstes: ", True, DARK), ("„Mein Hund Hat Meine Hausaufgaben Gefressen\" (DE) — Slot 10.07.", False, DARK)])
bullet([("Geparkt (off-brand): ", True, GREY), ("Seven More Weeks, KI-Wunder, Blazing Waves, Die Zukunft ist Jetzt, AI Symphony — passen nicht ins Feel-Good-Profil.", False, GREY)])

# ---------- Ausblick ----------
heading("Songs einzeln — wann du Bock hast", "🎈")
p = para(space_after=2)
run(p, "Kein Batch, kein Zeitplan-Zwang: Wenn dir ein Feel-Good-Song einfällt, machst du ihn in Suno fertig und "
        "drückst den grünen Knopf — der Rest passiert von selbst. So oft oder selten du magst.")
p = para(space_after=2)
run(p, "Ab 05.08. geht's wieder los (davor Urlaub, und Suno ist bis dahin pausiert).", color=GREY, size=9.5, italic=True)

# Footer
rule()
p = para(space_before=2)
run(p, "Erstellt von Bolla 🐾 · 06.07.2026 · bollawave = Chris + Bolla", size=8.5, color=GREY, italic=True)

doc.save(OUT)
print("Gespeichert:", OUT)
