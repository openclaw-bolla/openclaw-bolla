#!/usr/bin/env python3
"""
Mission Control API Server
Liefert Kalender, E-Mail und andere Daten an Mission Control (localhost:18790)
"""

import json
import os
import sys
import subprocess
import traceback
import urllib.request
import urllib.parse
from pathlib import Path
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from datetime import datetime, timezone, timedelta

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

# Sicherstellen dass ~/.local/bin im PATH ist (fehlt im Cron-Job-Kontext)
_local_bin = os.path.expanduser("~/.local/bin")
if _local_bin not in os.environ.get("PATH", ""):
    os.environ["PATH"] = _local_bin + ":" + os.environ.get("PATH", "")

CLAUDE_BIN = os.path.expanduser("~/.local/bin/claude")

WORKSPACE = os.path.expanduser("~/workspace")
TOKEN_FILE = os.path.join(WORKSPACE, "config/ms_token.json")
AZURE_SPEECH_FILE = os.path.join(WORKSPACE, "config/azure_speech.json")
CLIPBOARD_FILE = os.path.join(WORKSPACE, "config/clipboard.json")
CLIPBOARD_TRASH_FILE = os.path.join(WORKSPACE, "config/clipboard_trash.json")
CLIPBOARD_TRASH_DAYS = 14  # so lange bleiben gelöschte Einträge im Papierkorb
CLIPBOARD_IMAGES_DIR = os.path.join(WORKSPACE, "config/clipboard_images")
IMMO_BOOKMARKS_FILE  = Path(os.path.join(WORKSPACE, "config/immo_bookmarks.json"))
IMMO_CRITERIA_FILE   = Path(os.path.join(WORKSPACE, "config/immo_criteria.json"))
TRAVEL_FILE          = Path(os.path.join(WORKSPACE, "cache/travel.json"))
PHOTO_ANALYSIS_FILE  = os.path.join(WORKSPACE, "data/photo_analysis.json")
KORREKTUR_DIR        = os.path.join(WORKSPACE, "korrektur")

# Foto-Analyse Job-Status (global, threadsafe via GIL für einfache dict-ops)
_photo_job = {"running": False, "total": 0, "done": 0, "errors": 0, "stop": False, "folder": ""}

# Glücksrad Lehrer-Lösung (in-memory, reset bei Server-Neustart)
_gluecksrad_state = {"stil": None, "nummer": None}

# KI-Buch Async-Job (Hintergrund-Thread für Claude-Aufruf)
_ki_buch_job = {"status": "idle", "antwort": "", "inhalt": "", "inhalt_titel": "", "error": ""}

def get_clipboard():
    try:
        with open(CLIPBOARD_FILE, encoding="utf-8") as f:
            raw = json.load(f)
        # Migration: altes Format {"text":..., "ts":...} → neue Struktur
        if "entries" not in raw:
            entries = [{"text": raw["text"], "ts": raw.get("ts", ""), "source": "manual"}] if raw.get("text") else []
            return {"entries": entries}
        return raw
    except Exception:
        return {"entries": []}

def save_clipboard(text):
    data = {"entries": [{"text": text, "ts": datetime.now().isoformat(), "source": "manual"}]}
    with open(CLIPBOARD_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False)
    return data

def append_clipboard(text, source="voice"):
    try:
        with open(CLIPBOARD_FILE, encoding="utf-8") as f:
            raw = json.load(f)
        if "entries" not in raw:
            entries = [{"text": raw["text"], "ts": raw.get("ts", ""), "source": "manual"}] if raw.get("text") else []
        else:
            entries = raw["entries"]
    except Exception:
        entries = []
    entries.append({"text": text, "ts": datetime.now().isoformat(), "source": source})
    data = {"entries": entries}
    with open(CLIPBOARD_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False)
    return data

def _load_clipboard_trash():
    try:
        with open(CLIPBOARD_TRASH_FILE, encoding="utf-8") as f:
            raw = json.load(f)
        return raw.get("entries", [])
    except Exception:
        return []

def _save_clipboard_trash(entries):
    with open(CLIPBOARD_TRASH_FILE, "w", encoding="utf-8") as f:
        json.dump({"entries": entries}, f, ensure_ascii=False)

def _purge_clipboard_trash(entries):
    """Entfernt Papierkorb-Einträge, die älter als CLIPBOARD_TRASH_DAYS sind."""
    cutoff = datetime.now() - timedelta(days=CLIPBOARD_TRASH_DAYS)
    kept = []
    for e in entries:
        try:
            if datetime.fromisoformat(e.get("deleted_ts", "")) >= cutoff:
                kept.append(e)
        except Exception:
            kept.append(e)  # ohne gültigen Zeitstempel lieber behalten
    return kept

def delete_clipboard_entry(idx):
    """Soft-Delete: Eintrag wandert in den Papierkorb statt sofort verloren zu gehen."""
    try:
        with open(CLIPBOARD_FILE, encoding="utf-8") as f:
            raw = json.load(f)
        entries = raw.get("entries", [])
        if 0 <= idx < len(entries):
            removed = entries.pop(idx)
            removed = dict(removed)
            removed["deleted_ts"] = datetime.now().isoformat()
            trash = _purge_clipboard_trash(_load_clipboard_trash())
            trash.append(removed)
            _save_clipboard_trash(trash)
        data = {"entries": entries}
        with open(CLIPBOARD_FILE, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False)
        return data
    except Exception as e:
        return {"error": str(e)}

def get_clipboard_trash():
    """Liefert den (automatisch bereinigten) Papierkorb, neueste zuerst."""
    trash = _purge_clipboard_trash(_load_clipboard_trash())
    _save_clipboard_trash(trash)
    return {"entries": trash, "retention_days": CLIPBOARD_TRASH_DAYS}

def restore_clipboard_entry(idx):
    """Holt einen Eintrag aus dem Papierkorb zurück ins Clipboard."""
    try:
        trash = _purge_clipboard_trash(_load_clipboard_trash())
        if not (0 <= idx < len(trash)):
            return {"error": "Index ungültig"}
        item = dict(trash.pop(idx))
        item.pop("deleted_ts", None)
        _save_clipboard_trash(trash)
        try:
            with open(CLIPBOARD_FILE, encoding="utf-8") as f:
                raw = json.load(f)
            entries = raw.get("entries", [])
        except Exception:
            entries = []
        entries.append(item)
        with open(CLIPBOARD_FILE, "w", encoding="utf-8") as f:
            json.dump({"entries": entries}, f, ensure_ascii=False)
        return {"entries": entries}
    except Exception as e:
        return {"error": str(e)}

def delete_clipboard_trash_entry(idx):
    """Löscht einen Eintrag endgültig aus dem Papierkorb."""
    try:
        trash = _purge_clipboard_trash(_load_clipboard_trash())
        if 0 <= idx < len(trash):
            trash.pop(idx)
        _save_clipboard_trash(trash)
        return {"entries": trash}
    except Exception as e:
        return {"error": str(e)}

def _ensure_clipboard_images_dir():
    os.makedirs(CLIPBOARD_IMAGES_DIR, exist_ok=True)

def get_clipboard_images():
    _ensure_clipboard_images_dir()
    ALLOWED_EXT = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp"}
    images = []
    for fname in sorted(os.listdir(CLIPBOARD_IMAGES_DIR)):
        ext = os.path.splitext(fname)[1].lower()
        if ext not in ALLOWED_EXT:
            continue
        fpath = os.path.join(CLIPBOARD_IMAGES_DIR, fname)
        stat = os.stat(fpath)
        images.append({"filename": fname, "ts": datetime.fromtimestamp(stat.st_mtime).isoformat(), "size": stat.st_size})
    images.sort(key=lambda x: x["ts"], reverse=True)
    return {"images": images}

def save_clipboard_image(data_b64, mime):
    import base64, uuid
    _ensure_clipboard_images_dir()
    EXT_MAP = {"image/png": ".png", "image/jpeg": ".jpg", "image/jpg": ".jpg",
               "image/gif": ".gif", "image/webp": ".webp", "image/bmp": ".bmp"}
    ext = EXT_MAP.get(mime, ".png")
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    fname = f"clip_{ts}_{uuid.uuid4().hex[:6]}{ext}"
    fpath = os.path.join(CLIPBOARD_IMAGES_DIR, fname)
    with open(fpath, "wb") as f:
        f.write(base64.b64decode(data_b64))
    return {"filename": fname, "ts": datetime.now().isoformat()}

def delete_clipboard_image(filename):
    _ensure_clipboard_images_dir()
    # Sicherheitscheck: nur einfacher Dateiname, kein Pfad-Traversal
    if "/" in filename or "\\" in filename or ".." in filename:
        return {"error": "Ungültiger Dateiname"}
    fpath = os.path.join(CLIPBOARD_IMAGES_DIR, filename)
    if os.path.isfile(fpath):
        os.remove(fpath)
        return {"ok": True}
    return {"error": "Nicht gefunden"}


def azure_speech_config():
    """Lädt Azure Speech Config. None wenn nicht konfiguriert."""
    try:
        with open(AZURE_SPEECH_FILE) as f:
            cfg = json.load(f)
        if not cfg.get("key") or cfg["key"].startswith("REPLACE_"):
            return None
        return cfg
    except Exception:
        return None

GEMINI_API_FILE = os.path.join(WORKSPACE, "config/gemini_api.json")


def gemini_api_key():
    try:
        with open(GEMINI_API_FILE) as f:
            cfg = json.load(f)
        key = cfg.get("api_key", "")
        if not key or key.startswith("REPLACE_"):
            return None
        return key
    except Exception:
        return None


BILDGEN_LIMIT = 100
_bildgen_counter = {"date": "", "count": 0}

def bildgen_check_limit():
    today = datetime.now().strftime("%Y-%m-%d")
    if _bildgen_counter["date"] != today:
        _bildgen_counter["date"] = today
        _bildgen_counter["count"] = 0
    if _bildgen_counter["count"] >= BILDGEN_LIMIT:
        return False, f"Tageslimit von {BILDGEN_LIMIT} Bildern erreicht. Morgen wieder möglich."
    return True, None

def bildgen_generate(prompt, model="gemini-2.5-flash-image", aspect_ratio="1:1",
                     input_image_b64=None, input_mime_type="image/png"):
    import urllib.request, urllib.error, base64
    key = gemini_api_key()
    if not key:
        return None, "Kein Gemini API-Key konfiguriert. Bitte in config/gemini_api.json eintragen."
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent"
    parts = []
    if input_image_b64:
        parts.append({"inlineData": {"mimeType": input_mime_type, "data": input_image_b64}})
    parts.append({"text": prompt})
    gen_cfg = {"responseModalities": ["TEXT", "IMAGE"]}
    if not input_image_b64:
        gen_cfg["imageConfig"] = {"aspectRatio": aspect_ratio}
    payload = json.dumps({
        "contents": [{"parts": parts}],
        "generationConfig": gen_cfg
    }).encode()
    req = urllib.request.Request(url, data=payload,
                                  headers={"x-goog-api-key": key, "Content-Type": "application/json"})
    try:
        with urllib.request.urlopen(req, timeout=60) as resp:
            data = json.loads(resp.read())
    except urllib.error.HTTPError as e:
        msg = e.read().decode("utf-8", errors="replace")
        try:
            msg = json.loads(msg).get("error", {}).get("message", msg)
        except Exception:
            pass
        return None, f"API-Fehler: {msg}"
    except Exception as e:
        return None, str(e)
    for part in data.get("candidates", [{}])[0].get("content", {}).get("parts", []):
        if "inlineData" in part:
            img_b64 = part["inlineData"]["data"]
            mime = part["inlineData"].get("mimeType", "image/png")
            return img_b64, mime
    return None, "Kein Bild in der API-Antwort erhalten."


CLIENT_ID = "9e5f94bc-e8a4-4e73-b8be-63364c29d753"

_lms_host_cache = None
def lms_base_url():
    """Findet LM Studio auf dem Windows-Host (WSL2-Default-Gateway), Port 1234."""
    global _lms_host_cache
    if _lms_host_cache is None:
        try:
            with open("/proc/net/route") as f:
                for line in f.readlines()[1:]:
                    parts = line.split()
                    if parts[1] == "00000000":
                        ip = ".".join(str(int(parts[2][i:i+2], 16)) for i in (6, 4, 2, 0))
                        _lms_host_cache = ip
                        break
        except Exception:
            _lms_host_cache = "172.20.96.1"
    return f"http://{_lms_host_cache}:1234"

def get_token():
    """Gibt einen gültigen MS Graph Access Token zurück (refresht bei Bedarf)."""
    try:
        with open(TOKEN_FILE) as f:
            token_data = json.load(f)

        # Immer refreshen — Access Tokens laufen schnell ab
        import urllib.request, urllib.parse
        data = urllib.parse.urlencode({
            "client_id": CLIENT_ID,
            "grant_type": "refresh_token",
            "refresh_token": token_data["refresh_token"],
            "scope": "https://graph.microsoft.com/Mail.Read https://graph.microsoft.com/Mail.ReadWrite https://graph.microsoft.com/Calendars.Read https://graph.microsoft.com/Contacts.Read offline_access"
        }).encode()
        req = urllib.request.Request(
            "https://login.microsoftonline.com/common/oauth2/v2.0/token",
            data=data, method="POST",
            headers={"Content-Type": "application/x-www-form-urlencoded"},
        )
        with urllib.request.urlopen(req, timeout=10) as r:
            new_token = json.loads(r.read())
        with open(TOKEN_FILE, "w") as f:
            json.dump(new_token, f, indent=2)
        return new_token["access_token"]
    except Exception as e:
        print(f"Token error: {e}")
        # Fallback: alten Token versuchen
        try:
            with open(TOKEN_FILE) as f:
                return json.load(f).get("access_token")
        except:
            return None

def graph_get(path):
    import urllib.request, urllib.parse
    token = get_token()
    if not token:
        return None
    # URL-Teile splitten und Query-Parameter korrekt encodieren
    if '?' in path:
        base, query = path.split('?', 1)
        # Query-Parameter einzeln URL-encodieren (Leerzeichen → %20)
        encoded_query = urllib.parse.quote(query, safe='=&$,/%')
        full_url = f"https://graph.microsoft.com/v1.0{base}?{encoded_query}"
    else:
        full_url = f"https://graph.microsoft.com/v1.0{path}"
    req = urllib.request.Request(
        full_url,
        headers={"Authorization": f"Bearer {token}", "Content-Type": "application/json"}
    )
    try:
        with urllib.request.urlopen(req, timeout=10) as r:
            return json.loads(r.read())
    except Exception as e:
        print(f"Graph error {path}: {e}")
        return None

PRESET_COLORS = {
    "preset0":  "229,57,53",    # Rot
    "preset1":  "230,81,0",     # Orange
    "preset2":  "121,85,72",    # Braun
    "preset3":  "249,168,37",   # Gelb
    "preset4":  "67,160,71",    # Grün
    "preset5":  "0,172,193",    # Türkis
    "preset6":  "142,155,0",    # Olive
    "preset7":  "30,136,229",   # Blau
    "preset8":  "94,53,177",    # Lila
    "preset9":  "194,24,91",    # Cranberry/Pink
    "preset10": "84,110,122",   # Stahl
    "preset11": "55,71,79",     # Dunkelstahl
    "preset12": "117,117,117",  # Grau
    "preset13": "66,66,66",     # Dunkelgrau
    "preset14": "33,33,33",     # Schwarz
    "preset15": "183,28,28",    # Dunkelrot
    "preset16": "191,54,12",    # Dunkelorange
    "preset17": "78,52,46",     # Dunkelbraun
    "preset18": "245,127,23",   # Dunkelgelb
    "preset19": "27,94,32",     # Dunkelgrün
    "preset20": "0,96,100",     # Dunkeltürkis
    "preset21": "130,119,23",   # Dunkelolive
    "preset22": "13,71,161",    # Dunkelblau
    "preset23": "74,20,140",    # Dunkellila
    "preset24": "136,14,79",    # Dunkelcranberry
    "none":     "120,144,156",  # Fallback Grau
}

def get_category_colors():
    data = graph_get("/me/outlook/masterCategories")
    if not data:
        return {}
    result = {}
    for cat in data.get("value", []):
        name = cat.get("displayName", "")
        preset = cat.get("color", "none")
        result[name] = PRESET_COLORS.get(preset, PRESET_COLORS["none"])
    return result

def get_calendar():
    now = datetime.now(timezone.utc)
    end = now + timedelta(days=30)
    start_str = now.strftime("%Y-%m-%dT%H:%M:%SZ")
    end_str = end.strftime("%Y-%m-%dT%H:%M:%SZ")

    cat_colors = get_category_colors()

    data = graph_get(
        f"/me/calendarView?startDateTime={start_str}&endDateTime={end_str}"
        f"&$orderby=start/dateTime&$top=20"
        f"&$select=subject,start,end,categories,location"
    )
    if not data:
        return []

    events = []
    for ev in data.get("value", []):
        start_raw = ev.get("start", {}).get("dateTime", "")
        try:
            import zoneinfo
            berlin = zoneinfo.ZoneInfo("Europe/Berlin")
            dt_utc = datetime.fromisoformat(start_raw).replace(tzinfo=timezone.utc)
            dt_local = dt_utc.astimezone(berlin)
            date_str = dt_local.strftime("%d.%m.")
            time_str = dt_local.strftime("%H:%M")
            weekday = ["Mo", "Di", "Mi", "Do", "Fr", "Sa", "So"][dt_local.weekday()]
        except:
            date_str = start_raw[:10]
            time_str = ""
            weekday = ""

        cats = ev.get("categories", [])
        cat = cats[0] if cats else ""
        rgb = cat_colors.get(cat, PRESET_COLORS["none"])

        events.append({
            "title": ev.get("subject", ""),
            "date": date_str,
            "time": time_str,
            "weekday": weekday,
            "category": cat,
            "color": rgb,
            "location": ev.get("location", {}).get("displayName", "")
        })

    return events

def search_calendar_events(q, limit=5):
    from datetime import timezone, timedelta
    import zoneinfo, urllib.parse as _up
    berlin = zoneinfo.ZoneInfo("Europe/Berlin")
    cat_colors = get_category_colors()
    now_utc = datetime.now(timezone.utc)
    q_lower = q.lower()

    past2y = (now_utc - timedelta(days=730)).strftime("%Y-%m-%dT%H:%M:%SZ")
    fut3m  = (now_utc + timedelta(days=90)).strftime("%Y-%m-%dT%H:%M:%SZ")

    all_events = []
    url = (
        f"/me/calendarView?startDateTime={past2y}&endDateTime={fut3m}"
        f"&$top=200&$select=subject,start,categories,location"
    )
    pages = 0
    while url and pages < 3:
        data = graph_get(url)
        if not data:
            break
        all_events.extend(data.get("value", []))
        nxt = data.get("@odata.nextLink", "")
        # nextLink ist absolute URL — in relativen Pfad umwandeln
        if nxt:
            parsed = _up.urlparse(nxt)
            url = parsed.path + ("?" + parsed.query if parsed.query else "")
        else:
            url = ""
        pages += 1

    results = []
    for ev in all_events:
        subj = ev.get("subject", "")
        if q_lower not in subj.lower():
            continue
        start_raw = ev.get("start", {}).get("dateTime", "")
        try:
            dt_utc = datetime.fromisoformat(start_raw).replace(tzinfo=timezone.utc)
            dt_local = dt_utc.astimezone(berlin)
            date_str = dt_local.strftime("%d.%m.%Y")
            time_str = dt_local.strftime("%H:%M")
            weekday = ["Mo", "Di", "Mi", "Do", "Fr", "Sa", "So"][dt_local.weekday()]
            sort_key = dt_utc
        except Exception:
            date_str = start_raw[:10]
            time_str = ""
            weekday = ""
            sort_key = datetime.min.replace(tzinfo=timezone.utc)

        cats = ev.get("categories", [])
        cat = cats[0] if cats else ""
        rgb = cat_colors.get(cat, PRESET_COLORS["none"])
        past = sort_key < now_utc

        results.append({
            "title": subj,
            "date": date_str,
            "time": time_str,
            "weekday": weekday,
            "category": cat,
            "color": rgb,
            "location": ev.get("location", {}).get("displayName", ""),
            "past": past,
            "_sort": sort_key.timestamp()
        })

    # Vergangene: neueste zuerst — Zukünftige: nächste zuerst
    past_sorted   = sorted([r for r in results if r["past"]],      key=lambda x: x["_sort"], reverse=True)
    future_sorted = sorted([r for r in results if not r["past"]], key=lambda x: x["_sort"])
    ordered = past_sorted[:limit] + future_sorted[:limit]
    for r in ordered:
        del r["_sort"]
    return ordered[:limit]

SPAM_SUBJECT = ["newsletter", "unsubscribe", "abbestellen", "rabatt", "sonderangebot",
                "gutschein", "gewinnspiel", "angebot des tages", "% off", "% rabatt",
                "jetzt kaufen", "nur heute", "limited offer", "act now",
                "*** spam", "[spam]", "critical notice", "cloud storage plan",
                "your account", "verify your", "suspended", "unusual activity"]
SPAM_SENDER  = ["newsletter", "noreply", "no-reply", "donotreply", "mailer-daemon",
                "marketing@", "notification@", "bounce@", "alerts@", "info@newsletter"]

def _is_spam(subject, from_email):
    import re
    s = subject.lower()
    e = (from_email or "").lower()
    # Leet-speak Normalisierung (0→o, 1→l, 3→e)
    s_norm = s.translate(str.maketrans("013", "ole"))
    return (any(k in s for k in SPAM_SUBJECT)
            or any(k in s_norm for k in SPAM_SUBJECT)
            or any(k in e for k in SPAM_SENDER)
            or not e  # kein Absender = Spam
            )

def _parse_graph_mail(m, account="Outlook"):
    received = m.get("receivedDateTime", "")
    try:
        dt = datetime.fromisoformat(received.replace("Z", "+00:00"))
        dt_local = dt.astimezone(timezone(timedelta(hours=2)))
        now_local = datetime.now(timezone(timedelta(hours=2)))
        delta = (now_local.date() - dt_local.date()).days
        if delta == 0:   date_str = "Heute " + dt_local.strftime("%H:%M")
        elif delta == 1: date_str = "Gestern " + dt_local.strftime("%H:%M")
        elif delta < 7:  date_str = dt_local.strftime("%a %H:%M")
        else:            date_str = dt_local.strftime("%d.%m.")
        is_today = delta == 0
    except Exception:
        date_str = received[:16]; is_today = False
    return {
        "id": m.get("id", ""),
        "account": account,
        "from": m.get("from", {}).get("emailAddress", {}).get("name", "Unbekannt"),
        "from_email": m.get("from", {}).get("emailAddress", {}).get("address", ""),
        "subject": m.get("subject", "(kein Betreff)"),
        "date": date_str,
        "is_today": is_today,
        "isRead": m.get("isRead", True),
        "preview": m.get("bodyPreview", "")[:400]
    }

def get_emails_outlook():
    data = graph_get(
        "/me/mailFolders/inbox/messages?$filter=isRead%20eq%20false"
        "&$top=10&$select=id,subject,from,receivedDateTime,isRead,bodyPreview"
    )
    if not data: return []
    msgs = []
    for m in data.get("value", []):
        entry = _parse_graph_mail(m)
        if not _is_spam(entry["subject"], entry["from_email"]):
            msgs.append(entry)
    return msgs

def get_emails_recent_outlook():
    data = graph_get(
        "/me/mailFolders/inbox/messages?$top=15"
        "&$select=id,subject,from,receivedDateTime,isRead,bodyPreview"
    )
    if not data: return []
    msgs = []
    for m in data.get("value", []):
        entry = _parse_graph_mail(m)
        if not _is_spam(entry["subject"], entry["from_email"]):
            msgs.append(entry)
        if len(msgs) >= 6: break
    return msgs

def get_emails_sent_outlook():
    """Holt die letzten 5 gesendeten Mails."""
    data = graph_get(
        "/me/mailFolders/sentitems/messages"
        "?$top=5"
        "&$select=subject,toRecipients,sentDateTime,bodyPreview"
    )
    if not data:
        return []
    msgs = []
    for m in data.get("value", []):
        sent = m.get("sentDateTime", "")
        try:
            dt = datetime.fromisoformat(sent.replace("Z", "+00:00"))
            dt_local = dt.astimezone(timezone(timedelta(hours=2)))
            date_str = dt_local.strftime("%d.%m. %H:%M")
        except:
            date_str = sent[:16]
        recipients = m.get("toRecipients", [])
        to_name = recipients[0].get("emailAddress", {}).get("name", "Unbekannt") if recipients else "Unbekannt"
        msgs.append({
            "to": to_name,
            "subject": m.get("subject", "(kein Betreff)"),
            "date": date_str,
        })
    return msgs


def get_emails_wtnet():
    """Holt ungelesene Mails von chrismandel@wtnet.de via IMAP."""
    import imaplib, email as email_lib
    from email.header import decode_header

    try:
        wtnet_cfg = os.path.join(WORKSPACE, "config/wtnet_account.json")
        with open(wtnet_cfg) as f:
            cfg = json.load(f)

        mail = imaplib.IMAP4_SSL(cfg["imap_host"], cfg["imap_port"])
        mail.login(cfg["email"], cfg["password"])
        mail.select("INBOX")

        _, msg_ids = mail.search(None, "UNSEEN")
        ids = msg_ids[0].split()
        msgs = []

        for mid in reversed(ids[-10:]):  # max 10, neueste zuerst
            _, data = mail.fetch(mid, "(RFC822)")
            raw = data[0][1]
            msg = email_lib.message_from_bytes(raw)

            # Betreff dekodieren
            subj_raw = msg.get("Subject", "(kein Betreff)")
            subj_parts = decode_header(subj_raw)
            subj = ""
            for part, enc in subj_parts:
                if isinstance(part, bytes):
                    subj += part.decode(enc or "utf-8", errors="replace")
                else:
                    subj += part

            # Absender dekodieren
            from_raw = msg.get("From", "Unbekannt")
            from_parts = decode_header(from_raw)
            from_str = ""
            for part, enc in from_parts:
                if isinstance(part, bytes):
                    from_str += part.decode(enc or "utf-8", errors="replace")
                else:
                    from_str += part
            # Nur Name ohne E-Mail-Adresse
            from_name = from_str.split("<")[0].strip().strip('"')

            # Datum
            date_raw = msg.get("Date", "")
            try:
                from email.utils import parsedate_to_datetime
                dt = parsedate_to_datetime(date_raw)
                dt_local = dt.astimezone(timezone(timedelta(hours=2)))
                date_str = dt_local.strftime("%d.%m. %H:%M")
            except:
                date_str = date_raw[:16]

            # Spam überspringen
            if "*** SPAM" in subj.upper() or "[SPAM]" in subj.upper():
                continue

            msgs.append({
                "account": "wtnet",
                "from": from_name or "Unbekannt",
                "from_email": cfg["email"],
                "subject": subj,
                "date": date_str,
                "preview": ""
            })

        mail.logout()
        return msgs

    except Exception as e:
        print(f"wtnet IMAP Fehler: {e}")
        return []


def get_emails_recent_wtnet():
    """Holt die letzten Mails von wtnet (gelesen + ungelesen)."""
    import imaplib, email as email_lib
    from email.header import decode_header
    try:
        wtnet_cfg = os.path.join(WORKSPACE, "config/wtnet_account.json")
        with open(wtnet_cfg) as f:
            cfg = json.load(f)
        mail = imaplib.IMAP4_SSL(cfg["imap_host"], cfg["imap_port"])
        mail.login(cfg["email"], cfg["password"])
        mail.select("INBOX", readonly=True)
        _, uid_data = mail.uid("search", None, "ALL")
        ids = uid_data[0].split()
        msgs = []
        for uid in reversed(ids[-10:]):
            _, data = mail.uid("fetch", uid, "(RFC822 FLAGS)")
            raw_data = data[0]
            flags = data[0][0] if isinstance(data[0], tuple) else b""
            raw = raw_data[1] if isinstance(raw_data, tuple) else None
            if not raw:
                continue
            msg = email_lib.message_from_bytes(raw)
            subj_raw = msg.get("Subject", "(kein Betreff)")
            subj_parts = decode_header(subj_raw)
            subj = "".join(p.decode(enc or "utf-8", errors="replace") if isinstance(p, bytes) else p for p, enc in subj_parts)
            from_raw = msg.get("From", "Unbekannt")
            from_parts = decode_header(from_raw)
            from_str = "".join(p.decode(enc or "utf-8", errors="replace") if isinstance(p, bytes) else p for p, enc in from_parts)
            from_name = from_str.split("<")[0].strip().strip('"')
            date_raw = msg.get("Date", "")
            try:
                from email.utils import parsedate_to_datetime
                dt = parsedate_to_datetime(date_raw).astimezone(timezone(timedelta(hours=2)))
                date_str = dt.strftime("%d.%m. %H:%M")
                is_today = dt.date() == datetime.now().date()
            except Exception:
                date_str, is_today = date_raw[:16], False
            is_read = b"\\Seen" in (flags if isinstance(flags, bytes) else b"")
            if "*** SPAM" in subj.upper() or "[SPAM]" in subj.upper():
                continue
            msgs.append({"account": "wtnet", "id": uid.decode(),
                         "from": from_name or "Unbekannt",
                         "subject": subj, "date": date_str, "is_today": is_today,
                         "isRead": is_read, "preview": ""})
        mail.logout()
        return msgs
    except Exception as e:
        print(f"wtnet recent Fehler: {e}")
        return []


NEWSLETTER_WATCHLIST = Path(os.path.join(WORKSPACE, "config/newsletter_watchlist.json"))
NEWSLETTER_RESULTS   = Path(os.path.join(WORKSPACE, "cache/newsletter_results.json"))

MARKTGURU_KEY = "8Kk+pmbf7TgJ9nVj2cXeA7P5zBGv8iuutVVMRfOfvNE="
MARKTGURU_ZIP_FILE = Path(os.path.join(WORKSPACE, "config/marktguru_zip.txt"))

def _marktguru_zip():
    return MARKTGURU_ZIP_FILE.read_text().strip() if MARKTGURU_ZIP_FILE.exists() else "22844"

def offers_search(q, limit=20):
    import urllib.request as _ur
    zip_code = _marktguru_zip()
    url = (f"https://api.marktguru.de/api/v1/offers/search"
           f"?as=web&limit={limit}&offset=0&q={urllib.parse.quote(q)}&zipCode={zip_code}")
    req = _ur.Request(url, headers={
        "x-apikey": MARKTGURU_KEY,
        "Accept": "application/json",
        "User-Agent": "Mozilla/5.0",
    })
    try:
        with _ur.urlopen(req, timeout=8) as resp:
            data = json.loads(resp.read())
    except Exception as e:
        return {"error": str(e), "results": []}
    out = []
    for o in data.get("results", []):
        vd = o.get("validityDates", [{}])
        valid_from = vd[0].get("from", "")[:10] if vd else ""
        valid_to   = vd[0].get("to",   "")[:10] if vd else ""
        store = o["advertisers"][0]["name"] if o.get("advertisers") else "?"
        out.append({
            "store":      store,
            "brand":      o.get("brand", {}).get("name", "") if o.get("brand") else "",
            "name":       o.get("product", {}).get("name") or o.get("description", ""),
            "desc":       o.get("description", ""),
            "price":      o.get("price"),
            "old_price":  o.get("oldPrice"),
            "unit":       o.get("unit", {}).get("shortName", "") if o.get("unit") else "",
            "valid_from": valid_from,
            "valid_to":   valid_to,
        })
    return {"results": out, "total": data.get("filters", {}).get("retailers"), "zip": zip_code}

def offers_set_zip(zip_code):
    z = zip_code.strip()
    if not z.isdigit() or len(z) != 5:
        return {"error": "Ungültige PLZ"}
    MARKTGURU_ZIP_FILE.write_text(z)
    return {"ok": True, "zip": z}

MAKLER_FILE = Path(os.path.join(WORKSPACE, "data/makler_status.json"))

def get_travel():
    if TRAVEL_FILE.exists():
        return json.loads(TRAVEL_FILE.read_text(encoding="utf-8"))
    return _travel_default()

def _travel_default():
    return {
        "urlaube": [
            {
                "id": "sommer2026",
                "titel": "Sommer Badeurlaub",
                "emoji": "🌊",
                "von": "2026-07-11",
                "bis": "2026-07-30",
                "naechte": 19,
                "reisende": 2,
                "status": "geplant",
                "danach": {
                    "von": "2026-07-31",
                    "bis": "2026-08-09",
                    "titel": "Steffi & Enkel",
                    "ort": "Kaufering",
                    "adresse": "Thomas-Morus-Str. 12b, 86916 Kaufering"
                },
                "kriterien_pauschal": [
                    "Flug ab Hamburg (HAM)",
                    "mind. 4 Sterne",
                    "Halbpension",
                    "Pool + kostenlose Liegen & Schirmen",
                    "Warm & sonnig (Türkei, Griechenland, Ägypten...)",
                    "Bestes Preis-Leistungs-Verhältnis"
                ],
                "kriterien_eigenanreise": [
                    "VW ID3 1st Edition (≈350 km Reichweite)",
                    "Hotel mit Badeurlaub-Charakter",
                    "Warm, Schwimmen möglich",
                    "Evtl. Pool",
                    "Südroute → Heimweg über Kaufering ideal",
                    "Empfohlen: Gardasee, Bodensee, Adria"
                ],
                "pauschal": {"empfehlung": None, "aktualisiert": None},
                "eigenanreise": {"empfehlung": None, "aktualisiert": None}
            }
        ]
    }

def save_travel_recommendation(urlaub_id, typ, data):
    travel = get_travel()
    for u in travel.get("urlaube", []):
        if u["id"] == urlaub_id:
            u[typ] = {
                "empfehlung": data,
                "aktualisiert": datetime.now().isoformat()
            }
            break
    TRAVEL_FILE.write_text(json.dumps(travel, ensure_ascii=False, indent=2))
    return {"ok": True}

def get_makler():
    if not MAKLER_FILE.exists():
        return {"regionen": [], "statuses": {}}
    data = json.loads(MAKLER_FILE.read_text())
    # Emails von Maklern aus Outlook prüfen (letzten 30 Tage)
    emails = []
    try:
        all_emails = [m["email"] for r in data["regionen"]
                      for m in r["makler"] if m.get("email")]
        from mission_control_api import get_token as _gt
        token = _gt()
        import urllib.request as _ur
        req = _ur.Request(
            "https://graph.microsoft.com/v1.0/me/messages"
            "?$filter=receivedDateTime ge " +
            (datetime.utcnow() - timedelta(days=30)).strftime("%Y-%m-%dT00:00:00Z") +
            "&$select=id,subject,from,receivedDateTime,isRead"
            "&$top=50&$orderby=receivedDateTime desc",
            headers={"Authorization": f"Bearer {token}"}
        )
        with _ur.urlopen(req, timeout=15) as r:
            msgs = json.loads(r.read()).get("value", [])
        for m in msgs:
            addr = m.get("from", {}).get("emailAddress", {}).get("address", "").lower()
            if any(addr == e.lower() or addr.split("@")[-1] in e.lower()
                   for e in all_emails if e):
                emails.append({
                    "id":       m["id"],
                    "subject":  m.get("subject", ""),
                    "from":     addr,
                    "date":     m.get("receivedDateTime", "")[:10],
                    "isRead":   m.get("isRead", True),
                })
    except Exception as e:
        emails = []
    data["inbox"] = emails
    return data

def makler_set_status(makler_id, status, notiz=""):
    data = json.loads(MAKLER_FILE.read_text()) if MAKLER_FILE.exists() else {"regionen": [], "statuses": {}}
    data.setdefault("statuses", {})[makler_id] = {
        "status": status,
        "notiz":  notiz,
        "datum":  datetime.now().strftime("%d.%m.%Y"),
    }
    MAKLER_FILE.write_text(json.dumps(data, ensure_ascii=False, indent=2))
    return {"ok": True}


IMMO_NEWS_CACHE = Path(os.path.join(WORKSPACE, "cache/immo_news.json"))
IMMO_FEEDS = [
    {"url": "https://rss.sueddeutsche.de/rss/Wirtschaft",                  "source": "Süddeutsche",  "filter": True},
    {"url": "https://www.handelsblatt.com/contentexport/feed/schlagzeilen", "source": "Handelsblatt", "filter": True},
    {"url": "https://www.faz.net/rss/aktuell/wirtschaft/",                  "source": "FAZ",          "filter": True},
    {"url": "https://www.n-tv.de/wirtschaft/rss",                           "source": "n-tv",         "filter": True},
    {"url": "https://www.spiegel.de/wirtschaft/index.rss",                  "source": "Spiegel",      "filter": True},
    {"url": "https://newsfeed.zeit.de/wirtschaft/index",                    "source": "Zeit",         "filter": True},
]
_IMMO_KW = [
    "immobilien", "immobilienmarkt", "immobilienpreis", "immobilienkauf",
    "wohnungsmarkt", "wohnungsnot", "wohnraum", "wohnungskauf",
    "eigenheim", "eigentumsquote", "grundstück", "neubau",
    "baufinanzierung", "bauzinsen", "baukredit", "hypothek",
    "miete", "mieten", "mietpreise", "mietrecht",
    "makler", "kaufpreis", "kaufnebenkosten",
    "haus kaufen", "wohnung kaufen", "häuser", "neubauwohnungen",
    "immobilienfonds", "wohnimmobilien",
    "bayern", "münchen", "ammersee", "starnberg", "landsberg", "weilheim",
    "oberbayern", "fünf-seen", "fürstenfeldbruck",
]

def get_immo_news(force=False):
    import xml.etree.ElementTree as ET, re as _re
    # Cache prüfen (3h TTL)
    if not force and IMMO_NEWS_CACHE.exists():
        try:
            cached = json.loads(IMMO_NEWS_CACHE.read_text())
            age = (datetime.now() - datetime.fromisoformat(cached.get("fetched_at","2000-01-01"))).total_seconds()
            if age < 10800:
                return cached
        except Exception:
            pass

    results = []
    for feed in IMMO_FEEDS:
        try:
            req = urllib.request.Request(feed["url"],
                  headers={"User-Agent": "Mozilla/5.0 (Bolla-MC/1.0)"})
            with urllib.request.urlopen(req, timeout=10) as r:
                raw = r.read().decode("utf-8", errors="replace")
            root = ET.fromstring(raw)
            _ch = root.find("channel"); channel = _ch if _ch is not None else root
            count = 0
            for item in channel.findall("item"):
                title   = (item.findtext("title")        or "").strip()
                desc_r  = (item.findtext("description")  or "").strip()
                link    = (item.findtext("link")         or "").strip()
                pubdate = (item.findtext("pubDate")      or "")[:22].strip()
                # HTML aus Beschreibung entfernen
                desc = _re.sub(r"<[^>]+>", " ", desc_r)
                desc = _re.sub(r"&[a-z]+;|&#\d+;", " ", desc)
                desc = _re.sub(r"\s+", " ", desc).strip()[:220]
                # Keyword-Filter
                combined = (title + " " + desc).lower()
                if feed.get("filter") and not any(kw in combined for kw in _IMMO_KW):
                    continue
                if not title or not link:
                    continue
                results.append({
                    "title":   title[:130],
                    "summary": desc,
                    "url":     link,
                    "source":  feed["source"],
                    "date":    pubdate,
                })
                count += 1
                if count >= 4:
                    break
        except Exception as e:
            print(f"Immo-News Feed-Fehler ({feed['source']}): {e}")

    # Duplikate per URL raus, max 10 Artikel
    seen, deduped = set(), []
    for r in results:
        if r["url"] not in seen:
            seen.add(r["url"])
            deduped.append(r)
        if len(deduped) >= 10:
            break

    out = {"fetched_at": datetime.now().isoformat(), "news": deduped}
    IMMO_NEWS_CACHE.write_text(json.dumps(out, ensure_ascii=False, indent=2))
    return out

_SUMMARY_KEYWORDS = [
    "immobilien", "preis", "preis", "kaufpreis", "miete", "mieten", "zinsen", "bauzins",
    "baufinanzierung", "kredit", "hypothek", "rendite", "neubau", "grundstück",
    "eigenheim", "wohnungsmarkt", "wohnungsnot", "eigentumsquote",
    "markt", "angebot", "nachfrage", "leerstand", "förderung", "förderprogramm",
    "münchen", "bayern", "ammersee", "starnberg", "landsberg", "weilheim",
    "makler", "notarkosten", "grunderwerbsteuer",
]

def summarize_immo_news(title, summary_text):
    prompt = (
        "Du bist Bolla, Chris Mandels KI-Assistent. Chris hat einen Immobilienmarkt-Artikel angeklickt.\n"
        "Schreib eine kurze, prägnante Zusammenfassung auf Deutsch (3–5 Sätze, kein Bullshit).\n"
        "Hebe WICHTIGE Begriffe mit <mark>…</mark> hervor — Zahlen, Preise, Zinssätze, Orte, Trends.\n"
        "Antworte NUR mit dem HTML-Fließtext (keine Überschrift, kein <html>, kein <body>).\n\n"
        f"Titel: {title}\n"
        f"Teaser: {summary_text or '(kein Teaser)'}\n"
    )
    try:
        r = subprocess.run([CLAUDE_BIN, "-p", prompt], capture_output=True, text=True, timeout=45)
        raw = (r.stdout or "").strip()
        if not raw:
            return {"error": "Keine Antwort von Claude."}
        # Paragraphen in <p> wickeln wenn kein HTML vorhanden
        if "<p>" not in raw and "<mark>" not in raw:
            raw = "<p>" + raw.replace("\n\n", "</p><p>") + "</p>"
        elif "<p>" not in raw:
            raw = "<p>" + raw + "</p>"
        return {"html": raw}
    except subprocess.TimeoutExpired:
        return {"error": "Timeout — Claude hat zu lange gebraucht."}
    except Exception as e:
        return {"error": str(e)}

_IMMO_CRITERIA_DEFAULT = {
    "lage": [
        "Ort mit Gesundheitszentrum / Fitness + Sauna",
        "Großraum München West (Ammersee / Starnberg) – nicht München",
        "Einkaufsmöglichkeiten / Bäcker",
    ],
    "objekt": [
        "Penthouse inkl. Terrasse Südlage",
        "Neubau – oder günstig + Renovierung",
        "ca. 100 qm",
        "4 Zimmer (3,5)",
        "Max. 3-stöckig – keine Nachbar-Einsicht auf Terrasse",
        "Fußbodenheizung – Wärmepumpe / Fernwärme",
        "Terrasse: Südlage, Grill, Markise",
        "Parkett / Fliesen",
        "Bad: Dusche ebenerdig, Badewanne",
        "Keller / Abstellkammer",
        "Tiefgarage mit Ladestation",
        "Kein Straßenlärm / keine Autobahnnähe",
    ],
    "optional": [
        "Kaminofen (Einbau)",
        "Elektrische Rollläden",
        "Home-Steuerung Elektrik",
    ],
}

def get_crontab():
    import subprocess as _sp, re as _re3
    try:
        r = _sp.run(['crontab', '-l'], capture_output=True, text=True)
        entries = []
        for raw in r.stdout.splitlines():
            line = raw.strip()
            if not line or line.startswith('#'):
                continue
            parts = line.split(None, 5)
            if not parts:
                continue
            if parts[0] == '@reboot':
                schedule = '@reboot'
                cmd = ' '.join(parts[1:]) if len(parts) > 1 else ''
            elif len(parts) >= 6:
                schedule = ' '.join(parts[:5])
                cmd = parts[5]
            else:
                continue
            cmd = _re3.sub(r'\s*>>\s*\S+.*$', '', cmd).strip()
            cmd = _re3.sub(r'\s*2>&1$', '', cmd).strip()
            entries.append({'schedule': schedule, 'cmd': cmd})
        return entries
    except Exception as e:
        return []

def get_immo_criteria():
    if IMMO_CRITERIA_FILE.exists():
        return json.loads(IMMO_CRITERIA_FILE.read_text())
    return _IMMO_CRITERIA_DEFAULT

def save_immo_criteria(data):
    list_keys = {"lage", "objekt", "optional"}
    clean = {k: [str(x) for x in v] for k, v in data.items() if k in list_keys and isinstance(v, list)}
    if "priorities" in data and isinstance(data["priorities"], dict):
        clean["priorities"] = {str(k): str(v) for k, v in data["priorities"].items()}
    IMMO_CRITERIA_FILE.write_text(json.dumps(clean, ensure_ascii=False, indent=2))
    return {"ok": True}


def get_immo_bookmarks():
    if IMMO_BOOKMARKS_FILE.exists():
        return json.loads(IMMO_BOOKMARKS_FILE.read_text())
    return []

def save_immo_bookmark(data):
    bookmarks = get_immo_bookmarks()
    bid = str(int(datetime.now().timestamp() * 1000))
    entry = {
        "id":           bid,
        "title":        data.get("title", "").strip()[:200],
        "source":       data.get("source", "").strip(),
        "url":          data.get("url", "").strip(),
        "date":         data.get("date", "").strip(),
        "saved_at":     datetime.now().isoformat(),
        "category":     data.get("category", "Sonstiges").strip(),
        "summary_html": data.get("summary_html", "").strip(),
        "notiz":        data.get("notiz", "").strip(),
    }
    bookmarks.append(entry)
    IMMO_BOOKMARKS_FILE.write_text(json.dumps(bookmarks, ensure_ascii=False, indent=2))
    return {"ok": True, "id": bid}

def delete_immo_bookmark(bid):
    bookmarks = get_immo_bookmarks()
    bookmarks = [b for b in bookmarks if b.get("id") != bid]
    IMMO_BOOKMARKS_FILE.write_text(json.dumps(bookmarks, ensure_ascii=False, indent=2))
    return {"ok": True}

def update_immo_bookmark(bid, data):
    bookmarks = get_immo_bookmarks()
    for b in bookmarks:
        if b.get("id") == bid:
            if "notiz" in data:
                b["notiz"] = data["notiz"].strip()
            if "category" in data:
                b["category"] = data["category"].strip()
            break
    IMMO_BOOKMARKS_FILE.write_text(json.dumps(bookmarks, ensure_ascii=False, indent=2))
    return {"ok": True}


def get_newsletter():
    """Gibt Watchlist + letzte Scan-Ergebnisse zurück."""
    watchlist = json.loads(NEWSLETTER_WATCHLIST.read_text()) if NEWSLETTER_WATCHLIST.exists() else []
    results   = json.loads(NEWSLETTER_RESULTS.read_text())   if NEWSLETTER_RESULTS.exists()   else {"scanned_at": None, "results": []}
    return {"watchlist": watchlist, "scanned_at": results.get("scanned_at"), "results": results.get("results", [])}

def newsletter_watchlist_update(data):
    """Fügt einen Begriff hinzu oder entfernt ihn."""
    watchlist = json.loads(NEWSLETTER_WATCHLIST.read_text()) if NEWSLETTER_WATCHLIST.exists() else []
    if "add" in data:
        term = data["add"].strip().lower()
        if term and term not in watchlist:
            watchlist.append(term)
    if "remove" in data:
        watchlist = [w for w in watchlist if w != data["remove"]]
    NEWSLETTER_WATCHLIST.write_text(json.dumps(watchlist, ensure_ascii=False))
    return {"watchlist": watchlist}

def newsletter_scan_now():
    """Startet den Newsletter-Scanner sofort im Hintergrund."""
    _p = subprocess.Popen(["python3", os.path.join(WORKSPACE,"scripts/newsletter_scanner.py")],
              stdout=open(os.path.join(WORKSPACE,"logs/newsletter_scanner.log"),"a"),
              stderr=subprocess.STDOUT)
    return {"ok": True, "text": "Scan gestartet — dauert ca. 30–60 Sek."}

def newsletter_search(term):
    """Sofortsuche: scannt aktuelle Newsletter nach einem Begriff, speichert nichts."""
    import sys as _sys
    _sys.path.insert(0, os.path.join(WORKSPACE, "scripts"))
    from newsletter_scanner import fetch_outlook_newsletters, fetch_wtnet_newsletters, analyse_mail
    mails = fetch_outlook_newsletters() + fetch_wtnet_newsletters()
    results = []
    for mail in mails:
        hits = analyse_mail(mail, [term])
        for h in hits:
            results.append({
                "store": mail["store"]["name"], "color": mail["store"]["color"],
                "text_color": mail["store"].get("text","#fff"),
                "artikel": h.get("artikel",""), "preis": h.get("preis",""),
                "gueltig_von": h.get("gueltig_von",""), "gueltig_bis": h.get("gueltig_bis",""),
                "hinweis": h.get("hinweis",""), "mail_date": mail["date"], "subject": mail["subject"],
            })
    return {"results": results, "term": term, "mails_checked": len(mails)}


def mail_command(data):
    """Führt einen Sprachbefehl auf eine oder mehrere Mails aus."""
    import subprocess, smtplib, re, urllib.request
    from email.mime.text import MIMEText
    from pathlib import Path
    instruction = (data.get("instruction") or "").strip()
    instr_low   = instruction.lower()

    # Normalisierung: mails-Array (neu) oder einzelne mail (Fallback)
    mails = data.get("mails") or []
    if not mails and data.get("mail"):
        mails = [data["mail"]]
    if not mails:
        return {"type":"error","text":"Keine Mail ausgewählt."}
    mail = mails[0]  # Erste Mail für Einzeloperationen

    def _mail_block(m, idx=None):
        prefix = f"Mail {idx}: " if idx is not None else ""
        return (f"{prefix}Von: {m.get('from','')} <{m.get('from_email','')}>\n"
                f"Betreff: {m.get('subject','')}\nDatum: {m.get('date','')}\n"
                f"Inhalt: {m.get('preview','')}")

    # ── Löschen ────────────────────────────────────────────────────────────────
    if any(w in instr_low for w in ["lösch", "loesch", "delete", "entfern", "wegmach"]):
        deleted, errors = [], []
        ol_token = None
        wtnet_cfg = None
        for m in mails:
            mid, account = m.get("id",""), m.get("account","")
            if account == "Outlook" and mid:
                try:
                    if ol_token is None:
                        ol_token = get_token()
                    req = urllib.request.Request(
                        f"https://graph.microsoft.com/v1.0/me/messages/{mid}",
                        method="DELETE", headers={"Authorization": f"Bearer {ol_token}"})
                    urllib.request.urlopen(req, timeout=10)
                    deleted.append(m.get('from','?'))
                except Exception as e:
                    errors.append(str(e))
            elif account == "wtnet" and mid:
                try:
                    import imaplib
                    if wtnet_cfg is None:
                        wtnet_cfg = json.loads(Path(os.path.join(WORKSPACE,"config/wtnet_account.json")).read_text())
                    with imaplib.IMAP4_SSL(wtnet_cfg["imap_host"], wtnet_cfg["imap_port"]) as imap:
                        imap.login(wtnet_cfg["email"], wtnet_cfg["password"])
                        imap.select("INBOX")
                        imap.uid("store", mid.encode(), "+FLAGS", "\\Deleted")
                        imap.expunge()
                    deleted.append(m.get('from','?'))
                except Exception as e:
                    errors.append(str(e))
            else:
                errors.append(f"Kein Konto/ID für: {m.get('subject','?')}")
        parts = []
        if deleted:
            parts.append(f"✓ {len(deleted)} Mail(s) gelöscht: {', '.join(deleted)}")
        if errors:
            parts.append("⚠ Fehler: " + "; ".join(errors))
        return {"type":"deleted","text":"\n".join(parts) or "Nichts gelöscht."}

    # ── Antworten / Entwurf ────────────────────────────────────────────────────
    if any(w in instr_low for w in ["antwort","reply","schreib","verfass","sende","formulier"]):
        if len(mails) > 1:
            return {"type":"error","text":"Antworten geht nur für eine einzelne Mail — bitte nur eine auswählen."}
        prompt = f"""Du bist Bolla, Chris Mandels persönlicher KI-Assistent.
Verfasse eine E-Mail-Antwort gemäß der Anweisung.

Original-Mail:
Von: {mail.get('from','')} <{mail.get('from_email','')}>
Betreff: {mail.get('subject','')}
Inhalt: {mail.get('preview','')}

Anweisung: {instruction}

Schreibe NUR den E-Mail-Text (kein JSON, keine Erklärung).
Beginne direkt mit der Anrede. Unterschreibe als Chris Mandel."""
        try:
            r = subprocess.run([CLAUDE_BIN,"-p",prompt], capture_output=True, text=True, timeout=60)
            return {"type":"draft","text":r.stdout.strip(),
                    "draft_to": mail.get("from_email",""),
                    "draft_subject": "Re: " + mail.get("subject",""),
                    "draft_account": mail.get("account","Outlook")}
        except Exception as e:
            return {"type":"error","text":str(e)}

    # ── Senden (vorbereiteter Entwurf) ─────────────────────────────────────────
    if data.get("action") == "send":
        to      = data.get("draft_to","")
        subject = data.get("draft_subject","")
        text    = data.get("draft_text","")
        account = data.get("draft_account","Outlook")
        if account == "wtnet":
            try:
                cfg = json.loads(Path(os.path.join(WORKSPACE,"config/wtnet_account.json")).read_text())
                msg = MIMEText(text, "plain", "utf-8")
                msg["Subject"], msg["From"], msg["To"] = subject, cfg["email"], to
                with smtplib.SMTP(cfg["smtp_host"], cfg["smtp_port"]) as s:
                    s.starttls(); s.login(cfg["email"], cfg["password"]); s.send_message(msg)
                return {"type":"sent","text":f"✓ Gesendet an {to}"}
            except Exception as e:
                return {"type":"error","text":str(e)}
        else:
            try:
                token = get_token()
                body = {"message":{"subject":subject,
                    "body":{"contentType":"Text","content":text},
                    "toRecipients":[{"emailAddress":{"address":to}}]},
                    "saveToSentItems":True}
                req = urllib.request.Request(
                    "https://graph.microsoft.com/v1.0/me/sendMail",
                    data=json.dumps(body).encode(), method="POST",
                    headers={"Authorization":f"Bearer {token}","Content-Type":"application/json"})
                urllib.request.urlopen(req, timeout=15)
                return {"type":"sent","text":f"✓ Gesendet an {to}"}
            except Exception as e:
                return {"type":"error","text":str(e)}

    # ── Analyse / Prüfen (Standard) ────────────────────────────────────────────
    if len(mails) == 1:
        mail_context = _mail_block(mail)
    else:
        mail_context = "\n\n".join(_mail_block(m, i+1) for i, m in enumerate(mails))

    prompt = f"""Du bist Bolla, Chris Mandels KI-Assistent. Antworte kurz und direkt auf Deutsch.

{"E-Mail" if len(mails)==1 else f"{len(mails)} E-Mails"}:
{mail_context}

Aufgabe: {instruction or 'Fasse diese Mail(s) kurz zusammen.'}"""
    try:
        r = subprocess.run([CLAUDE_BIN,"-p",prompt], capture_output=True, text=True, timeout=60)
        return {"type":"analysis","text":r.stdout.strip()}
    except Exception as e:
        return {"type":"error","text":str(e)}


def get_emails():
    """Kombiniert Outlook + wtnet Mails."""
    outlook       = get_emails_outlook()
    wtnet         = get_emails_wtnet()
    recent_ol     = get_emails_recent_outlook()
    recent_wt     = get_emails_recent_wtnet()
    sent          = get_emails_sent_outlook()

    all_msgs = outlook + wtnet
    # Recent: Outlook + wtnet zusammen, nach Datum sortiert (neueste zuerst)
    recent_all = recent_ol + recent_wt
    return {
        "count": len(all_msgs),
        "messages": all_msgs,
        "recent": recent_all,
        "sent": sent,
        "accounts": [
            {"name": "Outlook", "count": len(outlook)},
            {"name": "wtnet",   "count": len(wtnet)}
        ]
    }


def get_quicknotes():
    """Zuletzt gespeicherte Schnellnotizen."""
    notes_file = os.path.join(WORKSPACE, "data/quicknotes.json")
    if os.path.exists(notes_file):
        with open(notes_file) as f:
            return json.load(f)
    return []

def get_sos_contacts():
    """SOS-Kontakte und Notfall-Infos."""
    cfg_file = os.path.join(WORKSPACE, "config/sos_contacts.json")
    if os.path.exists(cfg_file):
        with open(cfg_file) as f:
            return json.load(f)
    return {"contacts": [], "medical": {}}

# ── Bolla Chat-Verlauf ───────────────────────────────────────────────────────
CHAT_HISTORY_FILE = os.path.join(WORKSPACE, "data/chat_history.json")
CHAT_HISTORY_MAX = 30

def _chat_history_load():
    if os.path.exists(CHAT_HISTORY_FILE):
        try:
            with open(CHAT_HISTORY_FILE, encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return []

def _chat_history_save(lst):
    with open(CHAT_HISTORY_FILE, "w", encoding="utf-8") as f:
        json.dump(lst, f, ensure_ascii=False, indent=2)

def get_chat_history():
    return _chat_history_load()

def chat_history_add(entry):
    lst = _chat_history_load()
    lst = [e for e in lst if e.get("id") != entry.get("id")]
    lst.insert(0, entry)
    if len(lst) > CHAT_HISTORY_MAX:
        lst = lst[:CHAT_HISTORY_MAX]
    _chat_history_save(lst)
    return {"ok": True}

def chat_history_delete(entry_id):
    lst = _chat_history_load()
    lst = [e for e in lst if e.get("id") != entry_id]
    _chat_history_save(lst)
    return {"ok": True}

def chat_history_clear():
    _chat_history_save([])
    return {"ok": True}

# ── KI-Workshop Projekt-Seite (Legacy) ──────────────────────────────────────
WORKSHOP_MD = os.path.join(WORKSPACE, "projektwoche-ki-workshop/workshop-ideen.md")
WORKSHOP_AUFTRAEGE = os.path.join(WORKSPACE, "projektwoche-ki-workshop/auftraege.json")

# ── Projekte-Workspace ───────────────────────────────────────────────────────
PROJEKTE_FILE = os.path.join(WORKSPACE, "data/projekte.json")

def _projekte_load():
    if os.path.exists(PROJEKTE_FILE):
        try:
            with open(PROJEKTE_FILE, encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    # Erste Migration: KI-Workshop-Daten übernehmen
    md = ""
    if os.path.exists(WORKSHOP_MD):
        with open(WORKSHOP_MD, encoding="utf-8") as f:
            md = f.read()
    auftraege = []
    if os.path.exists(WORKSHOP_AUFTRAEGE):
        try:
            with open(WORKSHOP_AUFTRAEGE, encoding="utf-8") as f:
                auftraege = json.load(f)
        except Exception:
            pass
    default = {
        "current": "ki-shorttrack-2026",
        "projects": [{
            "id": "ki-shorttrack-2026",
            "title": "KI-ShortTrack für Lehrerkollegen",
            "icon": "🎓",
            "created": "2026-05-20",
            "updated": datetime.now().strftime("%Y-%m-%d"),
            "content": md,
            "auftraege": auftraege
        }]
    }
    _projekte_save_raw(default)
    return default

def _projekte_save_raw(data):
    os.makedirs(os.path.dirname(PROJEKTE_FILE), exist_ok=True)
    with open(PROJEKTE_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

def projekte_list():
    data = _projekte_load()
    return {
        "current": data.get("current"),
        "projects": [{"id": p["id"], "title": p["title"], "icon": p.get("icon","📁"),
                       "created": p.get("created",""), "updated": p.get("updated","")}
                     for p in data.get("projects", [])]
    }

def projekte_load(pid):
    data = _projekte_load()
    for p in data.get("projects", []):
        if p["id"] == pid:
            return p
    return None

def projekte_save(pid, title, icon, content, auftraege=None, items=None, prompts=None):
    data = _projekte_load()
    projects = data.get("projects", [])
    now = datetime.now().strftime("%Y-%m-%d")
    for p in projects:
        if p["id"] == pid:
            p["title"] = title
            p["icon"] = icon
            p["content"] = content
            p["updated"] = now
            if auftraege is not None:
                p["auftraege"] = auftraege
            if items is not None:
                p["items"] = items
            if prompts is not None:
                p["prompts"] = prompts
            data["current"] = pid
            _projekte_save_raw(data)
            return {"ok": True}
    projects.append({"id": pid, "title": title, "icon": icon,
                     "created": now, "updated": now,
                     "content": content, "auftraege": auftraege or [],
                     "items": items or [], "prompts": prompts or []})
    data["current"] = pid
    _projekte_save_raw(data)
    return {"ok": True, "new": True}

def projekte_add_auftrag(pid, text):
    data = _projekte_load()
    for p in data.get("projects", []):
        if p["id"] == pid:
            if "auftraege" not in p:
                p["auftraege"] = []
            p["auftraege"].append({
                "id": int(datetime.now().timestamp()),
                "ts": datetime.now().isoformat(timespec="minutes"),
                "text": text.strip(),
                "status": "offen"
            })
            _projekte_save_raw(data)
            return {"ok": True}
    return {"error": "Projekt nicht gefunden"}

def projekte_auftrag_status(pid, auftrag_id, status):
    data = _projekte_load()
    for p in data.get("projects", []):
        if p["id"] == pid:
            for a in p.get("auftraege", []):
                if str(a["id"]) == str(auftrag_id):
                    a["status"] = status
                    if status == "erledigt":
                        a["ts_done"] = datetime.now().isoformat(timespec="minutes")
                    _projekte_save_raw(data)
                    return {"ok": True}
    return {"error": "Nicht gefunden"}

def projekte_delete(pid):
    data = _projekte_load()
    data["projects"] = [p for p in data.get("projects", []) if p["id"] != pid]
    if data.get("current") == pid:
        data["current"] = data["projects"][0]["id"] if data["projects"] else ""
    _projekte_save_raw(data)
    return {"ok": True}
WORKSHOP_FORTSCHRITT = os.path.join(WORKSPACE, "projektwoche-ki-workshop/fortschritt.json")

def get_workshop_fortschritt():
    """Live-Fortschritt eines laufenden Workshop-Auftrags (von Bolla geschrieben)."""
    if os.path.exists(WORKSHOP_FORTSCHRITT):
        try:
            with open(WORKSHOP_FORTSCHRITT, encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return {"active": False, "pct": 0, "label": "", "ts": ""}

def get_workshop():
    """Liefert Markdown-Inhalt der Workshop-Ideensammlung + offene Aufträge."""
    md = ""
    if os.path.exists(WORKSHOP_MD):
        with open(WORKSHOP_MD, encoding="utf-8") as f:
            md = f.read()
    auftraege = []
    if os.path.exists(WORKSHOP_AUFTRAEGE):
        try:
            with open(WORKSHOP_AUFTRAEGE, encoding="utf-8") as f:
                auftraege = json.load(f)
        except Exception:
            auftraege = []
    return {"markdown": md, "auftraege": auftraege}

def save_workshop(markdown):
    """Speichert den bearbeiteten Markdown-Inhalt zurück (mit .bak-Sicherung)."""
    if not markdown or not isinstance(markdown, str):
        return {"error": "Kein Inhalt"}
    os.makedirs(os.path.dirname(WORKSHOP_MD), exist_ok=True)
    # Backup der vorigen Version
    if os.path.exists(WORKSHOP_MD):
        import shutil
        shutil.copy2(WORKSHOP_MD, WORKSHOP_MD + ".bak")
    with open(WORKSHOP_MD, "w", encoding="utf-8") as f:
        f.write(markdown)
    return {"ok": True, "bytes": len(markdown.encode("utf-8"))}

def add_workshop_auftrag(text):
    """Hängt einen Auftrag an die Queue (Bolla arbeitet ihn bei nächster Session ab)."""
    text = (text or "").strip()
    if not text:
        return {"error": "Leerer Auftrag"}
    auftraege = []
    if os.path.exists(WORKSHOP_AUFTRAEGE):
        try:
            with open(WORKSHOP_AUFTRAEGE, encoding="utf-8") as f:
                auftraege = json.load(f)
        except Exception:
            auftraege = []
    auftraege.append({
        "id": int(datetime.now().timestamp()),
        "ts": datetime.now().isoformat(timespec="minutes"),
        "text": text,
        "status": "offen"
    })
    os.makedirs(os.path.dirname(WORKSHOP_AUFTRAEGE), exist_ok=True)
    with open(WORKSHOP_AUFTRAEGE, "w", encoding="utf-8") as f:
        json.dump(auftraege, f, ensure_ascii=False, indent=2)
    return {"ok": True, "count": len([a for a in auftraege if a.get("status") == "offen"])}

# ── Charts ──────────────────────────────────────────────────────────────────
GEMINI_CONFIG = Path(os.path.join(WORKSPACE, "config/gemini_api.json"))
PARTY_CHARTS_CACHE = Path(os.path.join(WORKSPACE, "config/party_charts_cache.json"))
CREDENTIALS_FILE = Path(os.path.expanduser("~/.claude/.credentials.json"))
SUNO_API_BASE = "https://studio-api-prod.suno.com"
SUNO_ROUTENOTE_DIR = Path("/mnt/d/OneDrive/Dokumente/Bolla/Suno_RouteNote")

_charts_cache = {"data": None, "ts": 0}
CHARTS_TTL = 1800  # 30 Minuten

def _daily_sample(pool, n=10):
    """Täglich n Songs aus pool, Doppelungen mit Vortag vermeiden."""
    import random as _rnd, time as _t
    today = int(_t.time()) // 86400
    # Gestrige Auswahl berechnen
    rnd_y = _rnd.Random(today - 1)
    pool_y = pool[:]
    rnd_y.shuffle(pool_y)
    yesterday_titles = {s['title'] for s in pool_y[:n]}
    # Heutige Shufflefolge
    rnd_t = _rnd.Random(today)
    pool_t = pool[:]
    rnd_t.shuffle(pool_t)
    # Zuerst Songs nehmen, die gestern nicht drin waren
    result = [s for s in pool_t if s['title'] not in yesterday_titles]
    fallback = [s for s in pool_t if s['title'] in yesterday_titles]
    result = (result + fallback)[:n]
    return result

def _gemini_key():
    try:
        return json.loads(GEMINI_CONFIG.read_text()).get("api_key", "")
    except Exception:
        return ""

SUNO_TOKEN_FILE = Path(os.path.join(WORKSPACE, "config/suno_token.json"))

def _suno_token():
    try:
        tok = json.loads(SUNO_TOKEN_FILE.read_text()).get("token", "")
        return tok.removeprefix("TOKEN:").strip()
    except Exception:
        return ""

def _suno_token_save(token):
    token = token.removeprefix("TOKEN:").strip()
    SUNO_TOKEN_FILE.write_text(json.dumps({"token": token, "ts": datetime.now().isoformat()}))
    return {"ok": True}

def _fetch_kworb(kworb_slug, limit=10):
    """Spotify Charts via kworb.net (Spotify-Daten, täglich aktualisiert)."""
    import urllib.request as _ur2, re as _re2
    url = f"https://kworb.net/spotify/country/{kworb_slug}.html"
    req = _ur2.Request(url, headers={"User-Agent": "Mozilla/5.0 BollaMC/1.0"})
    try:
        html = _ur2.urlopen(req, timeout=12).read().decode("utf-8", "ignore")
        rows = _re2.findall(r'<tr[^>]*>(.*?)</tr>', html, _re2.DOTALL)
        results = []
        for row in rows[1:]:
            cells = _re2.findall(r'<td[^>]*>(.*?)</td>', row, _re2.DOTALL)
            cells = [_re2.sub(r'<[^>]+>', '', c).strip() for c in cells]
            if len(cells) >= 7 and cells[0].isdigit():
                combined = cells[2]
                if ' - ' in combined:
                    artist, title = combined.split(' - ', 1)
                else:
                    artist, title = '', combined
                results.append({"title": title.strip(), "artist": artist.strip(), "streams": cells[6]})
            if len(results) >= limit:
                break
        return results
    except Exception as e:
        return [{"error": str(e)}]

# Kuratierter Pool harmloser, sehr bekannter dt. Party-/Mitsing-Hits — familientauglich
# für ~13-Jährige (Schul-Party). Keine anzüglichen/Alkohol-/skandalträchtigen Titel.
# Wöchentliche Rotation: _fetch_party_charts() zieht per Wochen-Seed 10 davon.
_PARTY_HITS = [
    {"title": "Atemlos durch die Nacht",          "artist": "Helene Fischer",        "streams": "Partyklassiker"},
    {"title": "Ein Stern der deinen Namen trägt",  "artist": "DJ Ötzi & Nik P.",      "streams": "Partyklassiker"},
    {"title": "Wahnsinn",                          "artist": "Wolfgang Petry",        "streams": "Partyklassiker"},
    {"title": "Hey Baby",                          "artist": "DJ Ötzi",               "streams": "Partyklassiker"},
    {"title": "Anton aus Tirol",                   "artist": "DJ Ötzi",               "streams": "Partyklassiker"},
    {"title": "Cordula Grün",                      "artist": "Josh.",                 "streams": "Partyklassiker"},
    {"title": "Schwarz auf Weiß",                  "artist": "voXXclub",              "streams": "Partyklassiker"},
    {"title": "Hulapalu",                          "artist": "Andreas Gabalier",      "streams": "Partyklassiker"},
    {"title": "Das rote Pferd",                    "artist": "Markus Becker",         "streams": "Partyklassiker"},
    {"title": "Tage wie diese",                    "artist": "Die Toten Hosen",       "streams": "Mitsing-Hit"},
    {"title": "Auf uns",                           "artist": "Andreas Bourani",       "streams": "Mitsing-Hit"},
    {"title": "Applaus, Applaus",                  "artist": "Sportfreunde Stiller",  "streams": "Mitsing-Hit"},
    {"title": "80 Millionen",                      "artist": "Max Giesinger",         "streams": "Mitsing-Hit"},
    {"title": "Astronaut",                         "artist": "Sido & Andreas Bourani","streams": "Mitsing-Hit"},
    {"title": "Lieder",                            "artist": "Adel Tawil",            "streams": "Mitsing-Hit"},
    {"title": "99 Luftballons",                    "artist": "Nena",                  "streams": "Kult-Klassiker"},
    {"title": "Major Tom (Völlig losgelöst)",      "artist": "Peter Schilling",       "streams": "Kult-Klassiker"},
    {"title": "Marmor, Stein und Eisen bricht",    "artist": "Drafi Deutscher",       "streams": "Kult-Klassiker"},
    {"title": "Westerland",                        "artist": "Die Ärzte",             "streams": "Kult-Klassiker"},
    {"title": "Männer",                            "artist": "Herbert Grönemeyer",    "streams": "Kult-Klassiker"},
    {"title": "Ich war noch niemals in New York",  "artist": "Udo Jürgens",           "streams": "Kult-Klassiker"},
    {"title": "Aber bitte mit Sahne",              "artist": "Udo Jürgens",           "streams": "Kult-Klassiker"},
    {"title": "Verdammt, ich lieb dich",           "artist": "Matthias Reim",         "streams": "Kult-Klassiker"},
    {"title": "Cowboy und Indianer",              "artist": "Truck Stop",            "streams": "Partyklassiker"},
]

def _fetch_party_charts():
    """Kuratierte dt. Partyhits — täglich 10 zufällige, Doppelungen zum Vortag vermieden."""
    return _daily_sample(_PARTY_HITS, 10)

def _fetch_kworb_alltime(pick=10):
    """Top 100 meistgestreamte Songs aller Zeiten via kworb.net — täglich 10 zufällige, keine Vortags-Doppelungen."""
    import urllib.request as _ur2, re as _re2
    req = _ur2.Request("https://kworb.net/spotify/songs.html", headers={"User-Agent": "Mozilla/5.0 BollaMC/1.0"})
    try:
        html = _ur2.urlopen(req, timeout=12).read().decode("utf-8", "ignore")
        rows = _re2.findall(r'<tr[^>]*>(.*?)</tr>', html, _re2.DOTALL)
        pool = []
        for row in rows[1:]:
            cells = _re2.findall(r'<td[^>]*>(.*?)</td>', row, _re2.DOTALL)
            cells = [_re2.sub(r'<[^>]+>', '', c).strip() for c in cells]
            if len(cells) >= 2 and cells[0]:
                combined = cells[0]
                if ' - ' in combined:
                    artist, title = combined.split(' - ', 1)
                else:
                    artist, title = '', combined
                try:
                    streams_raw = int(cells[1].replace(',', '').replace('.', ''))
                    streams = f"{streams_raw/1_000_000_000:.1f} Mrd."
                except Exception:
                    streams = cells[1] if len(cells) > 1 else ''
                pool.append({"title": title.strip(), "artist": artist.strip(), "streams": streams})
            if len(pool) >= 100:
                break
        return _daily_sample(pool, pick) if len(pool) >= pick else pool
    except Exception as e:
        return [{"error": str(e)}]

def get_charts():
    """Streaming Charts: DE + Global (Spotify via kworb) + Party + Overall Alltime."""
    import time as _time
    now = _time.time()
    if _charts_cache["data"] and now - _charts_cache["ts"] < CHARTS_TTL:
        return _charts_cache["data"]
    de = _fetch_kworb("de_daily")
    gl = _fetch_kworb("global_daily")
    party = _fetch_party_charts()
    overall = _fetch_kworb_alltime()
    result = {"de": de, "global": gl, "party": party, "overall": overall}
    _charts_cache["data"] = result
    _charts_cache["ts"] = now
    return result


def get_chargers(lat: float, lon: float, radius: int = 12000):
    """Ladesäulen via OpenStreetMap Overpass API."""
    import urllib.request as _ureq, urllib.parse as _uparse, math
    query = (f'[out:json][timeout:20];'
             f'node["amenity"="charging_station"](around:{radius},{lat},{lon});'
             f'out body;')
    url = 'https://overpass-api.de/api/interpreter'
    req = _ureq.Request(url, data=_uparse.urlencode({'data': query}).encode(),
                        headers={'User-Agent': 'BollaMC/1.0', 'Accept': 'application/json'})
    try:
        with _ureq.urlopen(req, timeout=22) as resp:
            data = json.loads(resp.read().decode())
    except Exception as e:
        return {"error": str(e), "stations": []}

    TIER_MAP = [
        (['enbw','mobility+','e-wald'], 1, '✅ ~46 ct/kWh', '#34c759'),
        (['aral','pulse','total energies','avia'], 2, '🟡 ~55 ct/kWh', '#ff9500'),
    ]

    def tier(op):
        op_l = (op or '').lower()
        for names, t, label, color in TIER_MAP:
            if any(n in op_l for n in names):
                return t, label, color
        return 3, '⚠️ variabel', '#8e8e93'

    def dist_km(a_lat, a_lon):
        dlat = math.radians(a_lat - lat)
        dlon = math.radians(a_lon - lon)
        a = math.sin(dlat/2)**2 + math.cos(math.radians(lat))*math.cos(math.radians(a_lat))*math.sin(dlon/2)**2
        return 6371 * 2 * math.atan2(math.sqrt(a), math.sqrt(1-a))

    stations = []
    for el in data.get('elements', []):
        tags = el.get('tags', {})
        op = tags.get('operator') or tags.get('brand') or tags.get('name') or 'Unbekannt'
        name = tags.get('name') or op
        addr = ' '.join(filter(None, [tags.get('addr:street',''), tags.get('addr:housenumber',''), tags.get('addr:city','')]))
        kw_raw = tags.get('maxpower:ev') or tags.get('socket:type2_combo:output') or tags.get('socket:chademo:output') or ''
        try: kw = int(float(str(kw_raw).replace(' kW','').replace('kW','').strip()))
        except: kw = 0
        t, tlabel, tcolor = tier(op)
        stations.append({
            'name': name, 'operator': op, 'addr': addr.strip(),
            'lat': el['lat'], 'lon': el['lon'],
            'dist': round(dist_km(el['lat'], el['lon']), 2),
            'kw': kw, 'capacity': int(tags.get('capacity', 0) or 0),
            'tier': t, 'tier_label': tlabel, 'tier_color': tcolor,
            'maps': f"https://maps.google.com/?q={el['lat']},{el['lon']}"
        })
    stations.sort(key=lambda s: (s['tier'], s['dist']))
    return {"stations": stations, "count": len(stations)}

def get_birthdays():
    """Nächste Geburtstage aus Outlook-Kontakten."""
    from datetime import datetime
    import urllib.request
    token = get_token()
    if not token:
        return []
    
    today = datetime.now()
    contacts = []
    url = "https://graph.microsoft.com/v1.0/me/contacts?$select=displayName,birthday&$top=200"
    
    while url:
        data = graph_get(url.replace("https://graph.microsoft.com/v1.0", ""))
        if not data:
            break
        for c in data.get("value", []):
            bd = c.get("birthday")
            if not bd:
                continue
            try:
                dt = datetime.fromisoformat(bd[:10])
                next_bd = dt.replace(year=today.year)
                if next_bd.date() < today.date():
                    next_bd = next_bd.replace(year=today.year + 1)
                days_until = (next_bd.date() - today.date()).days
                birth_year = dt.year
                age = None if birth_year < 1900 else next_bd.year - birth_year
                contacts.append({
                    "days": days_until,
                    "name": c["displayName"],
                    "date": dt.strftime("%d.%m."),
                    "age": age,
                    "today": days_until == 0
                })
            except:
                pass
        url = data.get("@odata.nextLink")
    
    # Duplikate entfernen: gleiche Namen-Tokens + gleicher Geburtstag = Duplikat
    seen = set()
    unique = []
    for c in sorted(contacts, key=lambda x: x["days"]):
        # Name in Tokens aufteilen, sortieren und als Key nutzen
        name_tokens = frozenset(c["name"].lower().split())
        key = (name_tokens, c["date"])
        if key not in seen:
            seen.add(key)
            unique.append(c)
    
    return unique[:10]  # Nächste 10


def get_photo_of_day():
    """Holt ein zufälliges Foto aus OneDrive-Reiseordnern."""
    import random
    from datetime import date

    REISE_ORDNER = [
        "Ägypten 99", "Andalusien 2008", "Berlin", "Bormio 2001",
        "Côte d'Azur", "Damüls 2026", "Dänemark 2021", "Florida 2005",
        "Grindelwald 2000", "Harz 2024", "Hong Kong 1999", "Ischgl 2022",
        "Italien 2014", "Kreta 2012", "Kreuzfahrt 2025", "Kroatien 2018",
        "Malediven 2000", "Mallorca 2015", "New York 2000", "Norwegen 2005",
        "Norwegen 2007", "Paris 2002", "Prag 2015", "Rhodos 2020",
        "Sizilien 2013", "Sölden 2023", "Sylt 2012", "Thailand 2022",
        "Teneriffa 2011", "Venedig 2002", "Wien 2003", "Zillertal 2024",
        "Schladming 2024", "USA 2019", "Natur"
    ]

    # Täglich gleicher Ordner (per Datum-Seed)
    random.seed(date.today().toordinal())
    ordner = random.choice(REISE_ORDNER)

    try:
        # Fotos im Ordner holen (Leerzeichen URL-encoden)
        import urllib.parse
        ordner_enc = urllib.parse.quote(ordner)
        data = graph_get(f"/me/drive/root:/{ordner_enc}:/children?$select=name,id,file&$top=100")
        if not data:
            return None

        fotos = [
            item for item in data.get("value", [])
            if "file" in item and item["file"].get("mimeType", "").startswith("image")
        ]
        if not fotos:
            return None

        # 3 verschiedene zufällige Fotos
        auswahl = random.sample(fotos, min(3, len(fotos)))
        results = []
        for foto in auswahl:
            thumb = graph_get(f"/me/drive/items/{foto['id']}/thumbnails/0/large")
            if thumb and "url" in thumb:
                results.append({"url": thumb["url"], "name": foto["name"]})

        if not results:
            return None

        return {
            "photos": results,
            "album": ordner,
            "total_photos": len(fotos)
        }
    except Exception as e:
        print(f"Foto des Tages Fehler: {e}")
        return None


def get_recipe_of_day():
    """Generiert täglich ein deutsches Rezept via Claude + Bild via Gemini. Tages-Cache."""
    import subprocess
    from datetime import date

    today = date.today().isoformat()
    cache_file = os.path.join(WORKSPACE, f"cache/recipe_{today}.json")

    if os.path.exists(cache_file):
        with open(cache_file) as f:
            cached = json.load(f)
        # Cache ungültig wenn kein Bild vorhanden (z.B. API war beim ersten Aufruf noch nicht bereit)
        if cached.get('bild'):
            return cached
        os.remove(cache_file)

    # Letzte Rezepttitel sammeln um Wiederholungen zu vermeiden
    recent_titles = []
    cache_dir = os.path.join(WORKSPACE, "cache")
    for fname in sorted(os.listdir(cache_dir), reverse=True):
        if fname.startswith("recipe_") and fname.endswith(".json") and fname != f"recipe_{today}.json":
            try:
                with open(os.path.join(cache_dir, fname)) as rf:
                    t = json.load(rf).get("titel", "")
                if t:
                    recent_titles.append(t)
                if len(recent_titles) >= 14:
                    break
            except Exception:
                pass

    avoid_hint = ""
    if recent_titles:
        avoid_hint = (f" Diese Gerichte hatten wir zuletzt — bitte etwas anderes wählen: "
                      f"{', '.join(recent_titles)}.")

    prompt = (
        f"Heute ist {today}. Generiere ein leckeres Tagesrezept (deutsch oder international)."
        f"{avoid_hint} "
        "Antworte NUR als reines JSON ohne Markdown:\n"
        '{"titel":"Rezeptname","beschreibung":"1 Satz was das Gericht ist","'
        'zutaten":["2 Eier","100g Mehl"],"anweisungen":["Schritt 1","Schritt 2"],'
        '"bildprompt":"food photography prompt auf Englisch, appetitlich, weißer Hintergrund"}'
    )
    try:
        result = subprocess.run(
            ["claude", "-p", prompt],
            capture_output=True, text=True, timeout=60
        )
        text = result.stdout.strip()
        if "```" in text:
            text = text.split("```")[1]
            if text.startswith("json"):
                text = text[4:]
        data = json.loads(text.strip())
    except Exception as e:
        print(f"Rezept Claude-Fehler: {e}")
        return {"titel": "Pasta al Pomodoro", "beschreibung": "Klassische Tomatensauce mit frischem Basilikum.",
                "zutaten": ["400g Spaghetti", "800g Tomaten", "3 Knoblauchzehen", "Basilikum", "Olivenöl", "Salz"],
                "anweisungen": ["Spaghetti al dente kochen.", "Knoblauch in Öl anschwitzen.", "Tomaten dazugeben, 15 Min köcheln.", "Mit Basilikum und Salz abschmecken.", "Mit Pasta servieren."],
                "bildprompt": "spaghetti pomodoro, food photography, white background, appetizing"}

    bild_b64 = None
    try:
        bildprompt = data.get("bildprompt", f"appetizing {data.get('titel','food')} dish, food photography")
        img_b64, err = bildgen_generate(bildprompt, aspect_ratio="4:3")
        if img_b64:
            bild_b64 = img_b64
    except Exception as e:
        print(f"Rezept Bild-Fehler: {e}")

    data["bild"] = bild_b64

    try:
        with open(cache_file, "w") as f:
            json.dump(data, f, ensure_ascii=False)
    except Exception as e:
        print(f"Rezept Cache-Fehler: {e}")

    return data


REZEPTE_DIR = "/mnt/d/OneDrive/Dokumente/Rezepte-Cocktails"

def save_recipe_docx(data):
    """Speichert Rezept als Word-Datei — kompaktes 1-Seiten-Layout."""
    import base64, re, io
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import lxml.etree as etree

    def set_cell_border(cell, **kwargs):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ('top','left','bottom','right','insideH','insideV'):
            tag = OxmlElement(f'w:{edge}')
            tag.set(qn('w:val'), kwargs.get(edge, 'none'))
            tag.set(qn('w:sz'), '0')
            tag.set(qn('w:space'), '0')
            tag.set(qn('w:color'), 'auto')
            tcBorders.append(tag)
        tcPr.append(tcBorders)

    def set_para_spacing(para, before=0, after=0, line=None):
        pPr = para._p.get_or_add_pPr()
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), str(before))
        spacing.set(qn('w:after'), str(after))
        if line:
            spacing.set(qn('w:line'), str(line))
            spacing.set(qn('w:lineRule'), 'auto')
        pPr.append(spacing)

    ACCENT = RGBColor(0xC0, 0x39, 0x2B)
    GREY   = RGBColor(0x66, 0x66, 0x66)

    titel = data.get("titel", "Rezept").strip()
    fname = re.sub(r'[\\/:*?"<>|]', '', titel).strip() + ".docx"
    path = os.path.join(REZEPTE_DIR, fname)

    doc = Document()

    # Enge Ränder für mehr Platz
    sec = doc.sections[0]
    sec.top_margin    = Cm(1.5)
    sec.bottom_margin = Cm(1.5)
    sec.left_margin   = Cm(1.8)
    sec.right_margin  = Cm(1.8)

    # Basis-Schrift
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # ── Titel ──
    t = doc.add_heading(titel, level=1)
    t.runs[0].font.color.rgb = ACCENT
    t.runs[0].font.size = Pt(18)
    set_para_spacing(t, before=0, after=60)

    # Beschreibung (kursiv, kompakt)
    beschr = data.get("beschreibung", "")
    if beschr:
        p = doc.add_paragraph(beschr)
        r = p.runs[0]
        r.italic = True
        r.font.color.rgb = GREY
        r.font.size = Pt(10)
        set_para_spacing(p, before=0, after=80)

    # ── Trennlinie ──
    hr = doc.add_paragraph()
    hr_pPr = hr._p.get_or_add_pPr()
    pb = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1'); bot.set(qn('w:color'), 'C0392B')
    pb.append(bot); hr_pPr.append(pb)
    set_para_spacing(hr, before=0, after=80)

    # ── 2-Spalten-Tabelle: Bild links | Zutaten rechts ──
    bild_b64 = data.get("bild")
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    col_links  = tbl.columns[0]
    col_rechts = tbl.columns[1]
    # Spaltenbreiten: Bild 5.5cm, Zutaten Rest
    col_links.width  = Cm(5.5)
    col_rechts.width = Cm(11.2)

    cell_bild  = tbl.cell(0, 0)
    cell_zut   = tbl.cell(0, 1)
    set_cell_border(cell_bild)
    set_cell_border(cell_zut)

    # Bild-Zelle
    if bild_b64:
        try:
            img_bytes = base64.b64decode(bild_b64)
            bild_p = cell_bild.paragraphs[0]
            run = bild_p.add_run()
            run.add_picture(io.BytesIO(img_bytes), width=Cm(5.0))
            bild_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_para_spacing(bild_p, before=20, after=20)
        except Exception:
            cell_bild.paragraphs[0].add_run("(Bild fehlt)")
    else:
        cell_bild.paragraphs[0].add_run("")

    # Zutaten-Zelle
    h_zut = cell_zut.paragraphs[0]
    h_zut.clear()
    run_h = h_zut.add_run("Zutaten")
    run_h.bold = True
    run_h.font.color.rgb = ACCENT
    run_h.font.size = Pt(12)
    set_para_spacing(h_zut, before=20, after=60)

    for z in data.get("zutaten", []):
        p = cell_zut.add_paragraph(style="List Bullet")
        r = p.add_run(z)
        r.font.size = Pt(10)
        set_para_spacing(p, before=0, after=20)

    # Tipp (optional)
    if data.get("tipp"):
        p_tipp = cell_zut.add_paragraph()
        set_para_spacing(p_tipp, before=60, after=0)
        r = p_tipp.add_run(f"💡 {data['tipp']}")
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = GREY

    doc.add_paragraph()  # Abstand

    # ── Zubereitung (volle Breite) ──
    h_zub = doc.add_paragraph()
    r_h = h_zub.add_run("Zubereitung")
    r_h.bold = True
    r_h.font.color.rgb = ACCENT
    r_h.font.size = Pt(12)
    set_para_spacing(h_zub, before=40, after=60)

    for s in data.get("anweisungen", []):
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(s)
        r.font.size = Pt(10)
        set_para_spacing(p, before=0, after=30)

    # ── Fußzeile: Datum ──
    doc.add_paragraph()
    p_foot = doc.add_paragraph(f"🐾 Bolla · {datetime.now().strftime('%d.%m.%Y')}")
    p_foot.runs[0].font.size = Pt(8)
    p_foot.runs[0].font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
    p_foot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para_spacing(p_foot, before=60, after=0)

    os.makedirs(REZEPTE_DIR, exist_ok=True)
    doc.save(path)
    return fname


def get_robin_info():
    """Robin-Panel: Standort, Reise-Info, letzte Nachrichten."""
    from datetime import datetime
    today = datetime.now()
    
    # Famulatur Tansania: 01.08. - 08.09.2026
    tansania_start = datetime(2026, 8, 1)
    tansania_end   = datetime(2026, 9, 8)
    
    if today < tansania_start:
        days_until = (tansania_start - today).days
        location = "Ulm"
        location_detail = "Universität Ulm — 10. Semester Medizin"
        status = "studiert"
        travel_info = {
            "event": "Famulatur Tansania",
            "start": "01.08.2026",
            "end": "08.09.2026",
            "days_until": days_until,
            "flight": "Istanbul (RGXQU8)",
            "active": False
        }
        lat, lng = 48.4011, 9.9876  # Ulm
    elif today <= tansania_end:
        days_in = (today - tansania_start).days + 1
        location = "Tansania"
        location_detail = f"Famulatur — Tag {days_in} von 38"
        status = "auf Famulatur"
        travel_info = {
            "event": "Famulatur Tansania",
            "start": "01.08.2026",
            "end": "08.09.2026",
            "days_until": 0,
            "flight": "Istanbul (RGXQU8)",
            "active": True,
            "days_in": days_in
        }
        lat, lng = -6.369, 34.889  # Tansania (Mitte)
    else:
        location = "Ulm"
        location_detail = "Universität Ulm — Doktorarbeit"
        status = "schreibt Doktorarbeit"
        travel_info = None
        lat, lng = 48.4011, 9.9876
    
    return {
        "name": "Robin Mandel",
        "status": status,
        "location": location,
        "location_detail": location_detail,
        "lat": lat,
        "lng": lng,
        "travel": travel_info,
        "email": "robinmandel@outlook.de",
        "study": "Medizin, 10. Semester, Uni Ulm",
        "jarvis": "@RobinMandels_Jarvis_bot"
    }


def get_redesigns_meta():
    meta_path = os.path.expanduser("~/workspace/mission-control/redesign-meta.json")
    if not os.path.exists(meta_path):
        return {"date": None, "design1": None, "design2": None}
    with open(meta_path, encoding="utf-8") as f:
        return json.load(f)


def get_sysinfo():
    import subprocess, shutil, socket

    # RAM
    with open('/proc/meminfo') as f:
        mem = {}
        for line in f:
            k, v = line.split(':', 1)
            mem[k.strip()] = int(v.strip().split()[0])
    ram_total = mem['MemTotal'] // 1024
    ram_free = (mem['MemAvailable']) // 1024
    ram_used = ram_total - ram_free
    ram_pct = int(ram_used / ram_total * 100)

    # Swap
    swap_total = mem.get('SwapTotal', 0) // 1024
    swap_used = (mem.get('SwapTotal', 0) - mem.get('SwapFree', 0)) // 1024
    swap_pct = int(swap_used / swap_total * 100) if swap_total > 0 else 0

    # CPU Load + Kerne
    with open('/proc/loadavg') as f:
        load_parts = f.read().split()
    load1, load5, load15 = load_parts[0], load_parts[1], load_parts[2]
    cpu_count = os.cpu_count() or 1
    cpu_pct = min(100, int(float(load1) / cpu_count * 100))

    # Disk C:
    disk = os.statvfs('/mnt/c')
    disk_total = disk.f_blocks * disk.f_frsize // (1024**3)
    disk_free = disk.f_bavail * disk.f_frsize // (1024**3)
    disk_used = disk_total - disk_free
    disk_pct = int(disk_used / disk_total * 100)

    # Disk D: (OneDrive)
    disk_d = {"total_gb": 0, "free_gb": 0, "pct": 0}
    try:
        dd = os.statvfs('/mnt/d')
        dt = dd.f_blocks * dd.f_frsize // (1024**3)
        df = dd.f_bavail * dd.f_frsize // (1024**3)
        disk_d = {"total_gb": dt, "free_gb": df, "pct": int((dt - df) / dt * 100) if dt > 0 else 0}
    except Exception:
        pass

    # Uptime
    with open('/proc/uptime') as f:
        secs = float(f.read().split()[0])
    h = int(secs // 3600)
    m = int((secs % 3600) // 60)
    uptime = f"{h}h {m}m" if h > 0 else f"{m}m"

    # IP-Adresse (WSL)
    try:
        ip = subprocess.run(['hostname', '-I'], capture_output=True, text=True, timeout=3).stdout.strip().split()[0]
    except Exception:
        ip = '—'

    # Kernel
    try:
        kernel = subprocess.run(['uname', '-r'], capture_output=True, text=True, timeout=3).stdout.strip()
    except Exception:
        kernel = '—'

    # Prozesse
    try:
        ps_out = subprocess.run(['ps', 'ax', '--no-headers'], capture_output=True, text=True, timeout=3)
        proc_count = len(ps_out.stdout.strip().splitlines())
    except Exception:
        proc_count = 0

    # Git status
    git_info = {"commit": "unbekannt", "branch": "main", "dirty": 0, "commits_today": 0, "recent": []}
    try:
        branch = subprocess.run(
            ['git', '-C', WORKSPACE, 'rev-parse', '--abbrev-ref', 'HEAD'],
            capture_output=True, text=True, timeout=5
        )
        if branch.returncode == 0:
            git_info["branch"] = branch.stdout.strip()

        log3 = subprocess.run(
            ['git', '-C', WORKSPACE, 'log', '--oneline', '-5', '--format=%h|%s|%cr'],
            capture_output=True, text=True, timeout=5
        )
        if log3.returncode == 0 and log3.stdout.strip():
            lines = log3.stdout.strip().splitlines()
            recent = []
            for line in lines:
                parts = line.split('|')
                recent.append({
                    "commit": parts[0] if len(parts) > 0 else '',
                    "message": parts[1] if len(parts) > 1 else '',
                    "age": parts[2] if len(parts) > 2 else ''
                })
            git_info["recent"] = recent
            git_info["commit"] = recent[0]["commit"] if recent else ''
            git_info["message"] = recent[0]["message"] if recent else ''
            git_info["age"] = recent[0]["age"] if recent else ''

        today_log = subprocess.run(
            ['git', '-C', WORKSPACE, 'log', '--since=midnight', '--oneline'],
            capture_output=True, text=True, timeout=5
        )
        git_info["commits_today"] = len([l for l in today_log.stdout.strip().splitlines() if l.strip()])

        status = subprocess.run(
            ['git', '-C', WORKSPACE, 'status', '--short'],
            capture_output=True, text=True, timeout=5
        )
        git_info["dirty"] = len([l for l in status.stdout.strip().splitlines() if l.strip()])

        remote = subprocess.run(
            ['git', '-C', WORKSPACE, 'remote', 'get-url', 'origin'],
            capture_output=True, text=True, timeout=5
        )
        if remote.returncode == 0:
            url = remote.stdout.strip()
            git_info["repo"] = url.replace('https://github.com/', '').replace('.git', '')
    except Exception as e:
        git_info["error"] = str(e)

    def svc(name):
        return 'running' if subprocess.run(['pgrep', '-f', name],
                                           capture_output=True).returncode == 0 else 'stopped'

    claude_ver = '–'
    try:
        cv = subprocess.run([CLAUDE_BIN, '--version'], capture_output=True, text=True, timeout=5)
        if cv.returncode == 0:
            claude_ver = cv.stdout.strip().split('\n')[0]
    except Exception:
        pass

    last_backup = '–'
    try:
        import glob as _glob
        backups = sorted(_glob.glob(os.path.join(WORKSPACE, 'memory', '*.md')))
        if backups:
            last_backup = os.path.basename(backups[-1]).replace('.md', '')
    except Exception:
        pass

    return {
        "ram": {"used_mb": ram_used, "total_mb": ram_total, "pct": ram_pct},
        "swap": {"used_mb": swap_used, "total_mb": swap_total, "pct": swap_pct},
        "cpu": {"load1": load1, "load5": load5, "load15": load15, "cores": cpu_count, "pct": cpu_pct},
        "disk": {"used_gb": disk_used, "total_gb": disk_total, "pct": disk_pct, "free_gb": disk_free},
        "disk_d": disk_d,
        "uptime": uptime,
        "ip": ip,
        "kernel": kernel,
        "procs": proc_count,
        "gateway": 'claude-code',
        "git": git_info,
        "mission_control": svc('mission_control_api'),
        "cloudflared":     svc('cloudflared'),
        "telegram":        svc('telegram_bot'),
        "claude_version":  claude_ver,
        "last_backup":     last_backup,
    }


def get_book_health():
    import subprocess, base64, json as _json
    ps = (
        "[Console]::OutputEncoding=[System.Text.Encoding]::UTF8;"
        "$p=(Get-WmiObject Win32_Processor);"
        "$os=Get-WmiObject Win32_OperatingSystem;"
        "$cs=Get-WmiObject Win32_ComputerSystem;"
        "$rT=[math]::Round($os.TotalVisibleMemorySize/1KB,0);"
        "$rF=[math]::Round($os.FreePhysicalMemory/1KB,0);"
        "$rP=[math]::Round($cs.TotalPhysicalMemory/1GB,0);"
        "$dk=Get-WmiObject Win32_LogicalDisk -Filter \"DeviceID='C:'\";"
        "$dT=[math]::Round($dk.Size/1GB,0);"
        "$dF=[math]::Round($dk.FreeSpace/1GB,0);"
        "$gpus=(Get-WmiObject Win32_VideoController).Name -join ', ';"
        "$bs=Get-WmiObject -Namespace root/wmi -Class BatteryStatus;"
        "$b1=$bs|?{$_.InstanceName -like '*1_0'};"
        "$b2=$bs|?{$_.InstanceName -like '*2_0'};"
        "$fc=(Get-WmiObject -Namespace root/wmi -Class BatteryFullChargedCapacity|?{$_.InstanceName -like '*2_0'}).FullChargedCapacity;"
        "$up=(Get-Date)-(Get-CimInstance Win32_OperatingSystem).LastBootUpTime;"
        "@{cpu=$p.LoadPercentage;cpuName=$p.Name;ramUsed=($rT-$rF);ramTotal=$rT;ramPhys=$rP;"
        "diskTotal=$dT;diskFree=$dF;gpu=$gpus;winVer=$os.Caption;model=$cs.Model;"
        "bat1=if($b1){$b1.RemainingCapacity}else{-1};"
        "bat2=if($b2 -and $fc -gt 0){[math]::Round($b2.RemainingCapacity/$fc*100,0)}else{-1};"
        "bat2Charging=if($b2){[bool]$b2.Charging}else{$false};"
        "uptime=\"$([math]::Floor($up.TotalHours))h $($up.Minutes)m\"}|ConvertTo-Json"
    )
    enc = base64.b64encode(ps.encode('utf-16-le')).decode()
    try:
        r = subprocess.run(
            ['ssh', '-p', '2222', '-i', '/home/bolla/.ssh/id_ed25519',
             '-o', 'ConnectTimeout=5', '-o', 'StrictHostKeyChecking=no',
             'ernst@localhost', f'powershell -EncodedCommand {enc}'],
            capture_output=True, timeout=12
        )
        stdout = r.stdout.decode('utf-8', errors='replace') if r.stdout else ''
        if r.returncode == 0 and stdout.strip():
            d = _json.loads(stdout)
            d['online'] = True
            return d
    except Exception:
        pass
    return {'online': False}


def get_pro_health():
    import subprocess, base64, json as _json
    ps = (
        "[Console]::OutputEncoding=[System.Text.Encoding]::UTF8;"
        "$p=(Get-WmiObject Win32_Processor);"
        "$os=Get-WmiObject Win32_OperatingSystem;"
        "$cs=Get-WmiObject Win32_ComputerSystem;"
        "$rT=[math]::Round($os.TotalVisibleMemorySize/1KB,0);"
        "$rF=[math]::Round($os.FreePhysicalMemory/1KB,0);"
        "$rP=[math]::Round($cs.TotalPhysicalMemory/1GB,0);"
        "$dk=Get-WmiObject Win32_LogicalDisk -Filter \"DeviceID='C:'\";"
        "$dT=[math]::Round($dk.Size/1GB,0);"
        "$dF=[math]::Round($dk.FreeSpace/1GB,0);"
        "$gpus=(Get-WmiObject Win32_VideoController).Name -join ', ';"
        "$bs=Get-WmiObject -Namespace root/wmi -Class BatteryStatus;"
        "$b1=$bs|?{$_.InstanceName -like '*1_0'};"
        "$fc=(Get-WmiObject -Namespace root/wmi -Class BatteryFullChargedCapacity|?{$_.InstanceName -like '*1_0'}).FullChargedCapacity;"
        "$up=(Get-Date)-(Get-CimInstance Win32_OperatingSystem).LastBootUpTime;"
        "@{cpu=$p.LoadPercentage;cpuName=$p.Name;ramUsed=($rT-$rF);ramTotal=$rT;ramPhys=$rP;"
        "diskTotal=$dT;diskFree=$dF;gpu=$gpus;winVer=$os.Caption;model=$cs.Model;"
        "bat1=if($b1 -and $fc -gt 0){[math]::Round($b1.RemainingCapacity/$fc*100,0)}else{-1};"
        "bat1Charging=if($b1){[bool]$b1.Charging}else{$false};"
        "uptime=\"$([math]::Floor($up.TotalHours))h $($up.Minutes)m\"}|ConvertTo-Json"
    )
    enc = base64.b64encode(ps.encode('utf-16-le')).decode()
    try:
        r = subprocess.run(
            ['ssh', '-p', '2223', '-i', '/home/bolla/.ssh/id_ed25519',
             '-o', 'ConnectTimeout=5', '-o', 'StrictHostKeyChecking=no',
             'ernst@localhost', f'powershell -EncodedCommand {enc}'],
            capture_output=True, timeout=12
        )
        stdout = r.stdout.decode('utf-8', errors='replace') if r.stdout else ''
        if r.returncode == 0 and stdout.strip():
            d = _json.loads(stdout)
            d['online'] = True
            return d
    except Exception:
        pass
    return {'online': False}


def get_studio_health():
    import subprocess, base64, json as _json
    result = {'online': True}
    # Echte Windows-Hardware via powershell.exe (WSL2 → Windows direkt)
    ps = (
        "[Console]::OutputEncoding=[System.Text.Encoding]::UTF8;"
        "$p=(Get-WmiObject Win32_Processor);"
        "$os=Get-WmiObject Win32_OperatingSystem;"
        "$cs=Get-WmiObject Win32_ComputerSystem;"
        "$rT=[math]::Round($os.TotalVisibleMemorySize/1KB,0);"
        "$rF=[math]::Round($os.FreePhysicalMemory/1KB,0);"
        "$rP=[math]::Round($cs.TotalPhysicalMemory/1GB,0);"
        "$dk=Get-WmiObject Win32_LogicalDisk -Filter \"DeviceID='C:'\";"
        "$dT=[math]::Round($dk.Size/1GB,0);"
        "$dF=[math]::Round($dk.FreeSpace/1GB,0);"
        "$gpus=(Get-WmiObject Win32_VideoController).Name -join ', ';"
        "$up=(Get-Date)-(Get-CimInstance Win32_OperatingSystem).LastBootUpTime;"
        "@{cpu=$p.LoadPercentage;cpuName=$p.Name;ramUsed=($rT-$rF);ramTotal=$rT;ramPhys=$rP;"
        "diskTotal=$dT;diskFree=$dF;gpu=$gpus;winVer=$os.Caption;model=$cs.Model;"
        "uptime=\"$([math]::Floor($up.TotalHours))h $($up.Minutes)m\"}|ConvertTo-Json"
    )
    enc = base64.b64encode(ps.encode('utf-16-le')).decode()
    try:
        r = subprocess.run(
            ['/mnt/c/WINDOWS/System32/WindowsPowerShell/v1.0/powershell.exe', '-EncodedCommand', enc],
            capture_output=True, timeout=30
        )
        stdout = r.stdout.decode('utf-8', errors='replace') if r.stdout else ''
        if r.returncode == 0 and stdout.strip():
            result.update(_json.loads(stdout))
    except Exception as e:
        result['hwError'] = str(e)
    return result


def get_energie():
    import subprocess, base64, json as _json, threading

    PS_BAT = (
        "[Console]::OutputEncoding=[System.Text.Encoding]::UTF8;"
        "$bs=@(Get-WmiObject -Namespace root/wmi -Class BatteryStatus);"
        "$fcs=@(Get-WmiObject -Namespace root/wmi -Class BatteryFullChargedCapacity);"
        "$b1=if($bs.Count -ge 1){$bs[0]}else{$null};"
        "$b2=if($bs.Count -ge 2){$bs[1]}else{$null};"
        "$fc1=if($fcs.Count -ge 1){$fcs[0].FullChargedCapacity}else{0};"
        "$fc2=if($fcs.Count -ge 2){$fcs[1].FullChargedCapacity}else{0};"
        "$planStr=powercfg /GetActiveScheme|Out-String;"
        "$planName=if($planStr -match '\\((.+)\\)'){$matches[1].Trim()}else{'?'};"
        "$cpu=(Get-WmiObject Win32_Processor).LoadPercentage;"
        "$locked=[bool](Get-Process LogonUI -ErrorAction SilentlyContinue);"
        "$tz=Get-WmiObject -Namespace root/WMI -Class MSAcpi_ThermalZoneTemperature -EA SilentlyContinue|Select-Object -First 1;"
        "$cpuTempC=if($tz){[math]::Round(($tz.CurrentTemperature/10)-273.15,0)}else{-1};"
        "@{cpu=$cpu;cpuTempC=$cpuTempC;locked=$locked;powerPlan=$planName;"
        "powerOnline=if($b1){[bool]$b1.PowerOnline}else{$true};"
        "bat1Pct=if($b1 -and $fc1 -gt 0){[math]::Round($b1.RemainingCapacity/$fc1*100,0)}else{-1};"
        "bat2Pct=if($b2 -and $fc2 -gt 0){[math]::Round($b2.RemainingCapacity/$fc2*100,0)}else{-1};"
        "dischargeW=if($b1){[math]::Round($b1.DischargeRate/1000,1)}else{0};"
        "chargeW=if($b1){[math]::Round($b1.ChargeRate/1000,1)}else{0};"
        "voltV=if($b1){[math]::Round($b1.Voltage/1000,2)}else{0}}|ConvertTo-Json"
    )
    PS_STUDIO = (
        "[Console]::OutputEncoding=[System.Text.Encoding]::UTF8;"
        "$planStr=powercfg /GetActiveScheme|Out-String;"
        "$planName=if($planStr -match '\\((.+)\\)'){$matches[1].Trim()}else{'?'};"
        "$p=Get-WmiObject Win32_Processor;"
        "$cpu=$p.LoadPercentage;"
        "$cpuName=$p.Name;"
        "$tz=Get-WmiObject -Namespace root/WMI -Class MSAcpi_ThermalZoneTemperature -EA SilentlyContinue|Select-Object -First 1;"
        "$cpuTempC=if($tz){[math]::Round(($tz.CurrentTemperature/10)-273.15,0)}else{-1};"
        "$gpu=(Get-WmiObject Win32_VideoController -EA SilentlyContinue|Select-Object -First 1).Name;"
        "$os=Get-WmiObject Win32_OperatingSystem;"
        "$rT=[math]::Round($os.TotalVisibleMemorySize/1KB,0);"
        "$rF=[math]::Round($os.FreePhysicalMemory/1KB,0);"
        "@{cpu=$cpu;cpuName=$cpuName;cpuTempC=$cpuTempC;gpuName=$gpu;powerPlan=$planName;"
        "ramUsed=($rT-$rF);ramTotal=$rT;"
        "powerOnline=$true;bat1Pct=-1;bat2Pct=-1;dischargeW=0;chargeW=0;voltV=0}|ConvertTo-Json"
    )

    results = {}

    def query_studio():
        enc = base64.b64encode(PS_STUDIO.encode('utf-16-le')).decode()
        try:
            r = subprocess.run(
                ['/mnt/c/WINDOWS/System32/WindowsPowerShell/v1.0/powershell.exe', '-EncodedCommand', enc],
                capture_output=True, timeout=15)
            s = r.stdout.decode('utf-8', errors='replace') if r.stdout else ''
            if r.returncode == 0 and s.strip():
                data = {**_json.loads(s), 'online': True}
                # GPU-Watt via nvidia-smi (real measurement)
                try:
                    ng = subprocess.run(
                        ['/mnt/c/WINDOWS/System32/nvidia-smi.exe',
                         '--query-gpu=power.draw', '--format=csv,noheader,nounits'],
                        capture_output=True, timeout=5)
                    gpu_w = round(float(ng.stdout.decode().strip()), 1)
                    data['gpuW'] = gpu_w
                except Exception:
                    data['gpuW'] = None
                # CPU-Watt-Schätzung: Load% × 45W TDP (i7-13700H) + 5W Basis
                cpu_pct = data.get('cpu') or 0
                data['cpuW'] = round(cpu_pct * 45 / 100 + 5, 1)
                # Gesamtschätzung: CPU + GPU + ~28W Basis (Display, RAM, SSD, Board)
                base_w = 28
                data['totalW'] = round(data['cpuW'] + (data['gpuW'] or 8) + base_w, 0)
                results['studio'] = data
            else:
                results['studio'] = {'online': False}
        except Exception:
            results['studio'] = {'online': False}

    def query_remote(port, key):
        enc = base64.b64encode(PS_BAT.encode('utf-16-le')).decode()
        try:
            r = subprocess.run(
                ['ssh', '-p', str(port), '-i', '/home/bolla/.ssh/id_ed25519',
                 '-o', 'ConnectTimeout=5', '-o', 'StrictHostKeyChecking=no',
                 'ernst@localhost', f'powershell -EncodedCommand {enc}'],
                capture_output=True, timeout=15)
            s = r.stdout.decode('utf-8', errors='replace') if r.stdout else ''
            results[key] = {**_json.loads(s), 'online': True} if r.returncode == 0 and s.strip() else {'online': False}
        except Exception:
            results[key] = {'online': False}

    threads = [
        threading.Thread(target=query_studio),
        threading.Thread(target=query_remote, args=(2222, 'book')),
        threading.Thread(target=query_remote, args=(2223, 'pro')),
    ]
    for t in threads: t.start()
    for t in threads: t.join(timeout=16)
    return results


def do_book_action(action):
    import subprocess, base64
    if action == 'check':
        return get_book_health()
    if action == 'reboot':
        ps = 'Restart-Computer -Force'
        enc = base64.b64encode(ps.encode('utf-16-le')).decode()
        subprocess.run(
            ['ssh', '-p', '2222', '-i', '/home/bolla/.ssh/id_ed25519',
             '-o', 'ConnectTimeout=5', '-o', 'StrictHostKeyChecking=no',
             'ernst@localhost', f'powershell -EncodedCommand {enc}'],
            timeout=10
        )
        return {'ok': True, 'msg': 'Neustart gesendet'}
    if action == 'restart_tunnel':
        # Killt evtl. hängenden ssh und triggert den BollaTunnel-Task (= tunnel_hidden.vbs → tunnel.ps1)
        # Vorher zeigte der Aufruf auf C:\Users\ernst\surface_book_tunnel.ps1 — Datei existiert nicht.
        ps = ('Stop-Process -Name ssh -Force -ErrorAction SilentlyContinue;'
              'Start-Sleep 2;'
              'Start-Process wscript.exe -ArgumentList "C:\\\\ProgramData\\\\Bolla\\\\tunnel_hidden.vbs" -WindowStyle Hidden')
        enc = base64.b64encode(ps.encode('utf-16-le')).decode()
        subprocess.run(
            ['ssh', '-p', '2222', '-i', '/home/bolla/.ssh/id_ed25519',
             '-o', 'ConnectTimeout=5', '-o', 'StrictHostKeyChecking=no',
             'ernst@localhost', f'powershell -EncodedCommand {enc}'],
            timeout=10
        )
        return {'ok': True, 'msg': 'Tunnel-Neustart gesendet'}
    return {'error': 'Unbekannte Aktion'}


def do_pro_action(action):
    import subprocess, base64
    if action == 'check':
        return get_pro_health()
    if action == 'reboot':
        ps = 'Restart-Computer -Force'
        enc = base64.b64encode(ps.encode('utf-16-le')).decode()
        subprocess.run(
            ['ssh', '-p', '2223', '-i', '/home/bolla/.ssh/id_ed25519',
             '-o', 'ConnectTimeout=5', '-o', 'StrictHostKeyChecking=no',
             'ernst@localhost', f'powershell -EncodedCommand {enc}'],
            timeout=10
        )
        return {'ok': True, 'msg': 'Neustart gesendet'}
    if action == 'restart_tunnel':
        ps = ('Stop-Process -Name ssh -Force -ErrorAction SilentlyContinue;'
              'Start-Sleep 2;'
              'Start-Process powershell -ArgumentList "-WindowStyle Hidden -File C:\\\\Users\\\\renat\\\\surface_pro_tunnel.ps1" -WindowStyle Hidden')
        enc = base64.b64encode(ps.encode('utf-16-le')).decode()
        subprocess.run(
            ['ssh', '-p', '2223', '-i', '/home/bolla/.ssh/id_ed25519',
             '-o', 'ConnectTimeout=5', '-o', 'StrictHostKeyChecking=no',
             'ernst@localhost', f'powershell -EncodedCommand {enc}'],
            timeout=10
        )
        return {'ok': True, 'msg': 'Tunnel-Neustart gesendet'}
    return {'error': 'Unbekannte Aktion'}


def do_studio_action(action):
    import subprocess
    if action == 'restart_mc':
        subprocess.Popen(['pkill', '-f', 'mission_control_api'])
        return {'ok': True, 'msg': 'MC-Server wird neu gestartet...'}
    if action == 'restart_cloudflared':
        subprocess.run(['pkill', '-f', 'cloudflared'], capture_output=True)
        subprocess.Popen(['/home/bolla/.local/bin/cloudflared', 'tunnel', 'run', 'bolla-mc'])
        return {'ok': True, 'msg': 'Cloudflared neu gestartet'}
    if action == 'git_push':
        r = subprocess.run(
            ['git', '-C', WORKSPACE, 'add', '-A'],
            capture_output=True, text=True, timeout=15
        )
        r2 = subprocess.run(
            ['git', '-C', WORKSPACE, 'commit', '-m', 'auto: MC git push'],
            capture_output=True, text=True, timeout=15
        )
        r3 = subprocess.run(
            ['git', '-C', WORKSPACE, 'push'],
            capture_output=True, text=True, timeout=30
        )
        return {'ok': r3.returncode == 0, 'msg': r3.stdout.strip() or r3.stderr.strip() or 'Push abgeschlossen'}
    return {'error': 'Unbekannte Aktion'}


_token_cache = {"ts": 0, "data": None}
_halfday_cache = {"ts": 0, "data": None}

def get_token_halfdays(days=7):
    import time as _t
    if _halfday_cache["data"] and _t.time() - _halfday_cache["ts"] < 120:
        return _halfday_cache["data"]

    try:
        import zoneinfo
        tz = zoneinfo.ZoneInfo("Europe/Berlin")
    except Exception:
        tz = timezone.utc

    now_local = datetime.now(timezone.utc).astimezone(tz)
    cutoff = now_local - timedelta(days=days)
    buckets = {}
    proj_dir = "/home/bolla/.claude/projects"
    try:
        for root, _dirs, files in os.walk(proj_dir):
            for fn in files:
                if not fn.endswith(".jsonl"):
                    continue
                try:
                    with open(os.path.join(root, fn), encoding="utf-8", errors="ignore") as f:
                        for line in f:
                            try:
                                d = json.loads(line)
                            except Exception:
                                continue
                            msg = d.get("message")
                            if not isinstance(msg, dict):
                                continue
                            u = msg.get("usage")
                            if not isinstance(u, dict):
                                continue
                            ts = d.get("timestamp", "")
                            if not isinstance(ts, str):
                                continue
                            try:
                                dt = datetime.fromisoformat(ts.replace("Z", "+00:00"))
                            except Exception:
                                continue
                            local = dt.astimezone(tz)
                            if local < cutoff:
                                continue
                            date_str = local.strftime("%Y-%m-%d")
                            half = "AM" if local.hour < 12 else "PM"
                            key = (date_str, half)
                            b = buckets.setdefault(key, {"input":0,"output":0,"cache_read":0,"cache_creation":0})
                            b["input"] += u.get("input_tokens", 0) or 0
                            b["output"] += u.get("output_tokens", 0) or 0
                            b["cache_read"] += u.get("cache_read_input_tokens", 0) or 0
                            b["cache_creation"] += u.get("cache_creation_input_tokens", 0) or 0
                            # Modell-Zählung
                            mdl = msg.get("model", "") or ""
                            if "opus" in mdl.lower():
                                b["opus_calls"] = b.get("opus_calls", 0) + 1
                                b["opus_out"] = b.get("opus_out", 0) + (u.get("output_tokens", 0) or 0)
                            elif "sonnet" in mdl.lower():
                                b["sonnet_calls"] = b.get("sonnet_calls", 0) + 1
                except Exception:
                    continue
    except Exception as e:
        print(f"halfday error: {e}")

    result = []
    total_opus_calls = total_sonnet_calls = total_opus_out = total_out_7d = 0
    for i in range(days - 1, -1, -1):
        day = now_local - timedelta(days=i)
        date_str = day.strftime("%Y-%m-%d")
        weekday = ["Mo","Di","Mi","Do","Fr","Sa","So"][day.weekday()]
        for half in ["AM", "PM"]:
            b = buckets.get((date_str, half), {"input":0,"output":0,"cache_read":0,"cache_creation":0})
            total_opus_calls += b.get("opus_calls", 0)
            total_sonnet_calls += b.get("sonnet_calls", 0)
            total_opus_out += b.get("opus_out", 0)
            total_out_7d += b.get("output", 0)
            result.append({
                "date": date_str,
                "weekday": weekday,
                "label": f"{weekday} {day.strftime('%d.%m.')}",
                "half": half,
                "input": b.get("input", 0),
                "output": b.get("output", 0),
                "cache_read": b.get("cache_read", 0),
                "cache_creation": b.get("cache_creation", 0),
                "total_in": b.get("input", 0) + b.get("cache_read", 0) + b.get("cache_creation", 0),
                "total_out": b.get("output", 0),
                "opus_calls": b.get("opus_calls", 0),
                "sonnet_calls": b.get("sonnet_calls", 0),
            })
    avg_out_per_day = total_out_7d // max(days, 1)
    data = {
        "halfdays": result,
        "days": days,
        "models": {
            "opus_calls_7d": total_opus_calls,
            "sonnet_calls_7d": total_sonnet_calls,
            "opus_out_7d": total_opus_out,
            "avg_out_per_day": avg_out_per_day,
        }
    }
    _halfday_cache["data"] = data
    _halfday_cache["ts"] = _t.time()
    return data


def get_token_usage():
    """Summiert Claude-Code Token-Verbrauch aus den Session-JSONL-Dateien."""
    import time as _t
    if _token_cache["data"] and _t.time() - _token_cache["ts"] < 60:
        return _token_cache["data"]

    today = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    t_in = t_out = t_cr = t_ce = 0
    a_in = a_out = a_cr = a_ce = 0
    latest_ts = ""
    latest_model = ""
    proj_dir = "/home/bolla/.claude/projects"
    try:
        for root, _dirs, files in os.walk(proj_dir):
            for fn in files:
                if not fn.endswith(".jsonl"):
                    continue
                try:
                    with open(os.path.join(root, fn), encoding="utf-8", errors="ignore") as f:
                        for line in f:
                            try:
                                d = json.loads(line)
                            except Exception:
                                continue
                            msg = d.get("message")
                            if not isinstance(msg, dict):
                                continue
                            ts = d.get("timestamp", "")
                            mdl = msg.get("model")
                            if isinstance(mdl, str) and mdl and isinstance(ts, str) and ts > latest_ts:
                                latest_ts = ts
                                latest_model = mdl
                            u = msg.get("usage")
                            if not isinstance(u, dict):
                                continue
                            inp = u.get("input_tokens", 0) or 0
                            out = u.get("output_tokens", 0) or 0
                            cr = u.get("cache_read_input_tokens", 0) or 0
                            ce = u.get("cache_creation_input_tokens", 0) or 0
                            a_in += inp; a_out += out; a_cr += cr; a_ce += ce
                            if isinstance(ts, str) and ts.startswith(today):
                                t_in += inp; t_out += out; t_cr += cr; t_ce += ce
                except Exception:
                    continue
    except Exception as e:
        print(f"tokenusage error: {e}")

    def _pretty(m):
        if not m:
            return "Claude"
        p = m.split("-")
        if len(p) >= 4 and p[0] == "claude":
            return f"Claude {p[1].capitalize()} {p[2]}.{p[3]}"
        return m

    # Konfigurierter Default aus settings.json (z.B. "sonnet")
    _alias = {"sonnet": "Sonnet 4.6", "opus": "Opus 4.8", "haiku": "Haiku 4.5"}
    default_raw = ""
    try:
        with open("/home/bolla/.claude/settings.json", encoding="utf-8") as _sf:
            default_raw = (json.load(_sf).get("model") or "").strip()
    except Exception:
        default_raw = ""
    default_pretty = _alias.get(default_raw.lower(), _pretty(default_raw) if default_raw else "")

    # Kurzname des Live-Modells fürs Frontend (z.B. "Opus 4.8")
    live_short = _pretty(latest_model).replace("Claude ", "")
    # Mismatch nur melden, wenn beide bekannt sind und Familie abweicht
    mismatch = False
    if default_raw and latest_model:
        mismatch = default_raw.lower() not in latest_model.lower()

    data = {
        "model": f"{_pretty(latest_model)} (Pro Plan)",
        "model_live": live_short,
        "model_default": default_pretty,
        "model_mismatch": mismatch,
        "today": {"input": t_in, "output": t_out, "cache_read": t_cr, "cache_creation": t_ce},
        "total": {"input": a_in, "output": a_out, "cache_read": a_cr, "cache_creation": a_ce},
        "note": "Pro Plan — keine Kosten"
    }
    _token_cache["ts"] = _t.time()
    _token_cache["data"] = data
    return data


_TB_FILE = "/home/bolla/workspace/config/token_budget.json"

def get_token_budget():
    """Berechnet Anthropic-Wochenbudget-Prozentsatz aus gespeichertem Wochenstart."""
    import time as _t
    try:
        tok = get_token_usage()
        total_out = tok["total"]["output"]

        with open(_TB_FILE, encoding="utf-8") as f:
            cfg = json.load(f)

        week_start_out = cfg.get("week_start_output", 0)
        limit = cfg.get("estimated_limit_output", 2000000)
        week_start_date = cfg.get("week_start_date", "")

        used = max(0, total_out - week_start_out)
        pct  = min(100, round(used / limit * 100, 1))
        rem  = round(100 - pct, 1)

        # Countdown bis nächsten Samstag 23:59 Europe/Berlin
        from datetime import timezone as _tz
        import pytz
        berlin = pytz.timezone("Europe/Berlin")
        now_berlin = datetime.now(berlin)
        day = now_berlin.weekday()          # 0=Mo … 5=Sa, 6=So
        # Samstag = weekday 5
        days_until_sat = (5 - day) % 7
        if days_until_sat == 0:
            # heute Samstag: wenn vor 23:59 noch heute, sonst in 7 Tagen
            if now_berlin.hour < 23 or (now_berlin.hour == 23 and now_berlin.minute < 59):
                days_until_sat = 0
            else:
                days_until_sat = 7
        from datetime import timedelta
        reset_dt = (now_berlin + timedelta(days=days_until_sat)).replace(
            hour=23, minute=59, second=0, microsecond=0)
        diff_s = int((reset_dt - now_berlin).total_seconds())
        h, r   = divmod(diff_s, 3600)
        m      = r // 60

        # Tipp je nach Status
        if pct < 50:
            tip = "Entspannt — alles möglich."
        elif pct < 75:
            tip = "Moderat — große Features bewusst planen."
        elif pct < 90:
            tip = "Kritisch — nur kleine Fixes & kurze Sessions."
        else:
            tip = "Fast leer — bis Samstag auf das Nötigste beschränken."

        return {
            "pct": pct,
            "remaining": rem,
            "used_output": used,
            "limit_output": limit,
            "week_start_date": week_start_date,
            "reset_hours": h,
            "reset_minutes": m,
            "tip": tip
        }
    except Exception as e:
        return {"error": str(e), "pct": None}


_claude_quota_cache = {"ts": 0, "data": None}

# Öffentliche Claude-Code OAuth client_id (für Token-Refresh)
_CLAUDE_OAUTH_CLIENT_ID = "9d1c250a-e61b-44d9-88ed-5944d1962f5e"
_CLAUDE_CRED_PATH = os.path.expanduser("~/.claude/.credentials.json")


def _refresh_claude_token():
    """Erneuert den abgelaufenen OAuth-Token via refreshToken und speichert atomar.
    Gibt den neuen accessToken zurück oder None bei Fehler."""
    import time as _t
    import urllib.request as _ur
    try:
        creds = json.load(open(_CLAUDE_CRED_PATH))
        o = creds["claudeAiOauth"]
        rt = o.get("refreshToken")
        if not rt:
            return None
        body = json.dumps({
            "grant_type": "refresh_token",
            "refresh_token": rt,
            "client_id": _CLAUDE_OAUTH_CLIENT_ID,
        }).encode()
        req = _ur.Request(
            "https://console.anthropic.com/v1/oauth/token",
            data=body,
            headers={"Content-Type": "application/json", "Accept": "application/json",
                     "User-Agent": "claude-cli/1.0 (external, cli)"},
        )
        with _ur.urlopen(req, timeout=15) as r:
            tok = json.loads(r.read().decode())
        new_access = tok.get("access_token")
        if not new_access:
            return None
        o["accessToken"] = new_access
        o["refreshToken"] = tok.get("refresh_token", rt)
        if tok.get("expires_in"):
            o["expiresAt"] = int((_t.time() + tok["expires_in"]) * 1000)
        creds["claudeAiOauth"] = o
        # Sicherheitskopie + atomares Schreiben (chmod 600)
        try:
            import shutil
            shutil.copy2(_CLAUDE_CRED_PATH, _CLAUDE_CRED_PATH + ".bak-autorefresh")
        except Exception:
            pass
        tmp = _CLAUDE_CRED_PATH + ".tmp"
        with open(tmp, "w") as f:
            json.dump(creds, f, indent=2)
        os.chmod(tmp, 0o600)
        os.replace(tmp, _CLAUDE_CRED_PATH)
        return new_access
    except Exception:
        return None


def get_claude_quota():
    """Ruft echte Claude Pro Usage direkt von Anthropic API ab (OAuth).
    Erneuert bei 401 (abgelaufener Token) automatisch via refreshToken."""
    import time as _t
    import urllib.request as _ur
    import urllib.error as _ue
    if _claude_quota_cache["data"] and _t.time() - _claude_quota_cache["ts"] < 300:
        return _claude_quota_cache["data"]

    def _fetch(token):
        req = _ur.Request(
            "https://api.anthropic.com/api/oauth/usage",
            headers={"Authorization": f"Bearer {token}", "anthropic-beta": "oauth-2025-04-20", "Accept": "application/json"}
        )
        with _ur.urlopen(req, timeout=10) as r:
            return json.loads(r.read().decode())

    try:
        creds = json.load(open(_CLAUDE_CRED_PATH))
        token = creds["claudeAiOauth"]["accessToken"]
        try:
            raw = _fetch(token)
        except _ue.HTTPError as he:
            if he.code != 401:
                raise
            # Token abgelaufen -> erneuern und einmal erneut versuchen
            new_token = _refresh_claude_token()
            if not new_token:
                raise
            raw = _fetch(new_token)

        fh = raw.get("five_hour") or {}
        sd = raw.get("seven_day") or {}
        fh_pct = round(fh.get("utilization", 0), 1)
        sd_pct = round(sd.get("utilization", 0), 1)
        sd_rem = round(100 - sd_pct, 1)

        # Reset-Countdown aus resets_at
        import pytz
        from datetime import timezone as _tz
        berlin = pytz.timezone("Europe/Berlin")
        now_berlin = datetime.now(berlin)
        from datetime import datetime as _dt

        def _parse_reset(resets_at_str):
            if not resets_at_str:
                return 0, 0, "–"
            rt = _dt.fromisoformat(resets_at_str.replace("Z", "+00:00")).astimezone(berlin)
            diff_s = max(0, int((rt - now_berlin).total_seconds()))
            h, r2 = divmod(diff_s, 3600)
            m = r2 // 60
            label = rt.strftime("%a %d.%m. %H:%M")
            return h, m, label

        # 7-Tage Reset
        reset_h, reset_m, reset_label = _parse_reset(sd.get("resets_at"))
        # 5h-Fenster Reset
        fh_reset_h, fh_reset_m, fh_reset_label = _parse_reset(fh.get("resets_at"))

        if sd_pct < 50:
            tip = "Entspannt — alles möglich."
            tip_color = "#22c55e"
        elif sd_pct < 75:
            tip = "Moderat — große Features bewusst planen."
            tip_color = "#f59e0b"
        elif sd_pct < 90:
            tip = "Kritisch — nur kleine Fixes & kurze Sessions."
            tip_color = "#f97316"
        else:
            tip = "Fast leer — bis Reset auf das Nötigste beschränken."
            tip_color = "#ef4444"

        data = {
            "five_hour_pct": fh_pct,
            "five_hour_reset_hours": fh_reset_h,
            "five_hour_reset_minutes": fh_reset_m,
            "five_hour_reset_label": fh_reset_label,
            "seven_day_pct": sd_pct,
            "seven_day_rem": sd_rem,
            "reset_hours": reset_h,
            "reset_minutes": reset_m,
            "reset_label": reset_label,
            "tip": tip,
            "tip_color": tip_color,
            "source": "anthropic_oauth"
        }
        _claude_quota_cache["ts"] = _t.time()
        _claude_quota_cache["data"] = data
        return data
    except Exception as e:
        return {"error": str(e)}


_KOSTEN_FILE = "/home/bolla/workspace/config/kosten.json"

_PLAN_INFO = {
    "max":         {"label": "Max Plan",  "betrag": 100.0, "einheit": "$/Monat"},
    "max_5x":      {"label": "Max Plan",  "betrag": 100.0, "einheit": "$/Monat"},
    "pro":         {"label": "Pro Plan",  "betrag": 20.0,  "einheit": "$/Monat"},
    "free":        {"label": "Free Plan", "betrag": 0.0,   "einheit": ""},
}

def get_kosten():
    try:
        with open(_KOSTEN_FILE, encoding="utf-8") as f:
            data = json.load(f)
        return data
    except Exception as e:
        return {"error": str(e)}

def kosten_update_guthaben(name, betrag, info=None):
    try:
        with open(_KOSTEN_FILE, encoding="utf-8") as f:
            data = json.load(f)
        for k in data["konten"]:
            if k["name"] == name:
                k["betrag"] = betrag
                k["aktualisiert"] = datetime.now().strftime("%Y-%m-%d")
                if info is not None:
                    k["info"] = info
        data["zuletzt_geaendert"] = datetime.now().strftime("%Y-%m-%d")
        with open(_KOSTEN_FILE, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        return {"ok": True}
    except Exception as e:
        return {"error": str(e)}


def token_budget_snapshot():
    """Samstags-Reset: speichert aktuellen total_output als neuen Wochenstart."""
    try:
        tok = get_token_usage()
        total_out = tok["total"]["output"]
        with open(_TB_FILE, encoding="utf-8") as f:
            cfg = json.load(f)

        # Limit aus letzter Woche berechnen falls möglich
        old_start = cfg.get("week_start_output", 0)
        week_used = total_out - old_start
        history = cfg.get("history", [])
        history.append({
            "week_start": cfg.get("week_start_date"),
            "output_used": week_used,
            "pct_of_limit": round(week_used / cfg.get("estimated_limit_output", 2000000) * 100, 1)
        })
        # Limit nach 3+ Wochen aus Maximalwert der letzten 4 Wochen schätzen
        if len(history) >= 3:
            recent_used = [h["output_used"] for h in history[-4:]]
            new_limit = int(max(recent_used) * 1.1)  # 10% Puffer
            cfg["estimated_limit_output"] = new_limit

        cfg["week_start_output"] = total_out
        cfg["week_start_date"] = datetime.now().strftime("%Y-%m-%d")
        cfg["history"] = history[-12:]  # max 12 Wochen History

        with open(_TB_FILE, "w", encoding="utf-8") as f:
            json.dump(cfg, f, indent=2, ensure_ascii=False)

        return {"ok": True, "new_start": total_out, "week_used": week_used}
    except Exception as e:
        return {"error": str(e)}


def _run_claude_stream(cmd, tmp_path=None):
    """Hilfsgenerator: führt claude-Kommando aus und liefert JSON-Events."""
    import subprocess
    proc = None
    try:
        proc = subprocess.Popen(
            cmd,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            bufsize=1,
            cwd=os.path.expanduser("~"),
        )
        for line in proc.stdout:
            line = line.strip()
            if not line:
                continue
            try:
                yield json.loads(line)
            except json.JSONDecodeError:
                yield {"type": "raw", "text": line}
        proc.wait(timeout=5)
        if proc.returncode != 0:
            err = proc.stderr.read() if proc.stderr else ""
            yield {"type": "error", "error": err.strip() or f"Exit {proc.returncode}", "returncode": proc.returncode}
    except Exception as e:
        yield {"type": "error", "error": str(e)}
    finally:
        if proc and proc.poll() is None:
            proc.kill()
        if tmp_path:
            try:
                os.unlink(tmp_path)
            except:
                pass


def bolla_chat_stream(message, session_id=None, image_b64=None):
    """Streamt Claude-Code-Events als Generator. Jeder yield ist ein JSON-Objekt.
    Bei ungültiger session_id wird automatisch ohne Resume neu gestartet."""
    import tempfile, base64, shutil
    claude_bin = shutil.which("claude") or os.path.expanduser("~/.local/bin/claude")

    tmp_path = None
    if image_b64:
        try:
            img_data = base64.b64decode(image_b64)
            tmp = tempfile.NamedTemporaryFile(suffix=".png", prefix="bolla_img_", delete=False)
            tmp.write(img_data)
            tmp.close()
            tmp_path = tmp.name
            message = f"Schau dir das Bild an: {tmp_path}\n\n{message}" if message else f"Beschreibe was du in diesem Bild siehst: {tmp_path}"
        except Exception as e:
            print(f"Bild-Fehler: {e}")

    full_msg = f"[MC-UI] {message}" if message else message

    if session_id:
        cmd = [claude_bin, "-p", "--output-format", "stream-json", "--verbose", "--resume", session_id, full_msg]
        events = list(_run_claude_stream(cmd, tmp_path=None))
        # Wenn Resume fehlschlägt (session abgelaufen), Fallback ohne Resume
        if events and events[-1].get("type") == "error":
            err = events[-1].get("error", "")
            if "session" in err.lower() or "resume" in err.lower() or "not found" in err.lower() or events[-1].get("returncode", 0) != 0:
                yield {"type": "system", "text": "⚠️ Alte Session abgelaufen — starte frischen Chat"}
                cmd_fresh = [claude_bin, "-p", "--output-format", "stream-json", "--verbose", full_msg]
                yield from _run_claude_stream(cmd_fresh, tmp_path=tmp_path)
                return
        yield from iter(events)
        if tmp_path:
            try:
                os.unlink(tmp_path)
            except:
                pass
    else:
        cmd = [claude_bin, "-p", "--output-format", "stream-json", "--verbose", full_msg]
        yield from _run_claude_stream(cmd, tmp_path=tmp_path)


def azure_list_voices():
    """Holt die Stimmen-Liste von Azure."""
    cfg = azure_speech_config()
    if not cfg:
        return {"error": "Azure Speech nicht konfiguriert. Key in config/azure_speech.json eintragen."}
    import urllib.request
    url = f"https://{cfg['region']}.tts.speech.microsoft.com/cognitiveservices/voices/list"
    req = urllib.request.Request(url, headers={"Ocp-Apim-Subscription-Key": cfg["key"]})
    try:
        with urllib.request.urlopen(req, timeout=10) as r:
            voices = json.loads(r.read())
        # Filtern: nur DE + EN, neural, kompakte Form
        filtered = [
            {
                "name": v["ShortName"],
                "display": v.get("LocalName", v["ShortName"]),
                "locale": v["Locale"],
                "gender": v["Gender"],
                "styles": v.get("StyleList", []),
            }
            for v in voices
            if v.get("VoiceType") == "Neural" and (v["Locale"].startswith("de-") or v["Locale"].startswith("en-"))
        ]
        filtered.sort(key=lambda x: (not x["locale"].startswith("de-"), x["locale"], x["name"]))
        return {"voices": filtered, "default": cfg.get("default_voice", "de-DE-SeraphinaMultilingualNeural")}
    except Exception as e:
        return {"error": str(e)}


def azure_tts(text, voice, rate_percent=0):
    """Synthetisiert Text → MP3-Bytes. Gibt (bytes, error) zurück."""
    cfg = azure_speech_config()
    if not cfg:
        return None, "Azure Speech nicht konfiguriert"
    import urllib.request
    # SSML escapen
    def esc(s):
        return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    rate_sign = "+" if rate_percent >= 0 else ""
    locale = "-".join(voice.split("-")[:2]) if "-" in voice else "de-DE"
    ssml = (
        f"<speak version='1.0' xml:lang='{locale}'>"
        f"<voice name='{voice}'>"
        f"<prosody rate='{rate_sign}{int(rate_percent)}%'>{esc(text)}</prosody>"
        f"</voice></speak>"
    )
    url = f"https://{cfg['region']}.tts.speech.microsoft.com/cognitiveservices/v1"
    req = urllib.request.Request(
        url,
        data=ssml.encode("utf-8"),
        headers={
            "Ocp-Apim-Subscription-Key": cfg["key"],
            "Content-Type": "application/ssml+xml",
            "X-Microsoft-OutputFormat": "audio-24khz-48kbitrate-mono-mp3",
            "User-Agent": "bolla-mission-control",
        },
        method="POST",
    )
    try:
        with urllib.request.urlopen(req, timeout=20) as r:
            return r.read(), None
    except Exception as e:
        return None, str(e)


# ============ ADB / SURFACE ============
def _adb_run(args, timeout=15):
    import subprocess
    try:
        r = subprocess.run(["adb"] + list(args), capture_output=True, timeout=timeout)
        return r.returncode, r.stdout, r.stderr.decode("utf-8", "replace").strip()
    except subprocess.TimeoutExpired:
        return -1, b"", "Timeout — Gerät nicht erreichbar?"
    except FileNotFoundError:
        return -2, b"", "adb nicht installiert"

def adb_devices():
    rc, out, err = _adb_run(["devices"], timeout=5)
    if rc < 0:
        return {"devices": [], "error": err}
    devices = []
    for line in out.decode("utf-8", "replace").splitlines()[1:]:
        parts = line.strip().split()
        if len(parts) >= 2:
            devices.append({"id": parts[0], "state": parts[1]})
    return {"devices": devices, "error": None if rc == 0 else (err or "Fehler")}

def adb_packages(filter_str="", kind="all"):
    args = ["shell", "pm", "list", "packages"]
    if kind == "user":
        args.append("-3")
    elif kind == "system":
        args.append("-s")
    rc, out, err = _adb_run(args, timeout=15)
    if rc != 0:
        return {"packages": [], "error": err or "adb-Fehler"}
    pkgs = []
    f = (filter_str or "").lower().strip()
    for line in out.decode("utf-8", "replace").splitlines():
        line = line.strip()
        if line.startswith("package:"):
            name = line[8:]
            if not f or f in name.lower():
                pkgs.append(name)
    pkgs.sort()
    return {"packages": pkgs, "count": len(pkgs), "error": None}

def adb_pair(host, code):
    host = (host or "").strip()
    code = (code or "").strip()
    if not host or not code:
        return {"ok": False, "error": "IP:Port und Code erforderlich"}
    if " " in host or "\n" in host or not code.isdigit() or len(code) != 6:
        return {"ok": False, "error": "Ungültige Eingabe (IP:Port und 6-stelliger Code erwartet)"}
    import subprocess
    try:
        r = subprocess.run(
            ["adb", "pair", host],
            input=(code + "\n").encode(),
            capture_output=True, timeout=15
        )
        msg = (r.stdout.decode("utf-8", "replace") + " " + r.stderr.decode("utf-8", "replace")).strip()
        ok = r.returncode == 0 and ("successfully" in msg.lower() or "paired" in msg.lower())
        return {"ok": ok, "message": msg, "error": None if ok else (msg or "Kopplung fehlgeschlagen")}
    except subprocess.TimeoutExpired:
        return {"ok": False, "error": "Timeout — Duo nicht erreichbar?"}
    except FileNotFoundError:
        return {"ok": False, "error": "adb nicht installiert"}

def adb_connect(host):
    host = (host or "").strip()
    if not host or " " in host or "\n" in host:
        return {"ok": False, "error": "Ungültige Adresse"}
    rc, out, err = _adb_run(["connect", host], timeout=10)
    msg = (out.decode("utf-8", "replace") + " " + err).strip()
    ok = rc == 0 and ("connected" in msg.lower() or "already" in msg.lower())
    return {"ok": ok, "message": msg, "error": None if ok else (msg or "Verbindung fehlgeschlagen")}

def adb_disconnect(host=""):
    args = ["disconnect"]
    if host:
        args.append(host)
    rc, out, err = _adb_run(args, timeout=5)
    return {"ok": rc == 0, "message": (out.decode("utf-8","replace") + " " + err).strip()}

def adb_screenshot():
    rc, png, err = _adb_run(["exec-out", "screencap", "-p"], timeout=15)
    if rc != 0 or not png:
        return None, err or "Screenshot fehlgeschlagen"
    return png, None

def adb_push(filename, remote_dir, data_bytes):
    import tempfile
    filename = (filename or "").strip()
    remote_dir = (remote_dir or "").strip()
    if not filename or "/" in filename or "\\" in filename or "\0" in filename:
        return {"ok": False, "error": "Ungültiger Dateiname"}
    if not remote_dir or "\n" in remote_dir or "\0" in remote_dir:
        return {"ok": False, "error": "Ungültiger Zielpfad"}
    if not data_bytes:
        return {"ok": False, "error": "Keine Daten"}
    with tempfile.NamedTemporaryFile(delete=False, suffix="_" + filename) as tf:
        tf.write(data_bytes)
        local = tf.name
    try:
        remote = remote_dir.rstrip("/") + "/" + filename
        rc, out, err = _adb_run(["push", local, remote], timeout=300)
        if rc != 0:
            return {"ok": False, "error": err or "Push fehlgeschlagen"}
        return {"ok": True, "message": f"→ {remote} ({len(data_bytes)} Bytes)", "bytes": len(data_bytes)}
    finally:
        try: os.unlink(local)
        except Exception: pass

def adb_ls(path="/storage/emulated/0"):
    import shlex
    path = (path or "/storage/emulated/0").strip() or "/"
    if any(c in path for c in "\n\r\0"):
        return {"error": "Ungültiger Pfad"}
    rc, out, err = _adb_run(["shell", f"ls -laL {shlex.quote(path)} 2>/dev/null || ls -la {shlex.quote(path)}"], timeout=15)
    if rc != 0:
        return {"error": err or "Pfad nicht lesbar"}
    entries = []
    for line in out.decode("utf-8", "replace").splitlines():
        line = line.rstrip()
        if not line or line.startswith("total "):
            continue
        parts = line.split(None, 7)
        if len(parts) < 8:
            continue
        perms = parts[0]
        raw_name = parts[7]
        # Broken symlinks (l?????????) haben nur 7 Felder vor dem Namen statt 8
        if raw_name.startswith("-> ") and len(parts) > 6:
            name = parts[6]
        else:
            name = raw_name.split(" -> ", 1)[0]
        if name in (".", "..", ""):
            continue
        # Unleserliche Einträge überspringen
        if perms.startswith("?") or "?" in perms[:10] and perms.count("?") > 3:
            continue
        if perms.startswith("d"):
            kind = "dir"
        elif perms.startswith("l"):
            kind = "link"
        else:
            kind = "file"
        try:
            size = int(parts[4])
        except Exception:
            size = 0
        entries.append({"name": name, "type": kind, "size": size})
    entries.sort(key=lambda e: (e["type"] == "file", e["name"].lower()))
    normalized = path.rstrip("/") or "/"
    if normalized == "/":
        parent = None
    else:
        parent = "/".join(normalized.split("/")[:-1]) or "/"
    return {"path": normalized, "parent": parent, "entries": entries, "error": None}

def adb_info(kind="device"):
    kind = (kind or "device").strip().lower()
    if kind == "device":
        rc, out, err = _adb_run([
            "shell", "sh", "-c",
            "echo MODEL=$(getprop ro.product.model); "
            "echo MANUFACTURER=$(getprop ro.product.manufacturer); "
            "echo ANDROID=$(getprop ro.build.version.release); "
            "echo SDK=$(getprop ro.build.version.sdk); "
            "echo SIZE=$(wm size 2>/dev/null | grep -oE '[0-9]+x[0-9]+' | head -1)"
        ], timeout=10)
        if rc != 0:
            return {"error": err or "adb-Fehler"}
        props = {}
        for line in out.decode("utf-8", "replace").splitlines():
            if "=" in line:
                k, v = line.split("=", 1)
                props[k.strip()] = v.strip()
        return {"kind": "device", "props": props, "error": None}
    if kind == "battery":
        rc, out, err = _adb_run(["shell", "dumpsys", "battery"], timeout=10)
        if rc != 0:
            return {"error": err or "adb-Fehler"}
        props = {}
        for line in out.decode("utf-8", "replace").splitlines():
            line = line.strip()
            if ":" in line:
                k, v = line.split(":", 1)
                props[k.strip().replace(" ", "_")] = v.strip()
        status_map = {"1":"Unbekannt","2":"Lädt","3":"Entlädt","4":"Nicht lädt","5":"Voll"}
        if "status" in props:
            props["status_text"] = status_map.get(props["status"], props["status"])
        return {"kind": "battery", "props": props, "error": None}
    if kind == "foreground":
        rc, out, err = _adb_run(["shell", "dumpsys", "activity", "activities"], timeout=10)
        if rc != 0:
            return {"error": err or "adb-Fehler"}
        import re
        text = out.decode("utf-8", "replace")
        for pattern in (r'mResumedActivity[^\n]*\{[^}]*\s+(\S+)/(\S+)', r'topResumedActivity[^\n]*\{[^}]*\s+(\S+)/(\S+)'):
            m = re.search(pattern, text)
            if m:
                pkg, act = m.group(1), m.group(2)
                if act.startswith("."):
                    act = pkg + act
                return {"kind": "foreground", "package": pkg, "activity": act, "error": None}
        return {"error": "Keine aktive App erkennbar"}
    return {"error": f"Unbekannter Info-Typ: {kind}"}

def adb_pull(remote):
    import tempfile
    remote = (remote or "").strip()
    if not remote or "\0" in remote or "\n" in remote:
        return None, "Ungültiger Pfad"
    with tempfile.NamedTemporaryFile(delete=False) as tf:
        local = tf.name
    try:
        rc, _, err = _adb_run(["pull", remote, local], timeout=120)
        if rc != 0:
            return None, err or "Pull fehlgeschlagen"
        with open(local, "rb") as f:
            data = f.read()
        return data, None
    finally:
        try: os.unlink(local)
        except Exception: pass


# ═══════════════════════════════════════════════════════════
# FOTO-ANALYSE (LM Studio / Moondream)
# ═══════════════════════════════════════════════════════════

def _photo_load_results():
    os.makedirs(os.path.dirname(PHOTO_ANALYSIS_FILE), exist_ok=True)
    if os.path.exists(PHOTO_ANALYSIS_FILE):
        try:
            with open(PHOTO_ANALYSIS_FILE, encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return []

def _photo_save_results(results):
    os.makedirs(os.path.dirname(PHOTO_ANALYSIS_FILE), exist_ok=True)
    with open(PHOTO_ANALYSIS_FILE, "w", encoding="utf-8") as f:
        json.dump(results, f, ensure_ascii=False, indent=2)

GEMINI_API_KEY_FILE = os.path.join(WORKSPACE, "config/gemini_api.json")
GEMINI_DAILY_LIMIT = 1500

def _gemini_api_key():
    try:
        with open(GEMINI_API_KEY_FILE) as f:
            return json.load(f).get("api_key", "")
    except Exception:
        return ""

GEMINI_MODELS = ["gemini-2.5-flash", "gemini-2.5-flash-lite"]

def _photo_analyze_one_gemini(image_path, prompt):
    import base64 as _b64, io, urllib.request as _ur
    key = _gemini_api_key()
    if not key:
        raise RuntimeError("Kein Gemini API-Key gefunden")
    try:
        from PIL import Image as _Img
        img = _Img.open(image_path)
        img.thumbnail((1024, 1024), _Img.LANCZOS)
        if img.mode not in ("RGB", "L"):
            img = img.convert("RGB")
        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=85)
        img_b64 = _b64.b64encode(buf.getvalue()).decode()
    except ImportError:
        with open(image_path, "rb") as f:
            img_b64 = _b64.b64encode(f.read()).decode()
    payload = {
        "contents": [{"parts": [
            {"inline_data": {"mime_type": "image/jpeg", "data": img_b64}},
            {"text": prompt}
        ]}]
    }
    last_err = None
    for model in GEMINI_MODELS:
        try:
            req = _ur.Request(
                f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={key}",
                data=json.dumps(payload).encode(),
                headers={"Content-Type": "application/json"}
            )
            resp = json.loads(_ur.urlopen(req, timeout=60).read())
            return resp["candidates"][0]["content"]["parts"][0]["text"].strip()
        except _ur.HTTPError as e:
            last_err = e
            if e.code in (503, 429, 500):
                print(f"Gemini {model} → {e.code}, versuche nächstes Modell...")
                continue
            raise
    raise last_err

def _photo_analyze_one(image_path, prompt, model, lmstudio_url):
    import base64 as _b64, io
    ext = image_path.lower().rsplit(".", 1)[-1]
    mime = {"jpg": "image/jpeg", "jpeg": "image/jpeg", "png": "image/png",
            "gif": "image/gif", "webp": "image/webp", "bmp": "image/bmp"}.get(ext, "image/jpeg")

    def _load_img(max_px=1024):
        try:
            from PIL import Image as _Img
            img = _Img.open(image_path)
            img.thumbnail((max_px, max_px), _Img.LANCZOS)
            if img.mode not in ("RGB", "L"):
                img = img.convert("RGB")
            buf = io.BytesIO()
            img.save(buf, format="JPEG", quality=85)
            return _b64.b64encode(buf.getvalue()).decode(), "image/jpeg"
        except ImportError:
            with open(image_path, "rb") as f:
                return _b64.b64encode(f.read()).decode(), mime

    import urllib.request as _ur

    def _call(img_b64, img_mime):
        payload = {
            "model": model,
            "messages": [{"role": "user", "content": [
                {"type": "image_url", "image_url": {"url": f"data:{img_mime};base64,{img_b64}"}},
                {"type": "text", "text": prompt}
            ]}],
            "max_tokens": 300,
            "temperature": 0.1
        }
        data = json.dumps(payload).encode()
        req = _ur.Request(f"{lmstudio_url}/v1/chat/completions", data=data,
                          headers={"Content-Type": "application/json"})
        resp = json.loads(_ur.urlopen(req, timeout=60).read())
        return resp["choices"][0]["message"]["content"].strip()

    img_b64, img_mime = _load_img(1024)
    try:
        return _call(img_b64, img_mime)
    except _ur.HTTPError as e:
        if e.code == 500:
            # Bild verkleinern und nochmal versuchen
            img_b64, img_mime = _load_img(512)
            return _call(img_b64, img_mime)
        raise

def _photo_worker(folder, prompt, model, lmstudio_url, throttle=2.0, engine="lmstudio"):
    global _photo_job
    import time as _time
    IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".gif", ".webp", ".bmp", ".tif", ".tiff"}
    folders = folder if isinstance(folder, list) else [folder]
    all_files = []  # (fpath, base_folder)
    for f in folders:
        for root, _, files in os.walk(f):
            for fname in sorted(files):
                if os.path.splitext(fname.lower())[1] in IMAGE_EXTS:
                    all_files.append((os.path.join(root, fname), f))
    results = _photo_load_results()
    analyzed = {r["path"] for r in results}
    todo = [(fpath, base) for fpath, base in all_files if fpath not in analyzed]
    _photo_job["total"] = len(all_files)
    _photo_job["done"] = len(analyzed)
    _photo_job["errors"] = 0
    for fpath, base in todo:
        if _photo_job["stop"]:
            break
        try:
            if engine == "gemini":
                desc = _photo_analyze_one_gemini(fpath, prompt)
            else:
                desc = _photo_analyze_one(fpath, prompt, model, lmstudio_url)
            results.append({
                "path": fpath,
                "filename": os.path.basename(fpath),
                "folder": os.path.relpath(os.path.dirname(fpath), base),
                "description": desc,
                "analyzed_at": datetime.now().isoformat(),
                "engine": engine
            })
            _photo_save_results(results)
        except Exception as e:
            _photo_job["errors"] += 1
            err_detail = str(e)
            if hasattr(e, 'read'):
                try: err_detail += " | " + e.read().decode()[:200]
                except: pass
            print(f"Foto-Analyse Fehler [{engine}] {os.path.basename(fpath)}: {err_detail}")
            _photo_job["last_error"] = f"{os.path.basename(fpath)}: {err_detail[:120]}"
        _photo_job["done"] += 1
        if engine == "lmstudio" and throttle > 0 and not _photo_job["stop"]:
            _time.sleep(throttle)
    _photo_job["running"] = False

def get_lmstudio_config():
    """Findet die erreichbare LM Studio URL und das geladene Modell."""
    import subprocess as _sp, socket as _sock
    # Kandidaten: WSL2-Gateway + bekannte Windows-Interface-IPs
    candidates = []
    try:
        r = _sp.run(["ip", "route"], capture_output=True, text=True)
        for line in r.stdout.splitlines():
            if line.startswith("default"):
                candidates.append(line.split()[2])
    except Exception:
        pass
    candidates += ["10.5.0.2", "10.0.0.1", "192.168.1.1"]
    candidates = list(dict.fromkeys(candidates))  # deduplizieren

    for ip in candidates:
        try:
            s = _sock.create_connection((ip, 1234), timeout=1)
            s.close()
            base = f"http://{ip}:1234"
            # Geladenes Modell abfragen
            import urllib.request as _ur
            resp = json.loads(_ur.urlopen(f"{base}/v1/models", timeout=3).read())
            all_models = [m.get("id","") for m in resp.get("data", [])]
            PREFER = ["smolvlm", "moondream", "llava", "bakllava", "minicpm-v", "cogvlm", "internvl"]
            model = next((m for v in PREFER for m in all_models if v in m.lower()), \
                         all_models[0] if all_models else "moondream-2b-2025-04-14")
            return {"lmstudio_url": base, "model": model, "reachable": True, "all_models": all_models}
        except Exception:
            continue
    return {"lmstudio_url": "http://localhost:1234", "model": "smolvlm-500m-instruct", "reachable": False}

def photo_start(folder, prompt, model, lmstudio_url, throttle=2.0, engine="lmstudio"):
    global _photo_job
    if _photo_job["running"]:
        return {"error": "Analyse läuft bereits"}
    folders = folder if isinstance(folder, list) else [folder]
    folders = [os.path.expanduser(f) for f in folders]
    missing = [f for f in folders if not os.path.isdir(f)]
    if missing:
        return {"error": f"Ordner nicht gefunden: {missing[0]}"}
    try:
        throttle = float(throttle)
    except Exception:
        throttle = 2.0
    # Gemini: Limit-Check vor dem Start
    if engine == "gemini":
        IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".gif", ".webp", ".bmp", ".tif", ".tiff"}
        results = _photo_load_results()
        analyzed = {r["path"] for r in results}
        todo_count = 0
        for f in folders:
            for root, _, files in os.walk(f):
                for fname in files:
                    fpath = os.path.join(root, fname)
                    if os.path.splitext(fname.lower())[1] in IMAGE_EXTS and fpath not in analyzed:
                        todo_count += 1
        if todo_count > GEMINI_DAILY_LIMIT:
            return {"error": f"Gemini-Tageslimit: {todo_count} Fotos zu analysieren, aber max. {GEMINI_DAILY_LIMIT}/Tag erlaubt. Bitte Ordner aufteilen oder LM Studio verwenden."}
        if not _gemini_api_key():
            return {"error": "Kein Gemini API-Key in config/gemini_api.json gefunden"}
    label = folders[0] if len(folders) == 1 else f"{len(folders)} Ordner"
    _photo_job = {"running": True, "total": 0, "done": 0, "errors": 0, "stop": False, "folder": label, "engine": engine}
    import threading
    threading.Thread(target=_photo_worker, args=(folders, prompt, model, lmstudio_url, throttle, engine), daemon=True).start()
    return {"ok": True}

def photo_thumb(path):
    import base64 as _b64, io
    path = os.path.realpath(os.path.expanduser(path))
    allowed = ["/mnt/d/", "/mnt/c/Users/", os.path.expanduser("~/")]
    if not any(path.startswith(a) for a in allowed):
        return None, "Pfad nicht erlaubt"
    if not os.path.isfile(path):
        return None, "Datei nicht gefunden"
    try:
        from PIL import Image
        img = Image.open(path)
        img.thumbnail((240, 240), Image.LANCZOS)
        if img.mode in ("RGBA", "P"):
            img = img.convert("RGB")
        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=75)
        ext = "jpeg"
        return _b64.b64encode(buf.getvalue()).decode(), ext
    except ImportError:
        pass
    except Exception as e:
        return None, str(e)
    # Fallback: Datei direkt senden
    with open(path, "rb") as f:
        raw = f.read()
    ext = path.lower().rsplit(".", 1)[-1]
    if ext in ("jpg", "jpeg"):
        ext = "jpeg"
    elif ext == "png":
        ext = "png"
    else:
        ext = "jpeg"
    return _b64.b64encode(raw).decode(), ext


def photo_search(query):
    """German/any-language query → Claude expands to English terms → filter results."""
    if not query or not query.strip():
        return {"results": [], "terms": []}
    results = _photo_load_results()
    if not results:
        return {"results": [], "terms": [], "no_data": True}

    prompt = (
        "You are a photo search assistant. The user wants to find photos described in English.\n"
        "The user's search query may be in any language (often German).\n"
        "Return a JSON array of English search terms that would appear in a casual photo description.\n"
        "IMPORTANT rules:\n"
        "- Always include the BASIC/SIMPLE English word(s) for the concept first (e.g. 'wood' before 'timber')\n"
        "- Then add a few specific variants (e.g. 'firewood', 'wooden', 'log')\n"
        "- Photo descriptions use everyday language, not technical terms\n"
        "- 5-8 terms, simple first, specific last\n\n"
        f"User query: {query}\n\n"
        "Respond with ONLY a JSON array of strings. Example for 'Holz':\n"
        '[\"wood\", \"wooden\", \"firewood\", \"log\", \"timber\", \"pile of wood\", \"log cabin\"]'
    )
    try:
        r = subprocess.run([CLAUDE_BIN, "-p", prompt], capture_output=True, text=True, timeout=20)
        raw = (r.stdout or "").strip()
        import re as _re
        m = _re.search(r'\[.*?\]', raw, _re.DOTALL)
        if m:
            import json as _json
            terms = _json.loads(m.group())
            terms = [t.lower().strip() for t in terms if isinstance(t, str) and t.strip()]
        else:
            terms = [query.lower().strip()]
    except Exception:
        terms = [query.lower().strip()]

    import re as _re2
    matched = []
    # Kurze Begriffe (< 4 Zeichen) nur mit Word-Boundary matchen um False Positives zu vermeiden
    patterns = []
    for t in terms:
        if len(t) < 4:
            patterns.append(_re2.compile(r'\b' + _re2.escape(t) + r'\b', _re2.IGNORECASE))
        else:
            patterns.append(_re2.compile(_re2.escape(t), _re2.IGNORECASE))
    for r in results:
        desc = (r.get("description") or "")
        if any(p.search(desc) for p in patterns):
            matched.append(r)
    return {"results": matched, "terms": terms}


# ─── KI-Korrektur Excel-Klassenliste ─────────────────────────────────────────

def _excel_path(klasse, fach):
    import re as _re
    k = _re.sub(r'[^\w]', '_', klasse.strip())
    f = _re.sub(r'[^\w]', '_', fach.strip())
    return os.path.join(KORREKTUR_DIR, f"Klasse_{k}_{f}.xlsx")

def _parse_note_num(note_str):
    import re as _re
    m = _re.search(r'\b([1-6])\b', str(note_str or ''))
    return int(m.group(1)) if m else None

def _read_ka_students(ws):
    """Liest Schüler aus KA-Sheet (Zeilen 4–33)."""
    students = []
    for row_num in range(4, 34):
        vn = ws.cell(row=row_num, column=1).value
        nn = ws.cell(row=row_num, column=2).value
        nt = ws.cell(row=row_num, column=3).value
        if vn and str(vn).strip():
            students.append({
                'vorname': str(vn).strip(),
                'nachname': str(nn or '').strip(),
                'note': str(nt) if nt is not None else '',
            })
    return students

def _write_ka_sheet(ws, klasse, fach, ka_nr, thema, students):
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    HDR  = PatternFill("solid", fgColor="1E3A5F")
    SUB  = PatternFill("solid", fgColor="0F2538")
    COL  = PatternFill("solid", fgColor="16304F")
    AVG  = PatternFill("solid", fgColor="12243A")
    EVEN = PatternFill("solid", fgColor="0D1B2E")
    ODD  = PatternFill("solid", fgColor="0A1520")
    thin = Side(style='thin', color='2A4060')
    bd   = Border(left=thin, right=thin, top=thin, bottom=thin)
    GCOL = {1:"16A34A",2:"86EFAC",3:"FDE68A",4:"FCA5A5",5:"EF4444",6:"991B1B"}

    ws.delete_rows(1, ws.max_row or 1)

    # Zeile 1: Titel
    ws.merge_cells('A1:C1')
    c = ws['A1']
    c.value = f"Klasse {klasse}  ·  {fach}  ·  Klassenarbeit {ka_nr}"
    c.font = Font(bold=True, size=13, color="E2E8F0", name='Calibri')
    c.fill = HDR; c.alignment = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[1].height = 24

    # Zeile 2: Thema
    ws.merge_cells('A2:C2')
    c = ws['A2']
    c.value = f"Thema: {thema or '—'}"
    c.font = Font(italic=True, size=11, color="94A3B8", name='Calibri')
    c.fill = SUB; c.alignment = Alignment(horizontal='left', vertical='center', indent=1)
    ws.row_dimensions[2].height = 18

    # Zeile 3: Spaltenköpfe
    for col, lbl, align in [(1,'Vorname','left'),(2,'Nachname','left'),(3,'Note','center')]:
        c = ws.cell(row=3, column=col, value=lbl)
        c.font = Font(bold=True, size=11, color="FFFFFF", name='Calibri')
        c.fill = COL; c.border = bd
        c.alignment = Alignment(horizontal=align, vertical='center',
                                 indent=1 if align=='left' else 0)
    ws.row_dimensions[3].height = 20

    # Schülerzeilen 4–33
    notes = []
    for i in range(30):
        row = i + 4
        fill = EVEN if i % 2 == 0 else ODD
        s = students[i] if i < len(students) else {}
        vn, nn, nt_raw = s.get('vorname',''), s.get('nachname',''), s.get('note','')
        nt_num = _parse_note_num(nt_raw) if nt_raw else None
        if nt_num: notes.append(nt_num)

        ws.cell(row=row,column=1,value=vn or None).font = Font(size=11,name='Calibri',color="CBD5E1")
        ws.cell(row=row,column=1).fill = fill; ws.cell(row=row,column=1).border = bd
        ws.cell(row=row,column=1).alignment = Alignment(indent=1)
        ws.cell(row=row,column=2,value=nn or None).font = Font(size=11,name='Calibri',color="CBD5E1")
        ws.cell(row=row,column=2).fill = fill; ws.cell(row=row,column=2).border = bd
        ws.cell(row=row,column=2).alignment = Alignment(indent=1)

        c3 = ws.cell(row=row, column=3, value=nt_num or (nt_raw or None))
        c3.font = Font(size=12, bold=bool(nt_num), name='Calibri',
                       color="000000" if nt_num else "94A3B8")
        c3.fill = PatternFill("solid",fgColor=GCOL[nt_num]) if nt_num else fill
        c3.border = bd; c3.alignment = Alignment(horizontal='center')
        ws.row_dimensions[row].height = 18

    # Zeile 34: Durchschnitt
    avg = round(sum(notes)/len(notes), 2) if notes else None
    ws.cell(row=34,column=1,value=None).fill = AVG; ws.cell(row=34,column=1).border = bd
    ws.cell(row=34,column=2,value=f"Ø Durchschnitt  ({len(notes)} Schüler)").font = Font(bold=True,size=11,color="94A3B8",name='Calibri')
    ws.cell(row=34,column=2).fill = AVG; ws.cell(row=34,column=2).border = bd
    ws.cell(row=34,column=2).alignment = Alignment(horizontal='right')
    ws.cell(row=34,column=3,value=avg).font = Font(bold=True,size=13,color="F59E0B",name='Calibri')
    ws.cell(row=34,column=3).fill = AVG; ws.cell(row=34,column=3).border = bd
    ws.cell(row=34,column=3).alignment = Alignment(horizontal='center')
    ws.row_dimensions[34].height = 22

    ws.column_dimensions['A'].width = 18
    ws.column_dimensions['B'].width = 22
    ws.column_dimensions['C'].width = 9
    ws.freeze_panes = 'A4'

def _write_uebersicht_sheet(wb, klasse, fach):
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    ws = wb["Übersicht"]
    ws.delete_rows(1, ws.max_row or 1)
    HDR  = PatternFill("solid", fgColor="1E3A5F")
    COL  = PatternFill("solid", fgColor="16304F")
    AVG  = PatternFill("solid", fgColor="12243A")
    EVEN = PatternFill("solid", fgColor="0D1B2E")
    ODD  = PatternFill("solid", fgColor="0A1520")
    thin = Side(style='thin', color='2A4060')
    bd   = Border(left=thin, right=thin, top=thin, bottom=thin)
    GCOL = {1:"16A34A",2:"86EFAC",3:"FDE68A",4:"FCA5A5",5:"EF4444",6:"991B1B"}

    # Alle Schüler aus allen KA-Sheets zusammenführen
    all_students = {}
    for ka_nr in range(1, 5):
        ka_name = f"KA {ka_nr}"
        if ka_name not in wb.sheetnames:
            continue
        for s in _read_ka_students(wb[ka_name]):
            key = (s['vorname'].lower(), s['nachname'].lower())
            if key not in all_students:
                all_students[key] = {'vorname':s['vorname'],'nachname':s['nachname'],
                                     'ka1':None,'ka2':None,'ka3':None,'ka4':None}
            all_students[key][f'ka{ka_nr}'] = s.get('note') or None

    rows = sorted(all_students.values(), key=lambda s:(s['nachname'].lower(),s['vorname'].lower()))

    # Zeile 1: Titel
    ws.merge_cells('A1:G1')
    c = ws['A1']
    c.value = f"Klasse {klasse}  ·  {fach}  ·  Jahresübersicht"
    c.font = Font(bold=True,size=13,color="E2E8F0",name='Calibri')
    c.fill = HDR; c.alignment = Alignment(horizontal='center',vertical='center')
    ws.row_dimensions[1].height = 24

    # Zeile 2: Spaltenköpfe
    for col, lbl in enumerate(['Vorname','Nachname','KA 1','KA 2','KA 3','KA 4','Ø Gesamt'],1):
        c = ws.cell(row=2, column=col, value=lbl)
        c.font = Font(bold=True,size=11,color="FFFFFF",name='Calibri')
        c.fill = COL; c.border = bd
        c.alignment = Alignment(horizontal='center' if col>2 else 'left',
                                 vertical='center', indent=1 if col<=2 else 0)
    ws.row_dimensions[2].height = 20

    # Schülerzeilen
    for i, s in enumerate(rows[:30]):
        row = i + 3
        fill = EVEN if i%2==0 else ODD
        ws.cell(row=row,column=1,value=s['vorname']).font = Font(size=11,name='Calibri',color="CBD5E1")
        ws.cell(row=row,column=1).fill = fill; ws.cell(row=row,column=1).border = bd
        ws.cell(row=row,column=1).alignment = Alignment(indent=1)
        ws.cell(row=row,column=2,value=s['nachname']).font = Font(size=11,name='Calibri',color="CBD5E1")
        ws.cell(row=row,column=2).fill = fill; ws.cell(row=row,column=2).border = bd
        ws.cell(row=row,column=2).alignment = Alignment(indent=1)
        ka_notes = []
        for j, ka_key in enumerate(['ka1','ka2','ka3','ka4'],3):
            nt_raw = s.get(ka_key)
            nt_num = _parse_note_num(nt_raw) if nt_raw else None
            if nt_num: ka_notes.append(nt_num)
            c = ws.cell(row=row,column=j,value=nt_num or (nt_raw or None))
            c.font = Font(size=12,bold=bool(nt_num),name='Calibri')
            c.fill = PatternFill("solid",fgColor=GCOL[nt_num]) if nt_num else fill
            c.border = bd; c.alignment = Alignment(horizontal='center')
        avg_s = round(sum(ka_notes)/len(ka_notes),2) if ka_notes else None
        c = ws.cell(row=row,column=7,value=avg_s)
        c.font = Font(size=12,bold=True,name='Calibri',color="F59E0B" if avg_s else "94A3B8")
        c.fill = fill; c.border = bd; c.alignment = Alignment(horizontal='center')
        ws.row_dimensions[row].height = 18

    # Leere Zeilen bis Zeile 32
    for row in range(len(rows)+3, 33):
        fill = EVEN if (row-3)%2==0 else ODD
        for col in range(1,8):
            c = ws.cell(row=row,column=col,value=None)
            c.fill = fill; c.border = bd
        ws.row_dimensions[row].height = 18

    # Zeile 33: Klassendurchschnitt
    by_ka = {1:[],2:[],3:[],4:[]}
    for s in rows:
        for ka_nr in range(1,5):
            n = _parse_note_num(s.get(f'ka{ka_nr}'))
            if n: by_ka[ka_nr].append(n)
    ws.cell(row=33,column=1,value=None).fill = AVG; ws.cell(row=33,column=1).border = bd
    ws.cell(row=33,column=2,value="Ø Klassendurchschnitt").font = Font(bold=True,size=11,color="94A3B8",name='Calibri')
    ws.cell(row=33,column=2).fill = AVG; ws.cell(row=33,column=2).border = bd
    ws.cell(row=33,column=2).alignment = Alignment(horizontal='right')
    for j, ka_nr in enumerate(range(1,5),3):
        ns = by_ka[ka_nr]
        av = round(sum(ns)/len(ns),2) if ns else None
        c = ws.cell(row=33,column=j,value=av)
        c.font = Font(bold=True,size=12,color="F59E0B" if av else "94A3B8",name='Calibri')
        c.fill = AVG; c.border = bd; c.alignment = Alignment(horizontal='center')
    flat = [n for ns in by_ka.values() for n in ns]
    ov = round(sum(flat)/len(flat),2) if flat else None
    c = ws.cell(row=33,column=7,value=ov)
    c.font = Font(bold=True,size=13,color="F59E0B" if ov else "94A3B8",name='Calibri')
    c.fill = AVG; c.border = bd; c.alignment = Alignment(horizontal='center')
    ws.row_dimensions[33].height = 22

    ws.column_dimensions['A'].width = 18; ws.column_dimensions['B'].width = 22
    for col in ['C','D','E','F','G']: ws.column_dimensions[col].width = 9
    ws.freeze_panes = 'A3'

def excel_upsert_student(klasse, fach, ka_nr, thema, vorname, nachname, note_str):
    """Schüler in Klassenlisten-Excel eintragen oder aktualisieren."""
    from openpyxl import Workbook, load_workbook
    os.makedirs(KORREKTUR_DIR, exist_ok=True)
    fpath = _excel_path(klasse, fach)

    if os.path.exists(fpath):
        wb = load_workbook(fpath)
    else:
        wb = Workbook()
        if 'Sheet' in wb.sheetnames: del wb['Sheet']
        for n in range(1,5): wb.create_sheet(f"KA {n}")
        wb.create_sheet("Übersicht")

    for n in range(1,5):
        if f"KA {n}" not in wb.sheetnames: wb.create_sheet(f"KA {n}")
    if "Übersicht" not in wb.sheetnames: wb.create_sheet("Übersicht")

    ka_name = f"KA {ka_nr}"
    ws = wb[ka_name]
    students = _read_ka_students(ws)

    # Thema aus bestehendem Sheet lesen falls nicht übergeben
    if not thema:
        existing = str(ws.cell(row=2, column=1).value or "")
        thema = existing.replace('Thema: ','').replace('Thema:','').strip() or ""

    # Upsert
    key = (vorname.strip().lower(), nachname.strip().lower())
    updated = False
    for s in students:
        if (s['vorname'].lower(), s['nachname'].lower()) == key:
            s['note'] = note_str; updated = True; break
    if not updated:
        students.append({'vorname':vorname.strip(),'nachname':nachname.strip(),'note':note_str})
    students.sort(key=lambda s:(s['nachname'].lower(),s['vorname'].lower()))

    _write_ka_sheet(ws, klasse, fach, ka_nr, thema, students)
    _write_uebersicht_sheet(wb, klasse, fach)
    wb.active = wb[ka_name]
    wb.save(fpath)
    return fpath


class Handler(BaseHTTPRequestHandler):
    def log_message(self, format, *args):
        pass  # silent

    def _cors_headers(self):
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Methods", "GET, POST, OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "*")
        self.send_header("Content-Type", "application/json; charset=utf-8")

    def _send_json(self, data, status=200):
        try:
            self.send_response(status)
            self._cors_headers()
            self.end_headers()
            self.wfile.write(json.dumps(data, ensure_ascii=False).encode())
        except (BrokenPipeError, ConnectionResetError):
            pass

    def _proxy_lms(self, method, body_bytes=None):
        import urllib.request, urllib.error
        target = lms_base_url() + self.path[len("/api/lms"):]
        req = urllib.request.Request(target, data=body_bytes, method=method)
        if body_bytes is not None:
            req.add_header("Content-Type", "application/json")
        # Status-Check kurz halten (3s). Echte Chat-Anfragen brauchen länger.
        is_status = self.path.endswith("/v1/models") and method == "GET"
        timeout = 3 if is_status else 60
        try:
            with urllib.request.urlopen(req, timeout=timeout) as r:
                data = r.read()
                self.send_response(r.status)
                ct = r.headers.get("Content-Type", "application/json")
                self.send_header("Content-Type", ct)
                self._cors_headers()
                self.send_header("Content-Length", str(len(data)))
                self.end_headers()
                self.wfile.write(data)
        except (urllib.error.URLError, TimeoutError, OSError) as e:
            try:
                self._send_json({"error": f"LM Studio nicht erreichbar: {e}"}, status=502)
            except (BrokenPipeError, ConnectionResetError):
                pass  # Client weg, egal

    def do_GET(self):
        try:
            if self.path.startswith("/api/lms/"):
                self._proxy_lms("GET")
                return

            # Query-String abtrennen damit Routen wie /?reset-sidebar=1 weiterhin index.html liefern
            _path_only = self.path.split("?", 1)[0]
            if _path_only in ("/", "/index.html"):
                html_path = os.path.expanduser("~/workspace/mission-control/index.html")
                with open(html_path, "rb") as f:
                    body = f.read()
                self.send_response(200)
                self.send_header("Content-Type", "text/html; charset=utf-8")
                self.send_header("Cache-Control", "no-store")
                self.send_header("Permissions-Policy", "camera=*, microphone=*")
                self.send_header("Content-Length", str(len(body)))
                self.end_headers()
                self.wfile.write(body)
                return

            # ===== MOBILE / PWA (MC Chris) =====
            MOBILE_STATIC = {
                "/m":                        ("mobile.html",            "text/html; charset=utf-8", "no-store"),
                "/mobile.html":              ("mobile.html",            "text/html; charset=utf-8", "no-store"),
                "/manifest.webmanifest":     ("manifest.webmanifest",   "application/manifest+json", "max-age=3600"),
                "/sw.js":                    ("sw.js",                  "application/javascript", "no-store"),
                "/mc-icon-192.png":          ("mc-icon-192.png",        "image/png", "max-age=604800"),
                "/mc-icon-512.png":          ("mc-icon-512.png",        "image/png", "max-age=604800"),
                "/mc-apple-touch.png":       ("mc-apple-touch.png",     "image/png", "max-age=604800"),
                "/redesign-aurora.html":     ("redesign-aurora.html",   "text/html; charset=utf-8", "no-store"),
                "/redesign-cyber.html":      ("redesign-cyber.html",    "text/html; charset=utf-8", "no-store"),
                "/redesign-1.html":          ("redesign-1.html",        "text/html; charset=utf-8", "no-store"),
                "/redesign-2.html":          ("redesign-2.html",        "text/html; charset=utf-8", "no-store"),
                "/mmc_audio_problem.wav":    ("mmc_audio_problem.wav",  "audio/wav",               "no-store"),
            }
            if self.path in MOBILE_STATIC:
                fname, ctype, cache = MOBILE_STATIC[self.path]
                fpath = os.path.expanduser(f"~/workspace/mission-control/{fname}")
                with open(fpath, "rb") as f:
                    body = f.read()
                self.send_response(200)
                self.send_header("Content-Type", ctype)
                self.send_header("Cache-Control", cache)
                self.send_header("Content-Length", str(len(body)))
                self.end_headers()
                self.wfile.write(body)
                return

            # Wetter-Fotos: /img/weather/*.jpg
            if _path_only.startswith("/img/weather/") and _path_only.endswith(".jpg"):
                fname = os.path.basename(_path_only)
                img_path = os.path.expanduser(f"~/workspace/mission-control/img/weather/{fname}")
                if os.path.isfile(img_path):
                    with open(img_path, "rb") as f:
                        body = f.read()
                    self.send_response(200)
                    self.send_header("Content-Type", "image/jpeg")
                    self.send_header("Cache-Control", "public, max-age=86400")
                    self.send_header("Content-Length", str(len(body)))
                    self.end_headers()
                    self.wfile.write(body)
                    return

            # Generische Audio-Route: alle *.wav/.mp3/.ogg aus mission-control/ ausliefern
            # (für gclip-Player, Audio-Vergleiche, Sprachnotizen-Sharing)
            if _path_only.endswith((".wav", ".mp3", ".ogg", ".m4a", ".webm")):
                fname = os.path.basename(_path_only)  # Path-Traversal-Schutz
                audio_path = os.path.expanduser(f"~/workspace/mission-control/{fname}")
                if os.path.isfile(audio_path):
                    with open(audio_path, "rb") as f:
                        body = f.read()
                    ext_to_mime = {".wav":"audio/wav", ".mp3":"audio/mpeg", ".ogg":"audio/ogg",
                                   ".m4a":"audio/mp4", ".webm":"audio/webm"}
                    mime = ext_to_mime.get(os.path.splitext(fname)[1].lower(), "audio/wav")
                    self.send_response(200)
                    self.send_header("Content-Type", mime)
                    self.send_header("Cache-Control", "no-store")
                    self.send_header("Content-Length", str(len(body)))
                    self.end_headers()
                    self.wfile.write(body)
                    return

            if self.path == "/setup-pro.ps1":
                ps1_path = os.path.expanduser("~/workspace/mission-control/setup-pro.ps1")
                with open(ps1_path, "rb") as f:
                    body = f.read()
                self.send_response(200)
                self.send_header("Content-Type", "text/plain; charset=utf-8")
                self.send_header("Cache-Control", "no-store")
                self.send_header("Content-Length", str(len(body)))
                self.end_headers()
                self.wfile.write(body)
                return

            if self.path == "/favicon.ico":
                ico_path = os.path.expanduser("~/workspace/mission-control/favicon.ico")
                with open(ico_path, "rb") as f:
                    body = f.read()
                self.send_response(200)
                self.send_header("Content-Type", "image/x-icon")
                self.send_header("Cache-Control", "max-age=86400")
                self.send_header("Content-Length", str(len(body)))
                self.end_headers()
                self.wfile.write(body)
                return

            if self.path.startswith("/api/kiforum/img/"):
                fname = self.path.split("/api/kiforum/img/")[1]
                if "/" in fname or ".." in fname:
                    self._send_json({"error":"invalid"},status=400); return
                p = Path(os.path.join(WORKSPACE, "data", fname))
                if not p.exists(): self._send_json({"error":"not found"},status=404); return
                data = p.read_bytes()
                self.send_response(200); self.send_header("Access-Control-Allow-Origin","*")
                self.send_header("Content-Type","image/png"); self.send_header("Content-Length",str(len(data)))
                self.end_headers(); self.wfile.write(data); return

            if self.path == "/api/memory":
                import re as _re
                mem_file = os.path.expanduser("~/.claude/projects/-home-bolla/memory/MEMORY.md")
                try:
                    with open(mem_file) as f:
                        lines = f.read().splitlines()
                    entries, section = [], None
                    for line in lines:
                        line = line.strip()
                        if line.startswith("## "):
                            section = line[3:]
                        elif line.startswith("- ["):
                            m = _re.match(r'- \[([^\]]+)\]\(([^)]+)\)(?:\s+[—–-]+\s+(.+))?', line)
                            if m:
                                entries.append({"title": m.group(1), "file": m.group(2), "desc": m.group(3) or "", "section": section})
                    self._send_json({"entries": entries, "count": len(entries)})
                except Exception as e:
                    self._send_json({"error": str(e)}, status=500)
                return

            if self.path == "/api/kiforum/bg":
                p = Path(os.path.join(WORKSPACE, "data/kif_room_bg.jpg"))
                if not p.exists(): self._send_json({"error":"not found"},status=404); return
                data = p.read_bytes()
                self.send_response(200); self.send_header("Access-Control-Allow-Origin","*")
                self.send_header("Content-Type","image/png"); self.send_header("Content-Length",str(len(data)))
                self.end_headers(); self.wfile.write(data); return

            if self.path == "/api/kiforum/bolla":
                p = Path(os.path.join(WORKSPACE, "data/bolla_transparent.png"))
                if not p.exists(): self._send_json({"error":"not found"},status=404); return
                data = p.read_bytes()
                self.send_response(200); self.send_header("Access-Control-Allow-Origin","*")
                self.send_header("Content-Type","image/png"); self.send_header("Content-Length",str(len(data)))
                self.end_headers(); self.wfile.write(data); return

            if self.path == "/api/kiforum/chris":
                p = Path(os.path.join(WORKSPACE, "data/chris_transparent.png"))
                if not p.exists(): self._send_json({"error":"not found"},status=404); return
                data = p.read_bytes()
                self.send_response(200); self.send_header("Access-Control-Allow-Origin","*")
                self.send_header("Content-Type","image/png"); self.send_header("Content-Length",str(len(data)))
                self.end_headers(); self.wfile.write(data); return

            if self.path == "/api/chris/avatar/shrug/webm":
                p = os.path.expanduser("~/workspace/data/chris_shrug_transparent.webm")
                if not os.path.isfile(p):
                    self._send_json({"error": "WebM nicht gefunden"}, status=404); return
                with open(p, "rb") as f:
                    data = f.read()
                self.send_response(200)
                self.send_header("Access-Control-Allow-Origin", "*")
                self.send_header("Content-Type", "video/webm")
                self.send_header("Content-Length", str(len(data)))
                self.end_headers(); self.wfile.write(data); return

            if self.path == "/api/bolla/avatar/wave/webm":
                webm_path = os.path.expanduser("~/workspace/data/bolla_wave_transparent.webm")
                if not os.path.isfile(webm_path):
                    self._send_json({"error": "WebM nicht gefunden"}, status=404); return
                with open(webm_path, "rb") as f:
                    data = f.read()
                self.send_response(200)
                self.send_header("Access-Control-Allow-Origin", "*")
                self.send_header("Content-Type", "video/webm")
                self.send_header("Content-Length", str(len(data)))
                self.end_headers()
                self.wfile.write(data)
                return

            if self.path == "/api/bolla/avatar/wave":
                wave_path = os.path.expanduser("~/workspace/data/bolla_wave.gif")
                if not os.path.isfile(wave_path):
                    self._send_json({"error": "GIF nicht gefunden"}, status=404); return
                with open(wave_path, "rb") as f:
                    data = f.read()
                self.send_response(200)
                self.send_header("Access-Control-Allow-Origin", "*")
                self.send_header("Content-Type", "image/gif")
                self.send_header("Content-Length", str(len(data)))
                self.end_headers()
                self.wfile.write(data)
                return

            if self.path == "/api/bolla/avatar":
                avatar_path = os.path.expanduser("~/workspace/bolla_avatar.png")
                with open(avatar_path, "rb") as f:
                    body = f.read()
                self.send_response(200)
                self.send_header("Content-Type", "image/png")
                self.send_header("Content-Disposition", 'attachment; filename="bolla_avatar.png"')
                self.send_header("Content-Length", str(len(body)))
                self.end_headers()
                self.wfile.write(body)
                return

            if self.path == "/api/bolla/voices":
                self._send_json(azure_list_voices())
                return

            if self.path == "/api/speech/config":
                cfg = azure_speech_config()
                if cfg:
                    self._send_json({"key": cfg["key"], "region": cfg["region"]})
                else:
                    self._send_json({"error": "Azure Speech nicht konfiguriert"}, 503)
                return

            if self.path == "/api/adb/devices":
                self._send_json(adb_devices())
                return

            if self.path.startswith("/api/adb/packages"):
                import urllib.parse as _up
                qs = _up.urlparse(self.path).query
                params = dict(_up.parse_qsl(qs))
                self._send_json(adb_packages(params.get("filter", ""), params.get("kind", "all")))
                return

            if self.path.startswith("/api/fs/dirs"):
                import urllib.parse as _up
                qs = _up.urlparse(self.path).query
                params = dict(_up.parse_qsl(qs))
                req_path = params.get("path", "/mnt/d/OneDrive").strip()
                ALLOWED_ROOTS = ["/mnt/d/OneDrive", "/mnt/c/Users/ernst", os.path.expanduser("~")]
                if not any(req_path.startswith(r) for r in ALLOWED_ROOTS) or ".." in req_path:
                    self._send_json({"error": "Pfad nicht erlaubt"}); return
                IMG_EXTS = {".jpg", ".jpeg", ".png", ".gif", ".webp", ".bmp", ".tif", ".tiff"}
                def _count_photos(dirpath):
                    try:
                        return sum(1 for e in os.scandir(dirpath)
                                   if e.is_file() and os.path.splitext(e.name.lower())[1] in IMG_EXTS)
                    except Exception:
                        return 0
                try:
                    entries = sorted([
                        e for e in os.scandir(req_path)
                        if e.is_dir() and not e.name.startswith('.')
                    ], key=lambda e: e.name.lower())
                    dirs_info = [{"name": e.name, "photos": _count_photos(e.path)} for e in entries]
                    self._send_json({"path": req_path, "dirs": [e["name"] for e in dirs_info], "photos": dirs_info})
                except Exception as ex:
                    self._send_json({"error": str(ex)})
                return

            if self.path.startswith("/api/adb/info"):
                import urllib.parse as _up
                qs = _up.urlparse(self.path).query
                params = dict(_up.parse_qsl(qs))
                self._send_json(adb_info(params.get("kind", "device")))
                return

            if self.path.startswith("/api/adb/ls"):
                import urllib.parse as _up
                qs = _up.urlparse(self.path).query
                params = dict(_up.parse_qsl(qs))
                self._send_json(adb_ls(params.get("path", "/storage/emulated/0")))
                return

            if self.path.startswith("/api/offers/search"):
                import urllib.parse as _up
                qs = dict(_up.parse_qsl(_up.urlparse(self.path).query))
                q = qs.get("q", "").strip()
                limit = int(qs.get("limit", "20"))
                if not q:
                    self._send_json({"results": []}); return
                self._send_json(offers_search(q, limit)); return

            if self.path.startswith("/api/offers/zip"):
                self._send_json({"zip": _marktguru_zip()}); return

            if self.path.startswith("/api/calendar/search"):
                import urllib.parse as _up
                qs = _up.urlparse(self.path).query
                params = dict(_up.parse_qsl(qs))
                q = params.get("q", "").strip()
                if not q:
                    self._send_json([]); return
                self._send_json(search_calendar_events(q)); return

            simple = {
                "/api/calendar": get_calendar,
                "/api/email": get_emails,
                "/api/photo": lambda: get_photo_of_day() or {},
                "/api/recipe": get_recipe_of_day,
                "/api/birthdays": get_birthdays,
                "/api/robin": get_robin_info,
                "/api/sysinfo": get_sysinfo,
                "/api/surfaces/book": get_book_health,
                "/api/surfaces/pro": get_pro_health,
                "/api/surfaces/studio": get_studio_health,
                "/api/surfaces/energie": get_energie,
                "/api/tokenusage": get_token_usage,
                "/api/tokenusage/history": get_token_halfdays,
                "/api/tokenbudget": get_token_budget,
                "/api/tokenbudget/snapshot": token_budget_snapshot,
                "/api/claudequota": get_claude_quota,
                "/api/kosten": get_kosten,
                "/api/status": lambda: {"ok": True, "ts": datetime.now().isoformat()},
                "/api/redesigns-meta": get_redesigns_meta,
                "/api/clipboard": get_clipboard,
                "/api/clipboard/trash": get_clipboard_trash,
                "/api/clipboard/images": get_clipboard_images,
                "/api/newsletter": get_newsletter,
                "/api/makler": get_makler,
                "/api/immo-news": get_immo_news,
                "/api/immo-bookmarks": get_immo_bookmarks,
                "/api/immo-criteria":  get_immo_criteria,
                "/api/crontab":        get_crontab,
                "/api/travel":         get_travel,
                "/api/amadeus/config": amadeus_get_config,
                "/api/photos/status":  lambda: dict(_photo_job),
                "/api/photos/results": lambda: {"results": _photo_load_results()},
                "/api/photos/config":  get_lmstudio_config,
                "/api/quicknotes":     get_quicknotes,
                "/api/sos/contacts":   get_sos_contacts,
                "/api/chat-history":   get_chat_history,
                "/api/charts":         get_charts,
                "/api/workshop":       get_workshop,
                "/api/workshop/fortschritt": get_workshop_fortschritt,
                "/api/projekte/list":  projekte_list,
            }

            if self.path.startswith("/api/projekte/load"):
                import urllib.parse as _up
                pid = dict(_up.parse_qsl(_up.urlparse(self.path).query)).get("id","")
                p = projekte_load(pid) if pid else None
                if p is None:
                    self._send_json({"error": "Nicht gefunden"}, status=404); return
                self._send_json(p); return

            if self.path.startswith("/api/projekte/delete"):
                import urllib.parse as _up
                pid = dict(_up.parse_qsl(_up.urlparse(self.path).query)).get("id","")
                self._send_json(projekte_delete(pid) if pid else {"error": "Kein id"}); return

            # Chart Preview (iTunes)
            if self.path.startswith("/api/charts/preview"):
                import urllib.parse as _cup, urllib.request as _cur2, re as _re_p
                qs = _cup.parse_qs(_cup.urlparse(self.path).query)
                title = qs.get("title", [""])[0]
                artist = qs.get("artist", [""])[0]
                # Feature-Angaben bereinigen: (w/ ...), (feat. ...), (ft. ...)
                clean_title = _re_p.sub(r'\s*\((?:w/|feat\.?|ft\.?)[^)]*\)', '', title, flags=_re_p.IGNORECASE).strip()
                clean_artist = artist.split(',')[0].strip()  # nur ersten Künstler nehmen
                def _itunes_search(term):
                    q = _cup.urlencode({"term": term, "media": "music", "limit": "5", "entity": "song"})
                    req_it = _cur2.Request(f"https://itunes.apple.com/search?{q}", headers={"User-Agent": "Mozilla/5.0"})
                    d_it = json.loads(_cur2.urlopen(req_it, timeout=8).read())
                    for r in d_it.get("results", []):
                        if r.get("previewUrl"):
                            return r["previewUrl"]
                    return None
                try:
                    # Versuch 1: bereinigter Titel + Künstler
                    preview = _itunes_search(f"{clean_title} {clean_artist}")
                    # Versuch 2: nur bereinigter Titel
                    if not preview:
                        preview = _itunes_search(clean_title)
                    self._send_json({"preview_url": preview})
                except Exception as e:
                    self._send_json({"preview_url": None, "error": str(e)})
                return

            # Foto-Thumb
            if self.path.startswith("/api/photos/thumb"):
                import urllib.parse as _up
                qs = _up.parse_qs(_up.urlparse(self.path).query)
                path = qs.get("path", [""])[0]
                if not path:
                    self._send_json({"error": "Kein Pfad"}, status=400); return
                b64, ext = photo_thumb(path)
                if not b64:
                    self._send_json({"error": ext}, status=404); return
                import base64 as _b64
                body = _b64.b64decode(b64)
                self.send_response(200)
                self.send_header("Content-Type", f"image/{ext}")
                self.send_header("Cache-Control", "max-age=86400")
                self.send_header("Content-Length", str(len(body)))
                self.end_headers()
                self.wfile.write(body)
                return

            if self.path.startswith("/api/clipboard/image/"):
                fname = self.path[len("/api/clipboard/image/"):]
                if "/" in fname or "\\" in fname or ".." in fname:
                    self._send_json({"error": "Ungültig"}, status=400); return
                fpath = os.path.join(CLIPBOARD_IMAGES_DIR, fname)
                if not os.path.isfile(fpath):
                    self._send_json({"error": "Nicht gefunden"}, status=404); return
                ext = os.path.splitext(fname)[1].lower()
                MIME_MAP = {".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
                            ".gif": "image/gif", ".webp": "image/webp", ".bmp": "image/bmp"}
                ctype = MIME_MAP.get(ext, "application/octet-stream")
                with open(fpath, "rb") as f:
                    body = f.read()
                self.send_response(200)
                self.send_header("Content-Type", ctype)
                self.send_header("Cache-Control", "max-age=86400")
                self.send_header("Content-Length", str(len(body)))
                self.end_headers()
                self.wfile.write(body)
                return

            if self.path.startswith("/api/korrektur/excel-download"):
                import urllib.parse as _up
                qs = _up.parse_qs(_up.urlparse(self.path).query)
                fname = qs.get("file", [""])[0]
                if not fname or "/" in fname or "\\" in fname or ".." in fname:
                    self._send_json({"error": "Ungültig"}, status=400); return
                fpath = os.path.join(KORREKTUR_DIR, fname)
                if not os.path.isfile(fpath):
                    self._send_json({"error": "Nicht gefunden"}, status=404); return
                with open(fpath, "rb") as f:
                    data = f.read()
                self.send_response(200)
                self.send_header("Content-Type", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
                self.send_header("Content-Disposition", f'attachment; filename="{fname}"')
                self.send_header("Content-Length", str(len(data)))
                self.end_headers()
                self.wfile.write(data)
                return

            elif self.path == "/api/gluecksrad/stile":
                MUSIK_DIR = "/mnt/d/OneDrive/Dokumente/Office/Powerpoint/Beispiele/Spiele/Songs zum Glücksrad Musikstile"
                STIL_MAP = {
                    "Blues": "Blues - Lonely Echoes.mp3",
                    "Rock'n'Roll": "Rock´n´Roll - School of No Rules.mp3",
                    "Jazz": "Jazz - Night Whispers.mp3",
                    "Pop": "Pop - School Bells Ringing.mp3",
                    "Hip-Hop": "Hip-Hop - School Rules.mp3",
                    "Klassik": "Klassik - The Forgotten Star.mp3",
                    "Reggae": "Reggae - School Vibes.mp3",
                    "Metal": "Metal - Hellish Hallways.mp3",
                    "Country": "Country - Schoolhouse Days.mp3",
                    "Techno": "Techno - Classroom Chaos.mp3",
                    "Soul": "Soul - School of Life.mp3",
                    "Punk": "Punk - Classroom Chaos.mp3",
                    "Swing": "Swing - Schoolyard Swing.mp3",
                    "R&B": "Rythm & Blues - School Days Groove.mp3",
                    "Disco": "Dsico - Fever in the Classroom.mp3",
                    "Gospel": "Gospel - Faith in the Classroom.mp3",
                    "Rap": "Rap - Class in Session.mp3",
                    "Boogie-Woogie": "Boogie-Woogie - Shake It Loose.mp3",
                    "K-Pop": "K-Pop - Classroom Dreams.mp3",
                    "Latin": "Latin - Escuela de Vida.mp3",
                    "Indie": "Indie&Alternative - Classroom Daydreams.mp3",
                    "Funk": "Funk  -  School's Got Soul.mp3",
                    "Hard-Rock": "Hard-Rock - Classroom Chaos.mp3",
                    "Choral": "Choral - Echoes of Eternity.mp3",
                }
                stile = [{"name": k, "datei": v, "vorhanden": os.path.isfile(os.path.join(MUSIK_DIR, v))} for k, v in STIL_MAP.items()]
                self._send_json({"stile": stile, "dir": MUSIK_DIR})

            elif self.path.startswith("/api/gluecksrad/music"):
                import urllib.parse as _up
                qs = _up.parse_qs(self.path.split("?",1)[1] if "?" in self.path else "")
                stil = qs.get("stil", [""])[0]
                MUSIK_DIR = "/mnt/d/OneDrive/Dokumente/Office/Powerpoint/Beispiele/Spiele/Songs zum Glücksrad Musikstile"
                STIL_MAP = {
                    "Blues": "Blues - Lonely Echoes.mp3",
                    "Rock'n'Roll": "Rock´n´Roll - School of No Rules.mp3",
                    "Jazz": "Jazz - Night Whispers.mp3",
                    "Pop": "Pop - School Bells Ringing.mp3",
                    "Hip-Hop": "Hip-Hop - School Rules.mp3",
                    "Klassik": "Klassik - The Forgotten Star.mp3",
                    "Reggae": "Reggae - School Vibes.mp3",
                    "Metal": "Metal - Hellish Hallways.mp3",
                    "Country": "Country - Schoolhouse Days.mp3",
                    "Techno": "Techno - Classroom Chaos.mp3",
                    "Soul": "Soul - School of Life.mp3",
                    "Punk": "Punk - Classroom Chaos.mp3",
                    "Swing": "Swing - Schoolyard Swing.mp3",
                    "R&B": "Rythm & Blues - School Days Groove.mp3",
                    "Disco": "Dsico - Fever in the Classroom.mp3",
                    "Gospel": "Gospel - Faith in the Classroom.mp3",
                    "Rap": "Rap - Class in Session.mp3",
                    "Boogie-Woogie": "Boogie-Woogie - Shake It Loose.mp3",
                    "K-Pop": "K-Pop - Classroom Dreams.mp3",
                    "Latin": "Latin - Escuela de Vida.mp3",
                    "Indie": "Indie&Alternative - Classroom Daydreams.mp3",
                    "Funk": "Funk  -  School's Got Soul.mp3",
                    "Hard-Rock": "Hard-Rock - Classroom Chaos.mp3",
                    "Choral": "Choral - Echoes of Eternity.mp3",
                }
                fname = STIL_MAP.get(stil)
                if not fname:
                    self._send_json({"error": "Stil nicht gefunden"}, status=404); return
                fpath = os.path.join(MUSIK_DIR, fname)
                if not os.path.isfile(fpath):
                    self._send_json({"error": "Datei nicht gefunden"}, status=404); return
                with open(fpath, "rb") as f:
                    body = f.read()
                self.send_response(200)
                self.send_header("Content-Type", "audio/mpeg")
                self.send_header("Content-Length", str(len(body)))
                self.send_header("Cache-Control", "no-store")
                self.end_headers()
                self.wfile.write(body)

            elif self.path == "/api/gluecksrad/state":
                self._send_json(_gluecksrad_state)

            elif self.path == "/api/ki-news":
                f = os.path.join(WORKSPACE, "data/ki_news.json")
                if os.path.isfile(f):
                    with open(f) as fh:
                        self._send_json(json.load(fh))
                else:
                    self._send_json({"events": [], "news": [], "letztesUpdate": None})

            elif self.path == "/api/ki-buch/generiere-status":
                self._send_json(_ki_buch_job)

            elif self.path == "/api/ki-buch":
                f = os.path.join(WORKSPACE, "data/ki_buch.json")
                if os.path.isfile(f):
                    with open(f) as fh:
                        _buch_pub = json.load(fh)
                    # SPOILER-SCHUTZ: geheime Wendung NIE ans Frontend senden
                    _buch_pub.pop("geheim", None)
                    self._send_json(_buch_pub)
                else:
                    self._send_json({"titel": "Thriller-Projekt", "status": "Planung", "kriterien": {}, "kapitel": [], "kommentare": [], "naechsterSchritt": ""})

            elif self.path == "/api/sidebar-state":
                sf = os.path.join(WORKSPACE, "data/sidebar_state.json")
                if os.path.isfile(sf):
                    with open(sf) as f:
                        self._send_json(json.load(f))
                else:
                    self._send_json({})

            elif self.path == "/api/dokumente":
                BOLLA_DOCS = "/mnt/d/OneDrive/Dokumente/Bolla/claud code - openclaw Doku"
                ALLOWED_EXT = {".pdf", ".docx", ".html", ".txt", ".md"}
                EXT_ICON = {".pdf": "📄", ".docx": "📝", ".html": "🌐", ".txt": "📃", ".md": "📃"}
                docs = []
                if os.path.isdir(BOLLA_DOCS):
                    for fname in sorted(os.listdir(BOLLA_DOCS)):
                        ext = os.path.splitext(fname)[1].lower()
                        if ext not in ALLOWED_EXT:
                            continue
                        fpath = os.path.join(BOLLA_DOCS, fname)
                        stat = os.stat(fpath)
                        docs.append({
                            "name": fname,
                            "icon": EXT_ICON.get(ext, "📄"),
                            "size": stat.st_size,
                            "modified": datetime.fromtimestamp(stat.st_mtime).strftime("%d.%m.%Y"),
                            "download": f"/api/dokumente/download?file={urllib.parse.quote(fname)}"
                        })
                self._send_json({"docs": docs})
                return

            elif self.path.startswith("/api/dokumente/open"):
                import urllib.parse as _up
                BOLLA_DOCS = "/mnt/d/OneDrive/Dokumente/Bolla/claud code - openclaw Doku"
                qs = _up.parse_qs(_up.urlparse(self.path).query)
                fname = qs.get("file", [""])[0]
                if not fname or "/" in fname or "\\" in fname or ".." in fname:
                    self._send_json({"error": "Ungültig"}, status=400); return
                fpath = os.path.join(BOLLA_DOCS, fname)
                if not os.path.isfile(fpath):
                    self._send_json({"error": "Nicht gefunden"}, status=404); return
                win_path = fpath.replace("/mnt/d/", "D:\\").replace("/", "\\")
                subprocess.Popen(["/mnt/c/WINDOWS/System32/WindowsPowerShell/v1.0/powershell.exe", "-Command", f'Start-Process "{win_path}"'])
                self._send_json({"ok": True})
                return

            elif self.path.startswith("/api/dokumente/download"):
                import urllib.parse as _up
                BOLLA_DOCS = "/mnt/d/OneDrive/Dokumente/Bolla/claud code - openclaw Doku"
                qs = _up.parse_qs(_up.urlparse(self.path).query)
                fname = qs.get("file", [""])[0]
                if not fname or "/" in fname or "\\" in fname or ".." in fname:
                    self._send_json({"error": "Ungültig"}, status=400); return
                fpath = os.path.join(BOLLA_DOCS, fname)
                if not os.path.isfile(fpath):
                    self._send_json({"error": "Nicht gefunden"}, status=404); return
                ext = os.path.splitext(fname)[1].lower()
                MIME = {".pdf": "application/pdf",
                        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        ".html": "text/html",
                        ".txt": "text/plain",
                        ".md": "text/plain"}
                ctype = MIME.get(ext, "application/octet-stream")
                with open(fpath, "rb") as f:
                    data = f.read()
                self.send_response(200)
                self.send_header("Content-Type", ctype)
                disp = "inline" if ext == ".pdf" else "attachment"
                self.send_header("Content-Disposition", f'{disp}; filename="{fname}"')
                self.send_header("Content-Length", str(len(data)))
                self.end_headers()
                self.wfile.write(data)
                return

            elif self.path == "/api/kiforum":
                KIFORUM_FILE = os.path.join(WORKSPACE, "data/kiforum.json")
                posts = json.loads(Path(KIFORUM_FILE).read_text()) if Path(KIFORUM_FILE).exists() else []
                self._send_json(posts)

            elif self.path == "/api/kiforum/export":
                KIFORUM_FILE = os.path.join(WORKSPACE, "data/kiforum.json")
                data = Path(KIFORUM_FILE).read_bytes() if Path(KIFORUM_FILE).exists() else b"[]"
                self.send_response(200)
                self.send_header("Access-Control-Allow-Origin", "*")
                self.send_header("Content-Type", "application/json")
                self.send_header("Content-Disposition", 'attachment; filename="bolla-ki-forum.json"')
                self.send_header("Content-Length", str(len(data)))
                self.end_headers()
                self.wfile.write(data)

            elif self.path == "/api/kiforum/word":
                import io as _io
                from docx import Document as _DocxDoc
                from docx.shared import Pt as _Pt, RGBColor as _RGB, Inches as _Inches
                KIFORUM_FILE = Path(os.path.join(WORKSPACE, "data/kiforum.json"))
                posts = json.loads(KIFORUM_FILE.read_text()) if KIFORUM_FILE.exists() else []
                doc = _DocxDoc()
                for sec in doc.sections:
                    sec.left_margin = _Inches(1.1); sec.right_margin = _Inches(1.1)
                h = doc.add_heading("KI-Forum Diskussion", 0)
                h.runs[0].font.color.rgb = _RGB(0x1a,0x1a,0x2e)
                doc.add_paragraph(f"Exportiert: {datetime.now().strftime('%d.%m.%Y %H:%M')} · Bolla & Chris")
                doc.add_paragraph("")
                for p in reversed(posts):
                    t = p.get("type","")
                    if t == "user":
                        h2 = doc.add_heading(f"💬 Chris: {p.get('title','') or 'These'}", 2)
                        for r in h2.runs: r.font.color.rgb = _RGB(0xcc,0x77,0x00)
                    elif t == "bolla":
                        h2 = doc.add_heading(f"{p.get('emoji','🤖')} Bolla: {p.get('title','Reaktion')}", 2)
                        for r in h2.runs: r.font.color.rgb = _RGB(0x22,0x66,0xcc)
                    elif t == "conclusion":
                        h2 = doc.add_heading(f"{p.get('emoji','🎯')} Fazit: {p.get('title','')}", 2)
                        for r in h2.runs: r.font.color.rgb = _RGB(0x88,0x44,0x00)
                    else:
                        continue
                    content = p.get("content","")
                    img_path = None
                    if p.get("img"):
                        _candidate = Path(os.path.join(WORKSPACE, "data", p["img"]))
                        if _candidate.exists():
                            img_path = _candidate
                    if content and img_path:
                        from docx.oxml.ns import qn as _qn
                        from docx.oxml import OxmlElement as _OE
                        tbl = doc.add_table(rows=1, cols=2)
                        tbl.style = 'Table Grid'
                        # Rahmen entfernen
                        for cell in tbl.rows[0].cells:
                            tc = cell._tc
                            tcPr = tc.get_or_add_tcPr()
                            tcBorders = _OE('w:tcBorders')
                            for side in ('top','left','bottom','right','insideH','insideV'):
                                border = _OE(f'w:{side}')
                                border.set(_qn('w:val'), 'none')
                                tcBorders.append(border)
                            tcPr.append(tcBorders)
                        # Linke Zelle: Text (ca. 4.5")
                        cell_l = tbl.cell(0, 0)
                        cell_l.width = _Inches(4.5)
                        p_text = cell_l.paragraphs[0]
                        run = p_text.add_run(content)
                        run.font.size = _Pt(11)
                        # Rechte Zelle: Bild (ca. 1.3")
                        cell_r = tbl.cell(0, 1)
                        cell_r.width = _Inches(1.3)
                        p_img = cell_r.paragraphs[0]
                        p_img.paragraph_format.space_before = _Pt(2)
                        run_img = p_img.add_run()
                        run_img.add_picture(str(img_path), width=_Inches(1.2))
                    elif content:
                        para = doc.add_paragraph(content)
                        if para.runs: para.runs[0].font.size = _Pt(11)
                    doc.add_paragraph("")
                buf = _io.BytesIO()
                doc.save(buf)
                data = buf.getvalue()
                self.send_response(200)
                self.send_header("Access-Control-Allow-Origin", "*")
                self.send_header("Content-Type", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                self.send_header("Content-Disposition", 'attachment; filename="bolla-diskussion.docx"')
                self.send_header("Content-Length", str(len(data)))
                self.end_headers()
                self.wfile.write(data)

            elif self.path.startswith("/api/chargers"):
                import urllib.parse as _cup
                qs = dict(_cup.parse_qsl(_cup.urlparse(self.path).query))
                lat = qs.get("lat"); lon = qs.get("lon")
                radius = int(qs.get("radius", 12000))
                if not lat or not lon:
                    self._send_json({"error": "lat/lon required"}, status=400)
                else:
                    self._send_json(get_chargers(float(lat), float(lon), radius))
            elif self.path in simple:
                self._send_json(simple[self.path]())
            elif self.path.startswith("/api/amadeus/flights"):
                qs = dict(urllib.parse.parse_qsl(urllib.parse.urlparse(self.path).query))
                self._send_json(amadeus_search_flights(
                    qs.get("origin",""), qs.get("dest",""), qs.get("date",""),
                    qs.get("return_date",""), int(qs.get("adults",2))
                ))
            else:
                self._send_json({"error": "not found"}, status=404)
        except (BrokenPipeError, ConnectionResetError):
            pass
        except Exception as e:
            tb = traceback.format_exc()
            print(f"GET Error: {tb}")
            try:
                self._send_json({"error": str(e)}, status=500)
            except Exception:
                pass

    def do_POST(self):
        try:
            length = int(self.headers.get("Content-Length", 0))
            raw = self.rfile.read(length) if length else b""
        except Exception:
            raw = b""

        if self.path.startswith("/api/lms/"):
            try:
                self._proxy_lms("POST", raw)
            except Exception as e:
                tb = traceback.format_exc()
                print(f"LMS proxy error: {tb}")
                try:
                    self._send_json({"error": str(e)}, status=500)
                except Exception:
                    pass
            return

        try:
            body = json.loads(raw) if raw else {}
        except Exception:
            body = {}

        try:
            if self.path == "/api/gluecksrad/state":
                _gluecksrad_state["stil"] = body.get("stil")
                _gluecksrad_state["nummer"] = body.get("nummer")
                self._send_json({"ok": True})
            elif self.path == "/api/ki-buch/generiere":
                global _ki_buch_job
                if _ki_buch_job["status"] == "running":
                    self._send_json({"status": "running"}); return
                import threading as _thr3, shutil as _sh3, subprocess as _sp3, re as _re3, datetime as _dt3
                bf = os.path.join(WORKSPACE, "data/ki_buch.json")
                with open(bf) as fh:
                    buch = json.load(fh)
                kommentare = buch.get("kommentare", [])
                offene = [k for k in kommentare if not k.get("erledigt")]
                anweisung = offene[-1]["text"] if offene else "Schreibe das nächste Kapitel."
                prots = json.dumps(buch.get("protagonisten", []), ensure_ascii=False, indent=2)
                if buch.get("antagonist"):
                    prots += "\n\nANTAGONIST:\n" + json.dumps(buch["antagonist"], ensure_ascii=False, indent=2)
                if buch.get("nebenfiguren"):
                    prots += "\n\nNEBENFIGUREN (konsistent halten, plastisch einsetzen):\n" + json.dumps(buch["nebenfiguren"], ensure_ascii=False, indent=2)
                krit = json.dumps(buch.get("kriterien", {}), ensure_ascii=False, indent=2)
                # Produzenten-Steuerung (Rote Linien, Wunsch-Briefkasten, Richtungsimpuls)
                _st = buch.get("steuerung", {})
                _rl = _st.get("rote_linien", [])
                _eintraege = _st.get("eintraege", [])
                steuer_txt = ""
                if _rl:
                    steuer_txt += "\nROTE LINIEN (unbedingt einhalten — Dinge die NICHT passieren dürfen / Vorgaben):\n" + "\n".join(f"- {x}" for x in _rl) + "\n"
                # Steuerung & Wünsche — alle außer denen, von denen abgeraten wurde
                _aktive = [e for e in _eintraege if e.get("status") != "nicht_empfohlen"]
                if _aktive:
                    steuer_txt += "\nSTEUERUNG & WÜNSCHE (vom Autor; wo nicht anders vermerkt, liegt das Timing bei dir — nicht alles auf einmal, dramaturgisch einsetzen):\n"
                    _mark = {"in_arbeit": "[JETZT einbauen]", "umgesetzt": "[bereits umgesetzt — konsistent halten]", "geheim": "[umsetzen, dezent/unauffällig]", "vorgemerkt": "[für später vormerken]"}
                    for e in _aktive:
                        steuer_txt += f"- {_mark.get(e.get('status'),'')} {e.get('text','')}\n"
                # GEHEIME WENDUNG — nur intern, NIE im Überblick/Antwort verraten
                _gw = (buch.get("geheim", {}).get("wendung") or "").strip()
                geheim_txt = ""
                if _gw:
                    geheim_txt = f"""
GEHEIME MASTER-WENDUNG (NUR FÜR DICH — Chris darf sie NIEMALS erfahren, auch nicht andeutungsweise im ###ANTWORT### oder ###UEBERBLICK###):
{_gw}
Arbeite konsequent auf diese Wendung hin. Säe unauffällige Hinweise/Saatkörner, die sie später glaubwürdig und rückwirkend stimmig machen — aber so subtil, dass ein Erstleser sie nicht als Hinweis erkennt. Der ###UEBERBLICK### bleibt strikt pointenfrei.
"""
                # Vollständige Kapiteltexte für Konsistenz-Prüfung (Rückwirkungen erkennen)
                kap_voll = "\n\n".join([f"=== {k['titel']} ===\n{k['text']}" for k in buch.get("kapitel", [])])
                # Prüfen ob Chris eine Rückwirkungs-Änderung bereits bestätigt hat
                bestaetigt = any(w in anweisung.lower() for w in ["ja, ändere alle", "ja ändere alle", "bestätigt", "ja, alle anpassen", "ja alle anpassen", "ja bitte alle"])
                prompt = f"""Du bist Bolla, KI-Assistent von Chris Mandel, und schreibst gemeinsam einen deutschen KI-Thriller.

BUCH: {buch.get("titel","AURORA")} — {buch.get("untertitel","")}

BUCHKRITERIEN (UNBEDINGT BEACHTEN):
{krit}

PROTAGONISTEN:
{prots}

FIGURENFÜHRUNG (Ken-Follett-Prinzip): Nutze die Profile aktiv — zeige aussehen, detail, macke, wunde und stimme in Handlung und Dialog, statt sie nur zu kennen. Jede Figur (auch Neben- und Randfiguren, selbst die Antagonistin) soll ein klares, authentisches, interessantes Bild ergeben und konsistent zu ihrem Profil sprechen und handeln. Figuren wachsen durch konkrete Gesten und eigene Sprache, nicht durch Behauptungen.
{steuer_txt}{geheim_txt}
BISHERIGE KAPITEL (vollständig):
{kap_voll if kap_voll else "Noch keine Kapitel — fange frisch an."}

ANWEISUNG VON CHRIS:
{anweisung}

WICHTIGE REGEL ZU RÜCKWIRKUNGEN:
Wenn die Anweisung eine Änderung verlangt, die ZWINGEND auch in FRÜHEREN Kapiteln berücksichtigt werden müsste (z.B. ein Charakterzug, ein Handlungsdetail, ein neuer Name, eine geänderte Beziehung, die schon vorher vorkam), dann führe die Änderung NICHT sofort komplett aus. Stattdessen:
- Beschreibe im ANTWORT-Feld klar, WELCHE früheren Kapitel betroffen wären und WARUM.
- Stelle die Rückfrage: "Soll ich diese Änderung in allen betroffenen Kapiteln umsetzen? Antworte mit 'Ja, ändere alle'."
- Lass INHALT und TITEL_ABSCHNITT in diesem Fall LEER.
{"AUSNAHME: Chris hat mit 'Ja, ändere alle' bestätigt — setze die Änderung jetzt durchgängig in ALLEN betroffenen Kapiteln um und gib das/die überarbeitete(n) Kapitel zurück (bei mehreren: das wichtigste zuerst, erwähne im ANTWORT-Feld welche weiteren du noch anpasst)." if bestaetigt else ""}

Wenn die Anweisung KEINE Rückwirkung auf frühere Kapitel hat (z.B. neues Kapitel schreiben, aktuelles Kapitel isoliert überarbeiten), führe sie normal aus.

Antworte AUSSCHLIESSLICH in genau diesem Format mit den Trennmarken (kein JSON, kein Markdown). Lass Felder leer wenn nicht zutreffend:

###ANTWORT###
(Kurze Rückmeldung ODER Rückfrage bei Rückwirkungen, 1-3 Sätze)
###TITEL_ABSCHNITT###
(Titel des neuen/überarbeiteten Abschnitts, z.B. "Prolog" oder "Kapitel 1: ..." — exakt wie bestehendes Kapitel wenn überarbeitet)
###UEBERBLICK###
(STRIKT POINTENFREIER Kurz-Überblick des Kapitels, 3-5 Zeilen: wer, wo, welche Beziehung/Stimmung bewegt sich, welcher Strang kommt voran. NIEMALS Twists, Enthüllungen, Cliffhanger-Clou oder die geheime Wendung verraten. Dies liest Chris zum Steuern, OHNE sich zu spoilern. LEER bei Rückfrage.)
###INHALT###
(Der vollständige generierte Text — frei schreiben, Anführungszeichen, Absätze erlaubt. LEER bei Rückfrage.)
###BUCHTITEL_NEU###
(Nur wenn Buchtitel geändert werden soll, sonst leer)
###NAECHSTER_SCHRITT###
(Was als nächstes sinnvoll wäre)
###ENDE###"""
                _ki_buch_job = {"status": "running", "antwort": "", "inhalt": "", "inhalt_titel": "", "error": ""}
                def _run():
                    global _ki_buch_job
                    try:
                        cl = _sh3.which("claude") or os.path.expanduser("~/.local/bin/claude")
                        r = _sp3.run([cl, "-p", "--output-format", "json", "--model", "claude-sonnet-4-6", prompt],
                                     capture_output=True, text=True, timeout=900, stdin=_sp3.DEVNULL,
                                     cwd=os.path.expanduser("~"))
                        if r.returncode != 0:
                            _ki_buch_job = {"status": "error", "error": r.stderr[:200] or "Claude-Fehler", "antwort":"","inhalt":"","inhalt_titel":""}
                            return
                        raw = json.loads(r.stdout).get("result", "")
                        def _extract(tag_start, tag_end):
                            mm = _re3.search(_re3.escape(tag_start) + r'(.*?)' + _re3.escape(tag_end), raw, _re3.DOTALL)
                            return mm.group(1).strip() if mm else ""
                        gen = {
                            "antwort_text": _extract("###ANTWORT###", "###TITEL_ABSCHNITT###"),
                            "neuer_inhalt_titel": _extract("###TITEL_ABSCHNITT###", "###UEBERBLICK###"),
                            "ueberblick": _extract("###UEBERBLICK###", "###INHALT###"),
                            "neuer_inhalt": _extract("###INHALT###", "###BUCHTITEL_NEU###"),
                            "titel_update": _extract("###BUCHTITEL_NEU###", "###NAECHSTER_SCHRITT###"),
                            "naechster_schritt": _extract("###NAECHSTER_SCHRITT###", "###ENDE###"),
                        }
                        # Fallback wenn Format nicht erkannt
                        if not gen["antwort_text"] and not gen["neuer_inhalt"]:
                            gen["antwort_text"] = raw[:300]
                        with open(bf) as fh2: buch2 = json.load(fh2)
                        now = _dt3.datetime.now().strftime("%Y-%m-%d %H:%M")
                        if gen.get("titel_update"): buch2["titel"] = gen["titel_update"]
                        if gen.get("neuer_inhalt") and gen.get("neuer_inhalt_titel"):
                            neuer_titel = gen["neuer_inhalt_titel"]
                            # Überschreiben wenn Kapitel mit gleichem Titel existiert ODER Anweisung "nochmal/überarbeite" enthält
                            ueberschreibe_keywords = ["nochmal","noch einmal","überarbeit","ersetze","rewrite","verbessere kapitel"]
                            ist_ueberarbeitung = any(kw in anweisung.lower() for kw in ueberschreibe_keywords)
                            existing_idx = next((i for i,k in enumerate(buch2.get("kapitel",[])) if k["titel"]==neuer_titel), None)
                            kap_eintrag = {"titel":neuer_titel,"text":gen["neuer_inhalt"],"ueberblick":gen.get("ueberblick",""),"datum":now}
                            if existing_idx is not None:
                                buch2["kapitel"][existing_idx] = kap_eintrag
                            elif ist_ueberarbeitung:
                                # Findet das zuletzt erwähnte Kapitel und ersetzt es
                                import re as _re4
                                m4 = _re4.search(r'kapitel\s*(\d+)', anweisung.lower())
                                if m4:
                                    knum = int(m4.group(1))
                                    cidx = next((i for i,k in enumerate(buch2.get("kapitel",[])) if f"kapitel {knum}" in k["titel"].lower()), None)
                                    if cidx is not None: buch2["kapitel"][cidx] = kap_eintrag
                                    else: buch2.setdefault("kapitel",[]).append(kap_eintrag)
                                else:
                                    buch2.setdefault("kapitel",[]).append(kap_eintrag)
                            else:
                                buch2.setdefault("kapitel",[]).append(kap_eintrag)
                            buch2.setdefault("statistik",{})["kapitel_gesamt"] = len(buch2["kapitel"])
                            buch2["statistik"]["woerter_gesamt"] = sum(len(k["text"].split()) for k in buch2["kapitel"])
                            buch2["statistik"]["letzte_session"] = now
                        buch2["letzteAktion"] = now
                        if gen.get("naechster_schritt"): buch2["naechsterSchritt"] = gen["naechster_schritt"]
                        for k in [x for x in buch2.get("kommentare",[]) if not x.get("erledigt")]:
                            k["erledigt"]=True; k["erledigt_am"]=now; k["antwort"]=gen.get("antwort_text","")
                        with open(bf,"w") as fh2: json.dump(buch2, fh2, ensure_ascii=False, indent=2)
                        _ki_buch_job = {"status":"done","antwort":gen.get("antwort_text",""),"inhalt":gen.get("neuer_inhalt",""),"inhalt_titel":gen.get("neuer_inhalt_titel",""),"error":""}
                    except Exception as ex:
                        _ki_buch_job = {"status":"error","error":str(ex),"antwort":"","inhalt":"","inhalt_titel":""}
                _thr3.Thread(target=_run, daemon=True).start()
                self._send_json({"status": "started"})

            elif self.path == "/api/ki-buch/kommentar":
                bf = os.path.join(WORKSPACE, "data/ki_buch.json")
                with open(bf) as fh:
                    buch = json.load(fh)
                import datetime as _dt
                buch.setdefault("kommentare", []).append({
                    "text": body.get("text", ""),
                    "datum": _dt.datetime.now().strftime("%Y-%m-%d %H:%M")
                })
                with open(bf, "w") as fh:
                    json.dump(buch, fh, ensure_ascii=False, indent=2)
                self._send_json({"ok": True})

            elif self.path == "/api/ki-buch/steuerung":
                # Produzenten-Hebel: Wunsch-Briefkasten, Rote Linien, Richtungsimpuls
                bf = os.path.join(WORKSPACE, "data/ki_buch.json")
                with open(bf) as fh:
                    buch = json.load(fh)
                st = buch.setdefault("steuerung", {})
                st.setdefault("rote_linien", []); st.setdefault("wunsch_briefkasten", []); st.setdefault("richtungsimpuls", "")
                aktion = body.get("aktion", "")
                if aktion == "wunsch_add" and body.get("text", "").strip():
                    st["wunsch_briefkasten"].append(body["text"].strip())
                elif aktion == "wunsch_del":
                    i = body.get("index", -1)
                    if 0 <= i < len(st["wunsch_briefkasten"]): st["wunsch_briefkasten"].pop(i)
                elif aktion == "linie_add" and body.get("text", "").strip():
                    st["rote_linien"].append(body["text"].strip())
                elif aktion == "linie_del":
                    i = body.get("index", -1)
                    if 0 <= i < len(st["rote_linien"]): st["rote_linien"].pop(i)
                elif aktion == "richtung_set":
                    st["richtungsimpuls"] = body.get("text", "").strip()
                with open(bf, "w") as fh:
                    json.dump(buch, fh, ensure_ascii=False, indent=2)
                self._send_json({"ok": True, "steuerung": st})

            elif self.path == "/api/sidebar-state":
                sf = os.path.join(WORKSPACE, "data/sidebar_state.json")
                os.makedirs(os.path.dirname(sf), exist_ok=True)
                with open(sf, "w") as f:
                    json.dump(body, f)
                self._send_json({"ok": True})
            elif self.path == "/api/workshop/save":
                self._send_json(save_workshop(body.get("markdown", "")))
            elif self.path == "/api/workshop/auftrag":
                self._send_json(add_workshop_auftrag(body.get("text", "")))
            elif self.path == "/api/projekte/save":
                self._send_json(projekte_save(
                    body.get("id",""), body.get("title","Neues Projekt"),
                    body.get("icon","📁"), body.get("content",""),
                    body.get("auftraege"), body.get("items"), body.get("prompts")))
            elif self.path == "/api/projekte/new":
                import uuid, time
                new_id = "projekt-" + str(int(time.time()))
                self._send_json(projekte_save(new_id, body.get("title","Neues Projekt"), body.get("icon","📁"), "", []))
            elif self.path == "/api/projekte/auftrag/add":
                self._send_json(projekte_add_auftrag(body.get("pid",""), body.get("text","")))
            elif self.path == "/api/projekte/auftrag/status":
                self._send_json(projekte_auftrag_status(body.get("pid",""), body.get("id",""), body.get("status","")))
            elif self.path == "/api/kosten/guthaben":
                self._send_json(kosten_update_guthaben(body.get("name",""), body.get("betrag", 0), body.get("info")))
            elif self.path == "/api/travel/recommendation":
                uid = body.get("id", "sommer2026")
                typ = body.get("typ", "pauschal")
                self._send_json(save_travel_recommendation(uid, typ, body.get("empfehlung")))
            elif self.path == "/api/travel/request-refresh":
                self._send_json({"ok": True})
            elif self.path == "/api/recipe/save":
                try:
                    fname = save_recipe_docx(body)
                    self._send_json({"ok": True, "datei": fname})
                except Exception as e:
                    self._send_json({"ok": False, "error": str(e)}, status=500)
            elif self.path == "/api/quicknotes/add":
                # Nur speichern (kein Telegram-Send) — Speichern-Button in mmc
                try:
                    text = body.get("text", "").strip()
                    if not text:
                        self._send_json({"error": "kein Text"}, status=400)
                    else:
                        notes_file = os.path.join(WORKSPACE, "data/quicknotes.json")
                        notes = []
                        if os.path.exists(notes_file):
                            with open(notes_file) as _nf: notes = json.load(_nf)
                        notes.insert(0, {"text": text, "ts": datetime.now().isoformat()})
                        notes = notes[:50]
                        os.makedirs(os.path.dirname(notes_file), exist_ok=True)
                        with open(notes_file, "w") as _nf: json.dump(notes, _nf, ensure_ascii=False, indent=2)
                        self._send_json({"ok": True})
                except Exception as e:
                    self._send_json({"error": str(e)}, status=500)
            elif self.path == "/api/quicknotes/delete":
                try:
                    idx = int(body.get("idx", -1))
                    notes_file = os.path.join(WORKSPACE, "data/quicknotes.json")
                    if os.path.exists(notes_file):
                        with open(notes_file) as _nf: notes = json.load(_nf)
                        if 0 <= idx < len(notes):
                            notes.pop(idx)
                            with open(notes_file, "w") as _nf: json.dump(notes, _nf, ensure_ascii=False, indent=2)
                    self._send_json({"ok": True})
                except Exception as e:
                    self._send_json({"error": str(e)}, status=500)
            elif self.path == "/api/quickmsg":
                try:
                    text = body.get("text", "").strip()
                    save_note = body.get("save_note", False)
                    if not text:
                        self._send_json({"error": "kein Text"}, status=400)
                    else:
                        tg_cfg = json.load(open(os.path.join(WORKSPACE, "config/telegram_bot.json")))
                        import urllib.request as _ur_qm
                        msg = f"\U0001f4dd *Schnellnotiz*\n\n{text}"
                        qs = urllib.parse.urlencode({"chat_id": tg_cfg["chris_id"], "text": msg, "parse_mode": "Markdown"})
                        _ur_qm.urlopen(f"https://api.telegram.org/bot{tg_cfg['bot_token']}/sendMessage?{qs}", timeout=5)
                        if save_note:
                            notes_file = os.path.join(WORKSPACE, "data/quicknotes.json")
                            notes = []
                            if os.path.exists(notes_file):
                                with open(notes_file) as _nf: notes = json.load(_nf)
                            notes.insert(0, {"text": text, "ts": datetime.now().isoformat()})
                            notes = notes[:50]
                            with open(notes_file, "w") as _nf: json.dump(notes, _nf, ensure_ascii=False, indent=2)
                        self._send_json({"ok": True})
                except Exception as e:
                    self._send_json({"error": str(e)}, status=500)
            elif self.path == "/api/sos/send":
                try:
                    lat = body.get("lat")
                    lon = body.get("lon")
                    address = body.get("address", "Standort unbekannt")
                    tg_cfg = json.load(open(os.path.join(WORKSPACE, "config/telegram_bot.json")))
                    import urllib.request as _ur_sos
                    msg = f"\U0001f198 *SOS — Standort Chris*\n\n\U0001f4cd {address}"
                    if lat and lon:
                        msg += f"\n\n\U0001f5fa️ [Google Maps](https://maps.google.com/?q={lat},{lon})"
                    qs = urllib.parse.urlencode({"chat_id": tg_cfg["chris_id"], "text": msg, "parse_mode": "Markdown"})
                    _ur_sos.urlopen(f"https://api.telegram.org/bot{tg_cfg['bot_token']}/sendMessage?{qs}", timeout=5)
                    if lat and lon:
                        loc_qs = urllib.parse.urlencode({"chat_id": tg_cfg["chris_id"], "latitude": str(lat), "longitude": str(lon)})
                        _ur_sos.urlopen(f"https://api.telegram.org/bot{tg_cfg['bot_token']}/sendLocation?{loc_qs}", timeout=5)
                    self._send_json({"ok": True})
                except Exception as e:
                    self._send_json({"error": str(e)}, status=500)
            elif self.path == "/api/chat-history/add":
                self._send_json(chat_history_add(body))
            elif self.path == "/api/chat-history/delete":
                self._send_json(chat_history_delete(body.get("id")))
            elif self.path == "/api/chat-history/clear":
                self._send_json(chat_history_clear())
            elif self.path == "/api/bolla/chat":
                self._handle_bolla_stream(body)
            elif self.path == "/api/bolla/tts":
                self._handle_tts(body)
            elif self.path == "/api/mail-command":
                self._send_json(mail_command(body))
            elif self.path == "/api/offers/zip":
                self._send_json(offers_set_zip(body.get("zip","")))
            elif self.path == "/api/newsletter/watchlist":
                self._send_json(newsletter_watchlist_update(body))
            elif self.path == "/api/newsletter/scan":
                self._send_json(newsletter_scan_now())
            elif self.path == "/api/newsletter/search":
                term = body.get("term","").strip()
                if term:
                    self._send_json(newsletter_search(term))
                else:
                    self._send_json({"error":"Begriff fehlt"}, status=400)
            elif self.path == "/api/immo-criteria":
                self._send_json(save_immo_criteria(body))
            elif self.path == "/api/immo-bookmarks":
                self._send_json(save_immo_bookmark(body))
            elif self.path == "/api/immo-bookmarks/delete":
                bid = body.get("id", "").strip()
                if bid:
                    self._send_json(delete_immo_bookmark(bid))
                else:
                    self._send_json({"error": "id fehlt"}, status=400)
            elif self.path == "/api/immo-bookmarks/update":
                bid = body.get("id", "").strip()
                if bid:
                    self._send_json(update_immo_bookmark(bid, body))
                else:
                    self._send_json({"error": "id fehlt"}, status=400)
            elif self.path == "/api/immo-news/refresh":
                self._send_json(get_immo_news(force=True))
            elif self.path == "/api/immo-news/summarize":
                t = body.get("title","").strip()
                s = body.get("summary","").strip()
                if t:
                    self._send_json(summarize_immo_news(t, s))
                else:
                    self._send_json({"error":"title fehlt"}, status=400)
            elif self.path == "/api/makler/status":
                mid    = body.get("id","").strip()
                status = body.get("status","").strip()
                notiz  = body.get("notiz","").strip()
                if mid and status:
                    self._send_json(makler_set_status(mid, status, notiz))
                else:
                    self._send_json({"error":"id und status erforderlich"}, status=400)
            elif self.path == "/api/surfaces/book/action":
                self._send_json(do_book_action(body.get("action", "")))
            elif self.path == "/api/surfaces/pro/action":
                self._send_json(do_pro_action(body.get("action", "")))
            elif self.path == "/api/surfaces/studio/action":
                self._send_json(do_studio_action(body.get("action", "")))
            elif self.path == "/api/adb/pair":
                self._send_json(adb_pair(body.get("host", ""), body.get("code", "")))
            elif self.path == "/api/adb/connect":
                self._send_json(adb_connect(body.get("host", "")))
            elif self.path == "/api/adb/disconnect":
                self._send_json(adb_disconnect(body.get("host", "")))
            elif self.path == "/api/adb/screenshot":
                png, err = adb_screenshot()
                if err or not png:
                    self._send_json({"error": err or "Screenshot fehlgeschlagen"}, status=500)
                    return
                self.send_response(200)
                self.send_header("Access-Control-Allow-Origin", "*")
                self.send_header("Content-Type", "image/png")
                self.send_header("Content-Length", str(len(png)))
                self.end_headers()
                self.wfile.write(png)
            elif self.path == "/api/adb/push":
                import base64
                try:
                    data = base64.b64decode(body.get("data_b64", ""))
                except Exception:
                    self._send_json({"ok": False, "error": "Ungültige Datei-Daten"}, status=400)
                    return
                self._send_json(adb_push(body.get("filename", ""), body.get("remote_dir", "/sdcard/"), data))
            elif self.path == "/api/adb/pull":
                remote = body.get("remote", "")
                data, err = adb_pull(remote)
                if err or data is None:
                    self._send_json({"error": err or "Pull fehlgeschlagen"}, status=500)
                    return
                fname = os.path.basename(remote) or "download"
                self.send_response(200)
                self.send_header("Access-Control-Allow-Origin", "*")
                self.send_header("Content-Type", "application/octet-stream")
                self.send_header("Content-Disposition", f'attachment; filename="{fname}"')
                self.send_header("Content-Length", str(len(data)))
                self.end_headers()
                self.wfile.write(data)
            elif self.path == "/api/bildgen":
                prompt = body.get("prompt", "").strip()
                if not prompt:
                    self._send_json({"error": "Kein Prompt angegeben"}, status=400)
                    return
                ok, limit_err = bildgen_check_limit()
                if not ok:
                    self._send_json({"error": limit_err}, status=429)
                    return
                model = body.get("model", "gemini-2.5-flash-image")
                aspect = body.get("aspect_ratio", "1:1")
                in_img  = body.get("input_image_b64")
                in_mime = body.get("input_mime_type", "image/png")
                img_b64, mime_or_err = bildgen_generate(prompt, model, aspect, in_img, in_mime)
                if img_b64 is None:
                    self._send_json({"error": mime_or_err}, status=500)
                else:
                    _bildgen_counter["count"] += 1
                    remaining = BILDGEN_LIMIT - _bildgen_counter["count"]
                    self._send_json({"image_b64": img_b64, "mime_type": mime_or_err, "remaining": remaining})
            elif self.path == "/api/clipboard":
                text = body.get("text", "")
                if body.get("append"):
                    self._send_json(append_clipboard(text, source=body.get("source", "manual")))
                else:
                    self._send_json(save_clipboard(text))
            elif self.path == "/api/clipboard/delete":
                idx = body.get("idx", -1)
                self._send_json(delete_clipboard_entry(int(idx)))
            elif self.path == "/api/clipboard/restore":
                idx = body.get("idx", -1)
                self._send_json(restore_clipboard_entry(int(idx)))
            elif self.path == "/api/clipboard/trash/delete":
                idx = body.get("idx", -1)
                self._send_json(delete_clipboard_trash_entry(int(idx)))
            elif self.path == "/api/clipboard/image":
                img_b64 = body.get("image_b64", "")
                mime = body.get("mime", "image/png")
                if not img_b64:
                    self._send_json({"error": "Kein Bild"}, status=400); return
                self._send_json(save_clipboard_image(img_b64, mime))
            elif self.path == "/api/clipboard/image/delete":
                filename = body.get("filename", "")
                if not filename:
                    self._send_json({"error": "Kein Dateiname"}, status=400); return
                self._send_json(delete_clipboard_image(filename))
            elif self.path == "/api/photos/start":
                folder       = body.get("folder", "")
                prompt       = body.get("prompt", "Describe this photo briefly: who is visible, where, what activity, what mood.")
                model        = body.get("model", "moondream2")
                lmstudio_url = body.get("lmstudio_url", "http://localhost:1234")
                throttle     = body.get("throttle", 2.0)
                engine       = body.get("engine", "lmstudio")
                if not folder:
                    self._send_json({"error": "Kein Ordner angegeben"}, status=400); return
                self._send_json(photo_start(folder, prompt, model, lmstudio_url, throttle, engine))
            elif self.path == "/api/photos/stop":
                _photo_job["stop"] = True
                self._send_json({"ok": True})
            elif self.path == "/api/photos/clear":
                _photo_job["stop"] = True
                _photo_save_results([])
                self._send_json({"ok": True})
            elif self.path == "/api/photos/search":
                q = body.get("query", "").strip()
                self._send_json(photo_search(q))
            elif self.path == "/api/photos/remove":
                path = body.get("path", "")
                results = [r for r in _photo_load_results() if r.get("path") != path]
                _photo_save_results(results)
                self._send_json({"ok": True})
            elif self.path == "/api/photos/rename":
                old_path = body.get("path", "")
                new_name = body.get("new_name", "")
                if not old_path or not new_name or "/" in new_name or "\\" in new_name or ".." in new_name:
                    self._send_json({"error": "Ungültige Parameter"}, status=400); return
                new_path = os.path.join(os.path.dirname(old_path), new_name)
                if os.path.exists(new_path):
                    self._send_json({"error": "Datei existiert bereits"}); return
                os.rename(old_path, new_path)
                results = _photo_load_results()
                for r in results:
                    if r.get("path") == old_path:
                        r["path"] = new_path
                        r["filename"] = new_name
                _photo_save_results(results)
                self._send_json({"ok": True, "new_path": new_path, "new_name": new_name})
            elif self.path == "/api/photos/export":
                import shutil as _sh, re as _re
                paths = body.get("paths", [])
                query = (body.get("query", "") or "").strip()
                if not paths or not query:
                    self._send_json({"error": "Keine Treffer oder kein Suchbegriff"}); return
                safe = _re.sub(r'[^\w\- ]', '', query).strip()[:60] or "Suchergebnisse"
                dest_dir = os.path.join("/mnt/d/OneDrive", "Suchergebnisse", safe)
                os.makedirs(dest_dir, exist_ok=True)
                copied, errors = 0, 0
                for p in paths:
                    if not os.path.isfile(p):
                        errors += 1; continue
                    try:
                        _sh.copy2(p, os.path.join(dest_dir, os.path.basename(p)))
                        copied += 1
                    except Exception:
                        errors += 1
                self._send_json({"ok": True, "copied": copied, "errors": errors, "dest": dest_dir})
            elif self.path == "/api/korrektur/meta":
                import subprocess as _sp, shutil as _sh, tempfile as _tf, base64 as _b64, re as _re
                img_b64 = body.get("image_b64", "")
                if not img_b64:
                    self._send_json({"error": "Kein Bild"}, status=400); return
                tmp = _tf.NamedTemporaryFile(suffix=".jpg", prefix="korr_meta_", delete=False)
                tmp.write(_b64.b64decode(img_b64)); tmp.close()
                prompt = (
                    f"Schau dir diesen Klassenarbeit-Scan an: {tmp.name}\n\n"
                    "Extrahiere folgende Infos aus dem Scan-Header und antworte NUR mit JSON, "
                    "kein weiterer Text:\n"
                    '{\"klasse\": \"...\", \"fach\": \"...\", \"thema\": \"...\", \"vorname\": \"...\", \"nachname\": \"...\"}\n\n'
                    "klasse = Klassenstufe (z.B. \"10c\" oder \"9\")\n"
                    "fach = Unterrichtsfach (z.B. \"Biologie\", \"WiPo\")\n"
                    "thema = Thema aus der Kopfzeile (kurz, ohne Zeitangaben)\n"
                    "vorname = Vorname des Schuelers aus dem Name-Feld (leer falls nicht erkennbar)\n"
                    "nachname = Nachname des Schuelers aus dem Name-Feld (leer falls nicht erkennbar)"
                )
                claude_bin = _sh.which("claude") or os.path.expanduser("~/.local/bin/claude")
                try:
                    r = _sp.run([claude_bin, "-p", "--output-format", "text", prompt],
                                capture_output=True, text=True, timeout=30,
                                cwd=os.path.expanduser("~"))
                    m = _re.search(r'\{[^}]+\}', r.stdout or "", _re.DOTALL)
                    self._send_json(json.loads(m.group()) if m else {"error": "Parse-Fehler"})
                except Exception as ex:
                    self._send_json({"error": str(ex)})
                finally:
                    try: os.unlink(tmp.name)
                    except: pass
            elif self.path == "/api/korrektur/analyze":
                import subprocess as _sp
                img_b64    = body.get("image_b64", "")
                media_type = body.get("media_type", "image/jpeg")
                klassenstufe = body.get("klassenstufe", "").strip()
                fach         = body.get("fach", "").strip()
                thema        = body.get("thema", "").strip()
                erwartung    = body.get("erwartung", "").strip()
                if not img_b64 or not klassenstufe or not fach:
                    self._send_json({"error": "Bild, Klassenstufe und Fach sind Pflicht"}, status=400)
                    return
                erw_block = f"\n\nErwartungshorizont / Musterlösung:\n{erwartung}" if erwartung else ""
                prompt_text = f"""Du bist ein erfahrener Gymnasiallehrer und korrigierst eine Klassenarbeit.

Fach: {fach} | Klassenstufe: {klassenstufe} | Thema: {thema or '(nicht angegeben)'}{erw_block}

Bitte korrigiere sorgfältig und strukturiert:

1. Lies den handgeschriebenen Text so gut wie möglich. Weise auf unleserliche Stellen hin.

2. Gehe Aufgabe für Aufgabe durch:
   ✅ Richtig / ❌ Falsch oder unvollständig / ➕ Fehlt
   Kurze, konstruktive Anmerkung pro Aufgabe.

3. Bewertungsübersicht als Tabelle:
   | Aufgabe | Max. Punkte | Erreicht | Kommentar |
   (Punkte selbst schätzen wenn kein Erwartungshorizont angegeben)

4. Gesamtnote:
   - Gesamtpunkte: X von Y
   - Prozentzahl: XX%
   - Empfohlene Note: [X] (Schema: 1≥87%, 2≥73%, 3≥59%, 4≥45%, 5≥30%, 6<30%)
   - Kurze Begründung

5. Pädagogisches Feedback (3–4 Sätze): Was lief gut? Wo verbessern?
   Motivierend und konstruktiv formulieren.

Antworte auf Deutsch."""
                try:
                    cli_input = json.dumps({
                        "type": "user",
                        "message": {
                            "role": "user",
                            "content": [
                                {"type": "image", "source": {
                                    "type": "base64", "media_type": media_type, "data": img_b64}},
                                {"type": "text", "text": prompt_text}
                            ]
                        }
                    })
                    proc = _sp.run(
                        ["/home/bolla/.local/bin/claude", "-p",
                         "--input-format", "stream-json",
                         "--output-format", "stream-json",
                         "--verbose"],
                        input=cli_input.encode(),
                        capture_output=True,
                        timeout=120
                    )
                    result_text = ""
                    for line in proc.stdout.decode(errors="replace").splitlines():
                        try:
                            obj = json.loads(line)
                            if obj.get("type") == "assistant":
                                for block in obj.get("message", {}).get("content", []):
                                    if block.get("type") == "text":
                                        result_text += block["text"]
                        except Exception:
                            pass
                    if result_text:
                        self._send_json({"result": result_text})
                    else:
                        stderr = proc.stderr.decode(errors="replace")[:300]
                        self._send_json({"error": f"Keine Antwort vom CLI: {stderr}"}, status=500)
                except Exception as e:
                    self._send_json({"error": str(e)}, status=500)

            elif self.path == "/api/korrektur/docx":
                import base64 as _b64, io as _io, re as _re
                result_text  = body.get("result", "")
                klassenstufe = body.get("klassenstufe", "")
                fach         = body.get("fach", "Korrektur")
                thema        = body.get("thema", "")
                if not result_text:
                    self._send_json({"error": "Kein Korrektur-Text übergeben"}, status=400)
                    return
                try:
                    from docx import Document as _Doc
                    from docx.shared import Pt as _Pt, RGBColor as _RGB, Cm as _Cm
                    from docx.enum.text import WD_ALIGN_PARAGRAPH as _ALIGN
                    import datetime as _dt
                    doc = _Doc()
                    for sec in doc.sections:
                        sec.top_margin = _Cm(1.5); sec.bottom_margin = _Cm(1.5)
                        sec.left_margin = _Cm(2.0); sec.right_margin = _Cm(2.0)
                    # Kompakter Header
                    hdr_p = doc.add_paragraph()
                    hdr_p.alignment = _ALIGN.LEFT
                    hdr_r = hdr_p.add_run("KI-Korrektur")
                    hdr_r.bold = True; hdr_r.font.size = _Pt(13)
                    hdr_r.font.color.rgb = _RGB(0x7C, 0x3A, 0xED)
                    sep_r = hdr_p.add_run(f"   ·   Fach: {fach}   Klasse: {klassenstufe}   {('Thema: ' + thema) if thema else ''}   {_dt.date.today().strftime('%d.%m.%Y')}")
                    sep_r.font.size = _Pt(9)
                    sep_r.font.color.rgb = _RGB(0xA0, 0xA0, 0xA0)
                    # Inhalt
                    EMOJI_RE = _re.compile(r'^(#{1,3})\s*(.*)')
                    doc._korr_tbl = None
                    for raw_line in result_text.splitlines():
                        line = raw_line.rstrip()
                        if not line:
                            continue
                        hm = EMOJI_RE.match(line)
                        if hm:
                            lvl = len(hm.group(1))
                            hdr = doc.add_heading(hm.group(2).strip(), level=min(lvl+1, 3))
                            if hdr.runs:
                                hdr.runs[0].font.size = _Pt(11 if lvl == 1 else 10)
                                hdr.runs[0].font.color.rgb = _RGB(0x7C, 0x3A, 0xED) if lvl == 1 else _RGB(0x06, 0x80, 0xA0)
                            doc._korr_tbl = None
                        elif line.startswith('|') and '|' in line[1:]:
                            cells = [c.strip() for c in line.strip('|').split('|')]
                            if all(set(c.replace('-','').replace(' ','')) == set() for c in cells):
                                continue
                            if not hasattr(doc, '_korr_tbl') or doc._korr_tbl is None:
                                doc._korr_tbl = doc.add_table(rows=0, cols=len(cells))
                                doc._korr_tbl.style = 'Table Grid'
                            row = doc._korr_tbl.add_row()
                            for i, ct in enumerate(cells[:len(row.cells)]):
                                row.cells[i].text = ct
                                if row.cells[i].paragraphs[0].runs:
                                    row.cells[i].paragraphs[0].runs[0].font.size = _Pt(11)
                        else:
                            doc._korr_tbl = None
                            color = None
                            if '✅' in line: color = _RGB(0x16, 0xA3, 0x4A)
                            elif '❌' in line: color = _RGB(0xDC, 0x26, 0x26)
                            elif '➕' in line: color = _RGB(0xD9, 0x77, 0x06)
                            p = doc.add_paragraph()
                            p.paragraph_format.space_before = _Pt(0)
                            p.paragraph_format.space_after  = _Pt(1)
                            r = p.add_run(line)
                            r.font.size = _Pt(12)
                            if color: r.font.color.rgb = color
                    # Footer
                    foot = doc.add_paragraph()
                    foot.alignment = _ALIGN.RIGHT
                    fr = foot.add_run("Bolla · KI-Korrektur · Claude")
                    fr.font.size = _Pt(8); fr.font.italic = True
                    fr.font.color.rgb = _RGB(0xC0, 0xC0, 0xC0)
                    buf = _io.BytesIO()
                    doc.save(buf)
                    self._send_json({"docx_b64": _b64.standard_b64encode(buf.getvalue()).decode()})
                except Exception as e:
                    self._send_json({"error": str(e)}, status=500)

            elif self.path == "/api/korrektur/annotate":
                import base64 as _b64, io as _io, re as _re
                img_b64    = body.get("image_b64", "")
                media_type = body.get("media_type", "image/jpeg")
                result_text = body.get("result", "")
                if not img_b64 or not result_text:
                    self._send_json({"error": "Bild und Korrektur-Text erforderlich"}, status=400)
                    return
                try:
                    from PIL import Image as _Img, ImageDraw as _Draw, ImageFont as _Font
                    img_data = _b64.b64decode(img_b64)
                    img = _Img.open(_io.BytesIO(img_data)).convert("RGB")
                    W, H = img.size

                    # Rechte Randleiste anfügen (Lehrerrandnotizen)
                    MARG = max(210, int(W * 0.30))
                    canvas = _Img.new("RGB", (W + MARG, H), (252, 250, 248))
                    canvas.paste(img, (0, 0))
                    draw = _Draw.Draw(canvas)

                    # Gestrichelte Trennlinie Scan / Randnotizen
                    for yy in range(0, H, 8):
                        draw.line([(W, yy), (W, yy + 4)], fill=(190, 190, 190), width=1)

                    RED      = (196, 0, 0)
                    DARK_RED = (140, 0, 0)
                    GREEN    = (0, 130, 0)
                    ORANGE   = (190, 100, 0)

                    BOLD  = "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"
                    HAND  = "/mnt/c/Windows/Fonts/comic.ttf"
                    try:
                        fnt_note = _Font.truetype(BOLD, max(28, int(MARG * 0.18)))
                    except:
                        fnt_note = _Font.load_default()
                    try:
                        fnt_hd   = _Font.truetype(HAND, max(14, H // 58))
                        fnt_body = _Font.truetype(HAND, max(12, H // 72))
                        fnt_tiny = _Font.truetype(HAND, max(10, H // 90))
                    except:
                        try:
                            fnt_hd = fnt_body = fnt_tiny = _Font.truetype(BOLD, max(12, H // 72))
                        except:
                            fnt_hd = fnt_body = fnt_tiny = _Font.load_default()

                    mx = W + 10  # Rand-X-Start

                    # ── Note-Kreis ganz oben ──────────────────────────────
                    note_m = _re.search(r'Empfohlene Note[:\s*]+\[?([1-6](?:[,\.]\d)?)\]?', result_text, _re.IGNORECASE)
                    note   = note_m.group(1) if note_m else "?"
                    pct_m  = _re.search(r'(\d{1,3})\s*%', result_text)
                    pct    = pct_m.group(1) if pct_m else ""
                    pts_m  = _re.search(r'(\d+)\s*von\s*(\d+)\s*Punkt', result_text, _re.IGNORECASE)
                    nc     = RED if note in ('5','6') else (ORANGE if note in ('3','4') else GREEN)
                    cr     = max(36, int(MARG * 0.20))
                    cx     = W + (MARG - cr) // 2
                    draw.ellipse([(cx, 10), (cx+cr, 10+cr)], fill=nc, outline=DARK_RED, width=2)
                    draw.text((cx + cr//2, 10 + cr//2), note, font=fnt_note, fill=(255,255,255), anchor="mm")
                    y_after_note = 10 + cr + 4
                    if pct:
                        draw.text((cx + cr//2, y_after_note), f"{pct}%", font=fnt_body, fill=nc, anchor="mt")
                        y_after_note += 16
                    if pts_m:
                        draw.text((cx + cr//2, y_after_note), f"{pts_m.group(1)}/{pts_m.group(2)} Pkt", font=fnt_tiny, fill=(100,100,100), anchor="mt")
                        y_after_note += 14

                    # ── Aufgaben-Blöcke ───────────────────────────────────
                    # Markdown-Präfixe normalisieren damit ## Aufgabe 2 als Trenner erkannt wird
                    norm_text = _re.sub(r'^#{1,6}\s*', '', result_text, flags=_re.MULTILINE)
                    task_blocks = _re.findall(
                        r'(?:Aufgabe|Aufg\.?)\s*(\d+)[^\n]*\n((?:(?!(?:Aufgabe|Aufg\.?)\s*\d|\Z).*\n?)*)',
                        norm_text, _re.IGNORECASE)
                    point_lines = _re.findall(
                        r'\|\s*(Aufgabe\s*\d+[^|]*)\|[^|]*\|\s*(\d+(?:[,\.]\d+)?)\s*\|',
                        result_text, _re.IGNORECASE)

                    n_tasks  = max(len(task_blocks), 1)
                    y_start  = y_after_note + 10
                    usable_h = H - y_start - 16
                    # Mindest-Zeilenabstand pro Aufgabe
                    row_h    = max(int(usable_h / n_tasks), int(H * 0.12))

                    for i, (task_name, task_body) in enumerate(task_blocks[:8]):
                        y_base = y_start + i * row_h

                        is_ok  = bool(_re.search(r'✅', task_body))
                        is_bad = bool(_re.search(r'❌', task_body))
                        pts_str = next((f"  {pts}Pkt" for tname, pts in point_lines if task_name in tname), "")
                        color   = GREEN if (is_ok and not is_bad) else RED

                        # Gestrichelte Verbindungslinie Scan → Rand
                        for xc in range(W - 20, W, 5):
                            draw.line([(xc, y_base + 8), (xc + 3, y_base + 8)], fill=(210, 150, 150), width=1)

                        # Aufgaben-Kopfzeile
                        icon = "✓" if (is_ok and not is_bad) else "✗"
                        draw.text((mx, y_base), f"A{task_name} {icon}{pts_str}", font=fnt_hd, fill=color)

                        # Schlagworte aus Aufgaben-Body extrahieren
                        kws = []
                        for ln in task_body.splitlines():
                            ln = ln.strip()
                            if not ln or ln.startswith('#'): continue
                            clean = _re.sub(r'[✅❌➕\*#|]', '', ln).strip()
                            if len(clean) > 4:
                                prio = any(w in clean.lower() for w in ['fehlt','falsch','unvoll','rechts','fehler','nicht erwähnt','kein'])
                                kws.append((clean[:40], prio))
                            if len(kws) >= 6: break
                        kws.sort(key=lambda x: not x[1])
                        line_gap = max(13, fnt_tiny.size + 2) if hasattr(fnt_tiny, 'size') else 15
                        for j, (kw, _) in enumerate(kws[:4]):
                            draw.text((mx + 4, y_base + 18 + j * line_gap), f"→ {kw}", font=fnt_tiny, fill=DARK_RED)

                    # ── Kleiner "KI"-Stempel unten links im Scan ─────────
                    try:
                        fnt_ki = _Font.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 9)
                    except:
                        fnt_ki = _Font.load_default()
                    draw.text((4, H - 13), "KI-Korrektur", font=fnt_ki, fill=(195, 195, 195))

                    buf = _io.BytesIO()
                    canvas.save(buf, "JPEG", quality=90)
                    self._send_json({"image_b64": _b64.standard_b64encode(buf.getvalue()).decode()})
                except Exception as e:
                    self._send_json({"error": str(e)}, status=500)

            elif self.path == "/api/korrektur/excel":
                klasse   = body.get("klasse", "").strip()
                fach     = body.get("fach", "").strip()
                ka_nr    = int(body.get("ka_nr", 1))
                thema    = body.get("thema", "").strip()
                vorname  = body.get("vorname", "").strip()
                nachname = body.get("nachname", "").strip()
                note     = body.get("note", "").strip()
                if not klasse or not fach or not vorname or not nachname or not note:
                    self._send_json({"error": "klasse, fach, vorname, nachname, note sind Pflicht"}, status=400)
                    return
                if ka_nr not in range(1, 5):
                    self._send_json({"error": "ka_nr muss 1–4 sein"}, status=400)
                    return
                try:
                    fpath = excel_upsert_student(klasse, fach, ka_nr, thema, vorname, nachname, note)
                    fname = os.path.basename(fpath)
                    self._send_json({
                        "ok": True, "datei": fname,
                        "download": f"/api/korrektur/excel-download?file={urllib.parse.quote(fname)}"
                    })
                except Exception as e:
                    self._send_json({"error": str(e)}, status=500)

            elif self.path == "/api/elternbrief/entschaerfen":
                notizen    = body.get("notizen", "").strip()
                schueler   = body.get("schueler", "").strip()
                klasse     = body.get("klasse", "").strip()
                anlass     = body.get("anlass", "").strip()
                modus      = body.get("modus", "brief")  # "brief" oder "antwort"
                if not notizen:
                    self._send_json({"error": "Notizen fehlen"}, status=400); return

                import random as _rnd
                _seed = _rnd.randint(0, 9999)

                if modus == "antwort":
                    prompt = f"""Du bist Gymnasiallehrer und antwortest auf einen aufgebrachten Elternbrief.

Elternnachricht:
{notizen}

Schueler/in: {schueler or '(nicht angegeben)'} | Klasse: {klasse or '(nicht angegeben)'} | Anlass: {anlass or '(nicht angegeben)'}

Schreibe die Antwort so wie ein echter, erfahrener Lehrer sie schreiben wuerde — nicht wie ein Muster aus dem Internet.

Konkret:
- Kein Einstieg mit "Vielen Dank fuer Ihre Nachricht" oder "Ich verstehe Ihre Bedenken" oder aehnlichen Floskeln
- Kein gleichmaessiger Rhythmus — kurze und laengere Saetze mischen
- Keine Aufzaehlung von Punkten mit gleichem Satzbau
- Nicht entschuldigend, aber auch nicht arrogant — klar und auf Augenhoehe
- Eine persoenliche, warme Note darf durchscheinen — man soll merken dass du es gut meinst
- Einen konkreten naechsten Schritt nennen (Gespraechstermin o.ae.)
- Abschluss darf ruhig unkonventionell sein, kein "stehe Ihnen jederzeit zur Verfuegung"
- Zufallsseed fuer Variation: {_seed}

Nur der fertige Brief, kein Kommentar."""
                else:
                    prompt = f"""Du bist Gymnasiallehrer und schreibst einen Brief an Eltern.

Deine internen Notizen (NICHT woertlich uebernehmen, nur als Grundlage):
{notizen}

Schueler/in: {schueler or '(nicht angegeben)'} | Klasse: {klasse or '(nicht angegeben)'} | Anlass: {anlass or '(nicht angegeben)'}

Schreibe den Brief wie ein echter Lehrer — nicht wie eine KI-generierte Vorlage.

Konkret:
- Nicht mit "Sehr geehrte Familie X, ich wende mich an Sie bezueglich..." beginnen — finde einen direkteren Einstieg
- Keine drei gleichartig aufgebauten Abschaetze
- Kein "Darüber hinaus", "Abschliessend", "In diesem Sinne"
- Das Kind beim Namen nennen, nicht "Ihr Kind"
- Partnerschaftlich, aber nicht weichgespuelt — Dinge klar benennen
- Eine persoenliche, warme Note darf durchscheinen — man soll merken dass du das Kind magst und es gut meinst
- Schluss ohne "Ich stehe Ihnen jederzeit zur Verfuegung"
- Zufallsseed fuer Variation: {_seed}

Nur der fertige Brief, kein Kommentar."""

                import subprocess as _sp, shutil as _sh
                claude_bin = _sh.which("claude") or CLAUDE_BIN
                try:
                    r = _sp.run([claude_bin, "-p", "--output-format", "text", prompt],
                                capture_output=True, text=True, timeout=90,
                                cwd=os.path.expanduser("~"))
                    if r.returncode != 0:
                        self._send_json({"error": r.stderr[:300]}, status=500); return
                    self._send_json({"ok": True, "brief": r.stdout.strip()})
                except Exception as e:
                    self._send_json({"error": str(e)}, status=500)

            elif self.path == "/api/diff/generate":
                klasse  = body.get("klasse", "").strip()
                fach    = body.get("fach", "").strip()
                aufgabe = body.get("aufgabe", "").strip()
                if not aufgabe:
                    self._send_json({"error": "Aufgabe fehlt"}, status=400); return
                prompt = f"""Du bist Lehrer an einem Gymnasium und erstellst differenzierte Aufgaben.

Klasse: {klasse or 'nicht angegeben'}
Fach: {fach or 'nicht angegeben'}
Original-Aufgabe / Thema: {aufgabe}

Erstelle 3 differenzierte Versionen dieser Aufgabe — eine pro Niveau.
Jede Version enthält die Aufgabenstellung UND eine knappe Musterlösung.

Regeln:
- Fördern: vereinfacht, Hilfestellungen, weniger Komplexität
- Regelklasse: die Aufgabe wie gestellt (leicht aufbereitet)
- Hochbegabt: erweitert, anspruchsvoller, Transfer/Vertiefung

Antworte NUR als reines JSON ohne Markdown:
{{
  "foerdern": {{"aufgabe": "...", "loesung": "..."}},
  "regelklasse": {{"aufgabe": "...", "loesung": "..."}},
  "hochbegabt": {{"aufgabe": "...", "loesung": "..."}}
}}"""
                import subprocess, shutil
                claude_bin = shutil.which("claude") or os.path.expanduser("~/.local/bin/claude")
                result = subprocess.run(
                    [claude_bin, "-p", "--output-format", "json", prompt],
                    capture_output=True, text=True, timeout=120,
                    cwd=os.path.expanduser("~")
                )
                if result.returncode != 0:
                    self._send_json({"error": result.stderr[:300]}, status=500); return
                raw = json.loads(result.stdout).get("result", "").strip()
                if raw.startswith("```"):
                    raw = "\n".join(raw.split("\n")[1:])
                if raw.endswith("```"):
                    raw = raw.rsplit("```", 1)[0]
                raw = raw.strip()
                idx = raw.find('{')
                if idx < 0:
                    self._send_json({"error": "Kein JSON in Antwort"}, status=500); return
                diff_data, _ = json.JSONDecoder().raw_decode(raw, idx)
                self._send_json(diff_data)

            elif self.path == "/api/suno/generate":
                name = body.get("name", "").strip()
                klasse = body.get("klasse", "").strip()
                alter = body.get("alter", "").strip()
                geburtstag = body.get("geburtstag", "").strip()
                abspieltag = body.get("abspieltag", "").strip()
                sprache = body.get("sprache", "de")
                hit = body.get("hit", "").strip()
                style_hint = body.get("style_hint", "").strip()
                kontext = body.get("kontext", "").strip()
                feedback = body.get("feedback", "").strip()
                prev_lyrics = body.get("prev_lyrics", "").strip()
                jugendfrei = body.get("jugendfrei", True)
                jugendfrei_inst = ("WICHTIG: Songtext muss absolut jugendfrei und familienfreundlich sein. "
                                   "Keine zweideutigen Formulierungen, keine Anspielungen, keine suggestiven Ausdrücke. "
                                   "Clean lyrics, appropriate for all ages, family-friendly.\n") if jugendfrei else ""
                SCHOOL_KONTEXT = "Geburtstagssong für Schüler · Computerkurs Herrn Mandel · Lessing-Gymnasium"
                is_school = not kontext or kontext == SCHOOL_KONTEXT
                import re as _re2
                # Klassennamen erkennen: "7a", "10c", "Klasse 7b", "klasse 10" usw.
                is_class = bool(_re2.match(r'(?i)^(?:klasse\s*)?\d{1,2}\s*[a-zA-Z]?$', name))
                # Keine Einzelperson: Klassen, Gruppen oder leer
                is_personal = bool(name) and not is_class

                if sprache == "de" and is_personal:
                    lang_inst = ("auf Deutsch. WICHTIG — AUSSPRACHE DEUTSCHER NAMEN: Suno betont Vornamen oft englisch "
                                 "(Betonung auf letzter Silbe statt erster). Bevor du den Liedtext schreibst, prüfe '{name}': "
                                 "Wenn der Name auf -in, -en, -on, -an, -el, -er, -ig, -iv endet oder fremdsprachlich klingt, "
                                 "ersetze ihn im gesamten Liedtext durch eine lautgerechte Schreibweise die Suno zwingt, "
                                 "die erste Silbe zu betonen. Methoden (wähle die natürlichste): "
                                 "(a) Doppelvokal in der Betonungssilbe: 'Levin' → 'Leevin', 'Kevin' → 'Keevin' "
                                 "(b) Phonetische Umschreibung: 'Levin' → 'Lewien', 'Jason' → 'Jaison' "
                                 "(c) Trennung mit Bindestrich wenn Suno dadurch die erste Silbe betont. "
                                 "Beispiele: 'Levin' → 'Lewien' (LE-vin), 'Kevin' → 'Keewin', 'Justin' → 'Jussten', "
                                 "'Leon' → 'Leeon', 'Jason' → 'Jaison', 'Elias' → 'Eelias'. "
                                 "REGEL: Verwende im Liedtext NUR die phonetische Form. "
                                 "Im Titel und in allen Metadaten IMMER die originale Schreibweise.")
                elif sprache == "de":
                    lang_inst = "auf Deutsch"
                elif is_personal:
                    lang_inst = ("in English. IMPORTANT: Before writing the lyrics, analyze the name and find the single best "
                                 "English phonetic spelling that makes an English AI singer pronounce it exactly like a German "
                                 "speaker would. Rules: (1) Use hyphenated syllables if the name has multiple syllables "
                                 "(e.g. 'Jette' → 'Yet-teh', 'Jan' → 'Yahn', 'Jens' → 'Yens', 'Grete' → 'Greh-teh', "
                                 "'Heinz' → 'Hynts', 'Emilia' → 'Eh-MEE-lee-ah', 'Julia' → 'YOO-lee-ah', 'Maria' → 'Mah-REE-ah'). "
                                 "(2) Every syllable must be pronounceable by an English singer. "
                                 "(3) Use 'eh' for short German e, 'ah' for long German a, 'oo' for German u, 'y' for German j. "
                                 "CRITICAL: Never merge a consonant with 'y+vowel' into one syllable — always split: "
                                 "'-lia' → '-lee-ah', '-ria' → '-ree-ah', '-nia' → '-nee-ah', '-mia' → '-mee-ah'. "
                                 "(4) Replace EVERY occurrence of the name in the lyrics with this phonetic spelling. "
                                 "(5) Do NOT use parenthetical hints or footnotes — only the phonetic spelling in the text.")
                else:
                    lang_inst = "in English"
                style_hint_inst = (f"ZUSÄTZLICHE STIL-HINWEISE (vom Nutzer, haben höchste Priorität): {style_hint}\n"
                                   f"Diese Eigenschaften MÜSSEN den Style-Prompt dominieren — sie beschreiben den echten Sound des Songs.") if style_hint else ""

                def _fetch_song_info(query):
                    """Sucht im Web nach musikalischen Eigenschaften des Songs."""
                    try:
                        try:
                            from ddgs import DDGS
                        except ImportError:
                            from duckduckgo_search import DDGS
                        with DDGS() as ddgs:
                            results = list(ddgs.text(
                                f"{query} song genre tempo BPM instruments production style musical characteristics",
                                max_results=4
                            ))
                        if results:
                            snippets = []
                            for r in results[:4]:
                                body = r.get('body', '').strip()
                                if body and len(body) > 40:
                                    snippets.append(body[:300])
                            return "\n".join(snippets[:3])
                    except Exception:
                        pass
                    return ""

                web_song_info = _fetch_song_info(hit) if hit else ""
                web_song_inst = (f"SONG-RECHERCHE AUS DEM INTERNET (aktuell, verlässlich):\n{web_song_info}\n\n"
                                 f"Nutze diese Informationen um den Sound DIESES spezifischen Songs zu verstehen "
                                 f"und in einen präzisen Style-Prompt zu übersetzen.") if web_song_info else ""

                if hit:
                    hit_inst = (f"REFERENZ-SONG: '{hit}'. "
                                f"WICHTIG: Analysiere DIESEN SPEZIFISCHEN SONG — nicht den allgemeinen Stil des Künstlers. "
                                f"Viele Künstler haben verschiedene Songs mit sehr unterschiedlichem Klang. "
                                f"Extrahiere die konkreten Eigenschaften DIESES Songs: "
                                f"Tempo, Rhythmik, welche Instrumente dominieren (akustisch/elektrisch/elektronisch), "
                                f"Produktionscharakter (roh/weich/poliert/verzerrt), Energie (sanft/mittel/aggressiv), "
                                f"Vers/Chorus-Dynamik, Vocal-Charakter (ton, Ausdruck, Stärke, Textur). "
                                f"Falls du den Song gut kennst: bleib präzise. Falls unsicher: beschreibe was du weißt, "
                                f"aber übertrage NICHT den Klischee-Sound des Künstlers auf diesen Song. "
                                f"KRITISCH — LYRICS-STIL-HARMONIE: Ton und Energie der Lyrics MÜSSEN zum Song passen. "
                                f"Weicher melodischer Song → warme, fließende Lyrics. "
                                f"Aggressiver Beat → kraftvolle, pointierte Lyrics. "
                                f"Style-Prompt: Ausschließlich reine Musik-Eigenschaften — KEIN Künstlername, KEIN Songtitel.")
                else:
                    hit_inst = ""
                if feedback and prev_lyrics:
                    feedback_inst = f"Hier ist der vorherige Liedtext:\n\n{prev_lyrics}\n\nBitte diesen Liedtext gezielt verbessern: {feedback}\nStruktur und gute Passagen beibehalten, nur verbessern was nötig ist."
                elif feedback:
                    feedback_inst = f"Bitte beim Schreiben beachten: {feedback}"
                else:
                    feedback_inst = ""

                def _parse_date(s):
                    from datetime import date as _date
                    seg = [x.strip() for x in s.strip().rstrip(".").split(".") if x.strip()]
                    day, month = int(seg[0]), int(seg[1])
                    if len(seg) >= 3:
                        y = int(seg[2]); year = 2000 + y if y < 100 else y
                    else:
                        year = _date.today().year
                    return _date(year, month, day)

                def gb_kontext(gb_str, ref=None):
                    if not gb_str:
                        return ""
                    from datetime import date as _date
                    ref = ref or _date.today()
                    try:
                        bd = _parse_date(gb_str)
                        delta = (bd - ref).days
                        ds = bd.strftime("%d.%m.%Y")
                        ref_info = f" (Abspieltag: {ref.strftime('%d.%m.%Y')})" if ref != _date.today() else ""
                        if delta == 0:
                            return f"Geburtstag: heute ({ds}){ref_info}"
                        elif 1 <= delta <= 7:
                            return f"Geburtstag: in {delta} Tagen ({ds}){ref_info}"
                        elif delta > 7:
                            return f"Geburtstag: am {ds} (in {delta} Tagen){ref_info}"
                        elif -7 <= delta < 0:
                            return f"Geburtstag: vor {-delta} Tagen ({ds}){ref_info}"
                        else:
                            return f"Geburtstag: {ds} (vor {-delta} Tagen){ref_info}"
                    except Exception:
                        return f"Geburtstag: {gb_str}"

                def gb_lyrics_hint(gb_str, ref=None):
                    if not gb_str:
                        return ""
                    from datetime import date as _date
                    ref = ref or _date.today()
                    try:
                        bd = _parse_date(gb_str)
                        delta = (bd - ref).days
                        if delta == 0:
                            timing = "Der Geburtstag ist heute (am Abspieltag)."
                            tone = "Feier-Stimmung, Gegenwart."
                        elif delta == 1:
                            timing = "Der Geburtstag ist morgen (einen Tag nach dem Abspieltag)."
                            tone = "Vorfreude, Spannung — kreativ umsetzen, NICHT als wäre es heute."
                        elif 2 <= delta <= 7:
                            timing = f"Der Geburtstag ist in {delta} Tagen (nach dem Abspieltag)."
                            tone = "Vorfreude, Countdown-Gefühl — kreativ umsetzen, NICHT als wäre es heute."
                        elif delta > 7:
                            timing = f"Der Geburtstag ist in {delta} Tagen (nach dem Abspieltag)."
                            tone = "Ankündigung, Vorgeschmack — kreativ umsetzen, NICHT als wäre es heute."
                        elif delta == -1:
                            timing = "Der Geburtstag war gestern (einen Tag vor dem Abspieltag)."
                            tone = "Nachträgliche Gratulation, leicht selbstironisch — kreativ umsetzen, NICHT als wäre es heute."
                        elif -7 <= delta < 0:
                            timing = f"Der Geburtstag war vor {-delta} Tagen (vor dem Abspieltag)."
                            tone = "Verspätete Gratulation, Humor über die Verspätung — kreativ umsetzen, NICHT als wäre es heute."
                        else:
                            timing = f"Der Geburtstag war vor {-delta} Tagen (vor dem Abspieltag)."
                            tone = "Deutlich verspätete Gratulation, Augenzwinkern über die Verzögerung — NICHT als wäre es heute."
                        return f"ZEITLICHER KONTEXT FÜR DIE LYRICS: {timing} Stimmung: {tone} Die sprachliche Umsetzung ist frei — Hauptsache der zeitliche Bezug stimmt."
                    except Exception:
                        return ""

                from datetime import date as _dt_date
                ref_date = None
                if abspieltag:
                    try:
                        ref_date = _parse_date(abspieltag)
                    except Exception:
                        ref_date = None

                lehrer = "Mister Mandel" if sprache != "de" else "Herrn Mandel"
                STYLE_RULE = ("Suno style tags in English, comma-separated short tags only, MAX 120 characters total. "
                              "STRICT: NO artist names, NO song titles, NO full sentences — ONLY concise music tags. "
                              "Cover: genre, tempo feel, main instruments, energy, vocal style, production texture. "
                              "LESS IS MORE — 5-8 precise tags beat a long description. "
                              "Example: 'synth-pop, 80s, pulsing bassline, electronic drums, 120bpm, smooth male vocals, polished'")
                TITLE_RULE = ("kreativer Songtitel mit passenden Emojis"
                              + (" — verwende im Titel immer die ORIGINAL-Schreibweise des Namens, nicht die phonetische" if is_personal else ""))

                # Silbenanzahl grob berechnen (Vokalgruppen zählen)
                def _syllables(word):
                    return len(_re2.findall(r'[aeiouäöüyAEIOUÄÖÜY]+', word)) if word else 0

                # Rhythmus-Hinweis für Namen mit ≥3 Silben
                def _rhythm_hint(n):
                    s = _syllables(n)
                    if not n or s < 3:
                        return ""
                    return (f"RHYTHMUS-HINWEIS: '{n}' hat {s} Silben — passe den Namen so in die Verszeilen ein, "
                            f"dass er natürlich klingt und den Takt nicht bricht. "
                            f"Notfalls Kurzform oder Betonung anpassen.")

                # Hinweis wenn Hit-Stil + Zeitkontext kombiniert werden
                def _style_timing_hint(hit_str, gb_str, ref):
                    if not hit_str or not gb_str:
                        return ""
                    return ("STIL-TIMING: Stelle sicher, dass der zeitliche Ton der Lyrics (Vorfreude/Rückblick) "
                            "harmonisch mit dem Stil des Referenz-Songs zusammenpasst — kein Widerspruch zwischen Stimmung und Musik-Energie.")

                # Pflichtinhalt-Zeile je nach Name-Typ
                if is_personal:
                    name_inst = f"- Den Namen '{name}' mehrfach im Liedtext verwenden"
                elif is_class:
                    # Bei englischen Lyrics: "Klasse"-Prefix raus, englische Form "class 7A"
                    class_short = _re2.sub(r'(?i)^klasse\s*', '', name).strip()
                    if sprache != "de":
                        name_inst = f"- Address the class as 'class {class_short.upper()}' multiple times in the lyrics (NOT as 'Klasse {name}')"
                    else:
                        name_inst = f"- Die Klasse '{class_short}' mehrfach direkt ansprechen (z.B. 'Klasse {class_short}', 'ihr')"
                elif name:
                    name_inst = f"- '{name}' mehrfach im Liedtext erwähnen oder ansprechen"
                else:
                    name_inst = ""  # kein Name: nur anlassbezogen

                rhythm_hint = _rhythm_hint(name) if is_personal else ""
                style_timing_hint = _style_timing_hint(hit, geburtstag, ref_date)

                if is_school:
                    who = f"Schüler/in: {name}" if is_personal else (f"Gruppe/Klasse: {name}" if name else "Allgemeiner Klassen-Song")
                    prompt = f"""Du bist ein professioneller Songwriter für Suno AI. Erstelle einen Geburtstagssong {lang_inst}.

{jugendfrei_inst}{who}
Klasse: {klasse}
Alter: {alter}
{gb_kontext(geburtstag, ref_date)}
{gb_lyrics_hint(geburtstag, ref_date)}
{hit_inst}
{web_song_inst}
{style_hint_inst}
{rhythm_hint}
{style_timing_hint}
{feedback_inst}

Pflichtinhalte im Liedtext:
{name_inst}
- Lessing-Gymnasium erwähnen
- Computerkurs bei {lehrer} erwähnen

Struktur: [Intro], [Verse 1], [Chorus], [Verse 2], [Chorus], [Bridge], [Outro]

Gib deine Antwort als JSON zurück (kein Markdown, nur reines JSON):
{{
  "title": "{TITLE_RULE}",
  "lyrics": "vollständiger Liedtext mit Struktur-Tags",
  "style": "{STYLE_RULE}"
}}"""
                else:
                    who_line = f"Person/Gruppe: {name}" if name else "Kein spezifischer Adressat (nur Anlass)"
                    prompt = f"""Du bist ein professioneller Songwriter für Suno AI. Erstelle einen Song {lang_inst}.

{jugendfrei_inst}{who_line}
Anlass / Kontext: {kontext}
{gb_kontext(geburtstag, ref_date)}
{gb_lyrics_hint(geburtstag, ref_date)}
{hit_inst}
{web_song_inst}
{style_hint_inst}
{rhythm_hint}
{style_timing_hint}
{feedback_inst}
{('Pflichtinhalt: ' + name_inst) if name_inst else ''}

Struktur: [Intro], [Verse 1], [Chorus], [Verse 2], [Chorus], [Bridge], [Outro]

Gib deine Antwort als JSON zurück (kein Markdown, nur reines JSON):
{{
  "title": "{TITLE_RULE}",
  "lyrics": "vollständiger Liedtext mit Struktur-Tags",
  "style": "{STYLE_RULE}"
}}"""
                import subprocess, shutil
                claude_bin = shutil.which("claude") or os.path.expanduser("~/.local/bin/claude")
                try:
                    result = subprocess.run(
                        [claude_bin, "-p", "--model", "claude-sonnet-4-6", "--output-format", "json", prompt],
                        capture_output=True, text=True, timeout=240,
                        stdin=subprocess.DEVNULL,
                        cwd=os.path.expanduser("~")
                    )
                except subprocess.TimeoutExpired:
                    self._send_json({"error": "Song-Generierung hat zu lange gedauert (>240s). Bitte nochmal versuchen."}, status=500)
                    return
                if result.returncode != 0:
                    self._send_json({"error": result.stderr[:300] or "claude Fehler (kein stderr)"}, status=500)
                    return
                raw = json.loads(result.stdout).get("result", "")
                raw = raw.strip()
                if raw.startswith("```"):
                    raw = "\n".join(raw.split("\n")[1:])
                if raw.endswith("```"):
                    raw = raw.rsplit("```", 1)[0]
                raw = raw.strip()
                # raw_decode ignoriert Text nach dem ersten gültigen JSON-Objekt
                idx = raw.find('{')
                if idx < 0:
                    raise ValueError("Kein JSON-Objekt in Antwort gefunden")
                song_data, _ = json.JSONDecoder().raw_decode(raw, idx)
                if sprache == "de" and "style" in song_data:
                    song_data["style"] = song_data["style"].rstrip(", ") + ", german lyrics"
                # Englisch: Original-Name im Titel sicherstellen (nur bei Einzelpersonen)
                if sprache != "de" and is_personal and "title" in song_data:
                    if name.lower() not in song_data["title"].lower():
                        song_data["title"] = name + " — " + song_data["title"]
                # Deutsch Schul-Song: kompakten Titel bauen
                if sprache == "de" and is_school:
                    parts = ["Happy Birthday"]
                    if name: parts.append(name)
                    if alter: parts.append(alter)
                    if geburtstag:
                        gb = geburtstag.strip().rstrip(".")
                        segments = [s for s in gb.split(".") if s]
                        if len(segments) == 2:
                            gb = gb.rstrip(".") + ".26"
                        parts.append(gb)
                    song_data["title"] = " ".join(parts)
                self._send_json(song_data)

            elif self.path == "/api/suno/token":
                token = body.get("token", "").strip()
                if not token:
                    self._send_json({"error": "Kein Token"}, status=400); return
                self._send_json(_suno_token_save(token))
            elif self.path == "/api/suno/download-cover":
                title = body.get("title", "").strip()
                cover_engine = body.get("engine", "pollinations")  # "pollinations" oder "gemini"
                if not title:
                    self._send_json({"error": "Kein Titel angegeben"}, status=400); return
                # 1. Suno Feed abfragen
                token = _suno_token()
                if not token:
                    self._send_json({"error": "Kein Suno-Token gefunden"}, status=500); return
                feed_req = urllib.request.Request(
                    SUNO_API_BASE + "/api/feed/?page=0",
                    headers={"Authorization": f"Bearer {token}", "User-Agent": "BollaMC/1.0"}
                )
                try:
                    with urllib.request.urlopen(feed_req, timeout=20) as resp:
                        feed_data = json.loads(resp.read())
                except Exception as e:
                    self._send_json({"error": f"Suno Feed Fehler: {e}"}, status=500); return
                # 2. Song mit Titel suchen
                songs = feed_data if isinstance(feed_data, list) else feed_data.get("clips", feed_data.get("data", []))
                found = None
                for s in songs:
                    sname = (s.get("title") or s.get("display_name") or "")
                    if title.lower() in sname.lower():
                        found = s; break
                if not found:
                    self._send_json({"error": f"Song '{title}' nicht in Suno-Feed gefunden"}); return
                audio_url = found.get("audio_url") or found.get("mp3_url") or ""
                if not audio_url:
                    self._send_json({"error": "Kein audio_url im Song-Objekt"}); return
                # 3. MP3 herunterladen + auf 320kbps konvertieren (RouteNote-Anforderung)
                SUNO_ROUTENOTE_DIR.mkdir(parents=True, exist_ok=True)
                safe_title = "".join(c for c in title if c.isalnum() or c in " _-").strip()
                mp3_path = SUNO_ROUTENOTE_DIR / f"{safe_title}.mp3"
                try:
                    mp3_req = urllib.request.Request(audio_url, headers={"User-Agent": "BollaMC/1.0"})
                    with urllib.request.urlopen(mp3_req, timeout=60) as resp:
                        mp3_data = resp.read()
                    tmp_path = SUNO_ROUTENOTE_DIR / f"{safe_title}_tmp.mp3"
                    tmp_path.write_bytes(mp3_data)
                    # ffmpeg: auf 320kbps konvertieren (Suno liefert nur 64kbps)
                    import subprocess as _sp_ff
                    _ff = _sp_ff.run(
                        ["ffmpeg", "-y", "-i", str(tmp_path), "-b:a", "320k", "-ar", "44100", str(mp3_path)],
                        capture_output=True, timeout=120
                    )
                    tmp_path.unlink(missing_ok=True)
                    if _ff.returncode != 0 or not mp3_path.exists():
                        raise Exception("ffmpeg Konvertierung fehlgeschlagen")
                except Exception as e:
                    self._send_json({"error": f"MP3-Download Fehler: {e}"}); return
                # 4. Bildprompt mit Claude Haiku generieren
                try:
                    import subprocess as _sp2, shutil as _sh2
                    _claude_bin2 = _sh2.which("claude") or os.path.expanduser("~/.local/bin/claude")
                    _cover_prompt_instr = (
                        f"Create an ULTRA-SPECTACULAR, eye-catching album cover image prompt in English for the song '{title}'. "
                        f"This cover must be SO STUNNING that listeners STOP SCROLLING and MUST click play. "
                        f"Think: award-winning photography, viral visual impact, magazine cover quality. "
                        f"IMPORTANT: The artist is optimistic and life-affirming — always interpret the title POSITIVELY. "
                        f"Choose ONE of these visual power styles that fits the title: "
                        f"(1) EPIC GOLDEN HOUR — blazing sun rays, god rays through clouds, silhouettes, warm orange-gold atmosphere; "
                        f"(2) NEON NOIR — glowing neon reflections in rain puddles, dramatic urban night, electric blues and purples; "
                        f"(3) COSMIC WONDER — nebulas, galaxies, bioluminescent ocean meeting space, impossible scales; "
                        f"(4) HYPERREAL NATURE — macro textures, impossibly vivid flowers or waves, saturated dream colors; "
                        f"(5) CINEMATIC EXPLOSION — motion blur, particles, sparks, dust, dramatic action frozen in time. "
                        f"Be EXTREMELY specific: name exact colors (e.g. 'burnt sienna', 'electric cyan'), lighting direction, "
                        f"camera angle (low angle, bird's eye, extreme close-up), materials (chrome, glass, velvet, water droplets), "
                        f"mood (euphoric, mysterious, triumphant), and a unique focal element that anchors the composition. "
                        f"The result must look like it belongs on Spotify's editorial playlists. "
                        f"3000x3000px, square format, ultra high detail, professional color grading. "
                        f"End with: '— absolutely NO text, NO letters, NO words, NO typography anywhere in the image.' "
                        f"Reply with ONLY the image prompt, nothing else."
                    )
                    _cp_result = _sp2.run(
                        [_claude_bin2, "-p", "--model", "claude-sonnet-4-6", "--output-format", "json", _cover_prompt_instr],
                        capture_output=True, text=True, timeout=60, cwd=os.path.expanduser("~")
                    )
                    img_prompt = json.loads(_cp_result.stdout).get("result", "").strip() if _cp_result.returncode == 0 else ""
                    if not img_prompt:
                        raise Exception("Kein Prompt")
                except Exception:
                    img_prompt = f"Ultra-spectacular album cover for '{title}': epic golden hour light rays bursting through dramatic storm clouds over a vast landscape, god rays illuminating swirling particles of gold dust, deep burnt orange and electric violet sky, extreme low-angle shot, hyper-detailed textures, cinematic depth of field, Spotify editorial quality, viral visual impact, award-winning photography style — absolutely NO text, NO letters, NO words, NO typography in the image."
                # 5. Cover generieren — Pollinations (default) oder Gemini
                try:
                    import urllib.parse as _urlparse_cv, io as _io_cv
                    from PIL import Image as _PILImage
                    if cover_engine == "gemini":
                        import base64 as _b64_cv
                        gemini_key = _gemini_key()
                        if not gemini_key:
                            self._send_json({"error": "Kein Gemini API Key"}); return
                        _gem_url = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash-image:generateContent"
                        _gem_payload = json.dumps({
                            "contents": [{"parts": [{"text": img_prompt}]}],
                            "generationConfig": {"responseModalities": ["TEXT", "IMAGE"], "imageConfig": {"aspectRatio": "1:1"}}
                        }).encode()
                        _gem_req = urllib.request.Request(
                            _gem_url, data=_gem_payload,
                            headers={"x-goog-api-key": gemini_key, "Content-Type": "application/json"}
                        )
                        with urllib.request.urlopen(_gem_req, timeout=90) as resp:
                            _gem_data = json.loads(resp.read())
                        img_b64 = None
                        for part in _gem_data.get("candidates", [{}])[0].get("content", {}).get("parts", []):
                            if "inlineData" in part:
                                img_b64 = part["inlineData"]["data"]; break
                        if not img_b64:
                            self._send_json({"error": "Gemini lieferte kein Bild"}); return
                        img_bytes = _b64_cv.b64decode(img_b64)
                    else:
                        _poll_url = (
                            f"https://image.pollinations.ai/prompt/{_urlparse_cv.quote(img_prompt)}"
                            f"?width=3000&height=3000&nologo=true&enhance=true&model=flux-realism"
                        )
                        _poll_req = urllib.request.Request(_poll_url, headers={"User-Agent": "Bolla/1.0"})
                        with urllib.request.urlopen(_poll_req, timeout=120) as resp:
                            img_bytes = resp.read()
                    img = _PILImage.open(_io_cv.BytesIO(img_bytes)).convert("RGB")
                    if img.size != (3000, 3000):
                        img = img.resize((3000, 3000), _PILImage.LANCZOS)
                    cover_path = SUNO_ROUTENOTE_DIR / f"{safe_title}_cover.jpg"
                    img.save(str(cover_path), "JPEG", quality=95)
                except Exception as e:
                    self._send_json({"error": f"Cover-Generierung Fehler: {e}"}); return
                self._send_json({"ok": True, "mp3": str(mp3_path), "cover": str(cover_path)})

            elif self.path == "/api/kiforum/respond":
                import subprocess as _sp2, uuid as _uuid3, base64 as _b64
                thesis = body.get("thesis", "").strip()
                text_prompt = f"""Du bist Bolla, ein meinungsstarker KI-Assistent. Chris hat folgende These aufgestellt:

"{thesis}"

Bewerte diese These kurz und knackig:
- Nimm klar Stellung (dafür oder dagegen oder differenziert)
- Max. 3 Sätze, direkt und provokativ
- Extrahiere 3-5 Schlüsselbegriffe aus deiner Antwort
- Formuliere ein kurzes "Aber..."-Gegenargument (1 Satz) — auch wenn du zustimmst
- Schluss: ein detaillierter englischer Bildprompt der die Kernaussage stark visualisiert (konkrete Szene, Stimmung, Stil)

Antworte NUR als reines JSON:
{{"title": "Kurze Reaktion (max 6 Wörter)", "content": "Deine Bewertung in 2-3 Sätzen.", "emoji": "Emoji", "keywords": ["Begriff1", "Begriff2", "Begriff3"], "aber": "Aber: ein prägnanter Gegenpunkt in einem Satz.", "img_prompt": "Vivid digital illustration: [specific scene directly related to the argument], dramatic lighting, rich colors, detailed, high quality — NOT generic AI imagery"}}"""
                result = _sp2.run(
                    [CLAUDE_BIN, "-p", "--output-format", "json", text_prompt],
                    capture_output=True, text=True, timeout=60, cwd=os.path.expanduser("~")
                )
                if result.returncode != 0:
                    self._send_json({"error": result.stderr[:200]}, status=500); return
                raw = json.loads(result.stdout).get("result", "").strip()
                if "```" in raw:
                    raw = raw.split("```")[1]
                    if raw.startswith("json"): raw = raw[4:]
                raw = raw.strip()
                idx = raw.find('{')
                if idx < 0:
                    self._send_json({"error": "Kein JSON"}, status=500); return
                resp, _ = json.JSONDecoder().raw_decode(raw, idx)

                KIFORUM_FILE = Path(os.path.join(WORKSPACE, "data/kiforum.json"))
                posts = json.loads(KIFORUM_FILE.read_text()) if KIFORUM_FILE.exists() else []
                post_id = str(_uuid3.uuid4())
                img_file = None

                post = {
                    "id": post_id,
                    "type": "bolla",
                    "replyTo": thesis[:120],
                    "timestamp": datetime.now().isoformat(),
                    "title": resp.get("title", "Bollas Reaktion"),
                    "content": resp.get("content", ""),
                    "emoji": resp.get("emoji", "🐾"),
                    "keywords": resp.get("keywords", []),
                    "aber": resp.get("aber", ""),
                    "img_prompt": resp.get("img_prompt", ""),
                    "img": img_file,
                }
                posts.insert(0, post)
                KIFORUM_FILE.write_text(json.dumps(posts, ensure_ascii=False, indent=2))
                # chris_keywords: Bollas Keywords auch für Chris-These verwenden
                self._send_json({**post, "chris_keywords": resp.get("keywords", [])})

            elif self.path == "/api/kiforum/add":
                KIFORUM_FILE = Path(os.path.join(WORKSPACE, "data/kiforum.json"))
                posts = json.loads(KIFORUM_FILE.read_text()) if KIFORUM_FILE.exists() else []
                import uuid as _uuid
                post = {
                    "id": str(_uuid.uuid4()),
                    "type": body.get("type", "user"),
                    "timestamp": datetime.now().isoformat(),
                    "title": body.get("title", ""),
                    "content": body.get("content", ""),
                    "emoji": body.get("emoji", "👤"),
                }
                posts.insert(0, post)
                KIFORUM_FILE.write_text(json.dumps(posts, ensure_ascii=False, indent=2))
                self._send_json(post)

            elif self.path == "/api/kiforum/delete":
                KIFORUM_FILE = Path(os.path.join(WORKSPACE, "data/kiforum.json"))
                posts = json.loads(KIFORUM_FILE.read_text()) if KIFORUM_FILE.exists() else []
                post_id = body.get("id", "")
                posts = [p for p in posts if p.get("id") != post_id]
                KIFORUM_FILE.write_text(json.dumps(posts, ensure_ascii=False, indent=2))
                self._send_json({"ok": True})

            elif self.path == "/api/kiforum/clear":
                KIFORUM_FILE = Path(os.path.join(WORKSPACE, "data/kiforum.json"))
                KIFORUM_FILE.write_text("[]")
                self._send_json({"ok": True})

            elif self.path == "/api/kiforum/img-embed":
                import urllib.request as _urllib, uuid as _uuid_e
                KIFORUM_FILE = Path(os.path.join(WORKSPACE, "data/kiforum.json"))
                post_id = body.get("id", "")
                img_url  = body.get("url", "")
                if not post_id or not img_url:
                    self._send_json({"error": "id und url erforderlich"}, status=400); return
                posts = json.loads(KIFORUM_FILE.read_text()) if KIFORUM_FILE.exists() else []
                post = next((p for p in posts if p.get("id") == post_id), None)
                if not post:
                    self._send_json({"error": "Post nicht gefunden"}, status=404); return
                fname = f"kif_img_{_uuid_e.uuid4()}.png"
                dest  = Path(os.path.join(WORKSPACE, "data", fname))
                req = _urllib.Request(img_url, headers={"User-Agent": "Mozilla/5.0"})
                with _urllib.urlopen(req, timeout=30) as resp:
                    dest.write_bytes(resp.read())
                post["img"] = fname
                KIFORUM_FILE.write_text(json.dumps(posts, ensure_ascii=False, indent=2))
                self._send_json({"ok": True, "img": fname})

            elif self.path == "/api/kiforum/generate":
                import subprocess as _sp, uuid as _uuid2
                today = datetime.now().strftime("%d.%m.%Y")
                KIFORUM_FILE = Path(os.path.join(WORKSPACE, "data/kiforum.json"))
                posts = json.loads(KIFORUM_FILE.read_text()) if KIFORUM_FILE.exists() else []
                existing_titles = [p["title"] for p in posts if p.get("title")][:15]
                avoid_block = ""
                if existing_titles:
                    titles_str = "\n".join(f'- {t}' for t in existing_titles)
                    avoid_block = f"\nBereits behandelte Themen (NICHT wiederholen, auch inhaltlich nicht):\n{titles_str}\n"
                prompt = f"""Du bist Bolla, Chris Mandels KI-Assistent. Heute ist {today}.
Generiere ein spannendes, NEUES KI-Diskussionsthema für ein tägliches Forum.
Das Thema soll für einen technik-affinen, nicht-Entwickler interessant sein.
Provokativ, meinungsstark, zum Diskutieren einladend. Max. 3 Sätze Inhalt.
Schluss: ein detaillierter englischer Bildprompt der das Thema stark visualisiert (konkrete Szene, keine Abstraktion).
{avoid_block}
Antworte NUR als reines JSON:
{{
  "title": "Kurzer prägnanter Titel (max 8 Wörter)",
  "content": "2-3 spannende Sätze zum Thema. Regt zum Nachdenken an.",
  "emoji": "Ein passendes Emoji",
  "keywords": ["Begriff1", "Begriff2", "Begriff3"],
  "img_prompt": "Vivid digital illustration: [specific scene that powerfully visualizes the topic], dramatic lighting, rich colors, detailed, cinematic — NOT generic tech imagery"
}}"""
                claude_bin = CLAUDE_BIN
                result = _sp.run(
                    [claude_bin, "-p", "--output-format", "json", prompt],
                    capture_output=True, text=True, timeout=60,
                    cwd=os.path.expanduser("~")
                )
                if result.returncode != 0:
                    self._send_json({"error": result.stderr[:200]}, status=500); return
                raw = json.loads(result.stdout).get("result", "").strip()
                if "```" in raw:
                    raw = raw.split("```")[1]
                    if raw.startswith("json"): raw = raw[4:]
                raw = raw.strip()
                idx = raw.find('{')
                if idx < 0:
                    self._send_json({"error": "Kein JSON"}, status=500); return
                topic, _ = json.JSONDecoder().raw_decode(raw, idx)
                gen_img_file = None
                post_id_gen = str(_uuid2.uuid4())
                post = {
                    "id": post_id_gen,
                    "type": "bolla",
                    "timestamp": datetime.now().isoformat(),
                    "title": topic.get("title", "Tagesthema"),
                    "content": topic.get("content", ""),
                    "emoji": topic.get("emoji", "🐾"),
                    "keywords": topic.get("keywords", []),
                    "img_prompt": topic.get("img_prompt", ""),
                    "img": gen_img_file,
                }
                posts.insert(0, post)
                KIFORUM_FILE.write_text(json.dumps(posts, ensure_ascii=False, indent=2))
                self._send_json(post)

            elif self.path == "/api/kiforum/conclude":
                import subprocess as _sp5, uuid as _uuid5
                KIFORUM_FILE = Path(os.path.join(WORKSPACE, "data/kiforum.json"))
                posts = json.loads(KIFORUM_FILE.read_text()) if KIFORUM_FILE.exists() else []
                relevant = [p for p in reversed(posts) if p.get("type") in ("user","bolla") and not p.get("isConclusion")]
                lines = []
                for p in relevant:
                    who = "Chris" if p["type"] == "user" else "Bolla"
                    label = f"[{p.get('emoji','')} {p.get('title','')}] " if p.get("title") else ""
                    lines.append(f"{who}: {label}{p.get('content','')}")
                discussion = "\n\n".join(lines)
                prompt = f"""Du bist Bolla. Hier ist die bisherige Diskussion:\n\n{discussion}\n\nZiehe jetzt eine abschließende Schlussfolgerung:\n- Was war der Kern der Diskussion?\n- Zu welchem Ergebnis kommt man?\n- Dein persönliches Fazit (klar, direkt, meinungsstark)\n- Max. 4 Sätze gesamt\n\nAntworte NUR als reines JSON:\n{{"title": "Fazit: kurzer Satz", "content": "Deine Schlussfolgerung in max. 4 Sätzen.", "emoji": "🎯"}}"""
                result = _sp5.run(
                    [CLAUDE_BIN, "-p", "--output-format", "json", prompt],
                    capture_output=True, text=True, timeout=60, cwd=os.path.expanduser("~")
                )
                if result.returncode != 0:
                    self._send_json({"error": result.stderr[:200]}, status=500); return
                raw = json.loads(result.stdout).get("result", "").strip()
                if "```" in raw:
                    raw = raw.split("```")[1]
                    if raw.startswith("json"): raw = raw[4:]
                raw = raw.strip()
                idx = raw.find('{')
                if idx < 0:
                    self._send_json({"error": "Kein JSON"}, status=500); return
                resp, _ = json.JSONDecoder().raw_decode(raw, idx)
                post = {
                    "id": str(_uuid5.uuid4()),
                    "type": "conclusion",
                    "isConclusion": True,
                    "timestamp": datetime.now().isoformat(),
                    "title": resp.get("title", "Bollas Fazit"),
                    "content": resp.get("content", ""),
                    "emoji": resp.get("emoji", "🎯"),
                }
                posts.insert(0, post)
                KIFORUM_FILE.write_text(json.dumps(posts, ensure_ascii=False, indent=2))
                self._send_json(post)

            elif self.path == "/api/kiforum/word":
                from docx import Document as _DocxDoc
                from docx.shared import Pt as _Pt, RGBColor as _RGB, Inches as _Inches
                import uuid as _uuid6
                KIFORUM_FILE = Path(os.path.join(WORKSPACE, "data/kiforum.json"))
                posts = json.loads(KIFORUM_FILE.read_text()) if KIFORUM_FILE.exists() else []
                doc = _DocxDoc()
                for sec in doc.sections:
                    sec.left_margin = _Inches(1.1); sec.right_margin = _Inches(1.1)
                h = doc.add_heading("KI-Forum Diskussion", 0)
                h.runs[0].font.color.rgb = _RGB(0x1a,0x1a,0x2e)
                doc.add_paragraph(f"Exportiert: {datetime.now().strftime('%d.%m.%Y %H:%M')} · Bolla & Chris")
                doc.add_paragraph("")
                for p in reversed(posts):
                    t = p.get("type","")
                    if t == "user":
                        h2 = doc.add_heading(f"💬 Chris: {p.get('title','These')}", 2)
                        for r in h2.runs: r.font.color.rgb = _RGB(0xcc,0x77,0x00)
                    elif t == "bolla":
                        h2 = doc.add_heading(f"{p.get('emoji','🤖')} Bolla: {p.get('title','Reaktion')}", 2)
                        for r in h2.runs: r.font.color.rgb = _RGB(0x22,0x66,0xcc)
                    elif t == "conclusion":
                        h2 = doc.add_heading(f"{p.get('emoji','🎯')} Fazit: {p.get('title','')}", 2)
                        for r in h2.runs: r.font.color.rgb = _RGB(0x88,0x44,0x00)
                    else:
                        continue
                    para = doc.add_paragraph(p.get("content",""))
                    para.runs[0].font.size = _Pt(11) if para.runs else None
                    doc.add_paragraph("")
                fname = f"kif_diskussion_{_uuid6.uuid4().hex[:8]}.docx"
                fpath = Path(WORKSPACE) / "data" / fname
                doc.save(str(fpath))
                with open(fpath, "rb") as f:
                    data = f.read()
                self.send_response(200)
                self.send_header("Content-Type","application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                self.send_header("Content-Disposition",f'attachment; filename="{fname}"')
                self.send_header("Content-Length",str(len(data)))
                self.end_headers()
                self.wfile.write(data)

            elif self.path == "/api/trip/plan":
                self._send_json(trip_plan(
                    body.get("ziel",""), body.get("startort",""),
                    body.get("tage",2), body.get("datum",""),
                    body.get("fahrzeug","Auto"), body.get("interessen",""),
                    body.get("sprache","de")
                ))
            elif self.path == "/api/trip/docx":
                trip_data = body.get("trip")
                if not trip_data:
                    self._send_json({"error":"trip fehlt"}, status=400); return
                docx_bytes = trip_generate_docx(trip_data, body.get("sprache","de"))
                title = (trip_data.get("title") or "Trip").replace("/","_")[:60]
                fname = title + ".docx"
                self.send_response(200)
                self.send_header("Content-Type","application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                self.send_header("Content-Disposition",f'attachment; filename="{fname}"')
                self.send_header("Content-Length", str(len(docx_bytes)))
                self.end_headers()
                self.wfile.write(docx_bytes)
            elif self.path == "/api/amadeus/config":
                cid = body.get("client_id","").strip()
                sec = body.get("client_secret","").strip()
                if cid and sec:
                    self._send_json(amadeus_save_config(cid, sec))
                else:
                    self._send_json({"error":"client_id und client_secret erforderlich"}, status=400)
            else:
                self._send_json({"error": "not found"}, status=404)
        except (BrokenPipeError, ConnectionResetError):
            pass
        except Exception as e:
            tb = traceback.format_exc()
            print(f"POST Error: {tb}")
            try:
                self._send_json({"error": str(e)}, status=500)
            except Exception:
                pass

    def _handle_bolla_stream(self, body):
        msg = body.get("message", "").strip()
        if not msg and not body.get("image"):
            self._send_json({"error": "Keine Nachricht"}, status=400)
            return
        sid = body.get("session_id")
        img = body.get("image")

        self.send_response(200)
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Content-Type", "application/x-ndjson; charset=utf-8")
        self.send_header("Cache-Control", "no-cache")
        self.send_header("X-Accel-Buffering", "no")
        self.end_headers()

        try:
            for event in bolla_chat_stream(msg, sid, image_b64=img):
                line = json.dumps(event, ensure_ascii=False) + "\n"
                self.wfile.write(line.encode("utf-8"))
                self.wfile.flush()
        except (BrokenPipeError, ConnectionResetError):
            pass
        except Exception as e:
            try:
                err = json.dumps({"type": "error", "error": str(e)}) + "\n"
                self.wfile.write(err.encode())
                self.wfile.flush()
            except Exception:
                pass

    def _handle_tts(self, body):
        text = (body.get("text") or "").strip()
        if not text:
            self._send_json({"error": "Kein Text"}, status=400)
            return
        voice = body.get("voice") or "de-DE-SeraphinaMultilingualNeural"
        try:
            rate = int(body.get("rate", 0))
        except (TypeError, ValueError):
            rate = 0
        rate = max(-50, min(100, rate))

        # Text kürzen zur Sicherheit (Azure-Limit ist hoch, aber 2000 reicht für Sprachausgabe)
        if len(text) > 2000:
            text = text[:2000] + "..."

        audio, err = azure_tts(text, voice, rate)
        if err or not audio:
            self._send_json({"error": err or "TTS fehlgeschlagen"}, status=500)
            return
        self.send_response(200)
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Content-Type", "audio/mpeg")
        self.send_header("Content-Length", str(len(audio)))
        self.end_headers()
        self.wfile.write(audio)

    def do_OPTIONS(self):
        self.send_response(204)
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Methods", "GET, POST, OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "*")
        self.end_headers()


# ── Kurztrip-Planer ──────────────────────────────────────────────────────────

def trip_plan(ziel, startort, tage, datum, fahrzeug, interessen, sprache):
    if not ziel:
        return {"error": "Zielort fehlt"}
    import shutil, subprocess as _sp
    claude_bin = shutil.which("claude") or os.path.expanduser("~/.local/bin/claude")
    tage = max(1, min(3, int(tage) if str(tage).isdigit() else 2))
    lang_inst = "in German" if sprache == "de" else "in English"
    interessen_inst = f"Interests: {interessen}" if interessen else ""
    datum_str = datum or datetime.now().strftime("%Y-%m-%d")
    color_guide = (
        'Colors by category: parking/start: "#6366f1", sightseeing/culture: "#f59e0b", '
        'viewpoint/nature: "#16a34a", break/cafe/restaurant: "#0ea5e9", museum: "#8b5cf6", '
        'end/return: "#ef4444"'
    )
    # E-Auto (VW ID.3): saison-abhängige Reichweite + Ladestopps in die Reisezeit einrechnen.
    ev_inst = ""
    fz_low = (fahrzeug or "").lower()
    if "id.3" in fz_low or "id3" in fz_low or "e-auto" in fz_low or "elektro" in fz_low:
        try:
            _month = int(datum_str[5:7])
        except Exception:
            _month = datetime.now().month
        _summer = 4 <= _month <= 9
        _range = 200 if _summer else 160
        _season = "summer" if _summer else "winter"
        ev_inst = f"""

ELECTRIC VEHICLE — CHARGING (CRITICAL for time/distance estimates):
- The vehicle is a VW ID.3. Plan with a realistic, buffered range of ~{_range} km per charge in {_season}.
- After roughly every {_range} km of driving, the car MUST stop ~30 minutes to charge.
- For any leg or total route longer than ~{_range} km, insert charging stop(s) of ~30 min each
  at sensible locations with fast chargers (color "#0ea5e9", kategorie "break", name e.g. "Ladestopp ...").
- ADD the charging time (~30 min per stop) to "gesamt_zeit" and to per-leg times, and mention the charging
  stops in the route. Real-world reference: ~720 km took ~9 h in summer, ~10 h in winter — be generous and
  round travel times UP rather than down. Make the estimates realistic for an EV, not for a combustion car."""
    prompt = f"""You are an expert travel guide. Create a detailed {tage}-day trip plan for {ziel}.
Starting point: {startort or 'not specified (assume typical travel hub)'}
Vehicle: {fahrzeug}
Date: {datum_str}
{interessen_inst}
Language for output: {lang_inst}

IMPORTANT RULES:
- Provide PRECISE GPS coordinates (WGS84 decimal degrees) for every stop. Verify accuracy.
- Create a logical walking/driving route that minimizes backtracking.
- For {tage} day(s): distribute stops evenly across days with a logical daily theme.
- Each stop: realistic walking/driving distance and time from previous stop.
- Include 1 parking/start stop, at least 1 café/break stop per day, end back at start.
- {color_guide}{ev_inst}

Return ONLY this exact JSON (no markdown, no extra text):
{{
  "title": "CityName — {tage}-Day Trip",
  "subtitle": "Vehicle · Date · N stops · ~X km total · ~Y hours with breaks",
  "tage": [
    {{
      "tag": 1,
      "label": "Day 1 — Theme (e.g. Old Town & Harbor)",
      "stops": [
        {{
          "nr": 1,
          "name": "Parking / Start Point",
          "lat": 00.0000,
          "lon": 00.0000,
          "color": "#6366f1",
          "dist": "Start",
          "zeit": "",
          "kurz": "Short description (10-15 words)",
          "detail": "Detailed description with practical tips (3-5 sentences)",
          "kategorie": "parking"
        }}
      ]
    }}
  ],
  "gesamt_km": "~X km",
  "gesamt_zeit": "~Y hours including breaks and museums"
}}"""

    try:
        res = _sp.run(
            [claude_bin, "-p", "--model", "claude-sonnet-4-6", "--output-format", "json", prompt],
            capture_output=True, text=True, timeout=120, stdin=_sp.DEVNULL,
            cwd=os.path.expanduser("~")
        )
        if res.returncode != 0:
            return {"error": res.stderr[:300] or "Claude Fehler"}
        raw = json.loads(res.stdout).get("result", "").strip()
        if raw.startswith("```"): raw = "\n".join(raw.split("\n")[1:])
        if raw.endswith("```"): raw = raw.rsplit("```",1)[0]
        raw = raw.strip()
        idx = raw.find('{')
        if idx < 0: raise ValueError("Kein JSON")
        data, _ = json.JSONDecoder().raw_decode(raw, idx)
        return data
    except Exception as e:
        return {"error": str(e)}


def trip_generate_docx(trip_data, sprache="de"):
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from PIL import Image, ImageDraw, ImageFont
    import io, tempfile

    all_stops = [s for t in trip_data.get("tage", []) for s in t.get("stops", [])]

    def hex2rgb(h):
        h = h.lstrip('#')
        return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

    def set_bg(cell, hex_color):
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'), hex_color.lstrip('#'))
        tcPr.append(shd)

    def set_cell_text(cell, text, size=10, bold=False, color=None, italic=False,
                      space_before=2, space_after=2, align=WD_ALIGN_PARAGRAPH.LEFT):
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        p.alignment = align
        run = p.runs[0] if p.runs else p.add_run(text)
        if p.runs: run.text = text
        run.font.size = Pt(size); run.font.bold = bold; run.font.italic = italic
        if color: run.font.color.rgb = color

    # Karte generieren
    map_path = None
    try:
        from staticmap import StaticMap, CircleMarker
        from staticmap.staticmap import _lon_to_x, _lat_to_y
        m = StaticMap(900, 700, url_template="https://tile.openstreetmap.org/{z}/{x}/{y}.png")
        for s in all_stops:
            m.add_marker(CircleMarker((s["lon"], s["lat"]), "#ffffff", 2))
        img = m.render(zoom=14)
        draw = ImageDraw.Draw(img)
        try:
            font_nr = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 17)
            font_sm = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 11)
            font_hd = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 12)
        except Exception:
            font_nr = font_sm = font_hd = ImageFont.load_default()
        for s in all_stops:
            tx = _lon_to_x(s["lon"], m.zoom); ty = _lat_to_y(s["lat"], m.zoom)
            px = m._x_to_px(tx); py = m._y_to_px(ty)
            rgb = hex2rgb(s.get("color","#6366f1")); r = 16
            draw.ellipse([px-r-2,py-r-2,px+r+2,py+r+2], fill=(10,10,10))
            draw.ellipse([px-r,py-r,px+r,py+r], fill=rgb)
            txt = str(s["nr"]); bb = draw.textbbox((0,0),txt,font=font_nr)
            draw.text((px-(bb[2]-bb[0])//2-bb[0], py-(bb[3]-bb[1])//2-bb[1]), txt, font=font_nr, fill="white")
        # Legende
        lx, ly = 8, 8; lh = 28 + len(all_stops)*20 + 6
        draw.rectangle([lx,ly,lx+250,ly+lh], fill=(255,255,255,220), outline=(160,160,160))
        draw.text((lx+7,ly+6), trip_data.get("title","RUNDGANG")[:30], font=font_hd, fill=(20,20,20))
        for i, s in enumerate(all_stops):
            ry = ly+26+i*20; rgb = hex2rgb(s.get("color","#6366f1")); r2=8
            draw.ellipse([lx+7,ry+2,lx+7+r2*2,ry+2+r2*2], fill=rgb)
            draw.text((lx+30,ry+3), s["name"][:32], font=font_sm, fill=(20,20,20))
        tf = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
        img.save(tf.name, quality=92); map_path = tf.name
    except Exception as e:
        print(f"Karte fehlgeschlagen: {e}")

    doc = Document()
    for sec in doc.sections:
        sec.top_margin=Cm(1.5); sec.bottom_margin=Cm(1.5)
        sec.left_margin=Cm(1.8); sec.right_margin=Cm(1.8)

    t = doc.add_heading(trip_data.get("title","Reiseplan"), 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].font.size = Pt(20); t.runs[0].font.color.rgb = RGBColor(0x1e,0x40,0xaf)
    t.paragraph_format.space_before = Pt(0); t.paragraph_format.space_after = Pt(4)

    sub = doc.add_paragraph(trip_data.get("subtitle",""))
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(10); sub.runs[0].font.color.rgb = RGBColor(0x64,0x74,0x8b)
    sub.paragraph_format.space_after = Pt(8)

    if map_path:
        try:
            doc.add_picture(map_path, width=Inches(6.0))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.paragraphs[-1].paragraph_format.space_after = Pt(8)
        except Exception: pass

    BLUE = RGBColor(0x1e,0x40,0xaf)
    COL_W = [Cm(0.8), Cm(5.0), Cm(2.6), Cm(8.8)]

    for tag in trip_data.get("tage", []):
        if len(trip_data.get("tage",[])) > 1:
            lbl = doc.add_paragraph(tag.get("label",""))
            lbl.runs[0].font.size = Pt(12); lbl.runs[0].font.bold = True
            lbl.runs[0].font.color.rgb = BLUE; lbl.paragraph_format.space_after = Pt(4)

        tbl = doc.add_table(rows=1, cols=4); tbl.style = 'Table Grid'
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = tbl.rows[0].cells
        for i,(txt,w) in enumerate(zip(["#","Station","Weg/Zeit","Highlight"],COL_W)):
            hdr[i].width=w
            set_cell_text(hdr[i],txt,size=10,bold=True,color=RGBColor(0xff,0xff,0xff),align=WD_ALIGN_PARAGRAPH.CENTER)
            set_bg(hdr[i],"1e40af")
        for si, s in enumerate(tag.get("stops",[])):
            dist_txt = s.get("dist","") if s.get("dist")=="Start" else (s.get("dist","") + (" / "+s.get("zeit","") if s.get("zeit") else ""))
            row = tbl.add_row().cells
            for j,w in enumerate(COL_W): row[j].width=w
            set_cell_text(row[0],str(s["nr"]),size=10,bold=True,align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(row[1],s.get("name",""),size=10,bold=True)
            set_cell_text(row[2],dist_txt,size=9,color=RGBColor(0x44,0x55,0x6b),align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(row[3],s.get("kurz",""),size=9)
            if si % 2 == 1:
                for c in row: set_bg(c,"eef2ff")
        row = tbl.add_row().cells
        for j,w in enumerate(COL_W): row[j].width=w
        set_cell_text(row[0],"∑",size=10,bold=True,align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row[1],"Gesamt",size=10,bold=True)
        set_cell_text(row[2],trip_data.get("gesamt_km",""),size=10,bold=True,align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row[3],trip_data.get("gesamt_zeit",""),size=9,bold=True)
        for c in row: set_bg(c,"dbeafe")
        doc.add_paragraph("").paragraph_format.space_after = Pt(6)

    # Seite 2: Details
    doc.add_page_break()
    lbl2 = doc.add_paragraph("Stationen im Detail")
    lbl2.runs[0].font.size=Pt(14); lbl2.runs[0].font.bold=True
    lbl2.runs[0].font.color.rgb=BLUE; lbl2.paragraph_format.space_after=Pt(8)
    COL_W2=[Cm(0.8),Cm(4.4),Cm(12.0)]
    tbl2 = doc.add_table(rows=1, cols=3); tbl2.style='Table Grid'
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    h2 = tbl2.rows[0].cells
    for i,(txt,w) in enumerate(zip(["#","Station","Beschreibung & Tipps"],COL_W2)):
        h2[i].width=w
        set_cell_text(h2[i],txt,size=10,bold=True,color=RGBColor(0xff,0xff,0xff),align=WD_ALIGN_PARAGRAPH.CENTER)
        set_bg(h2[i],"1e40af")
    for si, s in enumerate(all_stops):
        row=tbl2.add_row().cells
        for j,w in enumerate(COL_W2): row[j].width=w
        set_cell_text(row[0],str(s["nr"]),size=11,bold=True,align=WD_ALIGN_PARAGRAPH.CENTER)
        p=row[1].paragraphs[0]; p.paragraph_format.space_before=Pt(3); p.paragraph_format.space_after=Pt(2)
        nr=p.add_run(s.get("name","")); nr.font.size=Pt(10); nr.font.bold=True
        if s.get("dist") and s["dist"]!="Start":
            dr=p.add_run(f"\n{s.get('dist','')} / {s.get('zeit','')}"); dr.font.size=Pt(8); dr.font.italic=True; dr.font.color.rgb=RGBColor(0x64,0x74,0x8b)
        dp=row[2].paragraphs[0]; dp.paragraph_format.space_before=Pt(3); dp.paragraph_format.space_after=Pt(3)
        dr2=dp.add_run(s.get("detail","")); dr2.font.size=Pt(10)
        if si % 2 == 1:
            for c in row: set_bg(c,"eef2ff")

    foot=doc.add_paragraph("🐾  Bolla · Viel Spaß und gutes Wetter!")
    foot.alignment=WD_ALIGN_PARAGRAPH.CENTER; foot.runs[0].font.size=Pt(9)
    foot.runs[0].font.color.rgb=RGBColor(0x94,0xa3,0xb8); foot.paragraph_format.space_before=Pt(10)

    buf = io.BytesIO(); doc.save(buf)
    if map_path:
        try: os.unlink(map_path)
        except Exception: pass
    return buf.getvalue()


# ── Amadeus API ──────────────────────────────────────────────────────────────

AMADEUS_CONFIG_FILE = Path(os.path.join(WORKSPACE, "config/amadeus_config.json"))
_amadeus_token_cache = {"token": None, "expires": 0}

def amadeus_get_config():
    if AMADEUS_CONFIG_FILE.exists():
        d = json.loads(AMADEUS_CONFIG_FILE.read_text())
        return {"client_id": d.get("client_id",""), "configured": bool(d.get("client_id"))}
    return {"client_id": "", "configured": False}

def amadeus_save_config(client_id, client_secret):
    AMADEUS_CONFIG_FILE.parent.mkdir(parents=True, exist_ok=True)
    AMADEUS_CONFIG_FILE.write_text(json.dumps({"client_id": client_id, "client_secret": client_secret}, indent=2))
    AMADEUS_CONFIG_FILE.chmod(0o600)
    _amadeus_token_cache["token"] = None
    return {"ok": True}

def _amadeus_token():
    import time, urllib.request as _ur, urllib.parse as _up2
    if _amadeus_token_cache["token"] and time.time() < _amadeus_token_cache["expires"] - 30:
        return _amadeus_token_cache["token"]
    if not AMADEUS_CONFIG_FILE.exists():
        raise ValueError("Amadeus API-Key nicht konfiguriert. Bitte ⚙️ API-Key klicken.")
    cfg = json.loads(AMADEUS_CONFIG_FILE.read_text())
    data = _up2.urlencode({"grant_type":"client_credentials","client_id":cfg["client_id"],"client_secret":cfg["client_secret"]}).encode()
    req = _ur.Request("https://test.api.amadeus.com/v1/security/oauth2/token",
                      data=data, headers={"Content-Type":"application/x-www-form-urlencoded"})
    with _ur.urlopen(req, timeout=15) as r:
        tok = json.loads(r.read())
    _amadeus_token_cache["token"] = tok["access_token"]
    _amadeus_token_cache["expires"] = time.time() + tok.get("expires_in", 1800)
    return _amadeus_token_cache["token"]

def amadeus_search_flights(origin, dest, date, return_date, adults=2):
    if not origin or not dest or not date:
        return {"error": "origin, dest und date sind Pflichtfelder"}
    import urllib.request as _ur, urllib.parse as _up2
    try:
        token = _amadeus_token()
        params = {"originLocationCode": origin, "destinationLocationCode": dest,
                  "departureDate": date, "adults": str(adults), "max": "15", "currencyCode": "EUR"}
        if return_date:
            params["returnDate"] = return_date
        url = "https://test.api.amadeus.com/v2/shopping/flight-offers?" + _up2.urlencode(params)
        req = _ur.Request(url, headers={"Authorization": f"Bearer {token}"})
        with _ur.urlopen(req, timeout=20) as r:
            data = json.loads(r.read())
        flights = data.get("data", [])
        return {"flights": flights, "meta": data.get("meta", {})}
    except Exception as e:
        return {"error": str(e)}


if __name__ == "__main__":
    # Spracherkennung läuft komplett im Handy-Browser (Web Speech API) — kein Whisper mehr,
    # kein /api/transcribe, kein Modell auf dem Server (2026-06-07 restlos entfernt).

    port = 18790
    # SO_REUSEADDR + SO_REUSEPORT damit Port-Bindung sofort nach Neustart klappt
    # (Linux: SO_REUSEPORT umgeht TIME_WAIT zuverlässig)
    import socket as _socket
    class _ReusableServer(ThreadingHTTPServer):
        allow_reuse_address = True
        def server_bind(self):
            self.socket.setsockopt(_socket.SOL_SOCKET, _socket.SO_REUSEADDR, 1)
            try:
                self.socket.setsockopt(_socket.SOL_SOCKET, _socket.SO_REUSEPORT, 1)
            except (AttributeError, OSError):
                pass  # SO_REUSEPORT nicht überall verfügbar
            super().server_bind()
    server = _ReusableServer(("0.0.0.0", port), Handler)
    print(f"Mission Control API läuft auf http://127.0.0.1:{port}")
    server.serve_forever()
