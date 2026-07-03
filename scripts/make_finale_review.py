#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""AURORA Finale-5★ Review-Doc aus finale_5stern.json (Aptos, kompakt)."""
import json, os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

D = json.load(open("/home/bolla/workspace/data/satzkur/finale_5stern.json", encoding="utf-8"))
AE = D["aenderungen"]

doc = Document()
st = doc.styles["Normal"]; st.font.name = "Aptos"; st.font.size = Pt(10.5)
st.element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
st.paragraph_format.space_after = Pt(0); st.paragraph_format.line_spacing = 1.13

def para(text="", size=10.5, bold=False, color=None, after=0, before=0, italic=False):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.line_spacing = 1.13
    r = p.add_run(text); r.font.name = "Aptos"; r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = RGBColor(*color)
    return p, r

para("AURORA — Finale auf 5★ (Fable)", 19, True, (0x1a,0x1a,0x3a), after=2)
para("17 gezielte Kürzungen · %d Wörter straffer · 4. Juli 2026" % D.get("woerter_entfernt_ca", 0), 11, False, (0x66,0x66,0x66), after=6, italic=True)

para("Was & warum (Fables Bilanz):", 12, True, (0x15,0x60,0x2d), after=2)
para(D.get("bilanz", ""), 10.5, after=6)

para("Sicherheitsnetz:", 11, True, after=1)
para("Alles ist angewendet, aber vollständig zurückrollbar — das Buch vor dieser Straffung liegt als "
     "Backup: backups/ki_buch_vor_finale5stern_*.json. Der komplette Diff steckt in data/satzkur/finale_5stern.json. "
     "Falls dir eine Kürzung nicht gefällt: sag mir die Nummer, ich hol die Stelle einzeln zurück.", 10, after=8)

para("Die 17 Änderungen im Detail", 14, True, (0x1a,0x1a,0x3a), after=3, before=2)

ACT = {"streichen": "STREICHEN", "ersetzen": "ERSETZEN / STRAFFEN"}
for i, a in enumerate(AE, 1):
    kap = "Prolog" if a["nr"] == 0 else f"Kap {a['nr']}"
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(4)
    r1 = p.add_run(f"{i}.  [{kap}] "); r1.bold = True; r1.font.name="Aptos"; r1.font.size=Pt(10.5); r1.font.color.rgb=RGBColor(0x2a,0x4a,0x8a)
    r2 = p.add_run(f"{ACT.get(a['aktion'], a['aktion'])} — {a.get('was','')}"); r2.bold=True; r2.font.name="Aptos"; r2.font.size=Pt(10.5)
    para("Warum: " + a.get("begruendung", ""), 9.5, italic=True, color=(0x55,0x55,0x55), after=1)
    # Vorschau des betroffenen Textes
    alt = (a.get("anker") or "").replace("\n", " ").strip()
    if len(alt) > 180: alt = alt[:180] + "…"
    para("Vorher: " + alt, 9.5, color=(0x80,0x40,0x40), after=1)
    if a["aktion"] == "ersetzen":
        neu = (a.get("ersatz") or "").replace("\n", " ").strip()
        if len(neu) > 180: neu = neu[:180] + "…"
        pv, rv = para("Nachher: " + neu, 9.5, after=2)
        rv.font.color.rgb = RGBColor(0x12,0x5a,0x28)
    else:
        para("Nachher: (ersatzlos gestrichen)", 9.5, color=(0x12,0x5a,0x28), after=2)

out = "/mnt/d/OneDrive/Desktop/AURORA_Finale_5Sterne_Review.docx"
doc.save(out)
print("Gespeichert:", out, "|", os.path.getsize(out), "bytes |", len(AE), "Änderungen")
