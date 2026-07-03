#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""AURORA Kanon-Entscheidungen - bündelt die 59 'vorlegen'-Funde in Grundsatz-Cluster.
   Fliesstext nutzt ausschliesslich typografische Anfuehrungszeichen, nie gerade."""
import json, os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

SRC = {f["id"]: f for f in json.load(open("/home/bolla/workspace/data/satzkur/review_B.json", encoding="utf-8"))}

# Cluster: (Titel, Grundsatzfrage, Empfehlung, [ids])
CLUSTERS = [
    ("① Zeitachse: Wie viele Wochen von AURORAs Erwachen bis zum Finale?",
     "Mehrere Stellen nennen unterschiedliche Spannen (zwei / vier / sechs Wochen; „drei Wochen nach Totalausfall“).",
     "EMPFEHLUNG: „vier Wochen“ als Kanon — Noah sagt es selbst mehrfach (Kap 21/26/46). Dann alle abweichenden Stellen darauf angleichen.",
     [22, 25, 35, 37, 38, 39, 42, 45, 75, 80, 81]),
    ("② „Elf Tage“ — festes Leitmotiv oder exakte Tageszählung?",
     "Das Motiv kollidiert mit einzelnen Tageszählungen (zwölf/elf).",
     "EMPFEHLUNG: Als bewusstes Leitmotiv behalten und die abweichenden Zahlen (Kap 18/9) daran angleichen.",
     [16, 19, 32]),
    ("③ Marias Alter — 45 oder ~50?",
     "Alter bzw. Firmengründungs-Zeitpunkt von Maria ist an mehreren Stellen widersprüchlich.",
     "EMPFEHLUNG: EIN Alter festlegen (z. B. 48), dann Funde 30/56/65/72 einheitlich nachziehen. Deine Wahl beim Alter.",
     [30, 56, 65, 72]),
    ("④ Tag/Nacht-Bruch Kap 30–35 (Kapitel-Header)",
     "Zwischen Kap 30 und 35 springen Header zwischen Tag und Nacht; der Prüfer schlägt eine Header-Umstellung mehrerer Kapitel vor.",
     "EMPFEHLUNG: Zeitleiste Kap 30–35 einmal glattziehen — ich baue dir die konkrete Header-Reihenfolge, du nickst sie ab.",
     [53, 57, 58, 60]),
    ("⑤ Jahres-Ketten (Karriere/Vergangenheit)",
     "Verschiedene Jahresangaben widersprechen sich: Theo 8 vs. 10 Jahre; 12 vs. 15 Jahre; vor 22 Jahren angelegt vs. gelöscht.",
     "EMPFEHLUNG: Pro Kette eine Zahl festlegen; bei Fund 40 (angelegt/gelöscht) ist es eine Plot-Frage — die brauche ich von dir.",
     [7, 18, 40, 54]),
    ("⑥ Plot-Lücken — hier fehlt neuer Text (kein mechanischer Fix)",
     "Diese Funde lösen sich nur durch 1–2 neue erklärende Sätze (Sicherungstausch, Plombierung, Notartermin, Wagen-Doppelung, Koffer-Übergabe, Abriss-Enthüllung).",
     "EMPFEHLUNG: Einzeln entscheiden, ob dir die Lücke wichtig ist. Wo ja, schreibe ich (bzw. Fable) einen passenden Überleitungssatz zur Freigabe.",
     [43, 44, 47, 77, 78, 83, 10]),
    ("⑦ Orte & Distanzen",
     "Ortsbenennungen/Distanzen uneinheitlich: Ottensen vs. Altona, Standort der getarnten Kiste, Bens Tochter-Wohnort, hundert Kilometer, AURORA-Transfer nach Lübeck.",
     "EMPFEHLUNG: Die einfachen (Ottensen/Altona einheitlich, 100 km als Rundung ok) mach ich auf dein Go; Lübeck-Transfer (67/68) ist Plot-Logik → deine Entscheidung.",
     [20, 46, 61, 64, 67, 68, 73]),
    ("⑧ Wie viele Personen im Raum?",
     "Zwei Stellen zählen Anwesende mehrdeutig (zählt sich die Figur selbst mit?).",
     "EMPFEHLUNG: Kurz klären, wer jeweils im Raum ist, dann exakte Zahl setzen.",
     [34, 66]),
    ("⑨ Einzelne Stil-/Formulierungsfragen (niedrige Priorität)",
     "Kleinere Stellen, wo der Fix eine freie Umformulierung wäre (präzise Stundenzahl vs. bewusste Rundung, Monatssetzung, Dialog-Nuancen).",
     "EMPFEHLUNG: Ein Sammel-„lass so“ ist völlig ok — oder ich gehe sie einzeln mit dir durch, wenn du magst.",
     [2, 4, 5, 11, 15, 31, 41, 49, 51, 52, 55, 69, 71]),
]
DONE = [0, 1]; NA = [48]; BUILD = [85]

doc = Document()
st = doc.styles["Normal"]; st.font.name = "Aptos"; st.font.size = Pt(10.5)
st.element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
st.paragraph_format.space_after = Pt(0); st.paragraph_format.line_spacing = 1.12

def para(text="", size=10.5, bold=False, color=None, after=0, before=0, italic=False):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run(text); r.font.name = "Aptos"; r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = RGBColor(*color)
    return p

para("AURORA — Kanon-Entscheidungen", 20, True, (0x1a,0x1a,0x3a), after=2)
para("Die 59 offenen Funde, gebündelt zu Grundsatzfragen · 3. Juli 2026", 11, False, (0x66,0x66,0x66), after=8, italic=True)

para("✅ Schon im Buch (mit Backup, verifiziert):", 12, True, (0x15,0x80,0x3d), after=2)
for t in ["83 Satz-Umbauten (Fable-kuratiert)", "2 glasklare Fixes (Yilmaz, Genus Mia)",
          "Auftakt-Datum + Timeline-Header auf Donnerstag, 1. März 2035",
          "33 weitere eindeutige Konsistenz-Fixes (Fable-geprüft, je 1x verifiziert)"]:
    para("   •  " + t, 10.5, after=1)
para("→ Von den 88 Konsistenz-Funden sind 35 erledigt. Bleiben die 59 unten — die brauchen DEIN Urteil.", 10.5, True, after=8)

para("So funktioniert's:", 11, True, after=2)
para("Ich habe die 59 Funde zu 9 Grundsatz-Clustern gebündelt. Pro Cluster triffst du EINE Entscheidung "
     "(oft reicht „Empfehlung übernehmen“), dann setze ich alle zugehörigen Stellen um. Die Fund-IDs "
     "stehen je Cluster dabei, falls du ins Detail willst.", 10.5, after=8)

for titel, frage, empf, ids in CLUSTERS:
    para(titel, 14, True, (0x1a,0x1a,0x3a), after=2, before=6)
    para("Worum geht's: " + frage, 10.5, after=1)
    para(empf, 10.5, bold=True, color=(0x15,0x60,0x2d), after=2)
    para("Betroffen (%d): " % len(ids) + ", ".join("Kap %s (id%d)" % (SRC[i]["nr"], i) for i in ids if i in SRC),
         9.5, italic=True, color=(0x55,0x55,0x55), after=2)
    for i in ids:
        if i not in SRC: continue
        prob = (SRC[i].get("problem") or "").replace("\n", " ").strip()
        if len(prob) > 135: prob = prob[:135] + "…"
        para("      · Kap %s: %s" % (SRC[i]["nr"], prob), 9.5, color=(0x40,0x40,0x40), after=0)

para("Sonderfälle", 14, True, (0x1a,0x1a,0x3a), after=2, before=8)
para("✔️ Bereits erledigt (März-Shift schon drin): id %s — nichts zu tun." % ", ".join(map(str, DONE)), 10, after=1)
para("➖ Fund trifft nicht zu: id %s (Kap 29, es sind wirklich nur vier anwesend)." % ", ".join(map(str, NA)), 10, after=1)
para("\U0001f6e0️ Build-Frage (kein Einzelfix): id %s — 14+ Kapitel führen ihre Überschrift/Ortsmarke IM Textfeld. "
     "Betrifft die EPUB/Print-Erzeugung, nicht den Inhalt. Klären wir separat beim Rebuild." % ", ".join(map(str, BUILD)), 10, after=1)

out = "/mnt/d/OneDrive/Desktop/AURORA_Kanon_Entscheidungen.docx"
doc.save(out)
covered = sum(len(ids) for *_, ids in CLUSTERS) + len(DONE) + len(NA) + len(BUILD)
print("Gespeichert:", out, "|", os.path.getsize(out), "bytes | abgedeckte Funde:", covered, "/ 59")
