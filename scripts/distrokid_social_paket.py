#!/usr/bin/env python3
"""
DistroKid Social-Media-Paket bauen — STANDARD-AKTION bei jedem neuen Song.

Aufruf:
    python3 distrokid_social_paket.py "Läuft, Sommer"        # Songname wie im Archiv
    python3 distrokid_social_paket.py "Läuft, Sommer" --title "Läuft, Sommer!"  # Anzeige-Titel abweichend

Was es tut (Cover-first, ohne DistroKid-Login):
  1. Findet Cover (.jpg) + MP3 im Archiv  D:/OneDrive/Dokumente/Bolla/Suno_DistroKid/
  2. Legt Zielordner  D:/OneDrive/Desktop/DistroKid/Social Media/<Song>/  an
  3. <Song>_PromoCard.jpg   = Cover (für Instagram; = das Promo-Bild)
  4. <Song>_TikTok-Video.mp4 = Cover 1080x1920 + 15s Song-Ausschnitt MIT TON
  5. <Song>_Instagram-Reel.mp4 = identische Kopie, seit 10.08.2026 dauerhafter Standard (Reels > reine Bildposts bei Reichweite)
  6. <Song>_Insta-Caption.txt + <Song>_TikTok-Caption.txt  (UTF-8 BOM, lebendige Caption via claude -p, Fallback-Template)
  7. ANLEITUNG_<Song>_Posten.docx  (Aptos, kompakt, KI-Label-Warnung fett, inkl. Reel-Schritt)

Chris postet selbst in den Desktop-Apps: TikTok-Video + Instagram-Bild + Instagram-Reel (3 Posts).
TikTok: KI-Label beim Upload AN (Mehr anzeigen -> KI-generierte Inhalte).
Danach Zielordner löschbar.
"""
import os, re, sys, shutil, subprocess, json, urllib.request

ARCHIVE = "/mnt/d/OneDrive/Dokumente/Bolla/Suno_DistroKid"
DEST_BASE = "/mnt/d/OneDrive/Desktop/DistroKid/Social Media"
SPOTIFY_CRED = "/home/bolla/workspace/config/spotify_credentials.json"
ARTIST = "Bollawave"
ARTIST_SLUG = "bollawave"
CLAUDE = os.path.expanduser("~/.local/bin/claude")
MODEL = "claude-sonnet-5"
LQ, RQ = "„", "“"  # deutsche Anführungszeichen (nie ASCII \" in Strings mit Umlaut-Text)

def slugify(name):
    s = name.lower()
    for k, v in {"ä": "ae", "ö": "oe", "ü": "ue", "ß": "ss"}.items():
        s = s.replace(k, v)
    return re.sub(r"[^a-z0-9]+", "-", s).strip("-")

def safe(name):
    return re.sub(r'[<>:"/\\|?*!]+', "", name).strip()

def find_asset(song):
    """Cover + MP3 im Archiv finden (toleranter Match)."""
    base_slug = slugify(song)
    cover = mp3 = None
    for f in os.listdir(ARCHIVE):
        stem, ext = os.path.splitext(f)
        if slugify(stem) == base_slug or slugify(stem).startswith(base_slug):
            if ext.lower() in (".jpg", ".jpeg", ".png"):
                cover = os.path.join(ARCHIVE, f)
            elif ext.lower() == ".mp3":
                mp3 = os.path.join(ARCHIVE, f)
    return cover, mp3

def make_video(cover, mp3, out, start=30, dur=15):
    """„Lebendiges" Musik-Video: geblurter Cover-Hintergrund füllt 1080x1920,
    scharfes Cover mit langsamem Ken-Burns-Zoom davor, Musik-Wellen (showwaves) unten,
    dazu Song-Ausschnitt als Ton. Fallback: stummes Standbild, falls MP3 fehlt."""
    frames = int(dur * 30)
    if mp3 and os.path.isfile(mp3):
        fc = (
            "[0:v]scale=1080:1920:force_original_aspect_ratio=increase,"
            "crop=1080:1920,boxblur=30:2,eq=brightness=-0.12[bg];"
            "[0:v]scale=1600:1600:flags=lanczos[big];"
            f"[big]zoompan=z='min(zoom+0.0007,1.18)':d={frames}:s=1000x1000:fps=30[fg];"
            "[bg][fg]overlay=x=(W-w)/2:y=(H-h)/2-140[comp];"
            "[1:a]showwaves=s=1080x220:mode=cline:rate=30:colors=0xFFFFFF|0xFFD447[wav];"
            "[comp][wav]overlay=x=0:y=H-300:format=auto[vid]"
        )
        cmd = ["ffmpeg", "-y", "-loop", "1", "-i", cover, "-ss", str(start), "-t", str(dur), "-i", mp3,
               "-filter_complex", fc, "-map", "[vid]", "-map", "1:a",
               "-c:v", "libx264", "-pix_fmt", "yuv420p", "-c:a", "aac", "-b:a", "192k",
               "-shortest", "-movflags", "+faststart", out]
        if subprocess.run(cmd, capture_output=True, text=True, timeout=300).returncode == 0:
            return True
    # Fallback: stummes Standbild mit Blur-Rahmen
    vf = ("scale=1080:1920:force_original_aspect_ratio=increase,crop=1080:1920,boxblur=30:2[bg];"
          "[0:v]scale=1000:1000[fg];[bg][fg]overlay=(W-w)/2:(H-h)/2-140")
    cmd2 = ["ffmpeg", "-y", "-loop", "1", "-i", cover, "-t", "12", "-r", "30",
            "-filter_complex", vf, "-c:v", "libx264", "-pix_fmt", "yuv420p",
            "-movflags", "+faststart", out]
    return subprocess.run(cmd2, capture_output=True, text=True, timeout=240).returncode == 0

def llm_caption(title, platform, link):
    plat_txt = "Instagram" if platform == "insta" else "TikTok"
    prompt = (
        f"Schreibe EINE fertige Social-Media-Caption fuer {plat_txt} auf Deutsch fuer den neuen Song "
        f'"{title}" des Kuenstlers Bollawave (Feel-Good-Musik mit Augenzwinkern). '
        "WICHTIG: Es ist NICHT der erste/Debut-Song — Bollawave hat schon mehrere Releases. Nicht als Erstveroeffentlichung framen. "
        "Stil: optimistisch, locker, warmherzig-humorvoll, Emojis, "
        f"{'3-5 knackige' if platform=='tiktok' else '4-6'} Hashtags am Ende (immer #{slugify(title).replace('-','')} passend dabei). "
        f"{'Kurz, 1-2 Saetze, mit Hook am Anfang, Hashtag #fyp.' if platform=='tiktok' else '2-3 Saetze.'} "
        "Keine KI-Floskeln, keine Anfuehrungszeichen um die ganze Caption. "
        "Gib AUSSCHLIESSLICH den Caption-Text aus, ohne Link und ohne Suno-Hinweis (die haenge ich selbst an)."
    )
    try:
        r = subprocess.run([CLAUDE, "-p", "--model", MODEL, prompt],
                           capture_output=True, text=True, timeout=180)
        out = (r.stdout or "").strip()
        if out and "unavailable" not in out.lower():
            return out
    except Exception:
        pass
    # Fallback-Template
    if platform == "tiktok":
        return f'Sommer an, Sorgen aus ☀️🎶 Neuer Song „{title}" ist jetzt ueberall! 🌻\n#{slugify(title).replace("-","")} #FeelGood #NeueMusik #fyp'
    return f'Neuer Song ist da! ☀️🎉 „{title}" — gute Laune zum Mitwippen, ab sofort ueberall streambar 🎶🌻\n#{slugify(title).replace("-","")} #NeuerSong #FeelGood #Bollawave'

def write_caption(path, caption, link, title):
    # Kein DistroKid-Link (in Caption nicht klickbar, Song nicht abspielbar) —
    # stattdessen Streaming-Dienste nennen + nach Künstler suchen lassen.
    body = (
        caption.strip()
        + "\n\n🎧 Überall streamen — Spotify, Apple Music, YouTube Music, Amazon & Co.:\n"
        + "einfach nach " + LQ + "Bollawave" + RQ + " suchen 🔎🎶\n"
        + "🤖 KI-unterstützt produziert (Suno).\n"
    )
    # UTF-8 mit BOM (Notepad-Umlaute)
    with open(path, "w", encoding="utf-8-sig") as f:
        f.write(body)

def make_anleitung(path, title, link):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    FONT = "Aptos"
    def sf(run, size=10.5, bold=False, color=None):
        run.font.name = FONT; run.font.size = Pt(size); run.font.bold = bold
        if color: run.font.color.rgb = RGBColor(*color)
        rPr = run._element.get_or_add_rPr()
        rF = rPr.find(qn('w:rFonts'))
        if rF is None:
            rF = rPr.makeelement(qn('w:rFonts'), {}); rPr.append(rF)
        for a in ('w:ascii','w:hAnsi','w:cs'): rF.set(qn(a), FONT)
    def par(doc, sa=1, sb=0):
        p = doc.add_paragraph(); pf = p.paragraph_format
        pf.space_after = Pt(sa); pf.space_before = Pt(sb); pf.line_spacing = 1.15
        return p
    def head(doc, t, sz=13, col=(0x1F,0x4E,0x79), sb=6):
        p = par(doc, sa=2, sb=sb); sf(p.add_run(t), size=sz, bold=True, color=col)
    def bul(doc, t, pre=""):
        p = par(doc); p.paragraph_format.left_indent = Pt(12)
        if pre: sf(p.add_run(pre), size=10.5, bold=True)
        sf(p.add_run(t), size=10.5)

    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Pt(28); s.left_margin = s.right_margin = Pt(36)
    p = par(doc, sa=1); sf(p.add_run("📣 Posten: „" + title + "“"), size=16, bold=True, color=(0x1F,0x4E,0x79))
    p = par(doc, sa=6); sf(p.add_run("Alle Dateien liegen in diesem Ordner. Danach Ordner löschbar."), size=10, color=(0x66,0x66,0x66))

    head(doc, "🎵 TikTok (Video)", sb=2)
    bul(doc, "TikTok-App/Studio → „+“ / „Hochladen“ → die _TikTok-Video.mp4 wählen.", pre="1.  ")
    bul(doc, "Beschreibung: Inhalt aus _TikTok-Caption.txt einfügen.", pre="2.  ")
    p = par(doc); p.paragraph_format.left_indent = Pt(12)
    sf(p.add_run("3.  ⚠️ „Mehr anzeigen ⌄“ → „KI-generierte Inhalte“ AN — NUR beim Upload möglich, NIE nachträglich!"), size=10.5, bold=True, color=(0xB0,0x30,0x00))
    bul(doc, "Veröffentlichen.", pre="4.  ")

    head(doc, "📸 Instagram (Bild)")
    bul(doc, "Erstellen → Beitrag → _PromoCard.jpg wählen → Weiter.", pre="1.  ")
    bul(doc, "Bildunterschrift: Inhalt aus _Insta-Caption.txt einfügen.", pre="2.  ")
    bul(doc, "Teilen. (KI-Kennzeichnung optional — Suno-Hinweis steht in der Caption.)", pre="3.  ")

    head(doc, "🎬 Instagram (Reel) — zusätzlich, dauerhafter Standard seit 10.08.2026")
    p = par(doc); p.paragraph_format.left_indent = Pt(12)
    sf(p.add_run("Gleiches Video wie TikTok (_Instagram-Reel.mp4, identisch zu _TikTok-Video.mp4) — Reels bekommen bei Instagram deutlich mehr Reichweite als reine Bildposts, kostet aber kein zusätzliches Material."), size=9.5, color=(0x66,0x66,0x66))
    bul(doc, "Erstellen → Reel → _Instagram-Reel.mp4 wählen.", pre="1.  ")
    bul(doc, "Bildunterschrift: Inhalt aus _Insta-Caption.txt einfügen (gleiche wie beim Bildpost).", pre="2.  ")
    bul(doc, "Teilen.", pre="3.  ")

    head(doc, "🔗 Link")
    bul(doc, link)

    p = par(doc, sb=8); sf(p.add_run("Hashtags & Suno-Hinweis stecken schon in den Caption-Dateien. — Bolla 🐾"), size=9, color=(0x88,0x88,0x88))
    doc.save(path)

def main():
    if len(sys.argv) < 2:
        print("Aufruf: distrokid_social_paket.py \"<Songname im Archiv>\" [--title \"Anzeige-Titel\"]"); sys.exit(1)
    song = sys.argv[1]
    title = song
    if "--title" in sys.argv:
        title = sys.argv[sys.argv.index("--title") + 1]

    cover, mp3 = find_asset(song)
    if not cover:
        print(f"❌ Kein Cover im Archiv gefunden für '{song}'. Vorhanden:")
        for f in sorted(os.listdir(ARCHIVE)):
            if f.lower().endswith((".jpg", ".jpeg", ".png")): print("   ", f)
        sys.exit(2)

    slug = slugify(song)
    link = f"https://distrokid.com/hyperfollow/{ARTIST_SLUG}/{slug}"
    dest = os.path.join(DEST_BASE, safe(song))
    os.makedirs(dest, exist_ok=True)
    base = os.path.join(dest, safe(song))

    print(f"🎵 {title}")
    print(f"   Cover: {os.path.basename(cover)}")
    print(f"   MP3:   {os.path.basename(mp3) if mp3 else '— (Video wird stumm)'}")
    print(f"   Link:  {link}")
    print(f"   Ziel:  {dest}")

    # 1) PromoCard (Cover)
    shutil.copy2(cover, base + "_PromoCard.jpg")
    print("   ✓ PromoCard.jpg")

    # 2) TikTok-Video (+ identische Kopie fürs Instagram-Reel, seit 10.08.2026 dauerhafter Standard)
    ok = make_video(cover, mp3, base + "_TikTok-Video.mp4")
    print("   ✓ TikTok-Video.mp4" if ok else "   ⚠️ Video fehlgeschlagen")
    if ok:
        shutil.copy2(base + "_TikTok-Video.mp4", base + "_Instagram-Reel.mp4")
        print("   ✓ Instagram-Reel.mp4 (Kopie)")

    # 3) Captions
    write_caption(base + "_Insta-Caption.txt", llm_caption(title, "insta", link), link, title)
    write_caption(base + "_TikTok-Caption.txt", llm_caption(title, "tiktok", link), link, title)
    print("   ✓ Insta-Caption.txt + TikTok-Caption.txt")

    # 4) Anleitung
    try:
        make_anleitung(os.path.join(dest, f"ANLEITUNG_{safe(song)}_Posten.docx"), title, link)
        print("   ✓ ANLEITUNG.docx")
    except Exception as e:
        print(f"   ⚠️ Anleitung: {e}")

    print(f"\n✅ Fertig. Ordner: {dest}")

if __name__ == "__main__":
    main()
