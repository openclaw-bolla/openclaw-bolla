#!/usr/bin/env python3
"""IT-Meilensteine Timeline – Word-Dokument Generator"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Seitenränder ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.2)
section.right_margin  = Cm(2.2)
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)

# ── Farben pro Ära ────────────────────────────────────────────────────────────
def hex2rgb(h): return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

ERAS = [
    {
        "title": "Pionierzeit",
        "years": "1941 – 1970",
        "hex":   "3B477A",
        "light": "E8EAF6",
        "dot":   "3B477A",
        "icon":  "⚙️",
        "entries": [
            ("1941", "Konrad Zuse – Z3",
             "Der erste funktionsfähige programmierbare Computer der Welt. Zuse baute ihn in seiner Berliner Wohnung – aus 2.600 Telefonrelais. Rechnete in Binär, hatte 64 Wörter Speicher.",
             "Niemand außer Militär & Forschung — Computer füllten ganze Hallen"),
            ("1947", "Transistor",
             "William Shockley, John Bardeen und Walter Brattain (Bell Labs) ersetzen Vakuumröhren durch winzige Halbleiter. Nobelpreis 1956. Alles danach wurde kleiner, schneller, billiger.",
             "Noch unsichtbar für die Gesellschaft — aber Grundlage für alles"),
            ("1954", "IBM 650 – Erster Massencomputer",
             "IBM verkauft über 1.000 Stück — ein Rekord. Preis: ca. 500.000 $ (heute ~5 Mio.). Zielgruppe: Versicherungen, Universitäten, das US-Militär.",
             "Nur Großkonzerne und Behörden — Computer als Prestige-Objekt"),
            ("1964", "IBM System/360 – Mainframe-Revolution",
             "Erstmals eine Computer-Familie mit kompatiblen Modellen. Banken, Airlines und Behörden rechnen damit. Lufthansa bucht Flüge, Banken verwalten Konten. Software wird zum Geschäftsmodell.",
             "Erste Berührung: Kontoauszüge kommen plötzlich vom Computer (~1968)"),
        ]
    },
    {
        "title": "Mikroelektronik & Personal Computer",
        "years": "1971 – 1990",
        "hex":   "1B5E20",
        "light": "E8F5E9",
        "dot":   "1B5E20",
        "icon":  "🖥️",
        "entries": [
            ("1971", "Intel 4004 – Der erste Mikroprozessor",
             "Alle Rechenfunktionen auf einem einzigen Chip. 2.300 Transistoren, so groß wie ein Fingernagel. Rechenleistung wie ein Raumfahrt-Rechner von 1960 — für 200 Dollar.",
             "Noch kein direkter Alltags-Kontakt — aber der Countdown läuft"),
            ("1977", "Apple II – Personal Computer",
             "Steve Wozniak baut den ersten wirklich massentauglichen Heimcomputer. Tastatur, Farbbild, BASIC — man kann sofort loslegen. VisiCalc (1979) macht ihn zur Büromaschine.",
             "Erste Hobby-Nutzer und Schulen — \"haben\" wollten ihn viele, leisten konnten es wenige"),
            ("1981", "IBM PC – Der Standard entsteht",
             "IBM öffnet die Architektur: Andere dürfen Clones bauen. Microsoft liefert MS-DOS. Ein Markt explodiert. Bis 1990 stehen 100 Millionen PCs auf Schreibtischen.",
             "Erste Sekretärinnen, Buchhalter, Ingenieure — der PC kommt ins Büro (~1985)"),
            ("1984", "Apple Macintosh – Grafische Revolution",
             "Maus, Fenster, Icons — kein Tippen von Befehlen mehr. Revolutionär, teuer (2.500 $). Zeigt der Welt, wie Computer sein könnten. Windows folgt 1985.",
             "Desktop Publishing verändert Verlage und Büros — \"benutzen\" lernt jeder"),
        ]
    },
    {
        "title": "Internet & World Wide Web",
        "years": "1991 – 2005",
        "hex":   "E65100",
        "light": "FFF3E0",
        "dot":   "E65100",
        "icon":  "🌐",
        "entries": [
            ("1969 / 1991", "ARPANET → World Wide Web",
             "1969 verbindet das US-Militär 4 Universitäten per Netzwerk. 1991 erfindet Tim Berners-Lee am CERN das Web: Hyperlinks, Browser, HTML. Aus einem Forschungsnetz wird die Infrastruktur der Welt.",
             "1995: Jeder mit Telefonanschluss kann sich einwählen — das Modem-Geräusch prägt eine Generation"),
            ("1993", "Mosaic / Netscape – Browser für alle",
             "Bilder und Text auf einer Seite, klickbare Links. Das Internet wird sichtbar. Innerhalb von 12 Monaten: Millionen Nutzer. Netscape IPO 1995 löst den Dot-com-Boom aus.",
             "~1996: Schulen und Firmen bekommen Internetzugang — E-Mail verdrängt Fax"),
            ("1998", "Google",
             "Larry Page und Sergey Brin starten mit einem Algorithmus der Links bewertet (PageRank). Aus der Informationsflut des Webs wird plötzlich Wissen abrufbar. Das Ende der Enzyklopädie beginnt.",
             "~2001: Wer googelt, findet. Wer nicht googelt, gilt als rückständig"),
            ("1999 / 2003", "Wi-Fi & Breitband",
             "802.11b (Wi-Fi) ermöglicht kabelloses Internet. DSL verdrängt das Modem. Aus Minuten-Downloads werden Sekunden. Streaming, VoIP und Online-Banking werden Alltag.",
             "~2005: Zuhause ohne Internet? Undenkbar für die Jüngeren"),
        ]
    },
    {
        "title": "Mobile Revolution & Cloud",
        "years": "2006 – 2015",
        "hex":   "4A148C",
        "light": "F3E5F5",
        "dot":   "4A148C",
        "icon":  "📱",
        "entries": [
            ("2006 / 2010", "Cloud Computing – AWS & Azure",
             "Amazon startet AWS: Server mieten statt kaufen. Startups brauchen keine eigene IT-Infrastruktur mehr. Microsoft Azure (2010), Google Cloud folgen. Netflix, Spotify, Slack existieren nur dank Cloud.",
             "~2012: Jeder nutzt Cloud — oft unwissentlich (Dropbox, iCloud, Google Drive)"),
            ("2007", "iPhone – Das Smartphone-Zeitalter",
             "Steve Jobs stellt drei Geräte vor: iPod, Telefon, Internet-Gerät — alles in einem. Multi-Touch, kein Stift, kein Handbuch. 2008 kommt Android. 2010 überholt das Smartphone den PC im Absatz.",
             "~2013: Mehr als die Hälfte der Deutschen hat ein Smartphone — immer online, immer erreichbar"),
            ("2008", "App Economy",
             "Apple App Store öffnet. In 3 Tagen: 10 Millionen Downloads. Eine neue Industrie entsteht. Angry Birds, Instagram, WhatsApp — Milliarden-Unternehmen aus einer Idee und einer Nacht Programmieren.",
             "~2015: Der App-Store ersetzt das Telefonbuch, den Stadtplan und den Reiseführer"),
            ("2012", "4G / LTE – Schnelles mobiles Internet",
             "Video-Streaming unterwegs, Echtzeit-Navigation, Video-Calls vom Zug. Die Trennung zwischen Zuhause-Online und Unterwegs-Offline verschwindet endgültig.",
             "~2014: Das Smartphone ist Kamera, Zeitung, Bank und Navigationssystem in einem"),
        ]
    },
    {
        "title": "Künstliche Intelligenz – Der Durchbruch",
        "years": "2012 – 2022",
        "hex":   "006064",
        "light": "E0F7FA",
        "dot":   "006064",
        "icon":  "🤖",
        "entries": [
            ("2012", "Deep Learning – AlexNet",
             "Geoffrey Hinton's Team gewinnt den ImageNet-Wettbewerb mit einem neuronalen Netz. Fehlerrate halbiert — ein Schock für die KI-Welt. Der Durchbruch: Bilderkennung besser als Menschen. Alles danach basiert darauf.",
             "Noch unsichtbar — läuft im Hintergrund (Facebook-Gesichtserkennung, Spam-Filter, Siri)"),
            ("2016", "AlphaGo schlägt den Weltmeister",
             "Google DeepMind's KI besiegt Lee Sedol 4:1 im Go — dem Spiel, das als zu komplex für Computer galt. Ein Wendepunkt: KI kann nicht nur rechnen, sie kann 'denken'. Die Fachwelt ist erschüttert.",
             "Breite Öffentlichkeit ahnt: Hier passiert etwas Historisches — aber noch kein Alltag"),
            ("2017", "Transformer – Die Architektur der KI-Zukunft",
             "Google-Forscher veröffentlichen 'Attention is All You Need'. Die Transformer-Architektur ermöglicht Sprachmodelle in ungeahnter Qualität. GPT, BERT, Claude — alle basieren darauf. Der wichtigste Paper seit Jahrzehnten.",
             "Noch Forschung — aber die Uhr tickt"),
            ("2020", "GPT-3 – Sprache wie ein Mensch",
             "OpenAI's Modell mit 175 Milliarden Parametern schreibt Texte, die von menschlichen kaum zu unterscheiden sind. Erster Schock für Journalisten, Texter, Programmierer. Aber: Nur per API, keine breite Nutzung.",
             "Tech-Industrie begreift die Dimension — Normalnutzer wissen noch nichts davon"),
        ]
    },
    {
        "title": "KI-Assistenten – Die Demokratisierung",
        "years": "2022 – heute",
        "hex":   "B71C1C",
        "light": "FFEBEE",
        "dot":   "B71C1C",
        "icon":  "✨",
        "entries": [
            ("Nov. 2022", "ChatGPT – Der Big Bang",
             "OpenAI macht GPT-3.5 für jeden kostenlos zugänglich. 1 Million Nutzer in 5 Tagen. 100 Millionen in 2 Monaten — schneller als jede Technologie zuvor. Jeder kann plötzlich mit KI reden, schreiben lassen, coden.",
             "JETZT — dieser Moment: Wer ChatGPT noch nicht probiert hat, wird gefragt warum"),
            ("März 2023", "GPT-4 & die Berufsreife",
             "Multimodal (Text + Bilder), besteht den Bar Exam im obersten Percentil, schreibt Code auf Senior-Niveau. Microsoft integriert es in Office, Windows, Bing. KI ist kein Spielzeug mehr — es ist ein Werkzeug.",
             "Bereits jetzt nutzen Millionen täglich Copilot, Gemini, ChatGPT für die Arbeit"),
            ("2023 / 2024", "Claude, Copilot, Gemini – KI überall",
             "Anthropic (Claude), Google (Gemini), Meta (Llama) bringen eigene Modelle. KI ist in jedem Betriebssystem, jeder Office-Suite, jedem Browser. Vom Nischenprodukt zur Infrastruktur — in 18 Monaten.",
             "2025: Wer KI nicht nutzt, hat einen Wettbewerbsnachteil — in fast jedem Beruf"),
            ("2025", "KI-Agenten – Die nächste Stufe",
             "KI führt nicht mehr nur aus — sie plant, entscheidet, handelt. Claude Code steuert Computer. Agenten buchen Meetings, schreiben Code, analysieren Daten, lösen mehrstufige Aufgaben autonom.",
             "Bereits Realität für Early Adopter — Mainstream-Welle: 2026–2028"),
        ]
    },
    {
        "title": "Quantencomputer – Die nächste Revolution",
        "years": "2019 – Vision 2050",
        "hex":   "1A237E",
        "light": "E8EAF6",
        "dot":   "1A237E",
        "icon":  "⚛️",
        "entries": [
            ("2019", "Google Sycamore – Quantum Supremacy",
             "Googles 54-Qubit-Prozessor löst ein Problem in 200 Sekunden — ein klassischer Supercomputer bräuchte 10.000 Jahre. Historischer Moment, aber: Das Problem war künstlich. Praxisrelevanz: noch keine.",
             "Wissenschaft und Rüstung beobachten — Normalnutzer: kein Kontakt"),
            ("2024", "IBM Condor & der 1000-Qubit-Durchbruch",
             "IBM überschreitet die 1.000-Qubit-Marke. Fehlerkorrektur wird besser. Erste echte Anwendungen in Pharma und Logistik-Optimierung. Die Kurve zeigt steil nach oben — analog zur Transistor-Kurve der 1960er.",
             "Erste Unternehmens-Pilotprojekte — noch kein Konsumenten-Kontakt"),
            ("~2030", "VISION: Erste praxisrelevante Quantum-Algorithmen",
             "Pharmaunternehmen simulieren Moleküle für neue Medikamente in Stunden statt Jahrzehnten. Finanzinstitute optimieren Portfolios in Echtzeit. Logistik-Probleme mit Millionen Variablen werden lösbar.",
             "Indirekt betroffen: günstigere Medikamente, bessere Versicherungen — aber kein direkter Kontakt"),
            ("~2035", "VISION: Q-Day – Ende der klassischen Verschlüsselung",
             "Ein ausreichend stabiler Quantencomputer bricht RSA-Verschlüsselung — die Basis von HTTPS, Banking, VPN. Regierungen und Banken bereiten sich bereits vor (Post-Quantum-Kryptographie). Historische Zäsur.",
             "ALLE sind betroffen — Banking, Kommunikation, Datenschutz müssen neu aufgestellt werden"),
            ("~2040–2050", "VISION: Quantencomputer als Standard",
             "Wie heute der klassische PC: Quantencomputing als Cloud-Service für Unternehmen (~2040), später für Privatnutzer (~2050). KI + Quantum: Medikamente werden individuell in Echtzeit synthetisiert. Klimamodelle in Minuten. Materialwissenschaft neu erfunden.",
             "Standard-Technologie — ähnlich wie das Smartphone heute. Jeder profitiert, kaum einer versteht es"),
        ]
    },
]

# ─────────────────────────────────────────────────────────────────────────────
# Hilfsfunktionen
# ─────────────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, spec in kwargs.items():
        border = OxmlElement(f'w:{edge}')
        for attr, val in spec.items():
            border.set(qn(f'w:{attr}'), val)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_para(cell, text, bold=False, size=11, color=None, italic=False, space_before=0, space_after=0, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p

def no_space_para(doc_or_cell, text, bold=False, size=11, color=None, italic=False):
    if hasattr(doc_or_cell, 'add_paragraph'):
        p = doc_or_cell.add_paragraph()
    else:
        p = doc_or_cell.paragraphs[-1] if doc_or_cell.paragraphs else doc_or_cell.add_paragraph()
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p

def remove_table_border(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        tblBorders.append(el)
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)

# ─────────────────────────────────────────────────────────────────────────────
# Titel-Seite / Header
# ─────────────────────────────────────────────────────────────────────────────

# Großer Titel in einer farbigen Box
header_table = doc.add_table(rows=1, cols=1)
header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
remove_table_border(header_table)
hdr_cell = header_table.rows[0].cells[0]
set_cell_bg(hdr_cell, "0D1B3E")
hdr_cell.width = Cm(16.6)

p = hdr_cell.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(18)
p.paragraph_format.space_after = Pt(4)
run = p.add_run("⚡ DIE BESCHLEUNIGUNG DER IT")
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

p2 = hdr_cell.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_before = Pt(0)
p2.paragraph_format.space_after = Pt(6)
run2 = p2.add_run("Von Konrad Zuses Z3 bis zum Quantencomputer — Meilensteine, die die Welt veränderten")
run2.font.size = Pt(12)
run2.font.italic = True
run2.font.color.rgb = RGBColor(0xBB, 0xCC, 0xFF)

p3 = hdr_cell.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_before = Pt(2)
p3.paragraph_format.space_after = Pt(18)
run3 = p3.add_run("1941 ·················· 1970 ·· 1990 · 2005 · 2015 · 2022 · 2025 → 2050")
run3.font.size = Pt(9)
run3.font.color.rgb = RGBColor(0x66, 0x88, 0xCC)

doc.add_paragraph()

# Beschleunigungshinweis
accel_table = doc.add_table(rows=1, cols=1)
remove_table_border(accel_table)
ac = accel_table.rows[0].cells[0]
set_cell_bg(ac, "FFF8E1")
p_ac = ac.add_paragraph()
p_ac.paragraph_format.space_before = Pt(8)
p_ac.paragraph_format.space_after = Pt(8)
p_ac.alignment = WD_ALIGN_PARAGRAPH.LEFT
r1 = p_ac.add_run("💡 Das Tempo der Veränderung: ")
r1.font.bold = True
r1.font.size = Pt(10)
r1.font.color.rgb = RGBColor(0x8B, 0x6A, 0x00)
r2 = p_ac.add_run(
    "Von Zuse (1941) bis zum Internet (1991) vergingen 50 Jahre. "
    "Vom Internet zum Smartphone (2007) nur 16 Jahre. "
    "Vom Smartphone zu ChatGPT (2022) nur 15 Jahre. "
    "Von ChatGPT zu KI-Agenten (2025): 3 Jahre. "
    "Die Kurve zeigt senkrecht nach oben."
)
r2.font.size = Pt(10)
r2.font.color.rgb = RGBColor(0x5D, 0x40, 0x00)

doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# Ären durchgehen
# ─────────────────────────────────────────────────────────────────────────────

for era in ERAS:
    hex_str = era["hex"]
    color   = hex2rgb(hex_str)
    light   = era["light"]
    dot_hex = era["dot"]

    # ── Ära-Header ──
    era_table = doc.add_table(rows=1, cols=1)
    remove_table_border(era_table)
    ec = era_table.rows[0].cells[0]
    set_cell_bg(ec, hex_str)

    p_era = ec.add_paragraph()
    p_era.paragraph_format.space_before = Pt(10)
    p_era.paragraph_format.space_after = Pt(2)
    r_icon = p_era.add_run(f"  {era['icon']}  ")
    r_icon.font.size = Pt(16)
    r_title = p_era.add_run(era["title"].upper())
    r_title.font.bold = True
    r_title.font.size = Pt(15)
    r_title.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    p_years = ec.add_paragraph()
    p_years.paragraph_format.space_before = Pt(0)
    p_years.paragraph_format.space_after = Pt(10)
    r_years = p_years.add_run(f"     {era['years']}")
    r_years.font.size = Pt(10)
    r_years.font.italic = True
    r_years.font.color.rgb = RGBColor(0xDD, 0xEE, 0xFF)

    doc.add_paragraph().paragraph_format.space_after = Pt(0)

    # ── Einträge ──
    for year, title, desc, affected in era["entries"]:
        # Tabelle: [Marker | Inhalt]
        row_table = doc.add_table(rows=1, cols=3)
        remove_table_border(row_table)
        row_table.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Spaltenbreiten
        widths = [Cm(3.0), Cm(0.5), Cm(13.1)]
        for i, cell in enumerate(row_table.rows[0].cells):
            cell.width = widths[i]

        # ── Linke Spalte: Jahr ──
        left = row_table.rows[0].cells[0]
        set_cell_bg(left, light)
        left.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        p_yr = left.add_paragraph()
        p_yr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_yr.paragraph_format.space_before = Pt(10)
        p_yr.paragraph_format.space_after = Pt(2)
        r_yr = p_yr.add_run(year)
        r_yr.font.bold = True
        r_yr.font.size = Pt(14)
        r_yr.font.color.rgb = color

        # ── Mittelspalte: vertikale Linie ──
        mid = row_table.rows[0].cells[1]
        set_cell_bg(mid, dot_hex)

        # ── Rechte Spalte: Inhalt ──
        right = row_table.rows[0].cells[2]
        right.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        p_title = right.add_paragraph()
        p_title.paragraph_format.space_before = Pt(10)
        p_title.paragraph_format.space_after = Pt(3)
        r_t = p_title.add_run(title)
        r_t.font.bold = True
        r_t.font.size = Pt(13)
        r_t.font.color.rgb = color

        p_desc = right.add_paragraph()
        p_desc.paragraph_format.space_before = Pt(0)
        p_desc.paragraph_format.space_after = Pt(5)
        r_d = p_desc.add_run(desc)
        r_d.font.size = Pt(10)
        r_d.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

        # Betroffenheit-Box
        p_aff = right.add_paragraph()
        p_aff.paragraph_format.space_before = Pt(2)
        p_aff.paragraph_format.space_after = Pt(10)
        r_label = p_aff.add_run("👥 Wann betraf es alle:  ")
        r_label.font.bold = True
        r_label.font.size = Pt(9.5)
        r_label.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        r_aff = p_aff.add_run(affected)
        r_aff.font.size = Pt(9.5)
        r_aff.font.italic = True
        r_aff.font.color.rgb = RGBColor(0x33, 0x55, 0x77)

        doc.add_paragraph().paragraph_format.space_after = Pt(1)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ─────────────────────────────────────────────────────────────────────────────
# Schluss-Box: Die Kurve
# ─────────────────────────────────────────────────────────────────────────────
doc.add_paragraph()
end_table = doc.add_table(rows=1, cols=1)
remove_table_border(end_table)
ec2 = end_table.rows[0].cells[0]
set_cell_bg(ec2, "0D1B3E")

p_end1 = ec2.add_paragraph()
p_end1.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_end1.paragraph_format.space_before = Pt(14)
p_end1.paragraph_format.space_after = Pt(6)
r_end1 = p_end1.add_run("🚀  Das Fazit")
r_end1.font.size = Pt(14)
r_end1.font.bold = True
r_end1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

p_end2 = ec2.add_paragraph()
p_end2.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_end2.paragraph_format.space_before = Pt(0)
p_end2.paragraph_format.space_after = Pt(14)
r_end2 = p_end2.add_run(
    "  Die IT-Geschichte ist keine lineare Entwicklung — sie ist eine Exponentialkurve. "
    "Was früher Jahrzehnte brauchte (Mainframe → PC: 30 Jahre), dauert heute Monate. "
    "ChatGPT brauchte 5 Tage bis zur Million Nutzer. Das Radio brauchte 38 Jahre, das Fernsehen 13, "
    "das Internet 4, Pokemon Go 19 Tage.\n\n"
    "  Der Quantencomputer wird nicht 50 Jahre brauchen. Wer die Muster kennt, rechnet mit 15–20 Jahren "
    "bis zur breiten Nutzung. Wer heute 50 ist, wird ihn noch als Alltags-Technologie erleben.\n\n"
    "  Die einzige Konstante: Die Veränderung kommt. Immer früher. Immer tiefer."
)
r_end2.font.size = Pt(10.5)
r_end2.font.color.rgb = RGBColor(0xCC, 0xDD, 0xFF)

# Speichern
out = "/mnt/d/OneDrive/Desktop/IT_Meilensteine_Timeline.docx"
doc.save(out)
print(f"✅ Gespeichert: {out}")
