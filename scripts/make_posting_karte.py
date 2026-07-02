#!/usr/bin/env python3
"""Erzeugt die Posting-Karte (TikTok/Insta + Caption) als kompakte .docx auf den OneDrive-Desktop."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

FONT = "Aptos"
# Deutsche Anführungszeichen als Unicode-Konstanten, damit ASCII-" nie den String schließt
LQ = "„"  # „
RQ = "“"  # "

def set_font(run, size=10.5, bold=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rFonts.set(qn(a), FONT)

def para(doc, space_after=0, space_before=0):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = 1.15
    return p

def heading(doc, text, size=13, color=(0x1F, 0x4E, 0x79), before=6):
    p = para(doc, space_after=2, space_before=before)
    set_font(p.add_run(text), size=size, bold=True, color=color)

def bullet(doc, text, prefix=""):
    p = para(doc, space_after=1)
    p.paragraph_format.left_indent = Pt(12)
    if prefix:
        set_font(p.add_run(prefix), size=10.5, bold=True)
    set_font(p.add_run(text), size=10.5)
    return p

def q(s):
    return LQ + s + RQ

doc = Document()
for sec in doc.sections:
    sec.top_margin = sec.bottom_margin = Pt(28)
    sec.left_margin = sec.right_margin = Pt(36)

# Titel
p = para(doc, space_after=1)
set_font(p.add_run("📣 Posting-Karte — Caption ins Video/Bild"), size=16, bold=True, color=(0x1F, 0x4E, 0x79))
p = para(doc, space_after=6)
set_font(p.add_run("Vom mmp-Clipboard in TikTok & Instagram — Schritt für Schritt"), size=10, color=(0x66, 0x66, 0x66))

# So kommt's zu dir
heading(doc, "So kommt's zu dir", before=2)
p = para(doc, space_after=1); p.paragraph_format.left_indent = Pt(12)
set_font(p.add_run("Caption tippen: in mmp ins Clipboard  "), size=10.5)
set_font(p.add_run("caption: "), size=10.5, bold=True, color=(0x0B, 0x71, 0x2B))
set_font(p.add_run("kurzer Kontext"), size=10.5, bold=True)
set_font(p.add_run("  → nach ~1 Min kommt die fertige 📣-Caption zurück ins Clipboard."), size=10.5)
bullet(doc, "Video/Bild: Bolla legt es ins mmp → grüner Knopf " + q("⬇ Auf Handy speichern") + " → landet in der Galerie.")

# TikTok
heading(doc, "📱 TikTok (Handy)")
bullet(doc, "Unten " + q("+") + " antippen.", prefix="1.  ")
bullet(doc, q("Hochladen") + " → dein Video aus der Galerie wählen.", prefix="2.  ")
bullet(doc, q("Weiter") + " (Schnitt/Effekte überspringen).", prefix="3.  ")
bullet(doc, "Ins große Beschreibungsfeld lange tippen → " + q("Einfügen") + " → Caption ist drin.", prefix="4.  ")
p = para(doc, space_after=1); p.paragraph_format.left_indent = Pt(12)
set_font(p.add_run("5.  ⚠️ Bei KI-Video: " + q("Mehr anzeigen ⌄") + " → " + q("KI-generierte Inhalte") + " AN  "), size=10.5, bold=True, color=(0xB0, 0x30, 0x00))
set_font(p.add_run("(geht NUR beim Upload!)"), size=10.5, bold=True, color=(0xB0, 0x30, 0x00))
bullet(doc, q("Posten") + ". Fertig.", prefix="6.  ")

# Instagram
heading(doc, "📸 Instagram (Handy)")
bullet(doc, "Unten " + q("+") + " → " + q("Beitrag") + ".", prefix="1.  ")
bullet(doc, "Foto/Video aus der Galerie wählen → " + q("Weiter") + " (Filter überspringen).", prefix="2.  ")
bullet(doc, "Feld " + q("Bildunterschrift schreiben…") + " antippen → " + q("Einfügen") + ".", prefix="3.  ")
bullet(doc, q("Teilen") + ". Fertig.", prefix="4.  ")

# PC
heading(doc, "💻 Am PC (kurz)")
bullet(doc, "TikTok/Insta im Browser → " + q("Hochladen") + "/" + q("Erstellen") + " → Datei reinziehen → ins Beschreibungsfeld Strg+V → Posten/Teilen.")

# Fuß
p = para(doc, space_before=8, space_after=0)
set_font(p.add_run("Hashtags stecken schon in der Caption. — Bolla 🐾"), size=9, color=(0x88, 0x88, 0x88))

out = "/mnt/d/OneDrive/Desktop/Posting-Karte_Caption.docx"
doc.save(out)
print("Gespeichert:", out)
