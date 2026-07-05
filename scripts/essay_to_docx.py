import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = "/home/bolla/workspace/data/aurora_essay.md"
DOCX = "/mnt/d/OneDrive/Dokumente/Bolla/AURORA/AURORA_Essay_Chris_Mandel.docx"
TXT = "/mnt/d/OneDrive/Desktop/AURORA_Essay_zum_Posten.txt"

lines = open(SRC, encoding="utf-8").read().split("\n")

FONT = "Aptos"
GLUT = RGBColor(0x8a, 0x4a, 0x6b)

doc = Document()
st = doc.styles["Normal"]
st.font.name = FONT
st.font.size = Pt(11.5)
pf = st.paragraph_format
pf.space_after = Pt(0)
pf.space_before = Pt(0)
pf.line_spacing = 1.15


def style_after(p, pts):
    p.paragraph_format.space_after = Pt(pts)


for ln in lines:
    s = ln.rstrip()
    if not s:
        continue
    if s.startswith("# "):
        p = doc.add_paragraph()
        r = p.add_run(s[2:])
        r.font.name = FONT; r.font.size = Pt(19); r.font.bold = True; r.font.color.rgb = GLUT
        style_after(p, 14)
    elif s.startswith("## "):
        p = doc.add_paragraph()
        r = p.add_run(s[3:])
        r.font.name = FONT; r.font.size = Pt(14); r.font.bold = True; r.font.color.rgb = GLUT
        p.paragraph_format.space_before = Pt(12); style_after(p, 6)
    elif s == "---":
        p = doc.add_paragraph(); style_after(p, 6)
    elif s.startswith("*") and s.endswith("*"):
        p = doc.add_paragraph()
        r = p.add_run(s.strip("*"))
        r.font.name = FONT; r.font.size = Pt(10.5); r.font.italic = True
        r.font.color.rgb = RGBColor(0x6b, 0x64, 0x78)
        p.paragraph_format.space_before = Pt(10)
    else:
        p = doc.add_paragraph(s)
        style_after(p, 8)

doc.save(DOCX)
print("DOCX:", DOCX)

# --- reine Textversion zum Posten (Markdown raus) ---
raw = open(SRC, encoding="utf-8").read()
txt = []
for ln in raw.split("\n"):
    s = ln.rstrip()
    if s == "---":
        txt.append("")
        continue
    s = re.sub(r"^#{1,6}\s*", "", s)      # Ueberschriften-Hashes weg
    s = s.strip("*") if s.startswith("*") and s.endswith("*") else s
    txt.append(s)
out = "\n".join(txt).strip() + "\n"
open(TXT, "w", encoding="utf-8").write(out)
print("TXT :", TXT)
