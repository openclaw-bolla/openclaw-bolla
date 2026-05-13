#!/usr/bin/env python3
"""Flensburg Rundgang — Word-Dokument Generator v3 (kein Routing, korrekte Marker)"""

from staticmap import StaticMap, CircleMarker
from staticmap.staticmap import _lon_to_x, _lat_to_y
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image, ImageDraw, ImageFont

OUTPUT   = "/mnt/d/OneDrive/Desktop/Flensburg_Rundgang.docx"
MAP_FILE = "/tmp/flensburg_map.png"

ZOOM  = 15
IMG_W = 1100
IMG_H = 880

STOPS = [
    {"nr": 1,
     "name": "Parkhaus Süderhofenden",
     "lat": 54.7861, "lon": 9.4365,
     "color": "#6366f1",
     "dist": "Start", "zeit": "",
     "kurz": "Ausgangspunkt im Zentrum",
     "detail": (
         "Zentralstes Parkhaus der Innenstadt, direkt am Süderhofenden. "
         "Normale Stadtparkgebühren, kurze Laufwege zu allen Sehenswürdigkeiten. "
         "Tipp: Ticket beim Einfahren ziehen — am Ende automatenlos zahlen."
     )},

    {"nr": 2,
     "name": "Südermarkt & St. Nikolai",
     "lat": 54.7828, "lon": 9.4363,
     "color": "#f59e0b",
     "dist": "350 m", "zeit": "4 Min",
     "kurz": "Gotische Hallenkirche (14. Jh.), historischer Marktplatz",
     "detail": (
         "Die gotische Nikolaikirche prägt seit dem 14. Jahrhundert die Flensburger Silhouette. "
         "Innen: Flügelaltar und mittelalterliche Ausstattung. "
         "Der Südermarkt davor ist der südliche historische Marktplatz — "
         "Do & Sa Wochenmarkt mit frischen Waren aus Schleswig-Holstein und Dänemark."
     )},

    {"nr": 3,
     "name": "Schifffahrts- & Rummuseum",
     "lat": 54.7931, "lon": 9.4333,
     "color": "#f59e0b",
     "dist": "500 m", "zeit": "7 Min",
     "kurz": "800 Jahre Seefahrt + Rum-Keller (Di–So 10–17 Uhr, 8 €)",
     "detail": (
         "Flensburg war Jahrhunderte lang eine der bedeutendsten Hafenstädte der Ostsee — "
         "und bis heute Deutschlands inoffizielle Rum-Hauptstadt. "
         "Das Museum zeigt prächtige Schiffsmodelle, Navigation, Handelsgeschichte "
         "und natürlich den legendären Flensburger Rum-Keller: "
         "original Fässer, Destillier-Technik und Verkostung. Unbedingt anschauen!"
     )},

    {"nr": 4,
     "name": "Museumshafen",
     "lat": 54.7937, "lon": 9.4342,
     "color": "#f59e0b",
     "dist": "250 m", "zeit": "3 Min",
     "kurz": "Historische Rahsegler & Schoner — Eintritt frei",
     "detail": (
         "Im Museumshafen liegen originalgetreu restaurierte Segelschiffe aus dem 19. Jahrhundert — "
         "Rahsegler, Schoner und Dampfschiffe, teils begehbar. "
         "Eintritt kostenlos. Bei schönem Wetter hervorragend für Fotos. "
         "Im Sommer finden hier regelmäßig Hafenfeste und Regatten statt."
     )},

    {"nr": 5,
     "name": "Hafenspitze — Kaffeepause",
     "lat": 54.7878, "lon": 9.4372,
     "color": "#0ea5e9",
     "dist": "400 m", "zeit": "5 Min",
     "kurz": "Biergarten mit Förde-Panorama bis nach Dänemark",
     "detail": (
         "Die Hafenspitze bietet den schönsten Weitblick der gesamten Tour: "
         "Über die Flensburger Förde bis Glücksburg und bei klarem Wetter bis nach Dänemark. "
         "Hier empfiehlt sich eine Kaffeepause im Biergarten — "
         "frischer Fisch, Krabbenbrötchen und dänische Backwaren. "
         "Der Ort trennt den deutschen und den früher dänischen Teil der Förde."
     )},

    {"nr": 6,
     "name": "Kaufmannshöfe Norderstr.",
     "lat": 54.7937, "lon": 9.4316,
     "color": "#f59e0b",
     "dist": "550 m", "zeit": "7 Min",
     "kurz": "18 historische Innenhöfe (18. Jh.) — Nr. 86 besonders empfohlen",
     "detail": (
         "Die Norderstraße ist eines der besterhaltenen Kaufmannshof-Ensembles Norddeutschlands. "
         "18 Höfe aus dem 17.–18. Jahrhundert, ursprünglich für Lagerung und Handel genutzt. "
         "Norderstraße 86 ist besonders beeindruckend — Durchgang einfach betreten, "
         "die Höfe sind öffentlich zugänglich. Manche beherbergen heute Galerien und Cafés."
     )},

    {"nr": 7,
     "name": "Nordermarkt & Neptunbrunnen",
     "lat": 54.7893, "lon": 9.4323,
     "color": "#f59e0b",
     "dist": "250 m", "zeit": "3 Min",
     "kurz": "Ältester Marktplatz (um 1200) — Neptunbrunnen von 1758",
     "detail": (
         "Der Nordermarkt ist der älteste Marktplatz Flensburgs, "
         "seit dem frühen Mittelalter (um 1200) Zentrum des Handels. "
         "Der Neptunbrunnen stammt aus dem Jahr 1758 und ist ein Wahrzeichen des Platzes. "
         "Di & Fr Wochenmarkt. Umgeben von historischen Bürgerhäusern "
         "aus dem 17.–19. Jahrhundert."
     )},

    {"nr": 8,
     "name": "Nordertor",
     "lat": 54.7955, "lon": 9.4302,
     "color": "#ef4444",
     "dist": "300 m", "zeit": "4 Min",
     "kurz": "Wahrzeichen Flensburgs — gotisches Stadttor von 1595",
     "detail": (
         "Das Nordertor von 1595 ist das bekannteste Wahrzeichen Flensburgs. "
         "Das gotische Backsteintor war einst Teil der mittelalterlichen Stadtmauer "
         "und ist heute einer der meistfotografierten Orte der Stadt. "
         "Der rote Backstein, die gotischen Spitzbögen und die Inschriften "
         "aus dem 17. Jahrhundert machen es zu einem perfekten Fotomotiv."
     )},

    {"nr": 9,
     "name": "Duburg-Ruine (Aussicht)",
     "lat": 54.7925, "lon": 9.4274,
     "color": "#16a34a",
     "dist": "700 m", "zeit": "9 Min",
     "kurz": "Mittelalterliche Burg (1411) — Panorama über Flensburg & Förde",
     "detail": (
         "Die Duburg wurde 1411 als herzogliche Residenz erbaut. "
         "Heute ist sie eine malerische Ruine auf einem Hügel über der Stadt — "
         "mit dem besten Panoramablick über Flensburg, die Förde und das dänische Ufer. "
         "Der Aufstieg dauert ca. 10 Minuten, lohnt sich aber sehr. "
         "Eintritt frei, immer geöffnet. Perfekter Abschluss des Rundgangs."
     )},
]

RETURN_ROW = ("→ 1", "Zurück zum Parkhaus", "1,1 km / ca. 14 Min",
              "Durch die Fußgängerzone zurück — ggf. noch shoppen")
TOTAL_ROW  = ("∑", "Gesamt", "~4,5 km",
              "ca. 56 Min reine Gehzeit · mit Pausen und Museum ca. 3–4 Stunden")


# ── Karte generieren ──────────────────────────────────────────────────────────
def generate_map():
    m = StaticMap(IMG_W, IMG_H,
                  url_template="https://tile.openstreetmap.org/{z}/{x}/{y}.png")

    # Minimale Marker damit staticmap die Bounds kennt
    for s in STOPS:
        m.add_marker(CircleMarker((s["lon"], s["lat"]), "#ffffff", 2))

    # Render (setzt m.x_center, m.y_center, m.zoom intern)
    img = m.render(zoom=ZOOM)
    draw = ImageDraw.Draw(img)

    try:
        font_nr  = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 18)
        font_sm  = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 12)
        font_hdr = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 13)
    except Exception:
        font_nr = font_hdr = font_sm = ImageFont.load_default()

    def hex2rgb(h):
        h = h.lstrip('#')
        return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

    def draw_marker(px, py, text, hex_col):
        r = 17
        rgb = hex2rgb(hex_col)
        draw.ellipse([px-r-2, py-r-2, px+r+2, py+r+2], fill=(10, 10, 10))
        draw.ellipse([px-r,   py-r,   px+r,   py+r],   fill=rgb)
        bb = draw.textbbox((0, 0), text, font=font_nr)
        tw, th = bb[2]-bb[0], bb[3]-bb[1]
        draw.text((px - tw//2 - bb[0], py - th//2 - bb[1]),
                  text, font=font_nr, fill="white")

    # Korrekte Pixel-Koordinaten via staticmap-interne API
    for s in STOPS:
        tile_x = _lon_to_x(s["lon"], m.zoom)
        tile_y = _lat_to_y(s["lat"], m.zoom)
        px = m._x_to_px(tile_x)
        py = m._y_to_px(tile_y)
        draw_marker(px, py, str(s["nr"]), s["color"])

    # Legende
    lx, ly = 8, 8
    lh = 30 + len(STOPS) * 21 + 8
    draw.rectangle([lx, ly, lx+255, ly+lh], fill=(255,255,255,230), outline=(160,160,160))
    draw.text((lx+8, ly+7), "RUNDGANG FLENSBURG", font=font_hdr, fill=(20,20,20))
    for i, s in enumerate(STOPS):
        ry = ly + 30 + i * 21
        rgb = hex2rgb(s["color"])
        r = 9
        draw.ellipse([lx+8, ry+2, lx+8+r*2, ry+2+r*2], fill=rgb)
        nb = draw.textbbox((0,0), str(s["nr"]), font=font_sm)
        nw, nh = nb[2]-nb[0], nb[3]-nb[1]
        draw.text((lx+8+r - nw//2 - nb[0], ry+2+r - nh//2 - nb[1]),
                  str(s["nr"]), font=font_sm, fill="white")
        draw.text((lx+32, ry+3), s["name"][:32], font=font_sm, fill=(20,20,20))

    img.save(MAP_FILE, quality=95)
    print(f"  Karte gespeichert: {MAP_FILE}")


# ── Word-Dokument ──────────────────────────────────────────────────────────────
def set_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_text(cell, text, size=10, bold=False, color=None, italic=False,
                  space_before=2, space_after=2, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.alignment = align
    run = p.add_run(text) if not p.runs else p.runs[0]
    if not p.runs:
        run = p.add_run(text)
    else:
        run = p.runs[0]
        run.text = text
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def add_cell_run(cell, text, size=10, bold=False, color=None, italic=False):
    """Weiteren Run zur ersten Paragraph hinzufügen."""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def generate_doc():
    doc = Document()

    for sec in doc.sections:
        sec.top_margin    = Cm(1.5)
        sec.bottom_margin = Cm(1.5)
        sec.left_margin   = Cm(1.8)
        sec.right_margin  = Cm(1.8)

    # ── Titel ──
    t = doc.add_heading("Flensburg — Stadtspaziergang", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].font.size = Pt(22)
    t.runs[0].font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
    t.paragraph_format.space_before = Pt(0)
    t.paragraph_format.space_after  = Pt(4)

    sub = doc.add_paragraph("Donnerstag  ·  VW ID.3  ·  9 Stopps  ·  ~4,5 km  ·  ca. 3–4 Stunden mit Pausen")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(11)
    sub.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
    sub.paragraph_format.space_after = Pt(6)

    # ── Info-Zeile ──
    p1 = doc.add_paragraph()
    r1 = p1.add_run("⚡  EnBW HPC 300 kW · BAUHAUS Schleswiger Str. 107–109 (2,5 km vor Zentrum)")
    r1.font.size = Pt(10); r1.font.bold = True
    r1.font.color.rgb = RGBColor(0x15, 0x80, 0x3d)
    p1.paragraph_format.space_after = Pt(2)

    p2 = doc.add_paragraph()
    r2 = p2.add_run("🅿  Parkhaus Süderhofenden · zentral · normale Stadtparkgebühren · Startpunkt")
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x44, 0x55, 0x6b)
    p2.paragraph_format.space_after = Pt(8)

    # ── Karte ──
    doc.add_picture(MAP_FILE, width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].paragraph_format.space_after = Pt(8)

    # ── Routen-Tabelle ──
    lbl = doc.add_paragraph("Route & Highlights")
    lbl.runs[0].font.size = Pt(12)
    lbl.runs[0].font.bold = True
    lbl.runs[0].font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
    lbl.paragraph_format.space_after = Pt(4)

    # Spaltenbreiten: # | Station | Weg/Zeit | Highlight
    COL_W = [Cm(0.9), Cm(5.0), Cm(2.8), Cm(8.6)]

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    tblPr = tbl._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl._tbl.insert(0, tblPr)
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)

    hdr = tbl.rows[0].cells
    for i, (txt, w) in enumerate(zip(["#", "Station", "Weg / Zeit", "Highlight"], COL_W)):
        hdr[i].width = w
        set_cell_text(hdr[i], txt, size=10, bold=True,
                      color=RGBColor(0xff, 0xff, 0xff), align=WD_ALIGN_PARAGRAPH.CENTER)
        set_bg(hdr[i], "1e40af")

    for i, s in enumerate(STOPS):
        dist_txt = s["dist"] if s["dist"] == "Start" else f"{s['dist']} / {s['zeit']}"
        row = tbl.add_row().cells
        for j, w in enumerate(COL_W):
            row[j].width = w
        set_cell_text(row[0], str(s["nr"]), size=10, bold=True,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row[1], s["name"], size=10, bold=True)
        set_cell_text(row[2], dist_txt, size=9,
                      color=RGBColor(0x44, 0x55, 0x6b), align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row[3], s["kurz"], size=9)
        if i % 2 == 1:
            for cell in row:
                set_bg(cell, "eef2ff")

    # Rückweg
    row = tbl.add_row().cells
    for j, w in enumerate(COL_W):
        row[j].width = w
    set_cell_text(row[0], "→1", size=9, color=RGBColor(0x64, 0x74, 0x8b),
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row[1], "Zurück zum Parkhaus", size=9)
    set_cell_text(row[2], "1,1 km / ~14 Min", size=9,
                  color=RGBColor(0x64, 0x74, 0x8b), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row[3], "Durch die Fußgängerzone zurück — noch shoppen?", size=9)

    # Gesamt
    row = tbl.add_row().cells
    for j, w in enumerate(COL_W):
        row[j].width = w
    set_cell_text(row[0], "∑", size=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row[1], "Gesamt", size=10, bold=True)
    set_cell_text(row[2], "~4,5 km", size=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row[3], "ca. 56 Min Gehzeit · mit Pausen & Museum ca. 3–4 Stunden", size=9, bold=True)
    for cell in row:
        set_bg(cell, "dbeafe")

    # ── Seite 2: Ausführliche Beschreibungen ──────────────────────────────────
    doc.add_page_break()

    lbl2 = doc.add_paragraph("Stationen im Detail")
    lbl2.runs[0].font.size = Pt(16)
    lbl2.runs[0].font.bold = True
    lbl2.runs[0].font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
    lbl2.paragraph_format.space_after = Pt(10)

    COL_W2 = [Cm(0.9), Cm(4.6), Cm(11.8)]
    tbl2 = doc.add_table(rows=1, cols=3)
    tbl2.style = 'Table Grid'
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER

    tblPr2 = tbl2._tbl.find(qn('w:tblPr'))
    if tblPr2 is None:
        tblPr2 = OxmlElement('w:tblPr')
        tbl2._tbl.insert(0, tblPr2)
    tblLayout2 = OxmlElement('w:tblLayout')
    tblLayout2.set(qn('w:type'), 'fixed')
    tblPr2.append(tblLayout2)

    hdr2 = tbl2.rows[0].cells
    for i, (txt, w) in enumerate(zip(["#", "Station", "Beschreibung & Tipps"], COL_W2)):
        hdr2[i].width = w
        set_cell_text(hdr2[i], txt, size=10, bold=True,
                      color=RGBColor(0xff, 0xff, 0xff), align=WD_ALIGN_PARAGRAPH.CENTER)
        set_bg(hdr2[i], "1e40af")

    for i, s in enumerate(STOPS):
        row = tbl2.add_row().cells
        for j, w in enumerate(COL_W2):
            row[j].width = w

        # Nummer (farbige Kugel simulieren mit farbigem Text)
        set_cell_text(row[0], str(s["nr"]), size=11, bold=True,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

        # Station + Weg
        p = row[1].paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(2)
        name_run = p.add_run(s["name"])
        name_run.font.size = Pt(10)
        name_run.font.bold = True
        if s["dist"] and s["dist"] != "Start":
            dist_run = p.add_run(f"\n{s['dist']} / {s['zeit']}")
            dist_run.font.size = Pt(8.5)
            dist_run.font.italic = True
            dist_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

        # Beschreibung
        dp = row[2].paragraphs[0]
        dp.paragraph_format.space_before = Pt(3)
        dp.paragraph_format.space_after  = Pt(3)
        dr = dp.add_run(s["detail"])
        dr.font.size = Pt(10)

        if i % 2 == 1:
            for cell in row:
                set_bg(cell, "eef2ff")

    # Footer
    foot = doc.add_paragraph("🐾  Bolla · Flensburg Ausflug · Viel Spaß und gutes Wetter!")
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    foot.runs[0].font.size = Pt(9)
    foot.runs[0].font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)
    foot.paragraph_format.space_before = Pt(10)

    doc.save(OUTPUT)
    print(f"  Word gespeichert: {OUTPUT}")


if __name__ == "__main__":
    print("Generiere Karte (ohne Routenlinie)...")
    generate_map()
    print("Erstelle Word-Dokument (2 Seiten)...")
    generate_doc()
    print("Fertig!")
