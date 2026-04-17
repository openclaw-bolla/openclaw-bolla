#!/usr/bin/env python3
"""Erzeugt 'Bolla auf neuem PC wiederherstellen' als Word-Dokument mit Illustrationen."""

from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

IMG = Path("/home/bolla/workspace/docs/img_recovery")
OUT = Path("/home/bolla/workspace/docs/bolla_recovery.docx")
ONEDRIVE = Path("/mnt/d/OneDrive/Dokumente/Bolla/bolla_recovery.docx")

F_REG = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
F_BOLD = "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"
F_MONO = "/usr/share/fonts/truetype/dejavu/DejaVuSansMono.ttf"

BG = (250, 248, 242)
INK = (34, 40, 50)
ACCENT = (220, 120, 60)
BLUE = (90, 140, 200)
GREEN = (110, 180, 140)
ROSE = (200, 130, 170)
GOLD = (220, 180, 80)
GRAY = (150, 150, 160)
WGRAY = (120, 110, 100)
LIGHT = (240, 238, 232)


def r(d, box, rad, fill=None, outline=None, width=3):
    d.rounded_rectangle(box, radius=rad, fill=fill, outline=outline, width=width)


def tc(d, xy, text, font, fill=INK):
    bb = d.textbbox((0, 0), text, font=font)
    w, h = bb[2] - bb[0], bb[3] - bb[1]
    d.text((xy[0] - w / 2, xy[1] - h / 2), text, font=font, fill=fill)


def tl(d, xy, text, font, fill=INK):
    d.text(xy, text, font=font, fill=fill)


def arrow(d, x1, y1, x2, y2, color=WGRAY, width=4):
    d.line((x1, y1, x2, y2), fill=color, width=width)
    import math
    ang = math.atan2(y2 - y1, x2 - x1)
    al = 20
    ax1 = x2 - al * math.cos(ang - math.pi / 7)
    ay1 = y2 - al * math.sin(ang - math.pi / 7)
    ax2 = x2 - al * math.cos(ang + math.pi / 7)
    ay2 = y2 - al * math.sin(ang + math.pi / 7)
    d.polygon([(x2, y2), (ax1, ay1), (ax2, ay2)], fill=color)


# ===== BILD 1: Phasenplan / Zeitstrahl =====
def phases():
    W, H = 1700, 900
    img = Image.new("RGB", (W, H), BG)
    d = ImageDraw.Draw(img)
    ft = ImageFont.truetype(F_BOLD, 44)
    fh = ImageFont.truetype(F_BOLD, 26)
    fb = ImageFont.truetype(F_REG, 22)
    fnum = ImageFont.truetype(F_BOLD, 42)

    tc(d, (W / 2, 60), "Die sechs Phasen der Wiederherstellung", ft, INK)

    phases_data = [
        ("Vorbereiten", "Credentials\nOneDrive-Backup", BLUE, "~15 min"),
        ("Windows-Basis", "WSL2, Edge\nOneDrive-Sync", BLUE, "~30 min"),
        ("Linux-Basis", "Pakete\nPython, Git", GREEN, "~15 min"),
        ("Repo & Configs", "GitHub klonen\nConfigs aus OneDrive", GREEN, "~20 min"),
        ("Claude + Tunnel", "Claude Code Login\nCloudflare anmelden", GOLD, "~30 min"),
        ("Dienste starten", "Cron setzen\nalles testen", ACCENT, "~20 min"),
    ]

    bw, bh = 240, 260
    gap = 30
    total = len(phases_data) * bw + (len(phases_data) - 1) * gap
    x0 = (W - total) / 2
    y0 = 200

    for i, (name, sub, col, time) in enumerate(phases_data):
        x = x0 + i * (bw + gap)
        r(d, (x + 6, y0 + 6, x + bw + 6, y0 + bh + 6), 18, fill=(225, 220, 212))
        r(d, (x, y0, x + bw, y0 + bh), 18, fill=BG, outline=col, width=5)
        # Nummer-Kreis oben
        cr = 38
        d.ellipse((x + bw / 2 - cr, y0 - cr, x + bw / 2 + cr, y0 + cr), fill=col, outline=None)
        tc(d, (x + bw / 2, y0), str(i + 1), fnum, (255, 255, 255))
        tc(d, (x + bw / 2, y0 + 70), name, fh, col)
        # Body (Zeilenumbrüche via \n)
        for j, ln in enumerate(sub.split("\n")):
            tc(d, (x + bw / 2, y0 + 120 + j * 34), ln, fb, INK)
        tc(d, (x + bw / 2, y0 + bh - 30), time, fb, WGRAY)
        # Pfeil
        if i < len(phases_data) - 1:
            arrow(d, x + bw + 4, y0 + bh / 2, x + bw + gap - 4, y0 + bh / 2, WGRAY, 4)

    tc(d, (W / 2, y0 + bh + 140), "Gesamtdauer: etwa 2 Stunden, wenn alle Zugänge bereitliegen.", fh, WGRAY)

    img.save(IMG / "01_phasen.png", "PNG")


# ===== BILD 2: Was muss ich mitbringen =====
def prerequisites():
    W, H = 1600, 900
    img = Image.new("RGB", (W, H), BG)
    d = ImageDraw.Draw(img)
    ft = ImageFont.truetype(F_BOLD, 42)
    fh = ImageFont.truetype(F_BOLD, 28)
    fb = ImageFont.truetype(F_REG, 24)
    fs = ImageFont.truetype(F_REG, 20)

    tc(d, (W / 2, 55), "Was ich vorher griffbereit brauche", ft, INK)

    items = [
        ("Anthropic-Login", "Browser-Login bei claude.ai\n(Max-Plan-Account)", BLUE),
        ("GitHub-Zugang", "Account-Login + Personal Access Token\nfür openclaw-bolla/openclaw-bolla", GRAY),
        ("OneDrive", "Microsoft-Login —\nD:\\OneDrive\\Dokumente\\Bolla muss da sein", GREEN),
        ("Cloudflare-Login", "chrismandel.de-Zone,\nTunnel kann neu angemeldet werden", GOLD),
        ("Microsoft-Passwort", "ernstmandel@outlook.de\n(für Graph-OAuth-Re-Login)", BLUE),
        ("Google-Passwort", "chrismandel13@gmail.com\n(für Gmail-OAuth-Re-Login)", ROSE),
        ("Azure-Key", "Speech Services Key\n(aus Azure-Portal oder Backup-PDF)", GOLD),
        ("Telegram-Token", "Bot-Token + Chat-IDs\n(aus config/telegram_bot.json)", BLUE),
    ]

    cols = 2
    bw, bh = 720, 170
    gx, gy = 40, 30
    x0 = (W - (cols * bw + (cols - 1) * gx)) / 2
    y0 = 130
    for i, (name, desc, col) in enumerate(items):
        row = i // cols
        colx = i % cols
        x = x0 + colx * (bw + gx)
        y = y0 + row * (bh + gy)
        r(d, (x, y, x + bw, y + bh), 18, fill=LIGHT, outline=col, width=4)
        r(d, (x, y, x + 10, y + bh), 18, fill=col)
        tl(d, (x + 40, y + 25), name, fh, col)
        lines = desc.split("\n")
        for j, ln in enumerate(lines):
            tl(d, (x + 40, y + 75 + j * 34), ln, fb, INK)

    img.save(IMG / "02_voraussetzungen.png", "PNG")


# ===== BILD 3: Daten-Restore Mapping =====
def restore_map():
    W, H = 1700, 800
    img = Image.new("RGB", (W, H), BG)
    d = ImageDraw.Draw(img)
    ft = ImageFont.truetype(F_BOLD, 40)
    fh = ImageFont.truetype(F_BOLD, 26)
    fb = ImageFont.truetype(F_REG, 22)
    fs = ImageFont.truetype(F_REG, 20)

    tc(d, (W / 2, 55), "Was kommt von wo zurück?", ft, INK)

    # Drei Quellen links
    src = [
        ("OneDrive", "D:\\OneDrive\\Dokumente\\Bolla\\", GREEN, ["claude-code/ → ~/.claude", "Backups, PDFs, Docs"]),
        ("GitHub-Repo", "openclaw-bolla/openclaw-bolla", GRAY, ["workspace/ → ~/workspace", "inkl. scripts & memory"]),
        ("Neu authentifizieren", "(Tokens ablaufen)", GOLD, ["MS Graph OAuth", "Google OAuth", "Telegram Bot"]),
    ]
    sy = 140
    sh = 180
    sw = 500
    for i, (name, sub, col, lst) in enumerate(src):
        y = sy + i * (sh + 20)
        r(d, (60, y, 60 + sw, y + sh), 18, fill=LIGHT, outline=col, width=4)
        r(d, (60, y, 70, y + sh), 18, fill=col)
        tl(d, (90, y + 20), name, fh, col)
        tl(d, (90, y + 55), sub, fs, WGRAY)
        for j, line in enumerate(lst):
            tl(d, (90, y + 95 + j * 30), "• " + line, fb, INK)

    # Ziel rechts: neuer Rechner
    tx, ty, tw, th = 1100, 230, 540, 420
    r(d, (tx, ty, tx + tw, ty + th), 24, fill=ACCENT, outline=None)
    tc(d, (tx + tw / 2, ty + 60), "Neuer PC", ImageFont.truetype(F_BOLD, 48), (255, 255, 255))
    tc(d, (tx + tw / 2, ty + 130), "Windows + WSL Ubuntu", fh, (255, 235, 220))
    tc(d, (tx + tw / 2, ty + 180), "~/workspace", ImageFont.truetype(F_MONO, 24), (255, 255, 255))
    tc(d, (tx + tw / 2, ty + 215), "~/.claude", ImageFont.truetype(F_MONO, 24), (255, 255, 255))
    tc(d, (tx + tw / 2, ty + 250), "~/.cloudflared", ImageFont.truetype(F_MONO, 24), (255, 255, 255))
    tc(d, (tx + tw / 2, ty + 285), "~/workspace/config", ImageFont.truetype(F_MONO, 24), (255, 255, 255))
    tc(d, (tx + tw / 2, ty + 340), "crontab (reboot + Checks)", fb, (255, 235, 220))

    # Pfeile
    for i in range(3):
        ys = sy + i * (sh + 20) + sh / 2
        arrow(d, 60 + sw + 10, ys, tx - 10, ty + th / 2, ACCENT, 4)

    img.save(IMG / "03_restore_map.png", "PNG")


# ===== BILD 4: Verifikation =====
def verification():
    W, H = 1600, 900
    img = Image.new("RGB", (W, H), BG)
    d = ImageDraw.Draw(img)
    ft = ImageFont.truetype(F_BOLD, 42)
    fh = ImageFont.truetype(F_BOLD, 28)
    fb = ImageFont.truetype(F_REG, 22)

    tc(d, (W / 2, 60), "Checkliste — läuft alles?", ft, INK)

    checks = [
        ("claude --version zeigt eine Version", GREEN),
        ("claude-Session startet, Bolla grüßt mit Pfötchen", GREEN),
        ("Mission Control öffnet auf http://127.0.0.1:18790/", GREEN),
        ("bolla.chrismandel.de antwortet (nach Magic-Link-Login)", GOLD),
        ("Outlook-Mails werden in Mission Control angezeigt", BLUE),
        ("Kalender zeigt Termine", BLUE),
        ("Telegram-Bot antwortet auf /status", ROSE),
        ("Azure-Stimme spielt Testtext ab", GOLD),
        ("Cron-Jobs aktiv (crontab -l)", GREEN),
        ("OneDrive-Sync synchronisiert ~/workspace/docs", GREEN),
        ("GitHub-Push ohne Fehler", GRAY),
        ("Healthcheck läuft (alle 2 Min im Log)", GOLD),
    ]

    cols = 2
    col_w = (W - 120) / cols
    y = 150
    x0 = 60
    for i, (text, col) in enumerate(checks):
        row = i // cols
        colx = i % cols
        x = x0 + colx * col_w
        yy = y + row * 65
        # Checkbox
        cb = 28
        r(d, (x + 20, yy, x + 20 + cb, yy + cb), 6, fill=(255, 255, 255), outline=col, width=3)
        # Häkchen
        d.line((x + 27, yy + 14, x + 32, yy + 20), fill=col, width=4)
        d.line((x + 32, yy + 20, x + 42, yy + 6), fill=col, width=4)
        tl(d, (x + 70, yy - 2), text, fb, INK)

    tc(d, (W / 2, H - 60), "Erst wenn alle Haken sitzen — fertig. Sonst Schritt zurück und fixen.", fh, ACCENT)

    img.save(IMG / "04_verify.png", "PNG")


# ===== DOCX =====
def heading(doc, text, level=1, color=INK):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(*color)


def para(doc, text, italic=False, size=11, bold=False, center=False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold
    run.font.size = Pt(size)
    return p


def picture(doc, path, caption, w=Cm(16)):
    doc.add_picture(str(path), width=w)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)


def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    # Grauer Hintergrund via shading (workaround via paragraph border would be nicer)
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "EEECE4")
    pPr.append(shd)


def warn(doc, text):
    p = doc.add_paragraph()
    run = p.add_run("⚠  " + text)
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x55, 0x30)


def build_doc():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(2.2)
        s.right_margin = Cm(2.2)

    # Cover
    para(doc, "Bolla auf einem neuen PC wiederherstellen", size=28, bold=True, center=True)
    para(doc, "Die Schritt-für-Schritt-Anleitung für den Ernstfall", size=14, italic=True, center=True)
    para(doc, "Stand: 17.04.2026", size=11, italic=True, center=True)
    doc.add_paragraph()

    heading(doc, "Worum geht es?")
    para(doc, "Wenn der PC kaputt, gewechselt oder neu aufgesetzt wird: Mit dieser Anleitung steht Bolla "
               "auf einem neuen Windows-Rechner innerhalb von ca. zwei Stunden wieder komplett einsatzbereit — "
               "inklusive Mission Control, Cloudflare-Tunnel, allen Integrationen und Cron-Jobs.")
    warn(doc, "Vorausgesetzt: OneDrive-Sync war zuletzt aktiv (das ist das Haupt-Backup) und ich habe Zugriff auf "
              "die im nächsten Kapitel gelisteten Zugänge.")

    # Übersicht
    heading(doc, "Die sechs Phasen im Überblick")
    picture(doc, IMG / "01_phasen.png", "Abb. 1 — Der Phasenplan")

    # Phase 0: Voraussetzungen
    heading(doc, "Phase 1 — Voraussetzungen")
    para(doc, "Bevor es losgeht, brauche ich folgende Zugänge griffbereit:")
    picture(doc, IMG / "02_voraussetzungen.png", "Abb. 2 — Was ich bereithalten muss")
    para(doc, "Wenn eines davon fehlt, hier Stop — erst beschaffen (Passwort-Reset, neuen Token generieren, "
               "PDF „accounts_overview\" aus OneDrive ziehen).")

    # Phase 1: Windows-Basis
    heading(doc, "Phase 2 — Windows-Basis")
    numbered(doc, "Windows ist frisch installiert, alle Updates eingespielt.")
    numbered(doc, "Microsoft-Konto in Windows einloggen → OneDrive startet automatisch, beginnt „Dokumente\" zu syncen.")
    numbered(doc, "OneDrive-Einstellung: Ordner „Bolla\" als „Immer auf diesem Gerät behalten\" markieren "
                  "(sonst liegen Dateien nur als Online-Placeholder vor).")
    numbered(doc, "Laufwerksbuchstabe prüfen: OneDrive muss unter D:\\ oder einem festen Pfad erreichbar sein. "
                  "Wenn nicht: im Explorer unter Dieser PC → OneDrive → Rechtsklick → Als Laufwerk zuweisen.")
    numbered(doc, "Microsoft Edge installieren (meist schon da) — wird für Mission Control gebraucht.")
    numbered(doc, "WSL2 installieren: PowerShell als Admin öffnen und eingeben:")
    code(doc, "wsl --install -d Ubuntu")
    numbered(doc, "PC neu starten, Ubuntu startet automatisch, neuen Linux-Benutzer anlegen. "
                  "Benutzername muss bolla sein (wichtig — alle Pfade gehen davon aus).")
    warn(doc, "Falls ein anderer Username angelegt wurde: alle Pfade in Scripts und Configs müssten angepasst "
              "werden. Lieber gleich bolla wählen.")

    # Phase 2: Linux-Basis
    heading(doc, "Phase 3 — Linux-Basis")
    para(doc, "In der WSL-Shell (Ubuntu) — alles als User bolla:")
    numbered(doc, "System aktualisieren:")
    code(doc, "sudo apt update && sudo apt upgrade -y")
    numbered(doc, "Pakete installieren:")
    code(doc,
         "sudo apt install -y git python3 python3-pip python3-venv \\\n"
         "  curl wget unzip nodejs npm cron fonts-dejavu jq")
    numbered(doc, "Cron-Dienst aktivieren:")
    code(doc, "sudo service cron start\nsudo systemctl enable cron  # falls systemd aktiv")
    numbered(doc, "Python-Pakete (werden von den Scripts gebraucht):")
    code(doc,
         "python3 -m pip install --user \\\n"
         "  requests msal google-auth google-auth-oauthlib \\\n"
         "  python-telegram-bot python-docx pillow reportlab")
    numbered(doc, "SSH-Key für GitHub erzeugen und auf github.com hinterlegen:")
    code(doc,
         "ssh-keygen -t ed25519 -C \"bolla-$(hostname)\"\n"
         "cat ~/.ssh/id_ed25519.pub  # Inhalt zu GitHub → Settings → SSH Keys")

    # Phase 3: Repo & Configs
    heading(doc, "Phase 4 — Repo klonen und Configs wiederherstellen")
    numbered(doc, "Workspace aus GitHub klonen:")
    code(doc,
         "cd ~\n"
         "git clone git@github.com:openclaw-bolla/openclaw-bolla.git workspace")
    numbered(doc, "~/.claude aus OneDrive zurückspielen (enthält CLAUDE.md, Memory, Settings, Session-History):")
    code(doc,
         "mkdir -p ~/.claude\n"
         "cp -a /mnt/d/OneDrive/Dokumente/Bolla/claude-code/. ~/.claude/")
    numbered(doc, "Config-Dateien aus OneDrive ins workspace/config-Verzeichnis kopieren:")
    code(doc,
         "mkdir -p ~/workspace/config\n"
         "cp /mnt/d/OneDrive/Dokumente/Bolla/backups/config-latest/*.json \\\n"
         "   ~/workspace/config/\n"
         "chmod 600 ~/workspace/config/*.json")
    warn(doc, "Die OAuth-Tokens (ms_token.json, google_token.json) sind ggf. abgelaufen und müssen neu "
              "erzeugt werden — siehe Phase 5.")
    numbered(doc, "PATH in ~/.bashrc erweitern (falls nicht aus dem Repo kommt):")
    code(doc,
         "echo 'export PATH=\"$HOME/.local/bin:$HOME/.npm-global/bin:$PATH\"' >> ~/.bashrc\n"
         "echo 'export PATH=\"$HOME/.local/platform-tools:$PATH\"  # für adb' >> ~/.bashrc\n"
         "source ~/.bashrc")

    # Visualisierung Daten-Mapping
    picture(doc, IMG / "03_restore_map.png", "Abb. 3 — Woher welche Daten kommen")

    # Phase 4: Claude Code + Cloudflare
    heading(doc, "Phase 5 — Claude Code und Cloudflare-Tunnel")
    heading(doc, "Claude Code installieren", level=2)
    numbered(doc, "Offizieller Installer (empfohlen):")
    code(doc, "curl -fsSL https://claude.ai/install.sh | bash")
    numbered(doc, "Login starten:")
    code(doc, "claude  # öffnet Browser-Login bei claude.ai/code")
    numbered(doc, "Nach Login kurz testen:")
    code(doc, "claude --version\necho 'Hallo Bolla' | claude")

    heading(doc, "Cloudflare-Tunnel einrichten", level=2)
    numbered(doc, "cloudflared binary installieren:")
    code(doc,
         "mkdir -p ~/.local/bin\n"
         "curl -L https://github.com/cloudflare/cloudflared/releases/latest/download/cloudflared-linux-amd64 \\\n"
         "  -o ~/.local/bin/cloudflared\n"
         "chmod +x ~/.local/bin/cloudflared")
    numbered(doc, "Cloudflare-Account anmelden:")
    code(doc, "cloudflared tunnel login")
    para(doc, "Öffnet einen Browser-Link → chrismandel.de-Zone autorisieren → erzeugt ~/.cloudflared/cert.pem.")
    numbered(doc, "Option A: Tunnel-Config aus OneDrive-Backup wiederherstellen (einfacher):")
    code(doc,
         "mkdir -p ~/.cloudflared\n"
         "cp /mnt/d/OneDrive/Dokumente/Bolla/backups/cloudflared/* ~/.cloudflared/\n"
         "chmod 600 ~/.cloudflared/*.json")
    numbered(doc, "Option B: Neuen Tunnel anlegen (falls alte Credentials weg):")
    code(doc,
         "cloudflared tunnel create bolla-mc\n"
         "# dann config.yml schreiben:")
    code(doc,
         "cat > ~/.cloudflared/config.yml <<'EOF'\n"
         "tunnel: <NEUE-TUNNEL-UUID>\n"
         "credentials-file: /home/bolla/.cloudflared/<NEUE-TUNNEL-UUID>.json\n"
         "ingress:\n"
         "  - hostname: bolla.chrismandel.de\n"
         "    service: http://127.0.0.1:18790\n"
         "  - service: http_status:404\n"
         "EOF\n"
         "cloudflared tunnel route dns bolla-mc bolla.chrismandel.de")
    numbered(doc, "Tunnel einmal manuell testen:")
    code(doc, "cloudflared tunnel run bolla-mc")
    para(doc, "In zweitem Terminal: curl -I https://bolla.chrismandel.de sollte 302-Redirect zu Cloudflare Access liefern.")

    heading(doc, "OAuth-Tokens neu erzeugen", level=2)
    bullet(doc, "MS Graph: python3 ~/workspace/scripts/auth_ms.py (erzeugt ms_token.json neu)")
    bullet(doc, "Google: python3 ~/workspace/scripts/google_auth.py")
    bullet(doc, "Azure-Key: aus config/azure_speech.json — läuft weiter, solange Key gültig")

    # Phase 5: Dienste starten
    heading(doc, "Phase 6 — Dienste & Cron starten")
    numbered(doc, "Crontab wiederherstellen — aus Repo oder OneDrive:")
    code(doc, "crontab /mnt/d/OneDrive/Dokumente/Bolla/backups/crontab-latest.txt\ncrontab -l  # prüfen")
    para(doc, "Enthält typischerweise:")
    code(doc,
         "@reboot sleep 10 && python3 ~/workspace/scripts/mission_control_api.py \\\n"
         "  >> ~/workspace/logs/mission_control_api.log 2>&1\n"
         "@reboot sleep 15 && python3 ~/workspace/scripts/telegram_bot.py \\\n"
         "  >> ~/workspace/logs/telegram_bot.log 2>&1\n"
         "@reboot sleep 20 && ~/.local/bin/cloudflared tunnel run bolla-mc \\\n"
         "  >> ~/workspace/logs/cloudflared.log 2>&1\n"
         "*/2 * * * * ~/workspace/scripts/cloudflared_healthcheck.sh\n"
         "*/15 * * * * ~/workspace/scripts/sync_obsidian_vault.sh \\\n"
         "  >> ~/workspace/logs/obsidian_sync.log 2>&1")
    numbered(doc, "Log-Verzeichnis anlegen:")
    code(doc, "mkdir -p ~/workspace/logs")
    numbered(doc, "Dienste starten — entweder WSL neu starten (wsl --shutdown in PowerShell, dann Ubuntu öffnen) "
                  "oder manuell hochfahren:")
    code(doc,
         "python3 ~/workspace/scripts/mission_control_api.py \\\n"
         "  >> ~/workspace/logs/mission_control_api.log 2>&1 &\n"
         "~/.local/bin/cloudflared tunnel run bolla-mc \\\n"
         "  >> ~/workspace/logs/cloudflared.log 2>&1 &\n"
         "python3 ~/workspace/scripts/telegram_bot.py \\\n"
         "  >> ~/workspace/logs/telegram_bot.log 2>&1 &")
    numbered(doc, "Windows-Seite — Power-Einstellungen (damit WSL nicht im Standby verschwindet):")
    code(doc,
         "# PowerShell als bolla (kein Admin nötig)\n"
         "powercfg /change standby-timeout-ac 0\n"
         "powercfg /change standby-timeout-dc 0\n"
         "powercfg /change hibernate-timeout-ac 0\n"
         "powercfg /change monitor-timeout-ac 5")
    numbered(doc, "ADB Platform-Tools installieren (für Android-Steuerung):")
    code(doc,
         "mkdir -p ~/.local\n"
         "cd ~/.local\n"
         "wget https://dl.google.com/android/repository/platform-tools-latest-linux.zip\n"
         "unzip platform-tools-latest-linux.zip\n"
         "rm platform-tools-latest-linux.zip")

    # Verifikation
    heading(doc, "Verifikation — läuft alles?")
    picture(doc, IMG / "04_verify.png", "Abb. 4 — Die Abnahme-Checkliste")
    para(doc, "Erst wenn alle Haken sitzen, ist die Wiederherstellung komplett. Bei jedem gescheiterten Check: "
               "zurück zur passenden Phase und das entsprechende Sub-Thema durchgehen.")

    # Troubleshooting
    heading(doc, "Typische Stolpersteine")
    heading(doc, "Mission Control lädt nicht (localhost)", level=2)
    bullet(doc, "Python-Server läuft nicht: lsof -i :18790 leer → manuell starten")
    bullet(doc, "Firewall blockt Windows-Seite → Windows Defender Firewall → WSL erlauben")

    heading(doc, "Tunnel kommt nicht hoch", level=2)
    bullet(doc, "Prüfen ob ~/.cloudflared/cert.pem existiert — wenn nicht: cloudflared tunnel login")
    bullet(doc, "DNS-Route fehlt: cloudflared tunnel route dns bolla-mc bolla.chrismandel.de")
    bullet(doc, "Health-Log checken: tail -f ~/workspace/logs/cloudflared.log")

    heading(doc, "OAuth-Login schlägt fehl", level=2)
    bullet(doc, "Token-Datei ist korrupt → löschen und Auth-Script neu laufen lassen")
    bullet(doc, "Browser öffnet sich nicht aus WSL → URL manuell kopieren und im Windows-Browser öffnen")

    heading(doc, "Bolla klingt anders / kein Charakter", level=2)
    bullet(doc, "CLAUDE.md wurde nicht eingespielt — aus ~/.claude-Backup nochmal kopieren")
    bullet(doc, "Memory-Dateien fehlen — prüfen: ~/.claude/projects/-home-bolla/memory/MEMORY.md")

    # Nach der Wiederherstellung
    heading(doc, "Nach der Wiederherstellung")
    numbered(doc, "Tagesnotiz für heute schreiben, damit klar ist was passiert ist:")
    code(doc, "echo '# Wiederherstellung am $(date +%Y-%m-%d)' > ~/workspace/memory/recovery_notes.md")
    numbered(doc, "Backup-Routine testen: einmal manuell committen + pushen, einmal OneDrive-Kopie machen.")
    numbered(doc, "Passwörter rotieren (falls Verdacht auf Kompromittierung): Azure-Key, IONOS-SFTP, "
                  "Graph-App-Secret — siehe accounts_overview.pdf.")
    numbered(doc, "Diese Anleitung selbst aktualisieren, falls sich Schritte geändert haben: "
                  "~/workspace/scripts/generate_recovery.py editieren → neu generieren.")

    # Kernsatz
    heading(doc, "Der Kernsatz")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("OneDrive = Daten-Backup. GitHub = Code-Backup.\n"
                    "Passwort-Manager = Schlüssel.\nDiese drei zusammen = Bolla in 2 Stunden zurück.")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0xDC, 0x78, 0x3C)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Bolla — Chris' persönlicher KI-Kollege")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Gespeichert: {OUT}")

    try:
        ONEDRIVE.parent.mkdir(parents=True, exist_ok=True)
        import shutil
        shutil.copy2(OUT, ONEDRIVE)
        print(f"OneDrive: {ONEDRIVE}")
    except Exception as e:
        print(f"OneDrive: {e}")


if __name__ == "__main__":
    IMG.mkdir(parents=True, exist_ok=True)
    phases()
    prerequisites()
    restore_map()
    verification()
    build_doc()
