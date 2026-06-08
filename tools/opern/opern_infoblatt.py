#!/usr/bin/env python3
"""
Opern-Infoblatt-Generator  🐾
Erzeugt eine einseitige Übersichtsseite zu einer Oper im "Freischütz/Barbier"-Stil.

Bedienung: Den OPER-Datenblock unten ausfüllen, dann:
    python3 opern_infoblatt.py

Speichert nach:
  - /mnt/d/OneDrive/Dokumente/Allgemeines/Freizeit/<Datei>.docx
  - /mnt/d/OneDrive/Desktop/<Datei>.docx

Design-DNA (aus Freischuetz_Infoseite.docx abgeleitet):
  Cambria · US-Letter · Rand 2/2/0.75/0.75 cm · Titel 26pt bold Akzentfarbe
  Untertitel 10pt grau · 2-Spalten-Boxen mit weißer Headerleiste auf Akzentfarbe
  Labels bold/Akzent, Fließtext 333333 · Handlung-Leiste · Akt-Überschriften + Bullets
  Pro Oper eine eigene Akzentfarbe (Hex ohne #).
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================== OPER-DATENBLOCK ==============================
OPER = {
    "datei":    "Barbier_von_Sevilla_Infoseite",
    "titel":    "Der Barbier von Sevilla",
    "untertitel": "Opera buffa in zwei Aufzügen  ·  Gioachino Rossini (1816)",
    "akzent":   "A8361E",   # warmer Sevilla-Terrakotta (Hex ohne #). Pro Oper eigene Farbe.
    "info": [
        ("Komponist:", "Gioachino Rossini (1792–1868)"),
        ("Originaltitel:", "Il barbiere di Siviglia"),
        ("Libretto:", "Cesare Sterbini (nach Beaumarchais)"),
        ("Uraufführung:", "20. Februar 1816, Rom"),
        ("Spieldauer:", "ca. 2,5 Stunden"),
        ("Gattung:", "Komische Oper / Opera buffa"),
    ],
    "personen": [
        ("Graf Almaviva (Tenor):", "Verliebt in Rosina — wirbt als armer Student „Lindoro“ um sie"),
        ("Figaro (Bariton):", "Barbier und Faktotum der Stadt — der gewitzte Strippenzieher"),
        ("Rosina (Mezzosopran):", "Bartolos Mündel — klug, eigensinnig, verliebt"),
        ("Doktor Bartolo (Bass):", "Rosinas Vormund — will sie selbst heiraten (Antagonist)"),
        ("Don Basilio (Bass):", "Musiklehrer und Intrigant — Meister der Verleumdung"),
        ("Berta (Sopran):", "Bartolos resolute Haushälterin"),
    ],
    "akte": [
        ("1. Aufzug", [
            "Vor Rosinas Fenster bringt Graf Almaviva ein Ständchen. Er will als mittelloser Student „Lindoro“ geliebt werden — nicht um seines Titels willen.",
            "Der Barbier Figaro, Tausendsassa der Stadt, bietet dem Grafen voller Selbstbewusstsein seine Dienste an (Auftrittsarie „Largo al factotum“).",
            "Rosina lebt streng bewacht im Haus des alten Doktor Bartolo, der sie selbst heiraten will. Heimlich verliebt sie sich in „Lindoro“ („Una voce poco fa“).",
            "Verkleidet als betrunkener Soldat verschafft sich Almaviva Zutritt — es endet in turbulentem Chaos mit der herbeigerufenen Wache.",
        ]),
        ("2. Aufzug", [
            "Almaviva kehrt als Musiklehrer „Don Alonso“ verkleidet zurück, angeblich Vertretung des erkrankten Basilio, um Rosina Unterricht zu geben.",
            "Während Figaro den Doktor zur Ablenkung rasiert, schmieden die Liebenden den Fluchtplan. Der echte Basilio platzt herein — und wird bestochen.",
            "Bartolo schöpft Verdacht und jagt alle hinaus. Basilio sät derweil das Gift der Verleumdung („La calunnia è un venticello“).",
            "In der Gewitternacht dringen Almaviva und Figaro ein. Bartolos bestellter Notar traut stattdessen die Liebenden. Bartolo fügt sich — alles endet heiter.",
        ]),
    ],
    "musik": [
        ("Ouvertüre:", "Sprühender Klassiker — eines der bekanntesten Opernvorspiele überhaupt"),
        ("„Largo al factotum“:", "Figaros rasante Auftrittsarie („Figaro hier, Figaro da“)"),
        ("„Una voce poco fa“:", "Rosinas glanzvolle Belcanto-Arie (1. Akt)"),
        ("„La calunnia“:", "Basilios berühmte Verleumdungsarie (2. Akt)"),
        ("„All’idea di quel metallo“:", "Funkelndes Duett Almaviva–Figaro"),
    ],
    "bedeutung": [
        "Der Barbier von Sevilla gilt als eine der vollkommensten komischen Opern überhaupt. Rossini schrieb das Werk mit nur 23 Jahren in knapp drei Wochen — sprudelnd vor Witz, Tempo und Belcanto-Virtuosität.",
        "Die Uraufführung 1816 in Rom geriet zum Fiasko, doch schon Tage später begann ein beispielloser Welterfolg. Die Handlung ist die Vorgeschichte zu Mozarts „Figaros Hochzeit“.",
    ],
}
# ===========================================================================

FONT = "Cambria"
GRAY = RGBColor(0x77, 0x77, 0x77)
BODY = RGBColor(0x33, 0x33, 0x33)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def build(oper):
    accent = oper["akzent"].upper()
    acc_rgb = RGBColor(int(accent[0:2], 16), int(accent[2:4], 16), int(accent[4:6], 16))

    d = Document()
    sec = d.sections[0]
    sec.page_width, sec.page_height = Cm(21.59), Cm(27.94)
    sec.left_margin = sec.right_margin = Cm(2)
    sec.top_margin = sec.bottom_margin = Cm(0.75)
    st = d.styles['Normal']
    st.font.name = FONT
    st.font.size = Pt(10)
    st._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

    def run_fmt(r, size=10, bold=False, color=BODY, fill=None):
        r.font.name = FONT
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.color.rgb = color
        if fill:
            rpr = r._r.get_or_add_rPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill)
            rpr.append(shd)

    # Titel + Untertitel
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run_fmt(p.add_run(oper["titel"]), 26, True, acc_rgb)
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run_fmt(p.add_run(oper["untertitel"]), 10, False, GRAY)

    def style_table(t):
        t.style = 'Table Grid'
        borders = OxmlElement('w:tblBorders')
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            e = OxmlElement('w:' + edge); e.set(qn('w:val'), 'none')
            borders.append(e)
        t._tbl.tblPr.append(borders)
        t.columns[0].width = t.columns[1].width = Cm(8.79)
        for c in t.rows[0].cells:
            c.width = Cm(8.79)

    def header_bar(para, text):
        para.paragraph_format.space_after = Pt(3)
        run_fmt(para.add_run("  " + text + "  "), 10, True, WHITE, fill=accent)

    def fill_cell_rows(cell, header, rows):
        cell.paragraphs[0].text = ''
        header_bar(cell.paragraphs[0], header)
        for label, value in rows:
            bp = cell.add_paragraph()
            bp.paragraph_format.space_after = Pt(1)
            bp.paragraph_format.line_spacing = 1.0
            if label:
                run_fmt(bp.add_run(label + " "), 10, True, acc_rgb)
            run_fmt(bp.add_run(value), 10, False, BODY)

    def fill_cell_text(cell, header, paras):
        cell.paragraphs[0].text = ''
        header_bar(cell.paragraphs[0], header)
        for para in paras:
            bp = cell.add_paragraph()
            bp.paragraph_format.space_after = Pt(4)
            bp.paragraph_format.line_spacing = 1.0
            run_fmt(bp.add_run(para), 10, False, BODY)

    # Box 1: Infos | Personen
    tA = d.add_table(rows=1, cols=2); style_table(tA)
    fill_cell_rows(tA.rows[0].cells[0], "Allgemeine Informationen", oper["info"])
    fill_cell_rows(tA.rows[0].cells[1], "Personen", oper["personen"])

    d.add_paragraph().paragraph_format.space_after = Pt(2)

    # Handlung-Leiste
    hb = d.add_paragraph()
    hb.paragraph_format.space_before = Pt(4); hb.paragraph_format.space_after = Pt(4)
    run_fmt(hb.add_run("  Handlung  "), 14, True, WHITE, fill=accent)

    for titel, bullets in oper["akte"]:
        ap = d.add_paragraph()
        ap.paragraph_format.space_before = Pt(4); ap.paragraph_format.space_after = Pt(2)
        run_fmt(ap.add_run(titel), 11, True, acc_rgb)
        for b in bullets:
            bp = d.add_paragraph(style='List Bullet')
            bp.paragraph_format.space_after = Pt(2); bp.paragraph_format.line_spacing = 1.0
            run_fmt(bp.add_run(b), 10, False, BODY)

    d.add_paragraph().paragraph_format.space_after = Pt(2)

    # Box 2: Musik | Bedeutung
    tB = d.add_table(rows=1, cols=2); style_table(tB)
    fill_cell_rows(tB.rows[0].cells[0], "Berühmte Musiknummern", oper["musik"])
    fill_cell_text(tB.rows[0].cells[1], "Bedeutung", oper["bedeutung"])

    return d


if __name__ == "__main__":
    import shutil
    d = build(OPER)
    freizeit = f"/mnt/d/OneDrive/Dokumente/Allgemeines/Freizeit/{OPER['datei']}.docx"
    desktop = f"/mnt/d/OneDrive/Desktop/{OPER['datei']}.docx"
    d.save(freizeit)
    shutil.copy(freizeit, desktop)
    print("Gespeichert:", freizeit)
    print("Auf Desktop:", desktop)
