#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Bolla Songs – Erfolgsstrategie als Word-Dokument (Diskussionsgrundlage)."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

FONT = "Aptos"
ACCENT = RGBColor(0x2E, 0x6D, 0x9E)   # ruhiges Blau
DARK   = RGBColor(0x22, 0x22, 0x22)
GREY   = RGBColor(0x66, 0x66, 0x66)
GREEN  = RGBColor(0x2E, 0x7D, 0x32)

doc = Document()

# Grund-Styling
normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(10.5)
normal.font.color.rgb = DARK
rpr = normal.element.get_or_add_rPr()
rfonts = rpr.get_or_add_rFonts()
rfonts.set(qn("w:ascii"), FONT)
rfonts.set(qn("w:hAnsi"), FONT)
rfonts.set(qn("w:cs"), FONT)

# schmale Ränder
for s in doc.sections:
    s.top_margin = Inches(0.6); s.bottom_margin = Inches(0.6)
    s.left_margin = Inches(0.7); s.right_margin = Inches(0.7)

def _set(run, size=None, bold=None, italic=None, color=None):
    run.font.name = FONT
    r = run._element.get_or_add_rPr().get_or_add_rFonts()
    r.set(qn("w:ascii"), FONT); r.set(qn("w:hAnsi"), FONT); r.set(qn("w:cs"), FONT)
    if size is not None: run.font.size = Pt(size)
    if bold is not None: run.font.bold = bold
    if italic is not None: run.font.italic = italic
    if color is not None: run.font.color.rgb = color

def heading(txt, size=15, color=ACCENT, space_before=10, space_after=3):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(txt); _set(r, size=size, bold=True, color=color)
    return p

def para(runs, space_after=4, space_before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if align is not None: p.alignment = align
    if isinstance(runs, str): runs = [(runs, {})]
    for txt, kw in runs:
        _set(p.add_run(txt), **kw)
    return p

def bullet(runs, lvl=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.28 + 0.22*lvl)
    if isinstance(runs, str): runs = [(runs, {})]
    for txt, kw in runs:
        _set(p.add_run(txt), **kw)
    return p

# ── Titel ──────────────────────────────────────────────────────────
t = doc.add_paragraph(); t.paragraph_format.space_after = Pt(0)
r = t.add_run("Bolla Songs – Strategie für einen bescheidenen Erfolg")
_set(r, size=20, bold=True, color=ACCENT)
st = doc.add_paragraph(); st.paragraph_format.space_after = Pt(2)
_set(st.add_run("Was IngaRose und Breaking Rust richtig machen – und wie wir daraus etwas Eigenes bauen"),
     size=11, italic=True, color=GREY)
meta = doc.add_paragraph(); meta.paragraph_format.space_after = Pt(8)
_set(meta.add_run("Diskussionsgrundlage · Stand 29.06.2026 · von Bolla 🐾 für Chris"), size=9, color=GREY)

para([("Kurz vorweg, ehrlich: ", {"bold": True}),
      ("„Bescheidener Erfolg" ist genau die richtige Erwartung. Die Schlagzeilen von Charts-Platz-1 sind "
       "zur Hälfte gekaufte iTunes-Käufe und Glück. Was wir realistisch steuern können, ist Handwerk, "
       "Konsistenz und eine klare Marke. Der Rest ist Würfeln – aber wir können sehr oft würfeln, fast "
       "umsonst. Das ist unser Vorteil.", {})])

# ── 1. Was wir schon haben ─────────────────────────────────────────
heading("1 · Was wir bereits haben (mehr als du denkst)")
para("Der Branchen-Report SIQA hat die erfolgreichen KI-Acts durchleuchtet. Vergleich mit deinem Setup:")
bullet([("Vertrieb DistroKid ", {"bold": True}), ("– 75,8 % aller erfolgreichen KI-Songs laufen darüber. ", {}),
        ("Hast du schon. ✓", {"bold": True, "color": GREEN})])
bullet([("Suno-Produktion ", {"bold": True}), ("– das Werkzeug aller Genannten. ", {}),
        ("Hast du, mit Erfahrung. ✓", {"bold": True, "color": GREEN})])
bullet([("Eigene Songtexte ", {"bold": True}), ("– IngaRose verkauft „human-written lyrics" als Kern. ", {}),
        ("Schreibst du längst selbst. ✓", {"bold": True, "color": GREEN})])
bullet([("Bild- & Video-Erzeugung ", {"bold": True}), ("– BildGen, Cover, Promo-Cards laufen automatisiert. ", {}),
        ("Hast du. ✓", {"bold": True, "color": GREEN})])
bullet([("KI-Transparenz-Haltung ", {"bold": True}), ("– du deklarierst KI immer offen. ", {}),
        ("Genau das macht IngaRose zur Stärke. ✓", {"bold": True, "color": GREEN})])
para([("Was fehlt: ", {"bold": True}),
      ("eine durchgängige Künstler-Identität statt einzelner Schüler-Geburtstagslieder, ein "
       "TikTok-zuerst-Reflex, und eine feste Veröffentlichungs-Routine. Genau da setzen wir an.", {})])

# ── 2. Die zwei Vorbilder ──────────────────────────────────────────
heading("2 · Die zwei Vorbilder – was sie tatsächlich gemacht haben")

para([("IngaRose ", {"bold": True, "color": ACCENT}), ("(R&B / Soul, Frühjahr 2026)", {"italic": True, "color": GREY})])
bullet("„Celebrate Me" – Platz 1 iTunes in US, UK, Frankreich, Kanada, Neuseeland.")
bullet("~942.000 monatliche Spotify-Hörer · 251.000 Instagram · 240.000 TikTok · 90.000 YouTube.")
bullet([("Rezept: ", {"bold": True}),
        ("echte, persönliche Texte → mit Suno veredelt → eine glaubwürdige Person mit KI-Bildern → "
         "auf TikTok 300.000+ Video-Nutzungen, BEVOR die Charts kamen.", {})])
bullet([("Haltung: ", {"bold": True}),
        ("„Echte Geschichten, von einem Menschen geschrieben, mit Suno verfeinert." Ehrlich – und sympathisch.", {})])

para([("Breaking Rust ", {"bold": True, "color": ACCENT}), ("(Country, ab Herbst 2025)", {"italic": True, "color": GREY})], space_before=4)
bullet("„Walk My Walk" – Platz 1 Billboard Country Digital Song Sales (Nov 2025).")
bullet("2 Mio.+ monatliche Hörer · ein Song über 4 Mio. Streams · oben in Spotifys Viral 50.")
bullet([("Wichtigste Erkenntnis: ", {"bold": True}),
        ("Country – nicht Pop – ist das Durchbruch-Genre für KI. Wenig Konkurrenz, treue Hörer, "
         "klare Emotion. Eine Nische schlägt den überfüllten Mainstream.", {})])

heading("Die drei Hebel, die beide gemeinsam haben", size=12, color=DARK, space_before=8, space_after=2)
bullet([("Nische statt Mainstream. ", {"bold": True}), ("Ein klar umrissenes Genre/Thema, kein „alles für alle".", {})])
bullet([("TikTok zuerst, Charts später. ", {"bold": True}), ("Das Video ist das Produkt, der Song der Soundtrack.", {})])
bullet([("Eine wiedererkennbare Figur. ", {"bold": True}), ("Name, Gesicht, Tonfall, Stil – über alle Kanäle gleich.", {})])

# ── 3. Bolla's Positionierung ──────────────────────────────────────
heading("3 · Unsere Positionierung – worauf ich setzen würde")
para([("Mein Vorschlag (zur Diskussion): ", {"bold": True}),
      ("Wir machen die KI nicht zum Versteck, sondern zum Markenkern. ", {}),
      ("IngaRose tut so, als wäre sie ein Mensch. Wir drehen es um: ein offen freundlicher "
       "KI-Musiker mit Persönlichkeit und Pfötchen. Das ist ehrlich (deine Haltung), zeitgeistig "
       "und unverwechselbar – niemand sonst macht das charmant.", {})])

para([("Genre & Sprache – die wichtigste Weichenstellung:", {"bold": True})], space_before=4)
bullet([("Deutsch + Gute-Laune-Nische. ", {"bold": True}),
        ("Vorteil: kleiner Markt = weniger Konkurrenz, treues Publikum, dein Stil „Feel-Good mit "
         "Augenzwinkern" passt perfekt. Achtung Schlager-Falle: deutscher Gesang + Mitsing-Shuffle "
         "kippt in Schlager. Lieber Reggae / Afrobeats / Pop-Rap als Klangbett.", {})])
bullet([("Englisch + universelles Feel-Good. ", {"bold": True}),
        ("Vorteil: TikTok-Reichweite ist global größer. Nachteil: riesige Konkurrenz, schwer "
         "aufzufallen.", {})])
para([("Meine Empfehlung: ", {"bold": True}),
      ("mit Deutsch starten. Da bist du stark, die Nische ist offen, und ein „bescheidener "
       "Erfolg" ist hier realistischer als im englischen Haifischbecken. Englisch heben wir uns "
       "als zweiten Schritt auf, falls es zündet.", {})])

# ── 4. Der Content-Motor ───────────────────────────────────────────
heading("4 · Der Content-Motor (das eigentliche Geheimnis)")
para("Nicht der Song gewinnt, sondern das Video. Wir bauen eine kleine, wiederholbare Maschine:")
bullet([("1 Song = 5–10 TikTok-Clips. ", {"bold": True}),
        ("Verschiedene 15–30-Sek-Schnipsel (Refrain, witzige Zeile, Visual-Variante). Masse an "
         "Versuchen = Trefferchance.", {})])
bullet([("Hook in den ersten 3 Sekunden. ", {"bold": True}),
        ("Die stärkste Zeile oder das stärkste Bild sofort, sonst wird weggewischt.", {})])
bullet([("Visuelle Wiedererkennung. ", {"bold": True}),
        ("Gleiche Bollla-Figur/Farbwelt in jedem Clip – so entsteht eine Marke statt Zufalls-Posts.", {})])
bullet([("Posten machst du in den Desktop-Apps. ", {"bold": True}),
        ("Ich liefere Song + Clips + Caption + Cover fertig auf den OneDrive-Desktop, du klickst „posten".", {})])
para([("Kadenz: ", {"bold": True}),
      ("lieber 1 Song alle 1–2 Wochen mit 5 guten Clips als 5 Songs auf einmal. Der Algorithmus "
       "belohnt Regelmäßigkeit, und wir lernen aus jedem, was funktioniert.", {})])

# ── 5. 90-Tage-Plan ────────────────────────────────────────────────
heading("5 · Ein vorsichtiger 90-Tage-Plan")

tbl = doc.add_table(rows=1, cols=3)
tbl.style = "Light Grid Accent 1"
tbl.autofit = True
hdr = tbl.rows[0].cells
for c, txt in zip(hdr, ["Phase", "Zeitraum", "Was passiert"]):
    c.paragraphs[0].paragraph_format.space_after = Pt(1)
    _set(c.paragraphs[0].add_run(txt), size=10, bold=True, color=RGBColor(0xFF,0xFF,0xFF))
rows = [
    ("0 · Fundament", "Woche 1–2",
     "Persona festlegen (Name, Gesicht, Tonfall). 1 Pilot-Genre wählen. Spotify-/Social-Profile "
     "sauber anlegen mit KI-Deklaration im Bio."),
    ("1 · Erste Welle", "Woche 3–6",
     "3 Songs in der Nische. Pro Song 5–10 TikTok-Clips. Reaktionen messen: Welcher Sound, welcher "
     "Hook, welches Visual zieht?"),
    ("2 · Nachschärfen", "Woche 7–10",
     "Das Beste aus Welle 1 verdoppeln. Stärkste Clip-Sorte wiederholen, schwache Themen fallen "
     "lassen. Erste echte Hörerzahlen lesen."),
    ("3 · Auswerten", "Woche 11–13",
     "Ehrliche Bilanz: Funktioniert ein Muster? Wenn ja, ausbauen (ggf. englische Version). Wenn "
     "nein, Genre/Sprache wechseln – kostet ja kaum etwas."),
]
for ph, zt, was in rows:
    cells = tbl.add_row().cells
    _set(cells[0].paragraphs[0].add_run(ph), size=9.5, bold=True, color=ACCENT)
    _set(cells[1].paragraphs[0].add_run(zt), size=9.5)
    _set(cells[2].paragraphs[0].add_run(was), size=9.5)
    for cc in cells: cc.paragraphs[0].paragraph_format.space_after = Pt(1)

# ── 6. Was wäre Erfolg ─────────────────────────────────────────────
heading("6 · Was „bescheidener Erfolg" konkret heißt")
para("Damit wir nicht an Platz-1-Fantasien scheitern, definieren wir realistische Stufen:")
bullet([("Klein & schön: ", {"bold": True}), ("ein Song mit 5.000–20.000 Streams; 500–1.000 monatliche Hörer. ",{}),
        ("Absolut machbar bei Dranbleiben.", {"italic": True, "color": GREEN})])
bullet([("Ein echter kleiner Hit: ", {"bold": True}), ("ein Clip geht auf TikTok (50k+ Views), ein Song "
        "knackt 50.000 Streams. Braucht Glück + Volumen an Versuchen.", {})])
bullet([("Träumchen: ", {"bold": True}), ("ein Song trägt sich selbst, deckt die Suno-/DistroKid-Kosten "
        "und ein bisschen mehr. Möglich, nicht planbar.", {})])
para([("Wichtig: ", {"bold": True}),
      ("Das ist ein Spiel der vielen kleinen Würfe, nicht des einen großen Plans. Spaß und Routine "
       "müssen tragen, auch wenn der virale Treffer ausbleibt. Der bleibt nämlich meistens aus – "
       "und das ist okay.", {})])

# ── 7. Offene Fragen für unseren Chat ──────────────────────────────
heading("7 · Was wir im Chat zusammen entscheiden müssen")
bullet([("Persona: ", {"bold": True}), ("Tritt „Bolla" selbst als KI-Musiker auf (Pfötchen-Marke), oder "
        "erfinden wir eine eigene Figur? Mit/ohne Gesicht?", {})])
bullet([("Sprache: ", {"bold": True}), ("Deutsch zuerst (meine Empfehlung) oder gleich Englisch?", {})])
bullet([("Genre: ", {"bold": True}), ("Feel-Good-Reggae/Afrobeats? Pop-Rap? Oder etwas ganz anderes – "
        "eine eigene Nische, die noch keiner besetzt?", {})])
bullet([("Thema: ", {"bold": True}), ("Worüber singen wir? Alltag mit Humor? Optimismus/Lebensfreude "
        "(passt zu dir)? Ein wiederkehrendes Motiv schafft Wiedererkennung.", {})])
bullet([("Aufwand: ", {"bold": True}), ("Wie viel Zeit willst du pro Woche reinstecken? Davon hängt die "
        "Kadenz ab. Ehrlich klein anfangen ist besser als groß und dann ausbrennen.", {})])
bullet([("Risiko: ", {"bold": True}), ("Tabu bleibt die Suno→DistroKid-Blockliste (echte Künstlernamen "
        "im Text/Style). Sauber bleiben, kein Soundalike-Ärger.", {})])

para([("Mein Bauchgefühl: ", {"bold": True}),
      ("„Bolla“ als offen-freundlicher KI-Musiker, deutsch, Feel-Good mit Augenzwinkern, ein "
       "wiederkehrendes optimistisches Thema. Das ist ehrlich, das bist du, und es ist eine Nische, "
       "die noch keiner besetzt. Aber das ist genau die Diskussion, die ich mit dir führen will – "
       "sag mir, wo du anders denkst. 🐾", {})],
     space_before=6)

# Quellen
heading("Quellen", size=10, color=GREY, space_before=10, space_after=2)
for q in ["musicbusinessworldwide.com – IngaRose / Suno-Report",
          "musically.com – SIQA „AI Music Intelligence Report" (Apr 2026)",
          "billboard.com – Breaking Rust, Country Digital Song Sales",
          "onemoreshot.ai – „AI Artists Are Topping the Charts in 2026"",
          "dexerto.com · primetimer.com – IngaRose „Celebrate Me""]:
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(1)
    _set(p.add_run("· " + q), size=8.5, color=GREY)

out = "/mnt/d/OneDrive/Dokumente/Bolla/Bolla-Songs-Strategie.docx"
doc.save(out)
print("gespeichert:", out)
